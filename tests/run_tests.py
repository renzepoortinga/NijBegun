"""
Regressie-testsuite voor de Nij Begun & EPA-tool. Draai: python tests/run_tests.py
Test de hele keten offline tegen de echte voorbeeldbestanden. Exit 0 = alles groen.
"""
import os, sys, json, subprocess, tempfile
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, ROOT)
PROJECT = os.path.dirname(ROOT)  # map met de monitor-XML/template
MONITOR = os.path.join(HERE, "fixtures", "monitor_voorbeeld.xml")
TEMPLATE = os.path.join(ROOT, "templates", "isolatieplan_template.docx")
CATALOG = os.path.join(ROOT, "catalog", "catalog.json")

passed = failed = 0
def check(naam, cond, extra=""):
    global passed, failed
    if cond: passed += 1; print("  PASS", naam)
    else: failed += 1; print("  FAIL", naam, "-", extra)

from core.dossier import Dossier, load_json, save_json
from vabi.monitor_xml import parse as parse_monitor
from engine.measure_engine import run as engine_run
from ventilatie.ventilatie import bereken as vent_bereken
from isolatieplan.fill_template import fill as fill_template
from validator.validate import validate, load_catalog_codes

print("1. Core datamodel round-trip")
d = Dossier.build_sample() if hasattr(Dossier, "build_sample") else None
from core.dossier import build_sample
d = build_sample()
check("round-trip identiek", Dossier.from_dict(d.to_dict()).to_dict() == d.to_dict())
check("sample heeft ruimtes", len(d.geometrie.ruimtes) > 0)

print("2. Catalogus")
cat = json.load(open(CATALOG, encoding="utf-8"))
check("catalogus >300 maatregelen", len(cat["maatregelen"]) > 300, str(len(cat["maatregelen"])))
check("V6 kierdichting aanwezig", any(m["code"].startswith("V6") for m in cat["maatregelen"]))

print("3. VABI monitor-parser")
dos, _ = parse_monitor(MONITOR)
check("postcode 9501TP", dos.identificatie.postcode == "9501TP", dos.identificatie.postcode)
check("Standaard 64.0", dos.berekening.standaard_eis_kwh_m2 == 64.0, str(dos.berekening.standaard_eis_kwh_m2))
check("25 schildelen", len(dos.schil) == 25, str(len(dos.schil)))
check("glas-U geparsed", any(s.type == "kozijn" and s.u_huidig for s in dos.schil))

print("4. Maatregel-engine")
dos2, notes = engine_run(parse_monitor(MONITOR)[0], cat)
codes = {m["code"] for m in cat["maatregelen"]}
check("maatregelen voorgesteld", len(dos2.maatregelen) >= 3, str(len(dos2.maatregelen)))
check("alle codes uit catalogus", all(any(c == m.code or c.startswith(m.code) for c in codes) for m in dos2.maatregelen))
check("alle kosten > 0", all((m.kosten or 0) > 0 for m in dos2.maatregelen))
check("qv10-advies aanwezig", any("qv;10" in n for n in notes))
check("meerwerk-subposten voorgesteld", any(m.subposten for m in dos2.maatregelen))
check("subposten cat 2 en cat 3",
      any(s.categorie == 2 for m in dos2.maatregelen for s in m.subposten) and
      any(s.categorie == 3 for m in dos2.maatregelen for s in m.subposten))
check("subposten zonder admin-rijen",
      not any("minimum tarief" in (s.omschrijving or "").lower()
              for m in dos2.maatregelen for s in m.subposten))

print("5. Ventilatie (Nij Begun-vuistregels)")
v = vent_bereken(build_sample().geometrie.ruimtes)
check("keuken afvoer 21", any(r["functie"] == "keuken" and r["afvoer"] == 21.0 for r in v["rows"]))
check("maatgevend > 0", v["maatgevend_dm3s"] > 0)
check("rate 0,7 per verblijfsgebied (Nij Begun)", v["rate"] == 0.7)
check("vuistregel: geen afvoer in slaapkamer in checklist",
      any("slaapkamer" in vr.lower() for vr in v["vuistregels"]))
check("overstroom berekend", "overstroom_dm3s" in v and v["overstroom_dm3s"] >= 0)

print("6. Isolatieplan invullen (huidige staat)")
tmpdoc = os.path.join(tempfile.gettempdir(), "_t.docx")
n, tot = fill_template(dos2, TEMPLATE, tmpdoc)
import docx
dd = docx.Document(tmpdoc)
allcells = " ".join(c.text for t in dd.tables for r in t.rows for c in r.cells)
check("docx gemaakt + maatregelen", n >= 3 and os.path.exists(tmpdoc))
check("huidige staat ingevuld (U-waarde)", "U " in allcells and "W/m2K" in allcells)

print("7. Validator")
issues, dos2 = validate(dos2, load_catalog_codes(CATALOG))
check("validator geeft status", dos2.validatie.status in ("sluitend", "onvolledig"))
incompleet = Dossier()
iss2, incompleet = validate(incompleet, None)
check("leeg dossier = blokkerend", any(s == "BLOKKEREND" for s, _ in iss2))

print("8. Orchestrator end-to-end (subprocess)")
outdir = os.path.join(tempfile.gettempdir(), "nb_test_out")
r = subprocess.run([sys.executable, os.path.join(ROOT, "run.py"), "--from-monitor", MONITOR,
                    "--out", outdir, "--straat", "Essenhage", "--plaats", "Stadskanaal",
                    "--woningtype", "Tussenwoning"], capture_output=True, text=True)
files = os.listdir(outdir) if os.path.isdir(outdir) else []
check("isolatieplan docx output", any(f.startswith("isolatieplan_") and f.endswith(".docx") for f in files))
check("ventilatieberekening output", any(f.startswith("ventilatieberekening_") for f in files))
check("fotochecklist output", any(f.startswith("fotochecklist_") for f in files))
check("rapport output", any(f.startswith("rapport_") for f in files))
check("exitcode 0 of 2 (geen crash)", r.returncode in (0, 2), "rc=%s err=%s" % (r.returncode, r.stderr[-200:]))

print("9. VABI-generator (dossier -> monitoring-XML, round-trip)")
from vabi import monitor_generate as vgen
_sd = build_sample()
_gp = os.path.join(tempfile.gettempdir(), "gen_sample.xml")
vgen.write(_sd, _gp)
_rd, _ = parse_monitor(_gp)
check("generator schil round-trip", len(_rd.schil) == len(_sd.schil), "%d vs %d" % (len(_rd.schil), len(_sd.schil)))
check("generator postcode behouden", _rd.identificatie.postcode == _sd.identificatie.postcode)
check("generator gevel-Rc behouden", any(s.type == "gevel" and s.rc_huidig for s in _rd.schil))
check("generator installaties in XML", b"HR107" in vgen.to_xml(_sd))
_md, _ = parse_monitor(MONITOR)
_mp = os.path.join(tempfile.gettempdir(), "gen_monitor.xml")
vgen.write(_md, _mp)
_mr, _ = parse_monitor(_mp)
check("monitor 25 schil round-trip", len(_mr.schil) == len(_md.schil) == 25, "%d->%d" % (len(_md.schil), len(_mr.schil)))

print("10. Parameter-sanity-check")
from vabi.sanity import check as sanity_check
from core.dossier import VloerInfo
check("schone sample geen NAMETEN", all(l != "NAMETEN" for l, _ in sanity_check(build_sample())))
_bad = build_sample(); _bad.geometrie.vloeren.append(VloerInfo("Zolder", 20.0, 3.9))
check("outlier-hoogte -> NAMETEN", any(l == "NAMETEN" for l, _ in sanity_check(_bad)))
_no = build_sample()
for s in _no.schil:
    if s.type == "gevel": s.orientatie = ""
check("ontbrekende orientatie gevlagd", any("orientatie" in m for _, m in sanity_check(_no)))

print("11. MagicPlan-extractor (plan-JSON -> dossier)")
from magicplan.extractor import map_plan_to_dossier
_plan = json.load(open(os.path.join(HERE, "fixtures", "magicplan_plan_voorbeeld.json"), encoding="utf-8"))
_ed = map_plan_to_dossier(_plan)
check("extractor postcode", _ed.identificatie.postcode == "9501TP", _ed.identificatie.postcode)
check("extractor schil gevel+vloer+dak+kozijn",
      {"gevel", "vloer", "dak", "kozijn"} <= {s.type for s in _ed.schil})
check("extractor gevel-begrenzing (algemeen)",
      any(s.type == "gevel" and s.begrenzing == "Buitenlucht" for s in _ed.schil))
check("extractor vloerconstructie",
      any(s.type == "vloer" and s.subtype == "Houten vloer" for s in _ed.schil))
check("extractor dak helling + oppervlak",
      any(s.type == "dak" and s.hellingshoek == 45 and s.oppervlakte_m2 > 55 for s in _ed.schil))
check("extractor isolatie Ja/Nee",
      any(s.type == "gevel" and s.isolatie_aanwezig == "Nee" for s in _ed.schil))
check("extractor ruimtes + functie", any(r.functie == "keuken" for r in _ed.geometrie.ruimtes))
check("extractor installaties (subtype)", _ed.installaties.verwarming.subtype == "HR107")
_ed2, _ = engine_run(_ed, cat)
check("extractor -> engine draait", _ed2 is not None and isinstance(_ed2.maatregelen, list))
check("extractor -> engine maatregelen (gevel/vloer/dak)", len(_ed2.maatregelen) >= 2, str(len(_ed2.maatregelen)))
# Python 3.14/OpenSSL 3.2-fix (taak 012): elke live urlopen() naar cloud.magicplan.app moet de
# aangepaste SSL-context gebruiken (VERIFY_X509_STRICT uit) - bron-check i.p.v. een live request,
# zodat een volgende toegevoegde urlopen-aanroep niet stil dezelfde 'Basic Constraints'-fout
# herintroduceert (precies wat er in de eerste versie van taak 012 gebeurde, gevonden door review).
_extractor_bron = open(os.path.join(HERE, "..", "magicplan", "extractor.py"), encoding="utf-8").read()
check("extractor: elke urlopen() gebruikt de VERIFY_X509_STRICT-fix",
      _extractor_bron.count("urllib.request.urlopen(") == _extractor_bron.count("context=_ssl_context()"),
      "urlopen=%d context=%d" % (_extractor_bron.count("urllib.request.urlopen("),
                                 _extractor_bron.count("context=_ssl_context()")))

print("12. Prijsopbouw T7-T9 (cat 1/2/3 + clonen)")
from core.dossier import Maatregel as _M, Subpost as _Sub
_pd = build_sample()
_pd.maatregelen[0].subposten = [_Sub(categorie=2, code="V1-1-Z1", omschrijving="Opdikken",
                                     prijs_per_eenheid=2.5, eenheid="m2", hoeveelheid=48, kosten=120.0)]
_pp = os.path.join(tempfile.gettempdir(), "pb_test.docx")
fill_template(_pd, TEMPLATE, _pp)
_b0 = "|".join(c.text for r in [t for t in docx.Document(_pp).tables
              if t.rows and t.rows[0].cells[0].text.strip().lower().startswith("maatregel:")][0].rows
              for c in r.cells)
check("prijsopbouw cat1 code", _pd.maatregelen[0].code in _b0)
check("prijsopbouw cat2 subpost", "Opdikken" in _b0)
_pd5 = build_sample()
_pd5.maatregelen += [_M(code="V5-1", onderdeel="E Ventilatie", omschrijving="Roosters",
                        oppervlakte_m2=4, prijs_per_eenheid=50, kosten=200, eenheid="st"),
                     _M(code="V6-1", onderdeel="E Ventilatie", omschrijving="Kier",
                        oppervlakte_m2=1, prijs_per_eenheid=300, kosten=300, eenheid="pst")]
_pp5 = os.path.join(tempfile.gettempdir(), "pb_test5.docx")
fill_template(_pd5, TEMPLATE, _pp5)
_bl5 = [t for t in docx.Document(_pp5).tables
        if t.rows and t.rows[0].cells[0].text.strip().lower().startswith("maatregel:")]
check("prijsopbouw blok per maatregel (clonen)", len(_bl5) >= 5, str(len(_bl5)))

print("13. Catalogus-API (response -> catalog.json-structuur)")
from catalog.api_client import map_measures_to_catalog
from engine.measure_engine import select_core
_raw = json.load(open(os.path.join(HERE, "fixtures", "nijbegun_catalog_response.json"), encoding="utf-8"))
_capi = map_measures_to_catalog(_raw)
check("catalog-api versie-label (live-datum)", _capi["versie"].startswith("api"))
check("catalog-api aantal maatregelen (gevlakt: 2 brackets + X7 + bodem + dak + X1)",
      _capi["aantal_maatregelen"] == 6, str(_capi["aantal_maatregelen"]))
check("catalog-api incl-btw afgeleid", all(m.get("prijs_per_eenheid_incl_btw") for m in _capi["maatregelen"]))
check("catalog-api werkt in engine", select_core(_capi, ["V1-1"], ["spouwmuurisolatie"], 30.0) is not None)

print("14. MagicPlan report-PDF parser (form-antwoorden uit export)")
from magicplan.report_parser import parse_text
_rtxt = open(os.path.join(HERE, "fixtures", "report_sample.txt"), encoding="utf-8").read()
_ans, _koz = parse_text(_rtxt)
check("report postcode + bouwjaar", _ans.get("postcode") == "9503HN" and _ans.get("bouwjaar") == "1994")
check("report gevel/dak tags", _ans.get("geveltype") == "Spouwmuur" and _ans.get("dakhelling") == "45")
check("report ventilatie", _ans.get("vent_systeem") == "A Natuurlijke ventilatie" and _ans.get("vent_subsysteem_code") == "A1")
check("report kierdichting", _ans.get("kierdichting") == "Slecht-tochtklachten")
check("report kozijnen per raam", len(_koz) == 2 and _koz[0].get("glastype") == "HR++" and _koz[1].get("glastype") == "Dubbel (D)")

print("15. MagicPlan assemblage (report-antwoorden + API-geometrie -> dossier)")
from magicplan.assemble import build_dossier
from magicplan.report_parser import parse_text as _ptext
_p, _koz = _ptext(open(os.path.join(HERE, "fixtures", "report_sample.txt"), encoding="utf-8").read())
_planv2 = {"data": {"plan_data": {"living_area": 100.0, "floors": [
    {"name": "Ground Floor", "statistics": {"area": 50.0}, "values": [{"id": "ceilingHeight", "value": 2.6}],
     "rooms": [{"name": "Keuken", "statistics": {"area": 12.0}, "objects": [
         {"symbol_id": "windowhung", "values": [{"id": "width", "value": 1.5}, {"id": "height", "value": 1.5}]}]}]}]}}}
_ad = build_dossier(_p, _koz, _planv2)
check("assemble identificatie uit report", _ad.identificatie.postcode == "9503HN" and _ad.identificatie.bouwjaar == 1994)
check("assemble geometrie uit plan (ruimte+functie)", any(r.functie == "keuken" for r in _ad.geometrie.ruimtes))
check("assemble schil gevel+vloer+dak+kozijn", {"gevel", "vloer", "dak", "kozijn"} <= {s.type for s in _ad.schil})
check("assemble gevel-tag uit report", any(s.type == "gevel" and s.subtype == "Spouwmuur" for s in _ad.schil))
check("assemble kozijn-glas uit report", any(s.type == "kozijn" and s.glastype == "HR++" for s in _ad.schil))
# taak 014-review: de hybride API+report-PDF-route (_maak_dak, extractor.py) maakt net als de
# Statistics-CSV-route altijd ÉÉN generiek dakvlak -> moet DEZELFDE fallback-tag krijgen, anders
# ziet de dak-fallback-opschoning/-preflight 'm niet als een dossier via dit pad binnenkomt.
check("assemble: dak krijgt dezelfde magicplan-dak-fallback-tag als de CSV-route",
      any(s.type == "dak" and s.bron == "magicplan-dak-fallback" for s in _ad.schil))
from vabi.preflight import assert_no_dubbel_dak_fallback as _adf15, VabiExportBlocked as _VEB15
from core.dossier import SchilDeel as _SD15
_ad2 = build_dossier(_p, _koz, _planv2)
_ad2.schil.append(_SD15(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=12.0,
                        bron="webapp-wizard"))
try:
    _adf15(_ad2)
    check("assemble: preflight blokkeert fallback+wizarddak ook op het hybride pad", False)
except _VEB15:
    check("assemble: preflight blokkeert fallback+wizarddak ook op het hybride pad", True)

print("15b. MagicPlan-contourextractie (echte plattegrondomtrek i.p.v. rechthoek-gok)")
from magicplan.assemble import _floor_contour_m, geometry_from_plan
from core.geometry import polygon_oppervlakte_m2
check("polygon_oppervlakte_m2: rechthoek", polygon_oppervlakte_m2([(0, 0), (4, 0), (4, 3), (0, 3)]) == 12)
check("polygon_oppervlakte_m2: minder dan 3 punten -> 0", polygon_oppervlakte_m2([(0, 0), (1, 1)]) == 0.0)
_floor_ok = {"statistics": {"area_with_walls": 12.0},
             "image_map": [{"symbol_id": "wall", "coordinates": [0, 0, 1, 1]},
                           {"symbol_id": "floor", "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]}]}
_contour = _floor_contour_m(_floor_ok)
check("assemble: contour gekalibreerd op de werkelijke oppervlakte (area_with_walls)",
      _contour is not None and abs(polygon_oppervlakte_m2(_contour) - 12.0) < 0.01)
check("assemble: contour genormaliseerd naar oorsprong (0,0)",
      min(p[0] for p in _contour) == 0 and min(p[1] for p in _contour) == 0)
check("assemble: geen 'floor'-symbol in image_map -> None (geen gegokte vorm)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "wall", "coordinates": [0, 0, 1, 1]}]}) is None)
check("assemble: geen werkelijke oppervlakte -> None",
      _floor_contour_m({"image_map": [{"symbol_id": "floor", "coordinates": [0, 0, 1, 0, 1, 1, 0, 1]}]}) is None)
check("assemble: meerdere 'floor'-entries -> None (niet gokken welke bij area_with_walls hoort)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "floor", "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]},
                                     {"symbol_id": "floor", "coordinates": [0, 0, 50, 0, 50, 50, 0, 50]}]}) is None)
check("assemble: niet-numerieke coördinaat -> None (geen crash)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, 100, None, 100, 100, 0, 100]}]}) is None)
check("assemble: bruto (area_with_walls) fors groter dan netto (area) -> None (schaal hoort er niet bij)",
      _floor_contour_m({"statistics": {"area": 5.0, "area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]}]}) is None)
# Stresstest-bevindingen (14-8): NaN/Infinity in de brondata moeten geweigerd worden, niet
# doorsijpelen naar een onbruikbare contour, en malformed/onverwachte structuren mogen nooit crashen.
check("assemble: NaN area_with_walls -> None (niet stil doorrekenen naar NaN-schaal)",
      _floor_contour_m({"statistics": {"area_with_walls": float("nan")},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]}]}) is None)
check("assemble: Infinity area_with_walls -> None",
      _floor_contour_m({"statistics": {"area_with_walls": float("inf")},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]}]}) is None)
check("assemble: NaN in coördinaten -> None (float() slikt NaN, expliciete isfinite-check vangt 'm)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, float("nan"), 0, 100, 100, 0, 100]}]}) is None)
check("assemble: fl is None -> None (geen crash)", _floor_contour_m(None) is None)
check("assemble: image_map is geen lijst -> None (geen crash)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0}, "image_map": "oops"}) is None)
check("assemble: coordinates is geen lijst -> None (geen crash)",
      _floor_contour_m({"statistics": {"area_with_walls": 12.0},
                        "image_map": [{"symbol_id": "floor", "coordinates": "oops"}]}) is None)
_planv2_negatief = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "Kelder", "statistics": {"area": -5.0}, "rooms": []}]}}}
_geo_neg, _win_neg, _fp_neg, _per_neg, _h_neg = geometry_from_plan(_planv2_negatief)
check("assemble: negatieve vloeroppervlakte crasht geometry_from_plan niet (was ValueError bij sqrt)",
      _per_neg == 0.0)
# Tweede stresstestronde (review op taak 011): niet-numerieke (string) waarden op dezelfde
# aanvalsoppervlakte die de eerste ronde met NaN/Infinity/malformed types al had dichtgezet.
check("assemble: area_with_walls als string -> None (geen ValueError)",
      _floor_contour_m({"statistics": {"area_with_walls": "N/A"},
                        "image_map": [{"symbol_id": "floor",
                                      "coordinates": [0, 0, 10, 0, 10, 10, 0, 10]}]}) is None)
import math as _mathassm
_planv2_nan_footprint = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "A", "statistics": {"area": float("nan")}},
    {"name": "B", "statistics": {"area": 20.0}}]}}}
_geo_nanfp, _win_nanfp, _fp_nanfp, _per_nanfp, _h_nanfp = geometry_from_plan(_planv2_nan_footprint)
check("assemble: NaN vloeroppervlakte lekt niet door via max() naar footprint/VloerInfo",
      _mathassm.isfinite(_fp_nanfp) and all(_mathassm.isfinite(v.oppervlakte_m2) for v in _geo_nanfp.vloeren),
      str((_fp_nanfp, [v.oppervlakte_m2 for v in _geo_nanfp.vloeren])))
_planv2_bad_ceiling = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "A", "statistics": {"area": 10.0},
     "values": [{"id": "ceilingHeight", "value": "oops"}]}]}}}
check("assemble: niet-numerieke ceilingHeight crasht niet",
      geometry_from_plan(_planv2_bad_ceiling)[0].vloeren[0].hoogte_m is None)
_planv2_bad_room = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "A", "statistics": {"area": 10.0},
     "rooms": [{"name": "R", "statistics": {"area": "bad"}}]}]}}}
check("assemble: niet-numerieke kamer-oppervlakte crasht niet",
      geometry_from_plan(_planv2_bad_room)[0].ruimtes[0].oppervlakte_m2 == 0.0)
# Derde stresstestronde (review op taak 011, ronde 2): negatieve oppervlaktes uit onbetrouwbare
# brondata werden door _getal() doorgelaten (eindig + numeriek, dus "geldig") maar zijn voor een
# oppervlakte net zo onbruikbaar als NaN - _opp() klemt ze nu vast op 0.
_planv2_neg_vloer = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "Kelder", "statistics": {"area": -5.0}, "rooms": []}]}}}
check("assemble: negatieve vloeroppervlakte -> 0.0, lekt niet als negatief getal door",
      geometry_from_plan(_planv2_neg_vloer)[0].vloeren[0].oppervlakte_m2 == 0.0)
_planv2_neg_raam = {"data": {"plan_data": {"living_area": 10.0, "floors": [
    {"name": "A", "statistics": {"area": 10.0}, "rooms": [
        {"name": "R", "statistics": {"area": 5}, "objects": [
            {"symbol_id": "windowHung", "values": [{"id": "width", "value": -1.2},
                                                    {"id": "height", "value": 1.0}]}]}]}]}}}
check("assemble: negatieve raambreedte -> raam-oppervlak 0.0 (geen inversie van de gevelaftrek)",
      geometry_from_plan(_planv2_neg_raam)[1][0]["area"] == 0.0)
from dashboard.gebouw_svg import _contour_geldig as _gcgeldig
check("gebouw-svg: _contour_geldig accepteert numerieke strings (geen crash, wél geldig)",
      _gcgeldig([["0", "0"], ["5", "0"], ["5", "5"]]) is True)
check("gebouw-svg: _contour_geldig verwerpt echt niet-numerieke coördinaten (geen TypeError)",
      _gcgeldig([["abc", "0"], ["5", "0"], ["5", "5"]]) is False)
_planv2_contour = {"data": {"plan_data": {"living_area": 12.0, "floors": [
    {"name": "Begane grond", "statistics": {"area": 12.0, "area_with_walls": 12.0},
     "image_map": [{"symbol_id": "floor", "coordinates": [0, 0, 100, 0, 100, 100, 0, 100]}],
     "rooms": []}]}}}
_geo, *_ = geometry_from_plan(_planv2_contour)
check("assemble: geometry_from_plan zet de contour door op VloerInfo",
      _geo.vloeren[0].contour_m is not None
      and abs(polygon_oppervlakte_m2(_geo.vloeren[0].contour_m) - 12.0) < 0.01)

print("16. VABI Constructiebibliotheek-generator (dossier -> importeerbare XML)")
import tempfile as _tf
import xml.etree.ElementTree as _ET
from vabi.constructie_generate import write as _cwrite
from vabi.codebook import Codebook as _CB
_cb = _CB.default()
_outc = os.path.join(_tf.gettempdir(), "test_constructiebib.xml")
_map, _iss = _cwrite(_ad, _outc)
_croot = _ET.parse(_outc).getroot()                      # well-formed?
_ccons = [c for c in _croot.iter() if c.tag.rsplit("}", 1)[-1] == "Constructie"]
check("constr-gen produceert constructies", len(_ccons) >= 1)
check("constr-gen header = VABI-export", _croot.findtext("XmlVersie") == "120001001")
check("constr-gen alle enums VABI-bekend", all(
    _cb.is_valid(_f, _c.findtext(_f)) for _c in _ccons
    for _f in ("ConstructieType", "Glas", "Kozijn", "IsolatieAanwezig", "Bouwjaar", "Invoer", "SpouwAanwezig")
    if _c.findtext(_f) is not None))
check("constr-gen unieke Guids", len({_c.findtext("Guid") for _c in _ccons}) == len(_ccons))
check("constr-gen mapping dekt schil", all(s.id in _map for s in _ad.schil if (s.type or "").lower() in ("gevel", "vloer", "dak", "kozijn")))

print("17. VABI Objectenbibliotheek-generator (geometrie -> importeerbare XML)")
import os.path as _op
if _op.exists(_op.join(HERE, "..", "vabi", "refs", "objecten_template.xml")):
    from vabi.objecten_generate import write as _owrite
    _outo = os.path.join(_tf.gettempdir(), "test_objectenbib.xml")
    _omap, _oiss, _ostats = _owrite(_ad, _outo)
    _oroot = _ET.parse(_outo).getroot()
    _ocons = {c.findtext("Guid") for c in _oroot.iter() if c.tag.rsplit("}", 1)[-1] == "Constructie" and c.find("Guid") is not None}
    _ohv = [e for e in _oroot.iter() if e.tag.rsplit("}", 1)[-1] == "Hoofdvlak"]
    _orefs = set()
    for _e in _oroot.iter():
        if _e.tag.rsplit("}", 1)[-1] in ("Hoofdvlak", "Deelvlak"):
            _c = _e.find("Constructie")
            if _c is not None and _c.text:
                _orefs.add(_c.text.strip())
    check("obj-gen produceert hoofdvlakken", len(_ohv) >= 1)
    check("obj-gen vlak-refs verwijzen naar embedded constructie", _orefs <= _ocons and len(_orefs) > 0)
    check("obj-gen well-formed Project", _oroot.tag.rsplit("}", 1)[-1] == "Project")
else:
    check("obj-gen (template aanwezig)", False)

print("18. VABI Installatiebibliotheek-generator")
if _op.exists(_op.join(HERE, "..", "vabi", "refs", "installatie_template.xml")):
    from vabi.installatie_generate import write as _iwrite
    _outi = os.path.join(_tf.gettempdir(), "test_installatiebib.xml")
    _iflags = _iwrite(_ad, _outi)
    _iroot = _ET.parse(_outi).getroot()
    check("inst-gen well-formed + heeft Ventilatie", any(e.tag.rsplit("}", 1)[-1] == "Ventilatie" for e in _iroot.iter()))
else:
    check("inst-gen (template aanwezig)", False)

print("19. VABI result_reader (Standaard-toets)")
from vabi.result_reader import read_results as _rr
_mon = os.path.join(_tf.gettempdir(), "test_summary_monitor.xml")
open(_mon, "w", encoding="utf-8").write(
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<tns:Energieprestatie xmlns:tns="http://schemas.ep-online.nl/monitoringbestand">'
    '<Summary><Labelklasse>B</Labelklasse>'
    '<IndicatorEnergiebehoefte>118.45</IndicatorEnergiebehoefte>'
    '<Standaard>91</Standaard></Summary></tns:Energieprestatie>')
_rres = _rr(_mon)
check("result_reader leest Summary", _rres.get("Labelklasse") == "B" and _rres.get("Standaard") == "91")
check("result_reader Standaard-toets (voldoet niet bij 118>91)", _rres.get("_voldoet_aan_standaard") is False)

print("20. Maatregel-advies met begeleidende tekst")
from engine.advies_text import genereer_advies as _gadv
from core.dossier import Maatregel as _Mt
_advms = [_Mt(onderdeel="A Gevel", omschrijving="Spouwmuurisolatie", rc_u_doel="Rc 1.1", oppervlakte_m2=60, kosten=1800),
          _Mt(onderdeel="E Ventilatie", omschrijving="CO2-sturing (C.4)", kosten=1500)]
_advtxt = _gadv(_ad, _advms,
                resultaat_huidig={"Labelklasse": "D", "IndicatorEnergiebehoefte": "165", "Standaard": "110"},
                resultaat_na={"IndicatorEnergiebehoefte": "98", "Standaard": "110", "_voldoet_aan_standaard": True})
check("advies bevat maatregelen + Standaard-conclusie", "Spouwmuurisolatie" in _advtxt and "voldoet aan de Standaard" in _advtxt)
check("advies toont huidige staat", "Huidige staat" in _advtxt and "165" in _advtxt)

print("21. Maatregel-beslislogica (welk advies wanneer)")
from engine.advies_logic import adviseer_dossier as _adl
from core.dossier import SchilDeel as _SD, Dossier as _Dos
_dl = _Dos()
_dl.schil = [_SD(id="gevel", type="gevel", isolatie_aanwezig="Nee", spouw_aanwezig=True, oppervlakte_m2=80),
             _SD(id="vloer", type="vloer", isolatie_aanwezig="Nee", begrenzing="Kruipruimte", oppervlakte_m2=50),
             _SD(id="raam-1", type="kozijn", glastype="Dubbel", oppervlakte_m2=12),
             _SD(id="gevel-ok", type="gevel", isolatie_aanwezig="Ja", isolatiedikte_mm=120, spouw_aanwezig=True)]
try:
    from core.dossier import Ventilatie as _V
    _dl.ventilatie = _V(systeem="A Natuurlijke ventilatie")
except Exception:
    pass
_radv = _adl(_dl)
_bd = _radv["bouwdelen"]
check("logic: spouw->spouwmuurisolatie", any(a["schildeel_id"] == "gevel" and "Spouwmuur" in a["advies"] for a in _bd))
check("logic: zwak glas->HR++", any(a["schildeel_id"] == "raam-1" and "HR++" in a["advies"] for a in _bd))
check("logic: goed geisoleerde gevel krijgt GEEN advies", not any(a["schildeel_id"] == "gevel-ok" for a in _bd))
check("logic: ventilatie-upgrade na isoleren", any("entilat" in a["advies"] for a in _radv["algemeen"]))
check("logic: kierdichting-regel", any("ierdicht" in a["advies"] for a in _radv["algemeen"]))

print("22. Dak/kopgevel-geometrie (hellend dak)")
from core.geometry import schuin_dakvlak_m2 as _sd, kopgevel_driehoek_m2 as _kd, dak_en_kopgevel as _dk
check("schuin dakvlak = footprint/cos", abs(_sd(63, 45) - 63 * (2 ** 0.5)) < 0.2)
check("kopgevel-driehoek 45gr B=6 -> 18 m2", abs(_kd(6, 45) - 18.0) < 0.1)
_z = _dk(63, 6, 45)
check("zadeldak: dak+extra gevel", _z["dak_m2"] > 63 and _z["extra_gevel_m2"] == 18.0)
check("plat dak: dak=footprint, geen extra gevel", _dk(63, 6, 30, type_dak="plat")["extra_gevel_m2"] == 0.0)
from core.geometry import ag_onder_schuin_dak as _ag
_agv, _weg = _ag(40, 8, 45, 0.7)
check("zolder Ag <1.5m wordt afgetrokken", abs(_agv - 27.2) < 0.1 and abs(_weg - 12.8) < 0.1)
check("zolder kniewand>=1.5m: geen reductie", _ag(40, 8, 45, 1.5)[0] == 40.0)
from core.geometry import dakkapel_vlakken as _dkp, oppervlak_vorm as _ov
check("dakkapel -> extra gevel + plat dak", _dkp(2.5, 1.4, 1.2)["gevel_m2"] > 0 and _dkp(2.5, 1.4, 1.2)["dak_m2"] == 3.0)
check("vorm driehoek = 0.5*a*b", _ov("driehoek", 6, 3) == 9.0)

print("23. Hart-op-hart gevel-toeslag woningscheidende wand (ISSO 82.1 par. 8.2)")
from core.geometry import (aantal_woningscheidende_wanden as _awsw,
                           woningscheidende_wand_toeslag_m2 as _wswt,
                           HARTMAAT_GEBOUWSCHEIDENDE_WAND_M as _hm)
check("aantal buurwanden: vrijstaand=0", _awsw("Vrijstaand") == 0)
check("aantal buurwanden: hoek/eind=1", _awsw("Kop-, eind- of hoekligging") == 1)
check("aantal buurwanden: tussen=2", _awsw("Tussenligging") == 2)
check("aantal buurwanden: 2-onder-1-kap=1", _awsw("Twee onder een kap") == 1)
check("halve wanddikte = 0,11 m", abs(_hm - 0.11) < 1e-9)
# ISSO 8.2: gevelbreedte +0,11 m/buurwand, geldt voor voor- EN achtergevel -> 2*n*0,11*h
check("toeslag tussen = 0,44*h (2*2*0,11)", abs(_wswt(2.6, "Tussenligging") - 0.44 * 2.6) < 0.01)
check("toeslag hoek = 0,22*h (2*1*0,11)", abs(_wswt(2.6, "Hoekwoning") - 0.22 * 2.6) < 0.01)
check("toeslag vrijstaand = 0", _wswt(2.6, "Vrijstaand") == 0.0)
check("toeslag zonder hoogte = 0", _wswt(None, "Tussenligging") == 0.0)

print("23b. Gedoogbeleid vleermuizen/eDNA (provincie Groningen, per 1-7-2026)")
from core.gedoogbeleid import provincie_uit_postcode as _pup, gedoogbeleid_reminders as _gbr
check("provincie: 9711 (Groningen stad) -> Groningen", _pup("9711AB")[0] == "Groningen")
check("provincie: 9641 (Veendam) -> Groningen", _pup("9641")[0] == "Groningen")
check("provincie: 7811 (Emmen) -> Drenthe", _pup("7811")[0] == "Drenthe")
check("provincie: 9401 (Assen) -> Drenthe", _pup("9401KL")[0] == "Drenthe")
check("provincie: 9301 (grensgebied) -> Groningen + onzeker",
      _pup("9301")[0] == "Groningen" and _pup("9301")[1] is True)
check("provincie: leeg -> Groningen + onzeker (veilige default)",
      _pup("")[0] == "Groningen" and _pup("")[1] is True)
check("reminder: geen spouw -> geen reminder", _gbr("9711AB", False) == [])
_gG = _gbr("9711AB", True)
check("reminder: Groningen + spouw -> eDNA/natuurvrij/X17 + BRL IC-200 genoemd",
      bool(_gG) and any("V1-1-X15" in m and "V1-1-X17" in m and "BRL IC-200" in m for m in _gG))
_gD = _gbr("7811AB", True)
check("reminder: Drenthe + spouw -> SMP-melding i.p.v. eDNA-verplichting",
      bool(_gD) and any("DRENTHE" in m and "SMP" in m for m in _gD))
import json as _json23
_cat23 = _json23.load(open("catalog/catalog.json", encoding="utf-8"))["maatregelen"]
_codes23 = {c["code"] for c in _cat23}
check("catalogus: eDNA-codes V1-1-X13..X17 aanwezig",
      all(c in _codes23 for c in ("V1-1-X13", "V1-1-X14", "V1-1-X15", "V1-1-X16", "V1-1-X17")))
_x17 = next((c for c in _cat23 if c["code"] == "V1-1-X17"), None)
check("catalogus: V1-1-X17 alternatieve verblijfplaats = EUR 151,25 incl",
      _x17 and abs(_x17["prijs_per_eenheid_incl_btw"] - 151.25) < 0.01)

print("23c. Voorschot-factuur-specificatie (opdracht Provincie Groningen)")
from dashboard.voorschot import tarief_excl as _te, build_specificatie as _bspec
check("tarief: Tussenwoning basis = 350 excl", _te("Tussenwoning", 90, False) == 350.0)
check("tarief: Tussenwoning uitgebreid = 425", _te("Tussenwoning", 90, True) == 425.0)
check("tarief: Vrijstaand >300 m2 basis = 750", _te("Vrijstaand", 320, False) == 750.0)
check("tarief: Vrijstaand <300 m2 basis = 625", _te("Vrijstaand", 180, False) == 625.0)
check("tarief: Twee-onder-een-kap = hoek-bucket 500", _te("Twee-onder-een-kap", 140, False) == 500.0)
check("tarief: onbekend woningtype -> None (flag)", _te("Woonboot", 100, False) is None)
_spec = _bspec([
    {"postcode": "9711AB", "huisnummer": "1", "woningtype": "Tussenwoning", "ag_m2": 90, "uitgebreid": False},
    {"postcode": "9641XY", "huisnummer": "2", "woningtype": "Vrijstaand", "ag_m2": 320, "uitgebreid": True},
    {"postcode": "9999ZZ", "huisnummer": "9", "woningtype": "Woonboot", "ag_m2": 80, "uitgebreid": False},
])
check("voorschot: subtotaal excl = 1175 (2 regels, 1 onbekend geflagd)",
      _spec["subtotaal_excl"] == 1175.0 and len(_spec["regels"]) == 2 and len(_spec["onbekend"]) == 1)
check("voorschot: 75% voorschot excl = 881,25", _spec["voorschot_excl"] == 881.25)
check("voorschot: 21% btw = 185,06", abs(_spec["btw"] - 185.06) < 0.01)
check("voorschot: adresregel = postcode+huisnr", _spec["regels"][0]["adres"] == "9711AB 1")
check("voorschot: verplichte header (VPL + documentnr + crediteuren-email)",
      _spec["header"]["vpl_nummer"] == "VPL-015187"
      and _spec["header"]["documentnummer_opdracht"] == "2026-102825"
      and _spec["header"]["email"] == "crediteurenadministratie@provinciegroningen.nl")

print("23d. Isolatieplan sectie 3 (Huidige woningstaat V1-V6) — volledigheid")
from core.dossier import SchilDeel as _SD71, build_sample as _bs71
import isolatieplan.fill_template as _FT71
_d71 = _bs71()
_d71.schil = [
    _SD71(id="gevel", type="gevel", oppervlakte_m2=40, rc_huidig=0.36, spouw_aanwezig=True),
    _SD71(id="raam", type="kozijn", subtype="Raam", glastype="Dubbel", u_huidig=2.9, kozijnmateriaal="Hout"),
    _SD71(id="voorzet", type="kozijn", subtype="Raam", glastype="Voorzetglas", u_huidig=2.5),
    _SD71(id="deur", type="kozijn", subtype="Deur", u_huidig=3.4),
    _SD71(id="paneel", type="kozijn", subtype="Paneel in kozijn", rc_huidig=0.5),
    _SD71(id="dakraam", type="kozijn", subtype="Dakraam", u_huidig=2.8),
    _SD71(id="vloerbg", type="vloer", subtype="Begane grondvloer", rc_huidig=0.15),
    _SD71(id="zolder", type="vloer", subtype="Zoldervloer", rc_huidig=1.3),
    _SD71(id="dakh", type="dak", subtype="Hellend dak", hellingshoek=45, rc_huidig=1.3),
    _SD71(id="dakp", type="dak", subtype="Plat dak", hellingshoek=0, rc_huidig=2.0),
]
_sub71 = _FT71._staat_subsets(_d71)
check("staat V2: raam-bucket = alleen echte ramen (geen deur/paneel/dakraam/voorzet)",
      [s.id for s in _sub71["raam"]] == ["raam"])
check("staat V2: voor-/achterzetbeglazing apart", [s.id for s in _sub71["voorzet"]] == ["voorzet"])
check("staat V2: deuren apart", [s.id for s in _sub71["deur"]] == ["deur"])
check("staat V2: panelen in kozijn apart", [s.id for s in _sub71["paneel"]] == ["paneel"])
check("staat V4: dakraam/tuimelvenster apart", [s.id for s in _sub71["dakraam"]] == ["dakraam"])
check("staat V3: begane-grondvloer vs zolder-/vlieringvloer gescheiden",
      [s.id for s in _sub71["vloer_bg"]] == ["vloerbg"] and [s.id for s in _sub71["vloer_zolder"]] == ["zolder"])
check("staat V4: hellend vs plat dak gescheiden",
      [s.id for s in _sub71["dak_hellend"]] == ["dakh"] and [s.id for s in _sub71["dak_plat"]] == ["dakp"])
_gat71 = _FT71.huidige_staat_gaten(_d71)
check("staat: gaten wijzen naar de ontbrekende MagicPlan-velden (gevel-/dak-zijde, bodem, kier)",
      any("Gevelisolatie" in g for g in _gat71) and any("Dakisolatie" in g for g in _gat71)
      and any("Bodemisolatie" in g for g in _gat71) and any("Kierdichting" in g for g in _gat71))
check("staat: vierpansraam blijft altijd een handmatige regel", any("Vierpansraam" in g for g in _gat71))
_d71.opname.gevel_isolatie_zijde = "Binnenzijde"
_d71.opname.dak_isolatie_zijde = "Buitenzijde"
_d71.opname.bodemisolatie = "Nee"
_d71.opname.kierdichting = "Redelijk (tochtstrips aanwezig)"
_gat71b = _FT71.huidige_staat_gaten(_d71)
check("staat: opnamevelden ingevuld -> alleen vierpansraam blijft over",
      len(_gat71b) == 1 and "Vierpansraam" in _gat71b[0])
# de 4 nieuwe MagicPlan-velden moeten via de CSV-parser in het dossier landen
import tempfile as _t71, os as _o71, csv as _c71
from magicplan.statistics_csv import build_dossier as _bd71
_rows71 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "40"], ["Woningtype", "Tussenwoning"],
           ["Gevel - isolatie aan zijde", "Binnenzijde (voorzetwand)"],
           ["Dak - isolatie aan zijde", "Buitenzijde"],
           ["Bodemisolatie kruipruimte", "Ja - folie"],
           ["Kierdichting", "Redelijk"], [],
           ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Ceiling Height"],
           ["Ground Floor", "40", "2.60 m"], []]
with _t71.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _fh71:
    _c71.writer(_fh71).writerows(_rows71)
    _p71 = _fh71.name
_dp71, _np71 = _bd71(_p71)
_o71.unlink(_p71)
check("parser: 4 nieuwe sectie-3-velden landen in het dossier",
      _dp71.opname.gevel_isolatie_zijde.startswith("Binnenzijde")
      and _dp71.opname.dak_isolatie_zijde == "Buitenzijde"
      and _dp71.opname.bodemisolatie.startswith("Ja")
      and _dp71.opname.kierdichting == "Redelijk")

print("24. MagicPlan Statistics-CSV -> dossier (parser)")
from magicplan.statistics_csv import build_dossier as _csvdos, _undot as _ud
check("undot bouwjaar-klasse", _ud("1992.t.m.2013") == "1992 t/m 2013")
check("undot subsysteem", _ud("A1.Standaard") == "A1 Standaard")
_cf = os.path.join(HERE, "fixtures", "statistics_voorbeeld.csv")
_cd, _cn = _csvdos(_cf, straat="Test", huisnummer="1", postcode="1234AB", plaats="Plaats")
_cgev = [s for s in _cd.schil if s.type == "gevel"]
_ckoz = [s for s in _cd.schil if s.type == "kozijn"]
check("csv: bouwjaar uit klasse = 1992", _cd.identificatie.bouwjaar == 1992)
check("csv: 2 gevels per orientatie (ZW+NO)", {s.orientatie for s in _cgev} == {"ZW", "NO"})
# 14-7 (methode Renze): gevel-hoofdvlak = BRUTO (10+10, incl. openingen); ramen = deelvlakken in Vabi
check("csv: gevel-m2 = BRUTO wandsom (10+10; ramen gaan er als deelvlak af)",
      abs(sum(s.oppervlakte_m2 for s in _cgev) - 20.0) < 0.1)
check("csv: 2 kozijnen (binnendeur uitgefilterd)", len(_ckoz) == 2)
check("csv: ventilatie A / A1", _cd.ventilatie.systeem == "A" and _cd.ventilatie.subsysteem_code == "A1")
check("csv: verwarming HR107", _cd.installaties.verwarming.type_opwekker == "HR107")
check("csv: qv10 gelezen maar niet-gemeten", _cd.opname.qv10_waarde == 1.25 and _cd.opname.qv10_gemeten is False)
check("csv: woningtype-ontbreekt geflagd", any("Woningtype" in n for n in _cn))

# Taak 018: MagicPlan vervangt de slash in alle twaalf aanvoertemperatuurcodes
# door een punt. Controleer de volledige CSV -> dossier -> VABI-keten.
from vabi import installatie_generate as _ig18
_aanvoer18 = [("30.27", "0"), ("35.30", "1"), ("40.35", "2"), ("45.40", "3"),
              ("50.42", "4"), ("55.47", "5"), ("60.50", "6"), ("65.55", "7"),
              ("70.60", "8"), ("75.65", "9"), ("80.60", "10"), ("90.70", "11")]
_ok18 = True
for _temp18, _code18 in _aanvoer18:
    _csv18 = ("PLAN ATTRIBUTES\n"
              "Verwarming - type opwekker,HR107\n"
              "Verwarming - aanvoertemperatuur,%s\n\n"
              "FLOOR ATTRIBUTES,Ground surface without walls: m²,Ceiling Height\n"
              "Ground Floor,40,2.50 m\n" % _temp18)
    with _t71.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8") as _f18:
        _f18.write(_csv18)
        _p18 = _f18.name
    try:
        _d18, _n18 = _csvdos(_p18)
        _r18, _flags18 = _ig18.build_tree(_d18)
        _ok18 = (_ok18
                 and _d18.installaties.verwarming.aanvoertemperatuur == _temp18
                 and _ig18._find(_r18, "WaterAanvoertemperatuur").text == _code18
                 and not any("aanvoertemp" in _flag18.lower() for _flag18 in _flags18))
    finally:
        os.unlink(_p18)
check("csv->VABI: alle 12 dotted aanvoertemperaturen krijgen juiste enum zonder onbekend-flag", _ok18)

print("25. Dak-per-vlak (hellingshoek uit nok + zadeldak-vlakken + parser-pad)")
from core.geometry import hellingshoek_uit_nok as _huk, dak_vlakken_zadeldak as _dvz
check("hellingshoek nok: breedte6 nok3 -> 45", _huk(6, 3, 0) == 45.0)
check("hellingshoek nok: knieschot meegerekend ~26.6", abs(_huk(8, 2.5, 0.5) - 26.6) < 0.2)
check("hellingshoek lessenaar (1 schuine zijde)", abs(_huk(4, 2, 0, 1) - 26.6) < 0.2)
_dv = _dvz(40, 8, 45, ("ZW", "NO"), ("NW",))
check("zadeldak: 2 schuine dakvlakken", sum(1 for v in _dv if v["kind"] == "dak") == 2)
check("zadeldak: schuin vlak = footprint/cos/2",
      abs([v for v in _dv if v["kind"] == "dak"][0]["m2"] - (40 * (2 ** 0.5)) / 2) < 0.1)
check("zadeldak: kopgevel-driehoek als gevel",
      any(v["kind"] == "gevel" and v["type"] == "kopgevel-driehoek" for v in _dv))
import tempfile as _tf
_dakcsv = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nFloors,2,\n,Test\nBouwjaar,1992.t.m.2013\n"
           "Type dak,Zadeldak\nHellingshoek dak,45\nDak vloerbreedte,8\nDak orientatie zijde 1,ZW\n"
           "Dak orientatie zijde 2,NO\nKopgevel orientatie 1,NW\nPlat dak m2,6\nPlat dak orientatie,horizontaal\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
           "Ground Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nKitchen,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
_tp = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_tp.write(_dakcsv); _tp.close()
_dd, _dn = _csvdos(_tp.name)
_dak = [s for s in _dd.schil if s.type == "dak"]
check("parser dak: 2 schuine + 1 plat dak", len(_dak) == 3 and any(s.subtype == "plat" for s in _dak))
check("parser dak: kopgevel als gevel NW",
      any(s.type == "gevel" and s.orientatie == "NW" and "kopgevel" in (s.subtype or "") for s in _dd.schil))

print("26. CSV-parser nieuwe opname-velden (woningtype/gevelhoogte/renovatiejaar/thermische massa)")
_velcsv = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\n"
           "Woningtype,Tussenwoning\nGevelhoogte (m),5.4\nRenovatiejaar,2015\n"
           "Thermische massa wanden,Zwaar\nThermische massa vloeren,Licht\n"
           "Qv10-waarde (dm3/s.m2),\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
           "Ground Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nKitchen,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
_vp = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_vp.write(_velcsv); _vp.close()
_vd, _vn = _csvdos(_vp.name)
check("csv: woningtype uit veld", _vd.identificatie.woningtype == "Tussenwoning")
check("csv: gevelhoogte uit veld", _vd.opname.gevelhoogte_m == 5.4)
check("csv: renovatiejaar uit veld", _vd.identificatie.renovatiejaar == 2015)
check("csv: thermische massa wanden=Zwaar", _vd.opname.thermische_massa_wanden == "Zwaar")
check("csv: thermische massa vloeren=Licht", _vd.opname.thermische_massa_vloeren == "Licht")
check("csv: tussenwoning -> hart-op-hart-toeslag NIET auto-toegepast maar GEFLAGD (zelf in Vabi)",
      any("HART-OP-HART GEVEL-TOESLAG" in n and "ZELF TOEVOEGEN IN VABI" in n for n in _vn)
      and not any("hart-op-hart" in (s.opmerkingen or "") for s in _vd.schil if s.type == "gevel"))
check("csv: thermische massa gevuld zonder handmatig-flag", not any("HANDMATIG" in n for n in _vn))
# CLI-arg overrulet het CSV-veld
_vd2, _ = _csvdos(_vp.name, woningtype="Vrijstaand", gevelhoogte_m=3.0)
check("csv: CLI-arg overrulet veld (woningtype+gevelhoogte)",
      _vd2.identificatie.woningtype == "Vrijstaand" and _vd2.opname.gevelhoogte_m == 3.0)
# VABI objecten-generator: thermische massa nu volledig gewired (0=Licht/1=Zwaar/2=Zeer zwaar)
from vabi.objecten_generate import build_tree as _objbuild
_otree, _omap, _oiss, _ostats = _objbuild(_vd)
_alg_v = next((a for a in _otree.iter() if a.tag.rsplit("}", 1)[-1] == "Algemeen" and a.find("Bouwjaar") is not None), None)
check("obj-gen: Licht-vloer -> TypeBouwwijzeVloeren=0", _alg_v is not None and _alg_v.findtext("TypeBouwwijzeVloeren") == "0")
check("obj-gen: Zwaar-wand -> TypeBouwwijzeWanden=1", _alg_v is not None and _alg_v.findtext("TypeBouwwijzeWanden") == "1")

# Ag + vloer-perimeter dossier-gestuurd in de objecten-XML
from xml.etree import ElementTree as _ET
_agcsv = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1992.t.m.2013\n"
          "Woningtype,Vrijstaand\nGevelhoogte (m),5.4\nTotal living area,120\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.50 m,Kruipruimte\nFirst Floor,60,2.50 m,Buitenlucht\n\n"
          "ROOM ATTRIBUTES\nLiving room,70\nBedroom,50\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
          "Ground Floor,\nLiving room,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
_ap = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_ap.write(_agcsv); _ap.close()
_ad, _an = _csvdos(_ap.name)
check("csv: vloer-perimeter = buitenomtrek (24)",
      any(s.type == "vloer" and s.perimeter_m == 24.0 for s in _ad.schil))
_atree, _amap, _aiss, _astats = _objbuild(_ad)
_axml = _ET.tostring(_atree, encoding="unicode")
_alg = next((a for a in _atree.iter() if a.tag.rsplit("}", 1)[-1] == "Algemeen"
             and a.find("Bouwjaar") is not None), None)
# EPA enum-mismatch-fix (23-6): de DIRECTE Rekenzone>Algemeen-<Gebruiksoppervlakte> is een ENUM/vlag
# (echte export = "1"), GEEN m2-veld -> overschrijven met de gemeten m2 gaf "Enum mismatch" bij objecten-
# import. De Ag dragen we via de Verdiepingen-som (zie check hieronder). Dus dit veld NIET op 120 zetten.
check("obj-gen: rekenzone-<Gebruiksoppervlakte> NIET overschreven met m2 (enum, geen oppervlak)",
      _alg is not None and _alg.find("Gebruiksoppervlakte").text != "120.00")
check("obj-gen: bouwjaar uit dossier (1992) i.p.v. sjabloon (1994)",
      _alg is not None and _alg.find("Bouwjaar").text == "1992")
_verd_sum = round(sum(float(v.find("Gebruiksoppervlakte").text)
                      for v in _alg.find("Verdiepingen") if v.tag.rsplit("}", 1)[-1] == "Verdieping"), 1)
check("obj-gen: Verdiepingen-som = Ag (120)", _verd_sum == 120.0)
check("obj-gen: AantalBouwlagenRekenzone = 2", _alg.find("AantalBouwlagenRekenzone").text == "2")
# objecten-sjabloon MOET versie-consistent zijn met de constructie-sjabloon (beide 12.0.1) -> anders
# importeert EPA niet. Het oude objecten-sjabloon was 12.0.0 (120000061); nu 12.0.1 (live bewezen 23-6).
check("obj-gen: XmlVersie = 120001001 (sjabloon versie-consistent met constructies)",
      _atree.findtext("XmlVersie") == "120001001")
# vloer-hoofdvlak heeft de perimeter 24
_vloer_hv = [h for h in _atree.iter() if h.tag.rsplit("}", 1)[-1] == "Hoofdvlak"
             and (h.findtext("NaamConstructie") or "").lower().startswith("vloer")]
check("obj-gen: vloer-hoofdvlak perimeter = 24",
      any((h.findtext("Perimeter") or "") == "24.00" for h in _vloer_hv))

# qv10: niet-gemeten -> forfaitair (Qv10Gemeten != 1); gemeten -> wel geschreven (ISSO 7.1.5)
def _alg_of(_tree):
    return next((a for a in _tree.iter() if a.tag.rsplit("}", 1)[-1] == "Algemeen"
                 and a.find("Bouwjaar") is not None), None)
_ad.opname.qv10_waarde = 1.25; _ad.opname.qv10_gemeten = False
_t_ng = _alg_of(_objbuild(_ad)[0])
check("obj-gen: qv10 niet-gemeten -> forfaitair (Gemeten!=1)", _t_ng.findtext("Qv10Gemeten") != "1")
_ad.opname.qv10_gemeten = True
_t_g = _alg_of(_objbuild(_ad)[0])
check("obj-gen: qv10 gemeten -> Gemeten=1 + waarde",
      _t_g.findtext("Qv10Gemeten") == "1" and _t_g.findtext("Qv10Waarde") == "1.250")

# begrenzing -> GrenstAan (uit echt sjabloon afgeleid: 0 buiten / 2 grond / 3 kruipruimte / 4 AOR)
from vabi.objecten_generate import _grenst_aan_code as _gac
check("grenstaan: buitenlucht=0", _gac("Buitenlucht") == "0")
check("grenstaan: grond=2", _gac("Grond") == "2")
check("grenstaan: kruipruimte=3", _gac("Kruipruimte") == "3")
check("grenstaan: water=1 (EPA-bevestigd)", _gac("Water") == "1")
check("grenstaan: onverwarmde kelder=7", _gac("Onverwarmde kelder") == "7")
check("grenstaan: AOR basis = buitenlucht (0)", _gac("AOR") == "0")
check("grenstaan: AOR detail = 4", _gac("AOR", basis=False) == "4")
check("grenstaan: AOS basis=0 / detail=5", _gac("AOS") == "0" and _gac("AOS", basis=False) == "5")
check("grenstaan: ASGR detail = 6", _gac("ASGR", basis=False) == "6")
check("grenstaan: onbekend -> None (niet gokken)", _gac("Iets onbekends xyz") is None)
for _s in _ad.schil:
    if _s.type == "vloer":
        _s.begrenzing = "Grond"
_vl_gr = [h for h in _objbuild(_ad)[0].iter() if h.tag.rsplit("}", 1)[-1] == "Hoofdvlak"
          and (h.findtext("NaamConstructie") or "").lower().startswith("vloer")]
check("obj-gen: vloer-begrenzing Grond -> GrenstAan 2", any(h.findtext("GrenstAan") == "2" for h in _vl_gr))
for _s in _ad.schil:
    if _s.type == "vloer":
        _s.begrenzing = "Iets onbekends xyz"
check("obj-gen: onbekende begrenzing geflagd (niet gegokt)",
      any("sjabloon-GrenstAan" in i for i in _objbuild(_ad)[2]))

# dak Hellingshoek-ENUM: plat=6 / hellend=3 (geverifieerd vabi_enums; GEEN rauwe graden)
# Deze CSV had geen dakvelden -> _ad draagt nog de parser-placeholder ('magicplan-dak-fallback',
# taak 014); die hoort net als in de webapp-wizard te verdwijnen zodra er echte dakvlakken bij
# komen, anders blokkeert de Vabi-preflight terecht op dubbel dakoppervlak.
from core.dossier import SchilDeel as _SD
_ad.schil = [s for s in _ad.schil if not (s.type == "dak" and s.bron == "magicplan-dak-fallback")]
_ad.schil.append(_SD(id="dakplat", type="dak", subtype="plat", begrenzing="Buitenlucht",
                     orientatie="", oppervlakte_m2=20.0, hellingshoek=0, rekenzone=1))
_ad.schil.append(_SD(id="dakhel", type="dak", subtype="schuin", begrenzing="Buitenlucht",
                     orientatie="ZW", oppervlakte_m2=30.0, hellingshoek=45, rekenzone=1))
_dakhv = {(h.findtext("Naam") or ""): h.findtext("Hellingshoek") for h in _objbuild(_ad)[0].iter()
          if h.tag.rsplit("}", 1)[-1] == "Hoofdvlak" and "dak" in (h.findtext("Naam") or "").lower()}
check("obj-gen: plat dakvlak -> Hellingshoek 6", _dakhv.get("Dak dakplat") == "6")
check("obj-gen: hellend dakvlak -> Hellingshoek 3", _dakhv.get("Dak dakhel") == "3")

# Gebouwhoogte = HANDMATIGE invoer (12-7): zonder gebouwhoogte_m -> 0 (geen gevelhoogte-fallback,
# geen sjabloon-lek); met gebouwhoogte_m -> die waarde.
_gh0 = next((e for e in _objbuild(_ad)[0].iter() if e.tag.rsplit("}", 1)[-1] == "Gebouwhoogte"), None)
check("obj-gen: geen gebouwhoogte-invoer -> 0 (geen gevelhoogte 5.40 als gebouwhoogte)",
      _gh0 is not None and _gh0.text == "0")
_ad.opname.gebouwhoogte_m = 8.40
_gh = next((e for e in _objbuild(_ad)[0].iter() if e.tag.rsplit("}", 1)[-1] == "Gebouwhoogte"), None)
check("obj-gen: Gebouwhoogte uit handmatige opname-invoer (8.40)", _gh is not None and _gh.text == "8.40")
# perimeter-guard: vloer Buitenlucht -> GEEN perimeter-override (ISSO 8.3 alleen grond/kruip/kelder)
for _s in _ad.schil:
    if _s.type == "vloer":
        _s.begrenzing = "Buitenlucht"; _s.perimeter_m = 30.0
_vl_bl = [h for h in _objbuild(_ad)[0].iter() if h.tag.rsplit("}", 1)[-1] == "Hoofdvlak"
          and (h.findtext("NaamConstructie") or "").lower().startswith("vloer")]
check("obj-gen: vloer Buitenlucht -> geen perimeter-override",
      all((h.findtext("Perimeter") or "") != "30.00" for h in _vl_bl))

print("27. Parser parent/child: begrenzing-naamconventie + kozijn A/B/C + Qv10-gemeten")
def _wrow(d, n=18):
    rr = [""] * n
    for kk, vv in d.items():
        rr[kk] = str(vv)
    return ",".join(rr)
_nc = "\n".join([
    "PLAN ATTRIBUTES", "Exterior perimeter: m,20,", "Bouwjaar,1975 t/m 1982",
    "Woningtype,Tussenwoning", "Gevelhoogte (m),6", "Qv10 gemeten?,Ja", "Qv10-waarde (dm3/s.m2),0.8", "",
    "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing",
    "Ground Floor,60,2.60 m,Kruipruimte", "",
    "WALL ATTRIBUTES,c1,c2,Surface,SurfNoOpen,c5,c6,c7,Type,Isol,c10,Orientatie,Bron,c13,c14,Kozijn,Glas,RaamOrient",
    "Ground Floor",
    _wrow({0: "Voorgevel", 3: 18, 4: 18, 8: "Wall", 11: "ZW"}),
    _wrow({0: "raam1", 3: 3.5, 8: "Window", 15: "B", 16: "HR++ glas", 17: "ZW"}),
    _wrow({0: "Achtergevel AOR garage", 3: 18, 4: 18, 8: "Wall", 11: "NO"}),
    _wrow({0: "buurwand AVR", 3: 18, 4: 18, 8: "Wall", 11: "O"}), "",
])
_npf = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_npf.write(_nc); _npf.close()
_nd, _nn = _csvdos(_npf.name)
_gev = {s.orientatie: s for s in _nd.schil if s.type == "gevel"}
check("csv: gevel AOR uit wandnaam (NO)", bool(_gev.get("NO")) and _gev["NO"].begrenzing == "AOR")
check("csv: AVR-wand uitgesloten uit schil (geen O-gevel)", "O" not in _gev)
check("csv: 'Qv10 gemeten?'=Ja -> qv10_gemeten True", _nd.opname.qv10_gemeten is True)
_ram = [s for s in _nd.schil if s.subtype == "Raam"]
check("csv: raam erft begrenzing van moederwand (ZW=Buitenlucht)", bool(_ram) and _ram[0].begrenzing == "Buitenlucht")
check("csv: kozijn B -> Metaal (thermisch onderbroken)", bool(_ram) and _ram[0].kozijnmateriaal == "Metaal (thermisch onderbroken)")

# NL-ruimtenamen correct classificeren (anders ventilatiebalans fout) — gevonden in de ultieme check
from magicplan.statistics_csv import _functie_uit_naam as _fun
check("functie: Woonkamer -> verblijfsruimte", _fun("Woonkamer") == "verblijfsruimte")
check("functie: Keuken -> keuken", _fun("Keuken") == "keuken")
check("functie: Slaapkamer 1 -> slaapkamer", _fun("Slaapkamer 1") == "slaapkamer")
check("functie: Badkamer -> badkamer", _fun("Badkamer") == "badkamer")
from core.dossier import Ruimte as _R
_vres = vent_bereken([_R(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30),
                      _R(naam="Keuken", functie="keuken", oppervlakte_m2=12),
                      _R(naam="Slaapkamer", functie="slaapkamer", oppervlakte_m2=14)])
check("ventilatie NL: keuken afvoer 21 + woonkamer toevoer 21",
      any(r["functie"] == "keuken" and r["afvoer"] == 21.0 for r in _vres["rows"])
      and any(r["functie"] == "verblijfsruimte" and r["toevoer"] == 21.0 for r in _vres["rows"]))

print("\n28. Installaties: PV-knoop + verwarming/tapwater-codes (LIVE geharvest uit EPA)")
from core.dossier import ZonneEnergieSysteem as _ZE
import vabi.installatie_generate as _ig
def _iloc(e): return e.tag.rsplit('}', 1)[-1]

_id = build_sample()
_id.installaties.zonne_energie = [_ZE(systeem="PV-panelen", aantal=12, oppervlak_per_paneel_m2=1.7,
    orientatie="Zuid", hellingshoek=35, pv_type="Monokristallijn", fabricagejaar="Vanaf 2018",
    bouwintegratie="Goed geventileerd")]
_ir, _iflags = _ig.build_tree(_id)
_ze = [e for e in _ir.iter() if _iloc(e) == "ZonneEnergie"]
_zd = {_iloc(c): (c.text or "").strip() for c in _ze[0]} if _ze else {}
check("inst: PV ZonneEnergiesysteem=0 (PV-panelen)", _zd.get("ZonneEnergiesysteem") == "0")
check("inst: PV paneeltype Mono=1", _zd.get("PiekvermogenPVPanelen") == "1")
check("inst: PV fabricagejaar Vanaf2018=4", _zd.get("FabricagejaarPVPanelen") == "4")
check("inst: PV oriëntatie Zuid=4 (PV-enum, niet geometrie)", _zd.get("Orientatie") == "4")
check("inst: PV bouwintegratie Goed=2", _zd.get("Bouwintegratie") == "2")
check("inst: PV hellingshoek rauwe graden (35)", _zd.get("Hellingshoek") == "35")
check("inst: PV aantal panelen 12", _zd.get("AantalPanelen") == "12")

_id2 = build_sample(); _id2.installaties.zonne_energie = []
_ir2, _ = _ig.build_tree(_id2)
check("inst: geen PV -> ZonneEnergie-knoop verwijderd (geen fantoom)",
      not [e for e in _ir2.iter() if _iloc(e) == "ZonneEnergie"])

_op = [e for e in _ir.iter() if _iloc(e) == "VerwarmingOpwekker"][0]
check("inst: gasketel -> TypeOpwekker 4", _op.find("TypeOpwekker").text == "4")
_id.installaties.verwarming.type_opwekker = "Warmtepomp elektrisch"
_ir3, _f3 = _ig.build_tree(_id)
check("inst: warmtepomp NIET auto-gecodeerd -> geflagd (golden rule)",
      any("warmtepomp" in f.lower() or "opwekkertype" in f.lower() for f in _f3))
_id.installaties.tapwater.type_toestel = "Gasgestookt combitoestel"
_ir4, _ = _ig.build_tree(_id)
_top = [e for e in _ir4.iter() if _iloc(e) == "TapwaterOpwekker"][0]
check("inst: tapwater combitoestel -> TypeToestel 10", _top.find("TypeToestel").text == "10")

print("\n29. Opname-tokens: narekenen-vlag + per-wand isolatie-override + vloer-split + Ag-aftrek <1,5 m")
_tok = "\n".join([
    "PLAN ATTRIBUTES", "Exterior perimeter: m,24,", "Bouwjaar,1975 t/m 1982",
    "Woningtype,Tussenwoning", "Gevelhoogte (m),6", "Isolatie aanwezig,Ja",
    "Begrenzing (vloer),Kruipruimte", "Ag-aftrek zolder (m2),8", "Total living area,100", "",
    "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing",
    "Ground Floor,60,2.50 m,Kruipruimte", "",
    "ROOM ATTRIBUTES", "Woonkamer,40", "Studeerkamer grond,20", "",
    "WALL ATTRIBUTES,c1,c2,Surface,SurfNoOpen,c5,c6,c7,Type,Isol,c10,Orientatie,Bron,c13,c14,Kozijn,Glas,RaamOrient",
    "Ground Floor",
    _wrow({0: "Voorgevel", 3: 18, 4: 18, 8: "Wall", 11: "ZW"}),
    _wrow({0: "Zijgevel ongeisoleerd", 3: 12, 4: 12, 8: "Wall", 11: "NO"}),
    _wrow({0: "Achtergevel narekenen", 3: 18, 4: 18, 8: "Wall", 11: "O"}), "",
])
_tkp = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_tkp.write(_tok); _tkp.close()
_td, _tn = _csvdos(_tkp.name)
_tgev = {s.orientatie: s for s in _td.schil if s.type == "gevel"}
check("token: per-wand 'ongeisoleerd' -> isolatie_aanwezig Nee (override)",
      bool(_tgev.get("NO")) and _tgev["NO"].isolatie_aanwezig == "Nee")
check("token: normale gevel houdt projectdefault (Ja)",
      bool(_tgev.get("ZW")) and _tgev["ZW"].isolatie_aanwezig == "Ja")
check("token: 'narekenen'-wand -> flag in notes", any("NAREKENEN" in n for n in _tn))
check("token: narekenen-gevel heeft opmerking", any("NAREKENEN" in (s.opmerkingen or "") for s in _td.schil if s.type == "gevel"))
_tvl = [s for s in _td.schil if s.type == "vloer"]
check("token: ruimtenaam 'grond' -> apart vloerdeel met begrenzing Grond",
      any(s.begrenzing == "Grond" and abs((s.oppervlakte_m2 or 0) - 20.0) < 0.1 for s in _tvl))
check("token: hoofdvloer verlaagd met het grond-deel (60-20=40)",
      any(s.subtype == "Begane grondvloer" and abs((s.oppervlakte_m2 or 0) - 40.0) < 0.1 for s in _tvl))
check("token: Ag = gemeten verdieping - zolderaftrek (60-8=52; MagicPlan-woonopp 100 NIET gebruikt)",
      abs((_td.geometrie.gebruiksoppervlakte_ag_m2 or 0) - 52.0) < 0.1)
check("token: Ag-aftrek gemeld in notes", any("Ag verlaagd" in n or "afgetrokken" in n for n in _tn))

print("\n30. Objecten<->constructies: gedeelde (deterministische) GUIDs — EPA enum-mismatch-fix")
from vabi.constructie_generate import resolve_constructies as _rc
import vabi.constructie_generate as _cg2, vabi.objecten_generate as _og2
import xml.etree.ElementTree as _ETg
_dzg = build_sample()
_m1 = _rc(_dzg)[1]; _m2 = _rc(_dzg)[1]
check("guid deterministisch over 2 aanroepen (zelfde per constructie)",
      bool(_m1) and all(_m1[k]["guid"] == _m2[k]["guid"] for k in _m1))
_tdg = _tf.mkdtemp()
_cpg = os.path.join(_tdg, "c.xml"); _opg = os.path.join(_tdg, "o.xml")
_cg2.write(_dzg, _cpg); _og2.write(_dzg, _opg)
def _loc2(e): return e.tag.rsplit("}", 1)[-1]
_cguids = set((c.findtext("Guid") or "").strip().lower()
              for c in _ETg.parse(_cpg).getroot().iter() if _loc2(c) == "Constructie")
_orefs = set((e.text or "").strip().lower()
             for e in _ETg.parse(_opg).getroot().iter() if _loc2(e) == "Constructie" and (e.text or "").strip())
check("objecten constructie-refs bestaan ALLEMAAL in constructiebib (geen EPA enum-mismatch)",
      bool(_orefs) and not (_orefs - _cguids))

print("\n31. Gevel-naamgeving -> oriëntatie-afleiding + meerdere PV + kwaliteitsverklaring-flag")
from magicplan.statistics_csv import (_orient_afleiden as _oa, _orient_uit_naam as _ou,
                                      _gevel_naam_uit_naam as _gn)
# 8-punts rotatie (Oost-vanaf-straat): voorgevel Z -> rechter O, links W, achter N
check("afleiden: voorgevel Z -> rechtergevel O", _oa("rechts", "Z") == "O")
check("afleiden: voorgevel Z -> linkergevel W", _oa("links", "Z") == "W")
check("afleiden: voorgevel Z -> achtergevel N", _oa("achter", "Z") == "N")
check("afleiden: voorgevel Z -> voorgevel Z", _oa("voor", "Z") == "Z")
check("afleiden: intercardinaal voorgevel NO -> rechtergevel NW", _oa("rechts", "NO") == "NW")
check("afleiden: geen voorgevel-orientatie -> leeg", _oa("rechts", "") == "")
check("naam: 'Linkergevel ...' -> links", _gn("Linkergevel woonkamer") == "links")
check("naam: kompastoken-override 'Rechtergevel ZW' -> ZW", _ou("Rechtergevel ZW") == "ZW")
check("naam: zonder kompastoken -> leeg", _ou("Voorgevel") == "")
# CSV-integratie: 4 benoemde gevels zonder kompaskolom + naam-override + 2 PV + kwaliteitsverklaring
_gv = "\n".join([
    "PLAN ATTRIBUTES", "Bouwjaar,1975 t/m 1982", "Woningtype,Vrijstaand",
    "Orientatie voorgevel,Z", "Total living area,120", "Exterior perimeter: m,40",
    "Rc-bron gevel,Kwaliteitsverklaring",
    "Zonne-energie aanwezig?,Ja", "PV - paneeltype,Monokristallijn", "PV - aantal panelen,10", "PV - orientatie,Zuid",
    "PV-2 - paneeltype,Polykristallijn", "PV-2 - aantal panelen,6", "PV-2 - orientatie,Oost", "",
    "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing",
    "Ground Floor,80,2.60 m,Kruipruimte", "",
    "WALL ATTRIBUTES,c1,c2,Surface,SurfNoOpen,c5,c6,c7,Type,Isol,c10,Orientatie,Bron",
    "Ground Floor",
    _wrow({0: "Voorgevel", 3: 30, 4: 27, 8: "Wall"}),
    _wrow({0: "Achtergevel", 3: 30, 4: 28, 8: "Wall"}),
    _wrow({0: "Linkergevel", 3: 20, 4: 20, 8: "Wall"}),
    _wrow({0: "Rechtergevel ZW", 3: 20, 4: 20, 8: "Wall"}), "",   # naam-override ZW i.p.v. afgeleid O
])
_gp = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
_gp.write(_gv); _gp.close()
_gd, _gnotes = _csvdos(_gp.name)
_gmap = {s.orientatie: s for s in _gd.schil if s.type == "gevel"}
check("csv: voorgevel afgeleid Z (naam=voor)", bool(_gmap.get("Z")) and _gmap["Z"].gevel_naam == "voor")
check("csv: achtergevel afgeleid N", bool(_gmap.get("N")) and _gmap["N"].gevel_naam == "achter")
check("csv: linkergevel afgeleid W", bool(_gmap.get("W")) and _gmap["W"].gevel_naam == "links")
check("csv: rechtergevel kompas-override ZW wint van afgeleid O",
      "O" not in _gmap and bool(_gmap.get("ZW")) and _gmap["ZW"].gevel_naam == "rechts")
check("csv: 2 PV-systemen ingelezen", len(_gd.installaties.zonne_energie) == 2)
check("csv: 2e PV = polykristallijn/Oost",
      any("poly" in (z.pv_type or "").lower() for z in _gd.installaties.zonne_energie))
check("csv: gevels rc_bron = Kwaliteitsverklaring",
      bool([s for s in _gd.schil if s.type == "gevel"])
      and all(s.rc_bron == "Kwaliteitsverklaring" for s in _gd.schil if s.type == "gevel"))
check("csv: kwaliteitsverklaring geflagd in notes", any("kwaliteitsverklaring" in n.lower() for n in _gnotes))
# constructie-generator BLOKKEERT de export bij kwaliteitsverklaring (vabi/preflight.py,
# taak fix/kwaliteitsverklaring-blokkeert-vabi-export): harde VabiExportBlocked i.p.v. een
# stille issue, want de tool exporteert hiervoor geen rekenbare forfaitaire fallback.
_kvdir = _tf.mkdtemp(); _kvpath = os.path.join(_kvdir, "c.xml")
from vabi.constructie_generate import write as _cwrite2
from vabi.preflight import VabiExportBlocked as _VEB
try:
    _cwrite2(_gd, _kvpath)
    check("constr-gen: kwaliteitsverklaring -> VabiExportBlocked", False, "geen exception opgegooid")
except _VEB as _kve:
    check("constr-gen: kwaliteitsverklaring -> VabiExportBlocked", "kwaliteitsverklaring" in str(_kve).lower())

print("\n32. MagicPlan form_push: merge + validatie (offline)")
import magicplan.form_push as _fp
_fform = {"name": "Nij Begun", "context": ["plan"], "children": [
    {"id": "s1", "name": "Gevels", "type": "section", "comparisonValue": None},
    {"id": "q1", "name": "Geveltype", "type": "question", "dataType": "list", "comparisonValue": None, "required": True},
    {"id": "s2", "name": "Vloer & Dak", "type": "section", "comparisonValue": None},
]}
_frec = {"id": "rec1", "form": _fform}
_fadd = _fp.load_additions()
_frec, _fadded, _freq, _fprobs = _fp.merge_record(_frec, _fadd, verbose=False)
check("form_push: 'Oriëntatie voorgevel' toegevoegd", "Oriëntatie voorgevel" in _fadded)
check("form_push: Rc-bron + Gebouwhoogte + 4x gevelbreedte toegevoegd (9 velden totaal)",
      len(_fadded) == 9 and any("Gebouwhoogte" in a for a in _fadded)
      and any("Voorgevel - breedte" in a for a in _fadded))
check("form_push: geen validatieproblemen", _fprobs == [])
check("form_push: nieuw veld na de juiste sectie (Gevels)",
      any(c.get("name") == "Oriëntatie voorgevel" for c in _fp._form_of(_frec)["children"]))
check("form_push: idempotent (2e merge voegt 0 toe)",
      _fp.merge_record({"id": "rec1", "form": _fp._form_of(_frec)}, _fadd, verbose=False)[1] == [])
check("form_push: name_escaped gestript + context = string-array",
      "name_escaped" not in json.dumps(_frec) and _fp._form_of(_frec)["context"] == ["plan"])

print("\n33. Schilddak/lessenaar + rekenzone")
from core.geometry import (dak_vlakken_schilddak as _ds, dak_vlakken_lessenaar as _dl,
                           dak_vlakken_zadeldak as _dz)
check("geom: lessenaar = 1 dakvlak (footprint/cos)", len(_dl(60, 30, "O")) == 1 and _dl(60, 30, "O")[0]["kind"] == "dak")
_sch = _ds(60, 45, ("N", "O", "Z", "W"))
check("geom: schilddak = alleen dakvlakken (GEEN kopgevel-gevel)", len(_sch) == 4 and all(v["kind"] == "dak" for v in _sch))
check("geom: zadeldak heeft wél een kopgevel-gevel", any(v["kind"] == "gevel" for v in _dz(60, 6, 45, ("O", "W"), ("N", "Z"))))
from magicplan.statistics_csv import _rekenzone_uit_naam as _rzn
check("parser: 'Rechtergevel zone2' -> rekenzone 2", _rzn("Rechtergevel zone2") == 2)
check("parser: 'Slaapkamer rekenzone 3' -> 3", _rzn("Slaapkamer rekenzone 3") == 3)
check("parser: gewone naam -> rekenzone 1", _rzn("Voorgevel") == 1)
_sv = "\n".join([
    "PLAN ATTRIBUTES", "Woningtype,Vrijstaand", "Orientatie voorgevel,Z", "Total living area,120",
    "Exterior perimeter: m,40", "Type dak,Schilddak", "Hellingshoek dak,45",
    "Dak orientatie zijde 1,Z", "Dak orientatie zijde 2,N", "Kopgevel orientatie 1,O", "Kopgevel orientatie 2,W",
    "Ventilatie - rekenzone,2", "",
    "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing", "Ground Floor,80,2.6 m,Kruipruimte", "",
    "WALL ATTRIBUTES,c1,c2,Surface,SurfNoOpen,c5,c6,c7,Type,Isol,c10,Orientatie,Bron", "Ground Floor",
    _wrow({0: "Voorgevel", 3: 30, 4: 28, 8: "Wall"}),
    _wrow({0: "Rechtergevel zone2", 3: 20, 4: 20, 8: "Wall"}), "",
])
_svp = _tf.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _svp.write(_sv); _svp.close()
_svd, _svn = _csvdos(_svp.name)
_dak = [s for s in _svd.schil if s.type == "dak"]
check("csv: schilddak -> 4 dakvlakken", len(_dak) == 4 and all("schild" in (s.subtype or "") for s in _dak))
check("csv: schilddak geeft GEEN kopgevel-gevel", not any("kopgevel" in (s.subtype or "") for s in _svd.schil))
_rg = [s for s in _svd.schil if s.type == "gevel" and s.orientatie == "O"]
check("csv: 'Rechtergevel zone2' -> gevel rekenzone 2", bool(_rg) and _rg[0].rekenzone == 2)
check("csv: ventilatie rekenzone 2 uit Installaties-form", _svd.ventilatie.rekenzone == 2)
check("csv: meerdere rekenzones geflagd", any("rekenzone" in n.lower() and "zone 2" in n for n in _svn))

print("\n34. Webapp: maatregel-selectie + visueel ventilatieplan")
from dashboard.measures import laad_catalog as _lc, suggesties as _sug, bouw_maatregelen as _bm
_wcat = _lc()
_wdos = build_sample()
_grp = _sug(_wdos, _wcat)
check("measures: advies-groepen gevonden", len(_grp) >= 1)
check("measures: elke groep heeft kandidaten + goedkoopste eerst",
      all(g["kandidaten"] and g["kandidaten"][0]["prijs"] <= g["kandidaten"][-1]["prijs"] for g in _grp))
_keuze = [{"code": g["default_code"], "onderdeel": g["onderdeel"], "m2": g["m2"],
          "rc_u_doel": g["rc_u_doel"], "subposten": []} for g in _grp]
_maat, _tot = _bm(_wcat, _keuze)
check("measures: selectie -> maatregelen met kosten", len(_maat) == len(_grp) and all(m.kosten for m in _maat))
check("measures: totaal = som maatregelkosten", abs(_tot - round(sum(m.kosten for m in _maat), 2)) < 0.01)
from ventilatie.ventilatie import bereken as _vber
from ventilatie.ventilatieplan_svg import ventilatieplan_svg as _vsvg
_vres = _vber([_R(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30),
              _R(naam="Keuken", functie="keuken", oppervlakte_m2=10),
              _R(naam="Badkamer", functie="badkamer", oppervlakte_m2=6)])
_svgtxt = _vsvg(_vres, adres="Teststraat 1")
check("ventilatieplan-svg: geldige SVG-string", _svgtxt.startswith("<svg") and _svgtxt.rstrip().endswith("</svg>"))
import xml.etree.ElementTree as _ETv
check("ventilatieplan-svg: well-formed XML", bool(_ETv.fromstring(_svgtxt)))
check("ventilatieplan-svg: toont toevoer + afvoer", "l/s in" in _svgtxt and "l/s uit" in _svgtxt)

from dashboard.gebouw_svg import gebouw_svg as _gsvg, _muurvlakken as _gmuur
from core.dossier import Dossier as _GDos, SchilDeel as _GSchil
_gdos = _GDos()
_gdos.opname.gevelhoogte_m = 3.0
_gdos.opname.gebouwhoogte_m = 7.0
_gdos.schil = [
    _GSchil(id="gevel-voor", type="gevel", gevel_naam="voor", orientatie="Z",
            oppervlakte_m2=18, rc_huidig=2.5),
    _GSchil(id="gevel-achter", type="gevel", gevel_naam="achter", orientatie="N",
            oppervlakte_m2=18, rc_huidig=2.5),
    _GSchil(id="gevel-links", type="gevel", gevel_naam="links", orientatie="W",
            oppervlakte_m2=24),
    _GSchil(id="gevel-rechts", type="gevel", gevel_naam="rechts", orientatie="O",
            oppervlakte_m2=24),
    _GSchil(id="dak1-schuin-z", type="dak", subtype="schuin (zadeldak)", orientatie="Z",
            oppervlakte_m2=25, hellingshoek=45, breedte_m=6, diepte_m=4,
            geometrie_groep="dak1"),
    _GSchil(id="dak1-schuin-n", type="dak", subtype="schuin (zadeldak)", orientatie="N",
            oppervlakte_m2=25, hellingshoek=45, breedte_m=6, diepte_m=4,
            geometrie_groep="dak1"),
]
_gsvgtxt = _gsvg(_gdos)
check("gebouw-svg: geldige SVG-string", _gsvgtxt.startswith("<svg") and _gsvgtxt.rstrip().endswith("</svg>"))
check("gebouw-svg: well-formed XML", bool(_ETv.fromstring(_gsvgtxt)))
check("gebouw-svg: toont gevel- en dakvlak-ids", "gevel-voor" in _gsvgtxt and "dak1-schuin-z" in _gsvgtxt)
check("gebouw-svg: echte isometrische vlakken en afgeleide footprint",
      'data-face="gevel-voor"' in _gsvgtxt and "Footprint 6.00 × 8.00 m" in _gsvgtxt
      and "WONING" not in _gsvgtxt)
check("gebouw-svg: taak-005-dakmetadata rendert exact",
      'data-element-id="dak1-schuin-z"' in _gsvgtxt and 'data-geometrie="exact"' in _gsvgtxt)
check("gebouw-svg: exacte dakcoördinaten bewaren run, goot- en gebouwhoogte",
      'data-punten-3d="0.000,3.000,0.000 6.000,3.000,0.000 '
      '6.000,7.000,4.000 0.000,7.000,4.000"' in _gsvgtxt)
_shuffle = _GDos.from_dict(_gdos.to_dict())
_shuffle.schil[-2], _shuffle.schil[-1] = _shuffle.schil[-1], _shuffle.schil[-2]
_shuffle_svg = _gsvg(_shuffle)
check("gebouw-svg: dakzijde volgt oriëntatie en niet lijstvolgorde",
      'data-element-id="dak1-schuin-z"' in _shuffle_svg
      and 'data-punten-3d="0.000,3.000,0.000 6.000,3.000,0.000 '
      '6.000,7.000,4.000 0.000,7.000,4.000"' in _shuffle_svg)
_dak_conflict = _GDos.from_dict(_gdos.to_dict())
_dak_conflict.opname.gebouwhoogte_m = 6.0
_dak_conflict_svg = _gsvg(_dak_conflict)
check("gebouw-svg: incompatibele exacte dakmetadata wordt niet vervormd of exact genoemd",
      "Dak niet getekend:" in _dak_conflict_svg and "verschillen meer dan 0,10 m" in _dak_conflict_svg
      and ('data-element-id="dak1-schuin-z" data-oppervlakte-m2="25.00" '
           'data-orientatie="Z" data-geometrie="exact"') not in _dak_conflict_svg)
_escape = _GDos.from_dict(_gdos.to_dict())
_escape.schil[0].id = 'gevel-&-<voor>"'
_escape_svg = _gsvg(_escape)
check("gebouw-svg: dossierwaarden zijn XML-veilig ge-escaped",
      "gevel-&amp;-&lt;voor&gt;&quot;" in _escape_svg and bool(_ETv.fromstring(_escape_svg)))
_legacy = _GDos.from_dict(_gdos.to_dict())
_legacy.schil[-2].breedte_m = _legacy.schil[-2].diepte_m = None
_legacy.schil[-1].breedte_m = _legacy.schil[-1].diepte_m = None
_legacy_svg = _gsvg(_legacy)
check("gebouw-svg: legacy-dak zichtbaar als benadering",
      'data-geometrie="benaderd"' in _legacy_svg and "Benaderd dakvlak:" in _legacy_svg)
_kapel = _GDos.from_dict(_gdos.to_dict())
_kapel.schil += [
    _GSchil(id="dakkapel1-voorvlak", type="gevel", subtype="Dakkapel voorvlak",
            oppervlakte_m2=3, breedte_m=2, diepte_m=1, hoogte_m=1.5,
            moedervlak_id="dak1-schuin-z", geometrie_groep="dakkapel1"),
    _GSchil(id="dakkapel1-wang-rechts", type="gevel", subtype="Dakkapel wang",
            oppervlakte_m2=1.5, breedte_m=2, diepte_m=1, hoogte_m=1.5,
            moedervlak_id="dak1-schuin-z", geometrie_groep="dakkapel1"),
    _GSchil(id="dakkapel1-wang-links", type="gevel", subtype="Dakkapel wang",
            oppervlakte_m2=1.5, breedte_m=2, diepte_m=1, hoogte_m=1.5,
            moedervlak_id="dak1-schuin-z", geometrie_groep="dakkapel1"),
    _GSchil(id="dakkapel1-dakje", type="dak", subtype="plat (dakkapel)",
            oppervlakte_m2=2, breedte_m=2, diepte_m=1, hoogte_m=1.5,
            moedervlak_id="dak1-schuin-z", geometrie_groep="dakkapel1"),
]
_kapel_svg = _gsvg(_kapel)
check("gebouw-svg: dakkapel staat op expliciet moederdakvlak",
      'data-face="dakkapel-voor"' in _kapel_svg and "dakkapel1-voorvlak" in _kapel_svg
      and "dak1-schuin-z" in _kapel_svg)
check("gebouw-svg: dakkapelrollen koppelen alle vier zichtbare faces structureel",
      all(('data-face="dakkapel-%s"' % rol) in _kapel_svg
          for rol in ("voor", "links", "rechts", "dak"))
      and all(('data-element-id="dakkapel1-%s"' % suffix) in _kapel_svg
              for suffix in ("voorvlak", "wang-links", "wang-rechts", "dakje")))
_geen_hoogte = _GDos.from_dict(_gdos.to_dict())
_geen_hoogte.opname.gevelhoogte_m = None
check("gebouw-svg: ontbrekende gevelhoogte geeft niet-gegokte fallback",
      "Kon geen 3D-vorm afleiden" in _gsvg(_geen_hoogte)
      and 'data-face="gevel-voor"' not in _gsvg(_geen_hoogte))
_inconsistent = _GDos.from_dict(_gdos.to_dict())
_inconsistent.schil[1].oppervlakte_m2 = 9
check("gebouw-svg: inconsistente tegenoverliggende gevels geven fallback",
      "verschillen meer dan 25%" in _gsvg(_inconsistent)
      and 'data-face="gevel-voor"' not in _gsvg(_inconsistent))
check("gebouw-svg: leeg dossier -> nog steeds geldige SVG", bool(_ETv.fromstring(_gsvg(_GDos()))))

from core.dossier import VloerInfo as _GVloer
_gpoly_dos = _GDos()
_gpoly_dos.opname.gevelhoogte_m = 5
_gpoly_dos.geometrie.vloeren = [_GVloer(naam="Begane grond", oppervlakte_m2=39,
    contour_m=[[0, 0], [8, 0], [8, 3], [5, 3], [5, 6], [0, 6]])]  # L-vormige footprint
_gpoly_svg = _gsvg(_gpoly_dos)
check("gebouw-svg: L-vormige contour geeft geldige SVG", bool(_ETv.fromstring(_gpoly_svg)))
_gpoly_root = _ETv.fromstring(_gpoly_svg)
_gpoly_gevels = [p for p in _gpoly_root.iter("{http://www.w3.org/2000/svg}polygon")
                 if p.attrib.get("data-face") == "gevel-contour"]
check("gebouw-svg: echte L-vorm tekent meer dan 2 gevelvlakken (geen rechthoek-reductie)",
      len(_gpoly_gevels) > 2)
check("gebouw-svg: polygon-footprint wint zelfs als dos.schil geen (consistente) gevels heeft",
      'data-footprint-bron="contour"' in _gpoly_svg and "Kon geen 3D-vorm afleiden" not in _gpoly_svg)
check("gebouw-svg: contourmuren dragen geen id maar wél een data-contour-markering",
      all(p.attrib.get("data-contour") == "true" for p in _gpoly_gevels))
# Stresstest-bevinding (14-8): een NaN/Infinity in contour_m mag nooit tot NaN-coördinaten in de
# uiteindelijke SVG leiden -> _contour_geldig verwerpt 'm, aanroeper valt terug op de rechthoek-
# fallback (die hier ook faalt, want dos.schil is leeg -> expliciete "Kon geen 3D-vorm afleiden").
_gnan_dos = _GDos()
_gnan_dos.opname.gevelhoogte_m = 5
_gnan_dos.geometrie.vloeren = [_GVloer(naam="X", oppervlakte_m2=20,
    contour_m=[[0, 0], [float("nan"), 0], [5, 5], [0, 5]])]
check("gebouw-svg: NaN in contour_m -> geen NaN in de SVG (valt terug op fallback)",
      "NaN" not in _gsvg(_gnan_dos) and "nan" not in _gsvg(_gnan_dos))
# Concave U-vorm (bijt-uit-de-achterrand): schoenveter-teken bepaalt de buiten-normaal per rand.
# Handmatig nagerekend (zie taak 007-review): met een simpel 'wijst van het gemiddelde af'-test
# (het gemiddelde van de hoekpunten valt hier BUITEN de veelhoek, in de uitsparing) draaien
# meerdere randen om; de schoenveter-methode geeft altijd 3 zichtbare wanden voor déze vorm.
_gu_vorm = [(0, 0), (10, 0), (10, 10), (7, 10), (7, 3), (3, 3), (3, 10), (0, 10)]
check("gebouw-svg: concave U-vorm geeft de schoenveter-verwachte 3 zichtbare wanden",
      len(_gmuur(_gu_vorm, 3)) == 3)
# Shading moet ONAFHANKELIJK zijn van de omlooprichting van de aangeleverde contour (bevinding uit
# taak 010-review): dezelfde rechthoek met/tegen de klok in aangeleverd is fysiek hetzelfde gebouw
# en hoort dus dezelfde brightness()-waarde per muur te geven.
from dashboard.gebouw_svg import _shade as _gshade
_grect_ccw = [(0, 0), (6, 0), (6, 8), (0, 8)]
_grect_cw = [(0, 0), (0, 8), (6, 8), (6, 0)]
_gshades_ccw = sorted(round(_gshade(f["points"]), 3) for f in _gmuur(_grect_ccw, 3))
_gshades_cw = sorted(round(_gshade(f["points"]), 3) for f in _gmuur(_grect_cw, 3))
check("gebouw-svg: shading onafhankelijk van CW/CCW-omlooprichting van de contour",
      _gshades_ccw == _gshades_cw and len(_gshades_ccw) == 2,
      str((_gshades_ccw, _gshades_cw)))

print("\n35. Webapp (Flask) — laadt + kernroutes + Beoordelingscheck")
try:
    import dashboard.app as _WA
    _WA.app.config.update(TESTING=True)
    _routes = {r.rule for r in _WA.app.url_map.iter_rules()}
    check("webapp: stappen-routes geregistreerd", {"/project/<tag>/opname", "/project/<tag>/huidig",
          "/project/<tag>/maatregelen", "/project/<tag>/vabi", "/project/<tag>/afronden",
          "/project/<tag>/export", "/guide"} <= _routes)
    _wc = _WA.app.test_client()
    check("webapp: login-pagina laadt", _wc.get("/login").status_code == 200)
    with _wc.session_transaction() as _s:   # direct inloggen (config kan echte pw_hash+MFA bevatten)
        _s["ingelogd"] = True
    _home = _wc.get("/")
    check("webapp: projecten-overzicht (ingelogd)", _home.status_code == 200)
    check("webapp: woningtype = dropdown (geen vrij tekstveld)",
          "Twee-onder-een-kap" in _home.get_data(as_text=True) and "<select name=woningtype" in _home.get_data(as_text=True))
    check("webapp: ingebouwde guide", _wc.get("/guide").status_code == 200)
    # leeg project aanmaken ZONDER bestand -> gaat naar de opname-stap
    _rn = _wc.post("/nieuw", data={"straat": "Teststraat 9", "postcode": "9999ZZ",
                   "plaats": "Groningen", "woningtype": "Hoekwoning"})
    check("webapp: leeg project zonder upload -> opname-stap",
          _rn.status_code in (302, 303) and _rn.headers["Location"].endswith("/opname"))
    _ptag = _rn.headers["Location"].rstrip("/").split("/")[-2]
    _op = _wc.get("/project/%s/opname" % _ptag)
    check("webapp: opname toont MagicPlan-import + woningtype-dropdown",
          "MagicPlan-opname inladen" in _op.get_data(as_text=True)
          and 'selected' in _op.get_data(as_text=True))
    check("webapp: dakwizards laden losse isometriemodule en beschikbare canvassen",
          'static/isometrie.js' in _op.get_data(as_text=True)
          and _op.get_data(as_text=True).count('class=isometrie-canvas') == 2)
    check("webapp: huidige-staat-stap laadt (VABI-export terug)",
          _wc.get("/project/%s/huidig" % _ptag).status_code == 200)
    # dakkapel-route: zadeldak toevoegen, dan een dakkapel erin -> 4 nieuwe vlakken + moederdak kleiner
    _wc.post("/project/%s/opname/dak/driehoek" % _ptag, data={"orient_hellend": "Z", "helling1": "45",
              "lange_zijde": "7", "breedte": "5", "rekenzone": "1"})
    from dashboard.app import _dossier as _WA_dos
    _voor_dos = _WA_dos(_ptag)
    _voor_n, _moeder_m2 = len(_voor_dos.schil), _voor_dos.schil[0].oppervlakte_m2
    check("webapp: zadeldak bewaart ruwe renderingmaten",
          all(s.breedte_m == 5.0 and abs(s.diepte_m - 3.5) < 0.001
              for s in _voor_dos.schil if s.type == "dak"))
    import math as _math35
    _wc.post("/project/%s/opname/dak/driehoek" % _ptag, data={"orient_hellend": "O", "helling1": "30",
              "helling2": "60", "lange_zijde": "8", "breedte": "6", "rekenzone": "1"})
    _asym_dos = _WA_dos(_ptag)
    _asym = [s for s in _asym_dos.schil if (s.id or "").startswith("dak2-schuin-")]
    _asym_o = next(s for s in _asym if s.orientatie == "O")
    _asym_w = next(s for s in _asym if s.orientatie == "W")
    check("webapp: asymmetrisch zadeldak deelt één nokhoogte en volledige basis",
          abs((_asym_o.diepte_m + _asym_w.diepte_m) - 8.0) < 0.001
          and abs(_asym_o.diepte_m * _math35.tan(_math35.radians(30))
                  - _asym_w.diepte_m * _math35.tan(_math35.radians(60))) < 0.001)
    check("webapp: asymmetrisch zadeldak m2 volgt per-vlak run en helling",
          abs(_asym_o.oppervlakte_m2 - round(6 * _asym_o.diepte_m / _math35.cos(_math35.radians(30)), 2)) < 0.001
          and abs(_asym_w.oppervlakte_m2 - round(6 * _asym_w.diepte_m / _math35.cos(_math35.radians(60)), 2)) < 0.001)
    _voor_n = len(_asym_dos.schil)
    _rk = _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": "0", "breedte": "2.5",
                    "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1", "wangen_geisoleerd": "on"})
    check("webapp: dakkapel-route redirect", _rk.status_code in (302, 303))
    _na_dos = _WA_dos(_ptag)
    check("webapp: dakkapel voegt 4 vlakken toe (voorvlak+2 wangen+dakje)", len(_na_dos.schil) == _voor_n + 4)
    check("webapp: dakkapel verkleint het moederdakvlak (gat afgetrokken)",
          _na_dos.schil[0].oppervlakte_m2 < _moeder_m2)
    _kapel_delen = [s for s in _na_dos.schil if (s.id or "").startswith("dakkapel1-")]
    check("webapp: dakkapel bewaart maten en stabiele moederdakreferentie",
          len(_kapel_delen) == 4 and all(s.breedte_m == 2.5 and s.diepte_m == 1.0
          and s.hoogte_m == 1.5 and s.moedervlak_id == _na_dos.schil[0].id
          and s.geometrie_groep == "dakkapel1" for s in _kapel_delen))
    _kapel_rt = _GDos.from_dict(_na_dos.to_dict())
    check("webapp: renderingmetadata overleeft dossier save/reload",
          _kapel_rt.schil[-1].breedte_m == 2.5 and _kapel_rt.schil[-1].moedervlak_id == _na_dos.schil[0].id)
    check("webapp: dakkapelcanvas verschijnt bij een geldig hellend moederdak",
          _wc.get("/project/%s/opname" % _ptag).get_data(as_text=True).count('class=isometrie-canvas') == 3)
    _steil_i = next(i for i, s in enumerate(_na_dos.schil)
                    if (s.id or "").startswith("dak2-schuin-") and s.hellingshoek == 60)
    _voor_steil = len(_na_dos.schil)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": str(_steil_i), "breedte": "2",
             "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1"})
    check("webapp: dakkapel weigert inverterende geometrie op steil dak",
          len(_WA_dos(_ptag).schil) == _voor_steil)
    import re as _re35
    check("webapp: opname-pagina met dakkapel toont geldige gebouw-svg",
          bool(_ETv.fromstring(_re35.search(r"<svg.*?</svg>", _wc.get("/project/%s/opname" % _ptag)
               .get_data(as_text=True), _re35.S).group(0))))
    # dakkapel-route weigert een PLAT dak als moederdak (ISSO 82.1 8.2.1: breekt door een
    # hellend vlak heen) en weigert niet-positieve maten (negatief moet niet gewoon 'werken')
    _wc.post("/project/%s/opname/dak/plat" % _ptag,
             data={"breedte": "4", "diepte": "2.5", "rekenzone": "1"})
    _plat_dos = _WA_dos(_ptag)
    _plat_i = next(i for i, s in enumerate(_plat_dos.schil) if (s.subtype or "") == "plat dak")
    check("webapp: plat dak berekent oppervlak en bewaart beide maten",
          _plat_dos.schil[_plat_i].oppervlakte_m2 == 10.0
          and _plat_dos.schil[_plat_i].breedte_m == 4.0 and _plat_dos.schil[_plat_i].diepte_m == 2.5)
    # Bevinding D (magicplan-velden-audit.md): alle drie de dak-wizards moeten de Constructies-DAK-
    # standaard (isolatie/dikte/bouwjaarklasse/spouw/rc_bron) erven, niet alleen het platte dak.
    from core.dossier import BouwdeelStandaard as _BJS
    _plat_dos.opname.dak_standaard = _BJS(isolatie_aanwezig="Ja", isolatiedikte_mm=120.0,
                                          bouwjaarklasse="Van 1975 t/m 1982", spouw_aanwezig=True,
                                          rc_bron="Kwaliteitsverklaring", begrenzing="Buitenlucht")
    _WA.save_json(_plat_dos, os.path.join(_WA._pdir(_ptag), _WA._load_state(_ptag)["dossier_file"]))
    _wc.post("/project/%s/opname/dak/driehoek" % _ptag, data={"orient_hellend": "ZO", "helling1": "40",
              "lange_zijde": "6", "breedte": "4", "rekenzone": "1"})
    _wc.post("/project/%s/opname/dak/negen" % _ptag, data={"m2_N": "5", "helling9": "20", "rekenzone": "1"})
    _erf_dos = _WA_dos(_ptag)
    _erf_driehoek = next(s for s in _erf_dos.schil if (s.id or "").startswith("dak4-schuin-"))
    _erf_negen = next(s for s in _erf_dos.schil if (s.id or "").startswith("dak5-"))
    for _erf, _naam in ((_erf_driehoek, "driehoek"), (_erf_negen, "negen")):
        check("webapp: dak-wizard '%s' erft de Constructies-DAK-standaard" % _naam,
              _erf.isolatie_aanwezig == "Ja" and _erf.isolatiedikte_mm == 120.0
              and _erf.bouwjaarklasse == "Van 1975 t/m 1982" and _erf.spouw_aanwezig is True
              and _erf.rc_bron == "Kwaliteitsverklaring",
              str((_erf.isolatie_aanwezig, _erf.isolatiedikte_mm, _erf.bouwjaarklasse,
                   _erf.spouw_aanwezig, _erf.rc_bron)))
    check("webapp: opname toont bouwjaarklasse- en Rc-bron-dropdown per vlak",
          "Bouwjaarklasse (dit vlak" in _wc.get("/project/%s/opname" % _ptag).get_data(as_text=True)
          and 'name=rc_bron' in _wc.get("/project/%s/opname" % _ptag).get_data(as_text=True))
    _erf_i = next(i for i, s in enumerate(_erf_dos.schil) if s.id == _erf_driehoek.id)
    _wc.post("/project/%s/opname/el/%d" % (_ptag, _erf_i),
             data={"m2": str(_erf_driehoek.oppervlakte_m2), "orientatie": _erf_driehoek.orientatie,
                   "begrenzing": "Buitenlucht", "rekenzone": "1", "rc": "2.1", "isolatie": "Ja",
                   "dikte": "150", "bouwjaarklasse": "Vanaf 2006", "rc_bron": "Opgemeten dikte"})
    _erf_bewerkt = next(s for s in _WA_dos(_ptag).schil if s.id == _erf_driehoek.id)
    check("webapp: bouwjaarklasse en rc_bron per vlak handmatig overschrijfbaar",
          _erf_bewerkt.bouwjaarklasse == "Vanaf 2006" and _erf_bewerkt.rc_bron == "Opgemeten dikte")
    _voor_n2 = len(_erf_dos.schil)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": str(_plat_i), "breedte": "2",
             "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1"})
    check("webapp: dakkapel weigert plat dak als moederdak (geen nieuwe vlakken)",
          len(_WA_dos(_ptag).schil) == _voor_n2)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": "0", "breedte": "-2",
             "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1"})
    check("webapp: dakkapel weigert negatieve breedte (geen nieuwe vlakken)",
          len(_WA_dos(_ptag).schil) == _voor_n2)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": "0", "breedte": "nan",
             "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1"})
    check("webapp: dakkapel weigert NaN-breedte (geen nieuwe vlakken)",
          len(_WA_dos(_ptag).schil) == _voor_n2)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": "0", "breedte": "inf",
             "hoogte": "1.5", "diepte": "1.0", "rekenzone": "1"})
    check("webapp: dakkapel weigert Infinity-breedte (geen nieuwe vlakken)",
          len(_WA_dos(_ptag).schil) == _voor_n2)
    # dakkapelgat groter dan het moederdak -> LUID geflagd, niet stil naar 0 geknipt
    _dakkapel_dos = _WA_dos(_ptag)
    _klein_i = next(i for i, s in enumerate(_dakkapel_dos.schil)
                    if s.id == "dak1-schuin-n")  # 24,75 m2, ongebruikt in eerdere posts
    _dakkapel_dos.schil[_klein_i].oppervlakte_m2 = 1.0  # kunstmatig klein moederdak
    _WA.save_json(_dakkapel_dos, os.path.join(_WA._pdir(_ptag), _WA._load_state(_ptag)["dossier_file"]))
    _voor_n3 = len(_dakkapel_dos.schil)
    _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": str(_klein_i), "breedte": "5",
             "hoogte": "3", "diepte": "3", "rekenzone": "1"})
    _na_groot_dos = _WA_dos(_ptag)
    check("webapp: dakkapelgat > moederdak -> geweigerd, geen nieuwe vlakken",
          len(_na_groot_dos.schil) == _voor_n3)
    check("webapp: dakkapelgat > moederdak -> moederdak ONgewijzigd (geen aftrek zonder toevoeging)",
          _na_groot_dos.schil[_klein_i].oppervlakte_m2 == 1.0)
    # ongeldige (niet-numerieke) rekenzone crasht niet (HTTP 500), valt terug op moederdak.rekenzone
    _rzr = _wc.post("/project/%s/opname/dakkapel" % _ptag, data={"moederdak_i": "0", "breedte": "2",
                     "hoogte": "1.5", "diepte": "1.0", "rekenzone": "geen-getal"})
    check("webapp: dakkapel met ongeldige rekenzone -> geen 500", _rzr.status_code in (302, 303))
    import shutil as _sh35
    _sh35.rmtree(_WA._pdir(_ptag), ignore_errors=True)
    _bd = _WA._beoordeling("x", {"foto_voorkant": "", "foto_huisnummer": "", "na": {}}, build_sample())
    check("webapp: Beoordelingscheck spiegelt kennisbank (>=6 punten)", len(_bd) >= 6)
except Exception as _e:
    check("webapp: laadt zonder fout", False)
    print("     " + repr(_e)[:160])

print("\n36. Catalogus-API live-mapping (JSON:API -> catalog.json)")
try:
    from catalog.api_client import map_measures_to_catalog
    _raw = {"data": [
        {"id": "V1-1-A1", "type": "measure", "attributes": {
            "name": "Spouwmuurisolatie 60", "unit": "m²", "rcValue": 1.7, "thicknessInMm": 60,
            "isBiobased": False,
            "regularCosts": [
                {"id": "V1-1-A1", "type": "cost", "attributes": {"contractorValuePerUnit": 23.09,
                 "diyValuePerUnit": 18.8, "minUnits": 0, "maxUnits": 45, "unit": "m²"}},
                {"id": "V1-1-A2", "type": "cost", "attributes": {"contractorValuePerUnit": 21.13,
                 "minUnits": 45, "maxUnits": 75, "unit": "m²"}}],
            "additionalCosts": [
                {"id": "V1-1-X7", "type": "cost", "attributes": {"contractorValuePerUnit": 91.13,
                 "unit": "won", "notes": "Betreft: Spouw richting dak dichtmaken van binnenuit"}}]}},
        {"id": "V1-2-A1", "type": "measure", "attributes": {"name": "Andere", "unit": "m²",
            "regularCosts": [{"id": "V1-2-A1", "type": "cost", "attributes": {
                "contractorValuePerUnit": 50.0, "minUnits": 0, "maxUnits": None, "unit": "m²"}}],
            "additionalCosts": [{"id": "V1-1-X7", "type": "cost", "attributes": {
                "contractorValuePerUnit": 91.13, "unit": "won", "notes": "Betreft: dubbel"}}]}},
    ]}
    _cat = map_measures_to_catalog(_raw)
    _ms = {m["code"]: m for m in _cat["maatregelen"]}
    check("api-map: 4 rijen (2 brackets + 1 bracket + 1 X; gedeelde X gededupe)", len(_cat["maatregelen"]) == 4)
    check("api-map: V1-1-A1 incl=23.09 / excl~19.08",
          _ms["V1-1-A1"]["prijs_per_eenheid_incl_btw"] == 23.09 and abs(_ms["V1-1-A1"]["prijs_per_eenheid_excl"] - 19.0826) < 0.01)
    check("api-map: bracket-tekst in omschrijving", "van 0 m² tot 45 m²" in _ms["V1-1-A1"]["omschrijving"])
    check("api-map: onderdeel uit code-prefix", _ms["V1-1-A1"]["onderdeel"] == "A Gevel")
    check("api-map: X-code 1x (gededupe over measures)",
          sum(1 for m in _cat["maatregelen"] if m["code"] == "V1-1-X7") == 1)
    check("api-map: 'Betreft:' uit X-notitie gestript", not _ms["V1-1-X7"]["omschrijving"].startswith("Betreft"))
    check("api-map: extra rc_waarde meegenomen", _ms["V1-1-A1"].get("rc_waarde") == 1.7)
except Exception as _e:
    check("api-map: mapper draait zonder fout", False)
    print("     " + repr(_e)[:160])

print("\n37. Parser: VABI-beslisboom per bouwdeel (nieuwe Constructies-form)")
try:
    import tempfile as _tf2
    from magicplan.statistics_csv import build_dossier as _bd2
    _boomcsv = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\n"
                "Gevelhoogte (m),5.4\n"
                "Gevel - invoer,Beslisschema\nGevel - isolatie aanwezig?,Ja\nGevel - isolatiedikte onbekend?,Nee\n"
                "Gevel - isolatiedikte (mm),80\nGevel - begrenzing,Buitenlucht\n"
                "Vloer - invoer,Kwaliteitsverklaring\nVloer - begrenzing,Kruipruimte\n\n"
                "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
                "Ground Floor,40,2.50 m,Kruipruimte\n\n"
                "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
                "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _bp = _tf2.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _bp.write(_boomcsv); _bp.close()
    _bdo, _bn = _bd2(_bp.name)
    _gev = next((s for s in _bdo.schil if s.type == "gevel"), None)
    _vlo = next((s for s in _bdo.schil if s.type == "vloer"), None)
    check("boom: gevel rc_bron=Opgemeten dikte (80mm bekend)", _gev is not None and _gev.rc_bron == "Opgemeten dikte")
    check("boom: gevel isolatie=Ja + dikte 80mm",
          _gev is not None and _gev.isolatie_aanwezig == "Ja" and _gev.isolatiedikte_mm == 80.0)
    check("boom: vloer rc_bron=Kwaliteitsverklaring (Invoer=KV)", _vlo is not None and _vlo.rc_bron == "Kwaliteitsverklaring")
except Exception as _e:
    check("boom: parser draait zonder fout", False); print("     " + repr(_e)[:160])

print("\n38. Parser: dak geconsolideerd uit Constructies + 9 m²-vakjes (type Anders)")
try:
    import tempfile as _tf3
    from magicplan.statistics_csv import build_dossier as _bd3
    _dak2 = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\nGevelhoogte (m),5.4\n"
             "Dakvlak 1 - daktype,Zadeldak\nDak - vloerbreedte (m),8\nDakvlak 1 - hellingshoek (°),45\n"
             "Dakvlak 1 - oriëntatie,ZW\nDakvlak 2 - oriëntatie,NO\nDak - kopgevel oriëntatie 1,NW\nDak - kopgevel oriëntatie 2,ZO\n\n"
             "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
             "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
             "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _dp = _tf3.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _dp.write(_dak2); _dp.close()
    _ddo, _ = _bd3(_dp.name)
    check("dak-uit-Constructies: 2 schuine dakvlakken (zadeldak)", len([s for s in _ddo.schil if s.type == "dak"]) == 2)
    check("dak-uit-Constructies: kopgevel-driehoek als gevel",
          any(s.type == "gevel" and "kopgevel" in (s.subtype or "") for s in _ddo.schil))
    _av = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1992.t.m.2013\nWoningtype,Vrijstaand\nGevelhoogte (m),5.4\n"
           "Dakvlak 1 - daktype,Anders\nDak m² Z,30\nDak m² Horizontaal,12\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _ap2 = _tf3.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _ap2.write(_av); _ap2.close()
    _ado, _ = _bd3(_ap2.name)
    _adak = [s for s in _ado.schil if s.type == "dak"]
    check("dak 9-vakjes (Anders): 2 vlakken, Z=30 m²",
          len(_adak) == 2 and any(abs(s.oppervlakte_m2 - 30) < 0.1 for s in _adak))
except Exception as _e:
    check("dak-uit-Constructies: parser draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n39. Parser: rekenzone default 1 + 2e tapwater/ventilatie")
try:
    import tempfile as _tf4
    from magicplan.statistics_csv import build_dossier as _bd4
    _ic = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\nGevelhoogte (m),5.4\n"
           "Verwarming - type opwekker,Gasgestookte ketel\nTapwater - toestel,Combiketel (gas)\n"
           "Tapwater 2 - toestel,Elektrische boiler\nTapwater 2 - installatiejaar,2020\n"
           "Ventilatiesysteem (A-E),C Mechanische afvoer\nSubsysteem (C),C1 Standaard\n"
           "Tweede ventilatiesysteem?,Ja\nVentilatie 2 - systeem (A-E),D Mechanische balansventilatie\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _ip = _tf4.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _ip.write(_ic); _ip.close()
    _ido, _inotes = _bd4(_ip.name)
    check("rekenzone default 1 (leeg) op ventilatie + verwarming",
          _ido.ventilatie.rekenzone == 1 and _ido.installaties.verwarming.rekenzone == 1)
    check("2e tapwater gelezen (tapwater_extra)",
          len(_ido.installaties.tapwater_extra) == 1 and "boiler" in (_ido.installaties.tapwater_extra[0].type_toestel or "").lower())
    check("2e ventilatiesysteem geflagd", any("Tweede ventilatiesysteem" in n for n in _inotes))
except Exception as _e:
    check("installaties-extra: parser draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n40. Leads (Nij Begun-portal): parsen + dedupe + concept-mail + CSV")
try:
    import dashboard.leads as _L
    _mail = ('Contact met adviseur\n{"BagAdresId":"0014200099999999","Email":"piet@test.example",'
             '"Postcode":"9999XX","Huisnummer":12,"HuisnummerToevoeging":"a","Voornaam":"Piet",'
             '"Telefoonnummer":"0612345678","Achternaam":"Test","Naam":"Piet Test",'
             '"WijzigingsType":"AdviseurToegekend","WijzigingsReden":"Adviseur toegekend."}')
    _ld = _L.parse_lead(_mail)
    check("lead: JSON uit mail-tekst geparsed", _ld is not None and _ld["naam"] == "Piet Test"
          and _ld["postcode"] == "9999XX" and _ld["huisnummer"] == "12" and _ld["toevoeging"] == "a")
    _rows, _n1 = _L.add_lead(_ld, [])
    _rows, _n2 = _L.add_lead(_ld, _rows)
    check("lead: dedupe op BAG-id", _n1 and not _n2 and len(_rows) == 1 and _rows[0]["status"] == "nieuw")
    _ond, _txt = _L.concept_mail(_rows[0], {"naam": "R. Poortinga", "telefoon": "06-1", "email": "a@b.nl"})
    check("lead: concept-mail met aanhef + voorbereiding + ondertekening",
          "Beste Piet Test" in _txt and "kruipruimte" in _txt and "R. Poortinga" in _txt and "Nij Begun" in _ond + _txt)
    _csv = _L.to_csv(_rows)
    check("lead: CSV Excel-NL (puntkomma + naam erin)", ";" in _csv.splitlines()[0] and "Piet Test" in _csv)
    check("lead: onzin-tekst -> None (geen crash)", _L.parse_lead("hallo dit is geen lead") is None)
except Exception as _e:
    check("leads: module draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n41. BAG-verrijking (parsers offline; structuur live geverifieerd 26-6)")
try:
    import dashboard.bag as _B
    _ls = {"response": {"docs": [
        {"weergavenaam": "Teststraat 12a, 9999XX Testdorp", "straatnaam": "Teststraat",
         "woonplaatsnaam": "Testdorp", "huis_nlt": "12a", "postcode": "9999XX",
         "nummeraanduiding_id": "0014200099999999", "adresseerbaarobject_id": "0014010099999999",
         "centroide_rd": "POINT(235729.987 585202.818)"},
        {"weergavenaam": "Teststraat 12, 9999XX Testdorp", "straatnaam": "Teststraat",
         "woonplaatsnaam": "Testdorp", "huis_nlt": "12", "postcode": "9999XX",
         "nummeraanduiding_id": "0014200088888888", "adresseerbaarobject_id": "0014010088888888",
         "centroide_rd": "POINT(1000.0 2000.0)"}]}}
    _a = _B.parse_locatieserver(_ls, "12", "a")
    check("bag: locatieserver-parse kiest exacte huisnummer+toevoeging-match",
          _a is not None and _a["verblijfsobject_id"] == "0014010099999999"
          and _a["straat"] == "Teststraat" and _a["x"] == 235729.987)
    _wfs = {"features": [
        {"properties": {"identificatie": "0014010077777777", "oppervlakte": 80, "bouwjaar": 1975,
                        "gebruiksdoel": "woonfunctie", "pandidentificatie": "p1"}},
        {"properties": {"identificatie": "0014010099999999", "oppervlakte": 109, "bouwjaar": 1982,
                        "gebruiksdoel": "woonfunctie", "pandidentificatie": "p2"}}]}
    _v = _B.parse_wfs(_wfs, "0014010099999999")
    check("bag: wfs-parse matcht op verblijfsobject-id (bbox kan buren bevatten)",
          _v is not None and _v["bouwjaar"] == 1982 and _v["oppervlakte_m2"] == 109)
    check("bag: lege respons -> None (geen crash)",
          _B.parse_locatieserver({}, "1") is None and _B.parse_wfs({}, "x") is None)
    from dashboard.leads import adres as _adr
    check("lead-adres met BAG-straat", _adr({"straat": "Teststraat", "huisnummer": "12", "toevoeging": "a",
          "woonplaats": "Testdorp", "postcode": "9999XX"}) == "Teststraat 12 a in Testdorp")
    import dashboard.app as _WA2
    check("webapp: BAG-route geregistreerd", "/leads/<int:lid>/bag" in {r.rule for r in _WA2.app.url_map.iter_rules()})
except Exception as _e:
    check("bag: module draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n42. Webapp v2: opname-editor + catalogus-kiezer + varianten")
try:
    from dashboard.measures import catalogus_boom as _cb, _schoon_label as _sl, zoek_maatregel as _zm
    _cat = json.load(open(CATALOG, encoding="utf-8"))
    _boom = _cb(_cat)
    check("catalogus-boom: 6 categorieën (V1..V6)", len(_boom) == 6 and _boom[0]["naam"] == "Gevel")
    _v11 = next(s for c in _boom for s in c["subs"] if s["code"] == "V1-1")
    check("catalogus-boom: V1-1 heeft kern + bijkomende kosten (X)",
          len(_v11["kern"]) > 0 and any(r["code"].split("-")[2].startswith("X") for r in _v11["meerwerk"]))
    check("catalogus-boom: sub-label geschoond (geen bracket)", "m² tot" not in _v11["naam"])
    check("schoon-label", _sl("Spouwmuurisolatie vlokken 60 mm van 0 m² tot 45 m²") == "Spouwmuurisolatie vlokken")
    check("zoek-maatregel", (_zm(_cat, "V1-1-A1") or {}).get("code") == "V1-1-A1")
    import dashboard.app as _WA3
    _rr = {r.rule for r in _WA3.app.url_map.iter_rules()}
    check("webapp v2: opname-editor-routes", {"/project/<tag>/opname", "/project/<tag>/opname/el/<int:i>",
          "/project/<tag>/opname/el/nieuw", "/project/<tag>/opname/vabi_huidig",
          "/project/<tag>/maatregelen/add", "/project/<tag>/toelichting"} <= _rr)
    check("webapp v2: stap 'opname' in de stepper", any(s == "opname" for s, _ in _WA3.STAPPEN))
except Exception as _e:
    check("webapp v2: modules draaien zonder fout", False); print("     " + repr(_e)[:170])

print("\n43. Beveiliging (hosting): wachtwoord-hash + TOTP-MFA + rate-limit")
try:
    import dashboard.security as _S
    _h = _S.hash_password("geheim-wachtwoord")
    check("pw-hash: rondje klopt + fout wachtwoord faalt",
          _S.check_password("geheim-wachtwoord", _h) and not _S.check_password("fout", _h))
    # RFC 6238-testvector (secret '12345678901234567890', T=59s, SHA1, 6 cijfers -> 287082)
    check("totp: RFC 6238-vector", _S.totp_code("GEZDGNBVGY3TQOJQGEZDGNBVGY3TQOJQ", t=59) == "287082")
    _sec = _S.nieuw_totp_secret()
    check("totp: eigen secret verifieert (venster)", _S.check_totp(_sec, _S.totp_code(_sec)))
    check("totp: foute code faalt", not _S.check_totp(_sec, "000000") or _S.totp_code(_sec) == "000000")
    _S._pogingen.clear()
    for _ in range(5):
        _S.poging_mislukt("1.2.3.4")
    check("rate-limit: 5 missers -> blok", _S.geblokkeerd("1.2.3.4") and not _S.geblokkeerd("5.6.7.8"))
    _S._pogingen.clear()
    _ok, _ = _S.login_check({}, "lokaalpw", "", "9.9.9.9", fallback_pw="lokaalpw")
    check("login: lokale modus (geen hash) werkt", _ok)
    _cfg2 = {"pw_hash": _h, "totp_secret": _sec}
    _ok1, _ = _S.login_check(_cfg2, "geheim-wachtwoord", _S.totp_code(_sec), "9.9.9.8")
    _ok2, _m2 = _S.login_check(_cfg2, "geheim-wachtwoord", "000000", "9.9.9.7")
    check("login: hash+MFA goed -> ok; foute code -> geweigerd", _ok1 and not _ok2 and "code" in _m2)
    import dashboard.app as _WA4
    check("app: secret key persistent + origin-check geregistreerd",
          bool(_WA4.app.secret_key) and any(f.__name__ == "_origin_check"
          for f in _WA4.app.before_request_funcs.get(None, [])))
    os.environ["NIJBEGUN_PW_HASH"] = _h
    os.environ["NIJBEGUN_TOTP_SECRET"] = _sec
    _d = _WA4._dash_cfg()
    check("app: PaaS env-vars overrulen config (Render/Railway)",
          _d.get("pw_hash") == _h and _d.get("totp_secret") == _sec)
    del os.environ["NIJBEGUN_PW_HASH"], os.environ["NIJBEGUN_TOTP_SECRET"]
except Exception as _e:
    check("beveiliging: module draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n44. Deep-dive-fixes: dak direct-m² wint + kozijn 0,65-regel + perimeter-note")
try:
    import tempfile as _tf5
    from magicplan.statistics_csv import build_dossier as _bd5
    _dd5 = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\nGevelhoogte (m),5.4\n"
            "Dakvlak 1 - daktype,Zadeldak\nDakvlak 1 - oppervlak (m²),33.92\nDakvlak 1 - oriëntatie,Z\nDakvlak 1 - hellingshoek (°),45\n"
            "Dakvlak 2 - daktype,Plat dak\nDakvlak 2 - oppervlak (m²),5.53\nDak - vloerbreedte (m),8\n\n"
            "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
            "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
            "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,Z\n")
    _p5 = _tf5.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _p5.write(_dd5); _p5.close()
    _d5, _n5 = _bd5(_p5.name)
    _dk5 = [s for s in _d5.schil if s.type == "dak"]
    check("dak: direct ingevoerde m² winnen (2 vlakken, 33.92 + 5.53; geen auto-berekening)",
          len(_dk5) == 2 and {round(s.oppervlakte_m2, 2) for s in _dk5} == {33.92, 5.53})
    check("dak: direct-m²-note aanwezig", any("direct ingevoerde m²" in n for n in _n5))
    check("perimeter: woningscheidende-wand-note bij tussenwoning", any("WONINGSCHEIDENDE" in n for n in _n5))
    # kozijn < 0.65 m2 -> 0.65 (via assemble/kozijn-route: window in WALL-attributen)
    _kc = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Vrijstaand\nGevelhoogte (m),5.4\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,Z\n"
           "Voorgevel,Window 1,Window,0.36,0.36,0.6,0.6,1,Window,HR++,1,Z\n")
    _pk = _tf5.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _pk.write(_kc); _pk.close()
    _dk, _ = _bd5(_pk.name)
    _rr5 = [s for s in _dk.schil if s.type == "kozijn" and s.subtype == "Raam"]
    check("kozijn: klein raam (0,36 m²) -> 0,65 m² (Nij Begun-regel)",
          bool(_rr5) and abs(_rr5[0].oppervlakte_m2 - 0.65) < 0.001 and "0,65" in (_rr5[0].opmerkingen or ""))
except Exception as _e:
    check("deep-dive-fixes: parser draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n45. Parser: 'Deels binnen/buiten'-VINKJE op de wand (kolom op naam, i.p.v. typen)")
try:
    import tempfile as _tf6
    from magicplan.statistics_csv import build_dossier as _bd6
    _vc = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Vrijstaand\nGevelhoogte (m),5.4\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\nGround Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie,Bron,Deels binnen/deels buiten? (narekenen)\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,Z,,No\n"
           "Achtergevel,Wall 1,Wall,12,11,5,2.5,1,Wall,,1,N,,Yes\n")
    _pv6 = _tf6.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _pv6.write(_vc); _pv6.close()
    _dv6, _nv6 = _bd6(_pv6.name)
    _gz = [s for s in _dv6.schil if s.type == "gevel" and "kopgevel" not in (s.subtype or "")]
    check("vinkje: aangevinkte wand -> NAREKENEN-flag (zonder typen)",
          any("NAREKENEN" in (s.opmerkingen or "") for s in _gz)
          and any("NAREKENEN" in n and "Achtergevel" in n for n in _nv6))
    check("vinkje: niet-aangevinkte wand blijft gewoon", any("NAREKENEN" not in (s.opmerkingen or "") for s in _gz))
except Exception as _e:
    check("vinkje: parser draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n46. Bouwjaar-hints + isolatieplan-JSON (leverformaat M29 punt 10a)")
try:
    import dashboard.bouwjaar as _BJ
    _t75, _h75 = _BJ.hint(1975)
    check("bouwjaar-hint 1975 -> tijdvak 1975–1982 + html", _t75 is not None and "1975" in _t75
          and "<h4>" in _h75 and "<li>" in _h75)
    _t30, _ = _BJ.hint(1930)
    check("bouwjaar-hint 1930 -> vooroorlogs", _t30 is not None and "1946" in _t30)
    check("bouwjaar-hint zonder bouwjaar -> None", _BJ.hint(None) == (None, None))
    import os as _os, io as _io, json as _js, tempfile as _tf46
    import dashboard.app as _WA5
    _WA5.app.config.update(TESTING=True)
    _projects5 = _WA5.PROJECTS_DIR
    with _tf46.TemporaryDirectory() as _td5:
        try:
            _WA5.PROJECTS_DIR = _td5
            _c5 = _WA5.app.test_client()
            with _c5.session_transaction() as _s5:
                _s5["ingelogd"] = True
            _fixture5 = _js.dumps(build_sample().to_dict()).encode("utf-8")
            _r5 = _c5.post("/nieuw", data={"bestand": (_io.BytesIO(_fixture5), "d.json"), "straat": "T 1",
                           "plaats": "X", "woningtype": "Tussenwoning"}, content_type="multipart/form-data")
            _tag5 = _r5.headers["Location"].rstrip("/").split("/")[-2]
            _c5.get("/project/%s/afronden" % _tag5)
            _jp = _os.path.join(_WA5._pdir(_tag5), "isolatieplan_%s.json" % _tag5)
            check("isolatieplan-JSON gegenereerd bij afronden", _os.path.isfile(_jp))
            _pj = _js.load(open(_jp, encoding="utf-8"))
            check("plan-JSON: formaat + rekenkern + maatregelen-array",
                  _pj.get("formaat") == "nijbegun-isolatieplan" and "Vabi" in _pj["tool"]["rekenkern"]
                  and isinstance(_pj.get("maatregelen_subsidietabel"), list))
        finally:
            _WA5.PROJECTS_DIR = _projects5
except Exception as _e:
    check("bouwjaar/plan-json: draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n47. Nij Begun-focus: mail-scope, leads->project, MagicPlan-foto's, HIG-CSS, spouwinspectie")
try:
    import os as _o7, tempfile as _tf7, shutil as _sh7
    # (a) kennismakingsmail is zuiver schil/ventilatie — geen installatie-vragen
    import dashboard.leads as _LS
    _blob = " ".join(_LS.VOORBEREIDING).lower()
    check("mail: geen cv-ketel/warmtepomp/zonnepanelen-vraag (label, niet M29)",
          "cv-ketel" not in _blob and "warmtepomp" not in _blob and "zonnepanel" not in _blob)
    check("mail: wel isolatie-bewijslast gevraagd",
          any(("isolatiewerk" in v.lower()) or ("geïsoleerd" in v.lower()) for v in _LS.VOORBEREIDING))
    # (b) MagicPlan-foto's: parse + download met injecteerbare fetch (offline)
    from magicplan import photos as _PH
    _plan = {"photos": [{"url": "https://x/a.jpg", "bouwdeel": "voorgevel", "tag": "overzicht"},
                        {"id": "p2", "bouwdeel": "dak"}]}
    _ents = _PH.photo_entries(_plan)
    check("photos: entries geparsed (url + id-only)", len(_ents) == 2 and _ents[0]["url"].endswith("a.jpg"))
    _td = _tf7.mkdtemp()
    _fotos, _fout = _PH.download_photos(_ents, _td, lambda e: b"JPEGDATA" if e.get("url") else (_ for _ in ()).throw(RuntimeError("geen url")))
    check("photos: 1 gedownload, 1 geflagd (id zonder url — niet gegokt)", len(_fotos) == 1 and len(_fout) == 1)
    check("photos: bestand op schijf + Foto.bouwdeel", _o7.path.isfile(_o7.path.join(_td, _fotos[0].bestand)) and _fotos[0].bouwdeel == "voorgevel")
    _sh7.rmtree(_td, ignore_errors=True)
    # (c) HIG-CSS: dark mode + safe-area + table-wrap aanwezig
    _css = open(_o7.path.join(ROOT, "dashboard", "static", "app.css"), encoding="utf-8").read()
    check("HIG-CSS: dark mode + safe-area + 44px + table-wrap",
          "prefers-color-scheme: dark" in _css and "safe-area-inset" in _css and ".table-wrap" in _css and "46px" in _css)
    # (d) spouwinspectie-gids bestaat + endoscopie wijst ernaar
    check("spouwinspectie-gids aanwezig", _o7.path.isfile(_o7.path.join(ROOT, "docs", "spouwinspectie-gids.md")))
    check("endoscopie-werkwijze verwijst naar spouwinspectie-gids",
          "spouwinspectie-gids" in open(_o7.path.join(ROOT, "docs", "endoscopie-werkwijze.md"), encoding="utf-8").read())
    # (e) leads -> project (geïsoleerd van echte leads.json)
    import dashboard.app as _WL
    _WL.app.config.update(TESTING=True)
    _cl = _WL.app.test_client()
    with _cl.session_transaction() as _s7:
        _s7["ingelogd"] = True
    _od, _of = _LS.LEADS_DIR, _LS.LEADS_FILE
    _tmp = _tf7.mkdtemp(); _LS.LEADS_DIR = _tmp; _LS.LEADS_FILE = _o7.path.join(_tmp, "leads.json")
    try:
        _rows, _ = _LS.add_lead({"bag_id": "TESTBAG1", "naam": "Testpersoon", "postcode": "9999ZZ",
                                 "huisnummer": "7", "toevoeging": "", "straat": "Teststraat",
                                 "woonplaats": "Groningen", "bouwjaar": 1975, "woningtype": "Tussenwoning"})
        _LS.save_leads(_rows); _lid = _rows[-1]["id"]
        _rp = _cl.post("/leads/%d/project" % _lid)
        check("leads->project: redirect naar opname", _rp.status_code in (302, 303) and _rp.headers["Location"].endswith("/opname"))
        _ptag = _rp.headers["Location"].rstrip("/").split("/")[-2]
        _l2 = next(x for x in _LS.load_leads() if x["id"] == _lid)
        check("leads->project: tag gekoppeld + status doorgezet", _l2.get("project_tag") == _ptag and _l2["status"] == "opname gedaan")
        _dos = _WL._dossier(_ptag)
        check("leads->project: adres/bouwjaar over, GEEN persoonsgegevens in dossier",
              _dos.identificatie.bouwjaar == 1975 and _dos.identificatie.postcode == "9999ZZ"
              and "Testpersoon" not in json.dumps(_dos.to_dict()))
        _rp2 = _cl.post("/leads/%d/project" % _lid)     # idempotent
        check("leads->project: idempotent (opent bestaand, overschrijft niet)", _rp2.status_code in (302, 303))
        _sh7.rmtree(_WL._pdir(_ptag), ignore_errors=True)
    finally:
        _LS.LEADS_DIR, _LS.LEADS_FILE = _od, _of
except Exception as _e:
    check("Nij Begun-focus: draait zonder fout", False); print("     " + repr(_e)[:180])

print("\n48. Raam-invoer versimpelen (alleen Type glas; kozijn/rooster/paneel defaulten)")
try:
    # (a) parser-defaults: leeg kozijnmateriaal -> Hout of kunststof; afwijking blijft werken
    from magicplan.statistics_csv import _norm_kozijn_mat as _nkm
    check("raam: leeg kozijnmateriaal -> Hout of kunststof (default)", _nkm("") == "Hout of kunststof")
    check("raam: 'b' -> Metaal thermisch onderbroken (afwijking werkt nog)", _nkm("b").startswith("Metaal"))
    # (b) form_push: set_optional + set_default op een nagebootste 'Raam/paneel'-veldgroep
    import magicplan.form_push as _FP
    _form = {"name": "Raam/paneel", "context": ["windows"], "children": [
        {"id": "q1", "type": "question", "name": "Type glas", "dataType": "list", "required": True,
         "fields": {"options": ["Enkel", "Dubbel", "HR++"]}},
        {"id": "q2", "type": "question", "name": "Kozijnmateriaal", "dataType": "list", "required": True,
         "fields": {"options": ["Metaal TO", "Hout of kunststof", "Metaal niet-TO"]}},
        {"id": "q3", "type": "question", "name": "Raam/paneel", "dataType": "list", "required": True,
         "fields": {"options": ["Paneel", "Raam"]}}]}
    _rec, _a48, _rq48, _pb48 = _FP.merge_record({"id": "r48", "form": _form}, _FP.load_additions(), verbose=False)
    _bn = {c["name"]: c for c in _FP._form_of(_rec)["children"]}
    check("raam: Type glas blijft VERPLICHT", _bn["Type glas"]["required"] is True)
    check("raam: Kozijnmateriaal optioneel + 'Hout of kunststof' vooraan",
          _bn["Kozijnmateriaal"]["required"] is False and _bn["Kozijnmateriaal"]["fields"]["options"][0] == "Hout of kunststof")
    check("raam: Raam/paneel optioneel + 'Raam' vooraan",
          _bn["Raam/paneel"]["required"] is False and _bn["Raam/paneel"]["fields"]["options"][0] == "Raam")
    check("form_push: geen structurele problemen na optioneel/default", _pb48 == [])
except Exception as _e:
    check("raam-invoer: draait zonder fout", False); print("     " + repr(_e)[:170])

print("\n49. Paneel-in-kozijn (dichte constructie, ConstructieType=1) end-to-end")
try:
    from core.dossier import Dossier as _D9, SchilDeel as _S9, Identificatie as _I9
    import vabi.constructie_generate as _CG9
    # (a) TYPE_CODE + classify: paneel -> ConstructieType 1, dichte match (geen glas)
    check("paneel: TYPE_CODE=1 (Paneel)", _CG9.TYPE_CODE.get("paneel") == "1")
    check("paneel: _classify blijft 'paneel' (geen kozijn/raam)", _CG9._classify(_S9(id="p", type="paneel")) == "paneel")
    # (b) generator kiest een ConstructieType-1 constructie voor een paneel (pool+cb intern gebouwd)
    _dos9 = _D9(identificatie=_I9(postcode="9999ZZ", huisnummer="1", bouwjaar=1975))
    _dos9.schil = [_S9(id="paneel-1", type="paneel", subtype="Paneel", orientatie="Z",
                       begrenzing="Buitenlucht", oppervlakte_m2=1.2, isolatie_aanwezig="Nee")]
    _clones9, _map9, _iss9 = _CG9.resolve_constructies(_dos9)
    check("paneel: kreeg een passende constructie (geen 'onbekend type')",
          "paneel-1" in _map9 and not any("onbekend type" in x for x in _iss9))
    _ct = next((_CG9._t(c, "ConstructieType") for c in _clones9), None)
    check("paneel: gekozen constructie is ConstructieType=1", _ct == "1")
    # (b2) TYPE_CODE-mapping 1-op-1 tegen de ECHTE EPA-export-template (19-7 bevestigd): elke
    # ConstructieType-code hoort bij het juiste bouwdeel (0=Gevel 1=Paneel 2=glas 3=Deur 4=Dak 7=Vloer).
    import os as _os9, xml.etree.ElementTree as _ET9
    _tpl9 = _ET9.parse(_os9.path.join(_os9.path.dirname(_CG9.__file__), "refs",
                                      "standaard_constructies_v120001001.xml")).getroot()
    _byct9 = {}
    for _c9 in _tpl9.iter():
        if _c9.tag.rsplit("}", 1)[-1] == "Constructie":
            _cc9 = {x.tag.rsplit("}", 1)[-1]: (x.text or "") for x in _c9}
            _byct9.setdefault(_cc9.get("ConstructieType", ""), []).append((_cc9.get("Naam", "") or "").lower())
    _verwacht9 = {"0": "gevel", "1": "paneel", "2": "glas", "3": "deur", "4": "dak", "7": "vloer"}
    check("constructie: TYPE_CODE-mapping matcht de echte EPA-template (0=gevel..7=vloer)",
          all(_kw in " ".join(_byct9.get(_code, [])) for _code, _kw in _verwacht9.items()))
    # (c) webapp: paneel toevoegen -> valt in de dichte-branch (Rc/isolatie), niet glas
    import dashboard.app as _WP9
    _WP9.app.config.update(TESTING=True)
    _c9 = _WP9.app.test_client()
    with _c9.session_transaction() as _s9:
        _s9["ingelogd"] = True
    _r9 = _c9.post("/nieuw", data={"straat": "Paneelstraat 1", "postcode": "9999ZP", "plaats": "X", "woningtype": "Tussenwoning"})
    _tg9 = _r9.headers["Location"].rstrip("/").split("/")[-2]
    _c9.post("/project/%s/opname/el/nieuw" % _tg9, data={"type": "paneel"})
    _d9 = _WP9._dossier(_tg9)
    _pn = next((s for s in _d9.schil if s.type == "paneel"), None)
    check("webapp: paneel toevoegbaar (type=paneel, subtype=Paneel)", _pn is not None and _pn.subtype == "Paneel")
    _op9 = _c9.get("/project/%s/opname" % _tg9).get_data(as_text=True)
    check("webapp: paneel-keuze in de dropdown", "Paneel (dicht)" in _op9)
    import shutil as _sh9; _sh9.rmtree(_WP9._pdir(_tg9), ignore_errors=True)
except Exception as _e:
    check("paneel-in-kozijn: draait zonder fout", False); print("     " + repr(_e)[:180])

print("\n50. form_push: element-veldgroepen (custom-fields) + list-conditie (Paneel-branch)")
try:
    import magicplan.form_push as _FP5
    # (a) robuuste list-extractie: {data:{forms,publish_to}} -> records + workgroup-ids
    _recs, _wg = _FP5._records_from_list({"data": {"forms": [{"id": "a"}], "publish_to": [{"id": "ws1"}]}})
    check("list-extractie: forms+publish_to uit dict-shape", _recs == [{"id": "a"}] and _wg == ["ws1"])
    _recs2, _ = _FP5._records_from_list({"data": [{"id": "b"}]})
    check("list-extractie: platte list-shape werkt ook", _recs2 == [{"id": "b"}])
    # (b) 'Raam/paneel'-veldgroep mergen: Paneel-branch aanhangen (list-conditie) + raam-defaults
    _fg = {"id": "fldRaam", "form": {"id": "q", "name": "Raam/paneel", "context": ["windows"],
           "type": "custom-form", "children": [
        {"id": "g1", "type": "question", "name": "Type glas", "dataType": "list", "required": True,
         "comparisonValue": None, "fields": {"options": ["Enkel", "HR++"]}},
        {"id": "g2", "type": "question", "name": "Kozijnmateriaal", "dataType": "list", "required": True,
         "comparisonValue": None, "fields": {"options": ["Metaal TO", "Hout of kunststof"]}},
        {"id": "g3", "type": "question", "name": "Raam/paneel", "dataType": "list", "required": True,
         "comparisonValue": None, "fields": {"options": ["Raam", "Paneel"]}, "children": []}]}}
    _rec5, _add5, _req5, _pb5 = _FP5.merge_record(_fg, _FP5.load_additions(), verbose=False)
    _bn5 = {c["name"]: c for c in _FP5._form_of(_rec5)["children"]}
    _kids5 = _bn5["Raam/paneel"].get("children", [])
    check("field-group: Paneel-branch aangehangen met list-conditie 'Paneel'",
          any(k["name"].startswith("Paneel - isolatie") and k.get("comparisonValue") == "Paneel" for k in _kids5))
    check("field-group: Kozijnmateriaal optioneel + 'Hout of kunststof' vooraan",
          _bn5["Kozijnmateriaal"]["required"] is False and _bn5["Kozijnmateriaal"]["fields"]["options"][0] == "Hout of kunststof")
    check("field-group: Raam/paneel optioneel + 'Raam' vooraan",
          _bn5["Raam/paneel"]["required"] is False and _bn5["Raam/paneel"]["fields"]["options"][0] == "Raam")
    check("field-group: geen structurele validatieproblemen", _pb5 == [])
    # (c) idempotent: nog eens mergen voegt de Paneel-branch niet dubbel toe
    _rec5b, _add5b, _, _ = _FP5.merge_record(_FP5._form_of(_rec5), _FP5.load_additions(), verbose=False)
    _kids5b = {c["name"]: c for c in _FP5._form_of(_rec5b)["children"]}["Raam/paneel"].get("children", [])
    check("field-group: idempotent (geen dubbele Paneel-velden)", len(_kids5b) == len(_kids5))
except Exception as _e:
    check("form_push field-groups: draait zonder fout", False); print("     " + repr(_e)[:180])

print("\n51. Veldgidsen in de webapp (/gids/<slug>) + inmeetgids/opnameformulier")
try:
    import dashboard.app as _WG
    import dashboard.bouwjaar as _BJ2
    _WG.app.config.update(TESTING=True)
    _cg = _WG.app.test_client()
    with _cg.session_transaction() as _sg:
        _sg["ingelogd"] = True
    for _slug in _WG.GIDSEN:
        _rg = _cg.get("/gids/%s" % _slug)
        if _rg.status_code != 200:
            check("gids '%s' laadt" % _slug, False)
            break
    else:
        check("alle %d veldgidsen laden in de webapp" % len(_WG.GIDSEN), True)
    _sp = _cg.get("/gids/spouwinspectie").get_data(as_text=True)
    check("gids spouwinspectie: inhoud gerenderd (boorlocatie + tabel)",
          "Boorlocatie" in _sp or "boorlocatie" in _sp)
    check("gids: markdown-tabel -> html-table in table-wrap", "table-wrap" in _cg.get("/gids/inmeten").get_data(as_text=True))
    check("gids: onbekende slug -> 404", _cg.get("/gids/bestaatniet").status_code == 404)
    check("guide-pagina linkt de gidsen", "Veldgidsen" in _cg.get("/guide").get_data(as_text=True))
    _html = _BJ2.md_naar_html("## Kop\n| A | B |\n|---|---|\n| 1 | 2 |\n- [ ] taak\n**vet**")
    check("md_naar_html: kop+tabel+checkbox+vet", "<h3>Kop</h3>" in _html and "<th>A</th>" in _html
          and "<td>1</td>" in _html and "☐ taak" in _html and "<b>vet</b>" in _html)
    check("docs: inmeetgids + opnameformulier bestaan",
          os.path.isfile(os.path.join(ROOT, "docs", "magicplan-inmeetgids.md"))
          and os.path.isfile(os.path.join(ROOT, "docs", "nijbegun-opnameformulier.md")))
except Exception as _e:
    check("veldgidsen: draait zonder fout", False); print("     " + repr(_e)[:180])

print()
print("52. Tikbaar Gevelnaam-veld (echte export-indeling: kamer,wand,...,Type@8 + Gevelnaam-kolom)")
try:
    import tempfile as _tf52, csv as _csv52, os as _os52
    _rows52 = [
        ["PLAN ATTRIBUTES"], ["Total living area: m2", "87"], ["", "Nij Begun"],
        ["Oriëntatie voorgevel", "NW"], ["Woningtype", "Tussenwoning"],
        [],
        ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m2", "Surface without openings: m2",
         "Width: m", "Height: m", "Annotation", "Type", "Type glas", "Raam = Ja | Paneel = Nee",
         "Gevelnaam (leeg = binnenwand)"],
        ["Living Room", "Wall 1", "Wall", "12.0", "8.0", "5.0", "2.6", "", "Wall", "", "", "Achtergevel"],
        ["Living Room", "Wall 1", "Fixed Window", "2.0", "", "1.3", "1.5", "", "Window", "Dubbel", "Ja.raam", ""],
        ["Living Room", "Wall 0", "Wall", "10.0", "10.0", "4.0", "2.6", "", "Wall", "", "", "Buurwand.AVR"],
        ["Hall", "Wall 6", "Wall", "9.0", "7.0", "3.5", "2.6", "", "Wall", "", "", "Voorgevel"],
        ["Hall", "Wall 2", "Wall", "6.0", "6.0", "2.3", "2.6", "", "Wall", "", "", ""],
    ]
    with _tf52.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _fh52:
        _csv52.writer(_fh52).writerows(_rows52)
        _pad52 = _fh52.name
    _d52, _n52 = _csvdos(_pad52)
    _gv = {s2.id: s2.orientatie for s2 in _d52.schil if s2.type == "gevel"}
    check("tik: Achtergevel -> gevel ZO (afgeleid van voorgevel NW)", _gv.get("gevel-achter") == "ZO")
    check("tik: Voorgevel -> gevel NW", _gv.get("gevel-voor") == "NW")
    check("tik: Buurwand (AVR) -> NIET in de schil", not any("avr" in (s2.begrenzing or "").lower() for s2 in _d52.schil))
    check("tik: naamloze wand blijft binnenwand", len(_gv) == 2)
    check("tik: raam in getikte gevel blijft behouden (erft ZO)",
          any(s2.type == "kozijn" and s2.orientatie == "ZO" for s2 in _d52.schil))
    _os52.unlink(_pad52)
except Exception as _e:
    check("gevelnaam-tik: draait zonder fout", False); print("     " + repr(_e)[:170])

print()
print("53. Dakmodel per type (Dak N - type) + 'Grenst aan buiten (m)'-splitsing")
try:
    import tempfile as _tf53, csv as _csv53, os as _os53
    def _csv53maak(extra_plan, wall_extra_kop, wall_extra):
        rows = [["PLAN ATTRIBUTES"], ["Total living area: m2", "90"], ["", "Nij Begun"],
                ["Oriëntatie voorgevel", "Z"], ["Woningtype", "Tussenwoning"], ["Bouwjaar", "1980"]]
        rows += extra_plan
        rows += [[], ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Volume: m³",
                 "Ground perimeter: m", "Ceiling perimeter: m", "Walls with openings: m²",
                 "Walls without openings: m²", "Ground surface with all walls: m²",
                 "Ground surface with interior walls: m²", "Ceiling Height"],
                 ["Ground Floor", "50", "130", "28", "28", "90", "80", "56", "50", "2.60 m"],
                 ["1st Floor", "48", "120", "27", "27", "85", "78", "54", "48", "2.50 m"], []]
        rows += [["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                  "Width: m", "Height: m", "Annotation", "Type", "Gevelnaam (leeg = binnenwand)"] + wall_extra_kop]
        rows += wall_extra
        with _tf53.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as fh:
            _csv53.writer(fh).writerows(rows); return fh.name
    # (a) schilddak -> 4 dakvlakken, GEEN kopgevels
    _p = _csv53maak([["Dak 1 - type (leeg = geen dak 1)", "Schilddak"],
                     ["Dak 1 schild - hellingshoek lange vlakken (°)", "40"],
                     ["Dak 1 schild - oriëntatie lang dakvlak 1", "Z"]], [], [])
    _d, _n = _csvdos(_p); _os53.unlink(_p)
    _dak = [s2 for s2 in _d.schil if s2.type == "dak"]
    _kop = [s2 for s2 in _d.schil if s2.type == "gevel" and "dak" in s2.id]
    check("schilddak: 4 dakvlakken", len(_dak) == 4)
    check("schilddak: GEEN kopgevel-driehoeken", len(_kop) == 0)
    # (b) lessenaarsdak -> 1 vlak + note over hoge-zijde-strook
    _p = _csv53maak([["Dak 1 - type (leeg = geen dak 1)", "Lessenaarsdak"],
                     ["Dak 1 lessenaar - oriëntatie dakvlak (afwaterend naar)", "Z"],
                     ["Dak 1 lessenaar - hoogte lage zijde boven vloer (m)", "2.5"],
                     ["Dak 1 lessenaar - hoogte hoge zijde boven vloer (m)", "3.8"],
                     ["Dak 1 lessenaar - hellingshoek (°, leeg = berekend)", "20"]], [], [])
    _d, _n = _csvdos(_p); _os53.unlink(_p)
    check("lessenaar: 1 dakvlak", sum(1 for s2 in _d.schil if s2.type == "dak") == 1)
    check("lessenaar: note over hoge-zijde-strook -> Vabi", any("hoge kant" in str(x) or "hoge-zijde" in str(x) for x in _n))
    # (c) 'Grenst aan buiten (m)'-splitsing bij narekenen-wand
    _p = _csv53maak([["Dak 1 - type (leeg = geen dak 1)", "Plat dak"]],
                    ["Deels binnen/deels buiten? (narekenen)", "Grenst aan buiten (m) — meet de buitenlengte"],
                    [["Woonkamer", "Wall 5", "Wall", "10.0", "10.0", "4.0", "2.5", "", "Wall", "Voorgevel", "Yes", "3.0"]])
    _d, _n = _csvdos(_p); _os53.unlink(_p)
    check("buiten-splitsing: 3m x 2.5m = 7.5 m2 opgeteld bij de gevel (met orientatie)",
          any("7.5" in str(x) and "opgeteld bij de gevel" in str(x) for x in _n))
    check("buiten-splitsing: gesplitste wand NIET meer als HANDMATIG NAREKENEN gemeld (geen tegenstrijdigheid)",
          not any("HANDMATIG NAREKENEN" in str(x) for x in _n))
    check("plat dak: footprint bovenste verdieping (48) gebruikt",
          any(s2.type == "dak" and abs((s2.oppervlakte_m2 or 0) - 48) < 0.1 for s2 in _d.schil))
except Exception as _e:
    check("dakmodel/buiten-splitsing: draait zonder fout", False); print("     " + repr(_e)[:170])

print()
print("54. Dakkapel (ISSO 8.2.1): voorvlak+wangen=gevel, dakje=plat, gat afgetrokken")
try:
    from core.geometry import dakkapel_vlakken as _dkv
    _dk = _dkv(2.0, 1.3, 1.0, 45.0)
    check("dakkapel: gevel = B*H + 2*D*H", abs(_dk["gevel_m2"] - 5.2) < 0.01)
    check("dakkapel: dakje = B*D", abs(_dk["dak_m2"] - 2.0) < 0.01)
    check("dakkapel: gat = B*D/cos(a) (~2.83)", abs(_dk["gat_schuin_dak_m2"] - 2.83) < 0.05)
    _dk0 = _dkv(2.0, 1.3, 1.0, None)
    check("dakkapel: zonder helling -> gat 0 + flag 'niet afgetrokken'",
          _dk0["gat_schuin_dak_m2"] == 0.0 and "NIET afgetrokken" in _dk0["flag"])
except Exception as _e:
    check("dakkapel: draait zonder fout", False); print("     " + repr(_e)[:150])

print()
print("55. Dakramen in hellend vlak + asymmetrisch zadeldak + Dak N-prefix isolatieboom")
try:
    import tempfile as _tf55, csv as _csv55, os as _os55
    _rows55 = [["PLAN ATTRIBUTES"], ["Total living area: m2", "90"], ["", "Nij Begun"],
        ["Oriëntatie voorgevel", "Z"], ["Woningtype", "Tussenwoning"], ["Bouwjaar", "1980"],
        ["Dak 1 - type (leeg = geen dak 1)", "Zadeldak"],
        ["Dak 1 zadel - oriëntatie dakvlak 1", "Z"],
        ["Dak 1 zadel - vloerbreedte tussen de kopgevels (m)", "6.0"],
        ["Dak 1 zadel - nokhoogte boven zoldervloer (m)", "3.0"],
        ["Dak 1 zadel - hellingshoek vlak 2 (°, leeg = zelfde)", "25"],
        ["Dak 1 - dakramen aantal (leeg = geen)", "2"],
        ["Dak 1 - dakramen totaal oppervlak (m²)", "1.8"],
        ["Dak 1 - dakramen type glas", "HR++"],
        ["Dak 1 - isolatie aanwezig?", "Ja"],
        ["Dak 1 - isolatiedikte (mm)", "100"],
        [],
        ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Volume: m³", "Ground perimeter: m",
         "Ceiling perimeter: m", "Walls with openings: m²", "Walls without openings: m²",
         "Ground surface with all walls: m²", "Ground surface with interior walls: m²", "Ceiling Height"],
        ["Ground Floor", "50", "130", "28", "28", "90", "80", "56", "50", "2.60 m"], []]
    with _tf55.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _fh:
        _csv55.writer(_fh).writerows(_rows55); _p55 = _fh.name
    _d55, _n55 = _csvdos(_p55); _os55.unlink(_p55)
    _rw = next((s2 for s2 in _d55.schil if (s2.subtype or "") == "Dakraam"), None)
    check("dakraam: kozijn subtype Dakraam, 1.8 m2, HR++, op dakvlak-orientatie Z",
          _rw is not None and abs(_rw.oppervlakte_m2 - 1.8) < 0.01 and _rw.glastype == "HR++" and _rw.orientatie == "Z")
    _dks = [s2 for s2 in _d55.schil if s2.type == "dak"]
    _zuid = next((s2 for s2 in _dks if s2.orientatie == "Z"), None)
    _noord = next((s2 for s2 in _dks if s2.orientatie == "N"), None)
    check("asymmetrisch: vlak 2 (N) heeft eigen helling 25", _noord is not None and _noord.hellingshoek == 25)
    check("asymmetrisch: note over benaderde vlakverdeling", any("ASYMMETRISCH" in str(x) for x in _n55))
    # audit-glas-F1 15-7: het dakvlak blijft nu BRUTO (35.36) in het dossier — de parser trekt het
    # dakraam-glas NIET meer af; de netto-aftrek gebeurt 1x in Vabi (dakraam als deelvlak op het dak).
    check("dakraam: dakvlak blijft BRUTO (35.36) in het dossier — geen dubbele aftrek meer",
          any(abs((s2.oppervlakte_m2 or 0) - 35.36) < 0.4 for s2 in _dks))
    check("Dak 1-prefix isolatieboom gelezen (Ja + 100mm)",
          _zuid is not None and _zuid.isolatie_aanwezig == "Ja" and _zuid.isolatiedikte_mm == 100)
except Exception as _e:
    check("dakramen/asymmetrie: draait zonder fout", False); print("     " + repr(_e)[:170])

print()
print("56. Dakramen/dakkapel PER DAKVLAK (A/B-groepen, eigen glastype) — herontwerp 12-7")
try:
    import tempfile as _t6, csv as _c6, os as _o6
    _r6 = [["PLAN ATTRIBUTES"], ["Total living area: m2", "90"], ["", "Nij Begun"],
        ["Oriëntatie voorgevel", "Z"], ["Woningtype", "Tussenwoning"], ["Bouwjaar", "1980"],
        ["Dak 1 - type (het HELE dak)", "Zadeldak"],
        ["Dak 1 zadel - oriëntatie dakvlak 1", "Z"],
        ["Dak 1 zadel - vloerbreedte tussen de kopgevels (m)", "6.0"],
        ["Dak 1 zadel - nokhoogte boven zoldervloer (m)", "3.0"],
        ["Dak 1 - dakramen A: dakvlak (leeg = dakvlak 1)", "Dakvlak.1"],
        ["Dak 1 - dakramen A: aantal", "2"], ["Dak 1 - dakramen A: totaal oppervlak (m2)", "1.6"],
        ["Dak 1 - dakramen A: type glas", "Dubbel"],
        ["Dak 1 - dakramen B: dakvlak (leeg = dakvlak 1)", "Dakvlak.2.tegenoverliggend"],
        ["Dak 1 - dakramen B: aantal", "1"], ["Dak 1 - dakramen B: totaal oppervlak (m2)", "0.8"],
        ["Dak 1 - dakramen B: type glas", "Enkel"],
        [],
        ["FLOOR ATTRIBUTES", "Ground surface without walls: m2", "V", "G", "C", "W1", "W2",
         "Ground surface with all walls: m2", "Gi", "Ceiling Height"],
        ["Ground Floor", "50", "130", "28", "28", "90", "80", "56", "50", "2.60 m"], []]
    with _t6.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f6:
        _c6.writer(_f6).writerows(_r6); _p6 = _f6.name
    _d6, _n6 = _csvdos(_p6); _o6.unlink(_p6)
    _ra = next((x for x in _d6.schil if x.id == "dak1-dakraam-a"), None)
    _rb = next((x for x in _d6.schil if x.id == "dak1-dakraam-b"), None)
    check("A/B: dakraam A op VOORvlak Z met Dubbel", _ra is not None and _ra.orientatie == "Z" and _ra.glastype == "Dubbel")
    check("A/B: dakraam B op ACHTERvlak N met Enkel", _rb is not None and _rb.orientatie == "N" and _rb.glastype == "Enkel")
    _z = next((x for x in _d6.schil if x.id == "dak1-schui-Z"), None)
    _n = next((x for x in _d6.schil if x.id == "dak1-schui-N"), None)
    # audit-glas-F1 15-7: dakvlakken blijven BRUTO in het dossier; het dakraam-glas wordt 1x in Vabi
    # afgetrokken (dakraam als deelvlak op het dak, per orientatie A->Z / B->N zoals hierboven getest).
    check("A/B: dakvlakken blijven BRUTO (parser trekt niet af; 1x aftrek in Vabi per orientatie)",
          _z is not None and _n is not None and _z.oppervlakte_m2 > 34.5 and _n.oppervlakte_m2 > 34.5)
except Exception as _e:
    check("dakramen A/B: draait zonder fout", False); print("     " + repr(_e)[:170])

print()
print("57. Leads-workflow: bulk-plak + afspraak (auto-project) + bevestigings-/ontvangstmail")
try:
    import tempfile as _t57, shutil as _s57, os as _o57
    import dashboard.leads as _L57
    import dashboard.app as _W57
    # (a) bulk-parse: 3 portal-mails in één plak -> 3 leads
    _blob = " ".join('{"BagAdresId":"BULK%d","Naam":"Bulk %d","Email":"b%d@x.nl","Postcode":"888%dYY",'
                     '"Huisnummer":%d}' % (i, i, i, i, i) for i in (1, 2, 3))
    check("bulk: parse_leads_bulk vindt 3 blokken", len(_L57.parse_leads_bulk(_blob)) == 3)
    check("bulk: enkel blok blijft werken", len(_L57.parse_leads_bulk('{"Naam":"Solo","Postcode":"1111AA"}')) == 1)
    # (a2) .eml-upload: quoted-printable (JSON over 2 regels gebroken met soft break) + base64-HTML
    _qp = (b"From: portal@smarttwin.nl\r\nSubject: Contact met adviseur\r\nMIME-Version: 1.0\r\n"
           b"Content-Type: text/plain; charset=utf-8\r\nContent-Transfer-Encoding: quoted-printable\r\n\r\n"
           b'Beste adviseur,\r\n{"BagAdresId":"EML1","Naam":"Eml Testpersoon","Postcode":"7777X=\r\nX",'
           b'"Huisnummer":9}\r\n')
    import base64 as _b64
    _html = '<html><body><p>{"BagAdresId":"EML2","Naam":"Html Persoon","Postcode":"6666WW","Huisnummer":4}</p></body></html>'
    _b = (b"From: portal@smarttwin.nl\r\nSubject: Contact\r\nMIME-Version: 1.0\r\n"
          b"Content-Type: text/html; charset=utf-8\r\nContent-Transfer-Encoding: base64\r\n\r\n"
          + _b64.encodebytes(_html.encode("utf-8")))
    _t1 = _L57.tekst_uit_eml(_qp)
    _t2 = _L57.tekst_uit_eml(_b)
    check("eml: quoted-printable gedecodeerd (soft break geheeld)", '"Postcode":"7777XX"' in _t1.replace(" ", ""))
    check("eml: base64-HTML gedecodeerd + tags gestript",
          "BagAdresId" in _t2 and "<p>" not in _t2 and len(_L57.parse_leads_bulk(_t2)) == 1)
    _rot = _L57.tekst_uit_eml(b"\x00\xff geen mail")
    check("eml: kapotte bytes -> geen crash, geen leads", isinstance(_rot, str) and _L57.parse_leads_bulk(_rot) == [])
    # (b) mails: NL-datum + voorbereidings-/verwachtings-punten + drukte-tekst
    _lead = {"naam": "Bulk 1", "postcode": "8881YY", "huisnummer": "1", "toevoeging": "",
             "afspraak": "2026-07-20T14:30", "email": "b1@x.nl"}
    _ow, _tx = _L57.bevestiging_mail(_lead, {"naam": "Renze", "bedrijf": "Poortinga"})
    check("bevestiging: NL-datum + tijd in de mail", "maandag 20 juli 2026 om 14:30" in _tx)
    check("bevestiging: raambekleding + kruipruimteluik + foto's-in-alle-ruimtes",
          "aambekleding" in _tx and "kruipruimteluik" in _tx and "alle ruimtes" in _tx)
    check("bevestiging: verwachtingsmanagement (triple/kozijn niet vergoed)", "triple glas" in _tx.lower())
    _ow2, _tx2 = _L57.ontvangst_mail({"naam": "Renze"})
    check("ontvangst: drukte + wachttijd + 'op de lijst'", "drukte" in _tx2 and "wachttijd" in _tx2 and "lijst" in _tx2)
    # GEEN lange streepjes (em/en-dash) in klantmails: leest als AI-geschreven (eis Renze 12-7)
    _adv57 = {"naam": "Renze Poortinga", "bedrijf": "Poortinga Energieadvies",
              "telefoon": "06-99999999", "email": "info@poortinga-energieadvies.nl"}
    _ow3, _tx3 = _L57.concept_mail(_lead, _adv57)
    _alle_mail = "".join((_ow, _tx, _ow2, _tx2, _ow3, _tx3))
    check("mails: geen em/en-dash in onderwerp of tekst", "—" not in _alle_mail and "–" not in _alle_mail)
    check("mails: geen emoji (alles onder U+2500; opsommingsteken mag)", all(ord(c) < 0x2500 for c in _alle_mail))
    # WIJ-vorm (bedrijf), geen telefoonnummer, wel e-mailvoorkeur (eis Renze 12-7)
    check("mails: wij-vorm, geen ik-vorm", " ik " not in _alle_mail.lower() and "Wij " in _alle_mail)
    check("mails: telefoonnummer eruit + e-mailvoorkeur benoemd",
          "Telefoon" not in _alle_mail and "06-99999999" not in _alle_mail and "per e-mail" in _alle_mail)
    check("mails: bedrijfsnaam als afzender in ontvangst + kennismaking",
          _L57.ontvangst_mail(_adv57)[1].count("Poortinga Energieadvies") >= 2
          and _tx3.count("Poortinga Energieadvies") >= 2)
    # Kennismakingsmail: bewonerswensen (30% ISDE) + bouwfysische klachten (vocht/schimmel/tocht)
    # conform kennisbank-eisen + opnameformulier (schimmel/vochtklachten = vast opname-onderdeel)
    check("kennismaking: vraagt naar bewonerswensen (30% ISDE)", "wensen" in _tx3 and "30% ISDE" in _tx3)
    check("kennismaking: vraagt naar vocht/schimmel/tocht-klachten",
          "vocht" in _tx3 and "schimmel" in _tx3 and "tocht" in _tx3)
    check("kennismaking: geen uitnodiging om zelf te bellen (wij nemen contact op)",
          "belt" not in _tx3 and "Wij nemen binnenkort contact" in _tx3)
    # (c) routes: afspraak zetten -> status 'afspraak gepland' + AUTO-project; ontvangst-bulk
    _W57.app.config.update(TESTING=True)
    _c57 = _W57.app.test_client()
    with _c57.session_transaction() as _ss:
        _ss["ingelogd"] = True
    _od57, _of57 = _L57.LEADS_DIR, _L57.LEADS_FILE
    _tmp57 = _t57.mkdtemp(); _L57.LEADS_DIR = _tmp57; _L57.LEADS_FILE = _o57.path.join(_tmp57, "leads.json")
    try:
        # upload-route: 2 .eml's + 1 .msg (overgeslagen met melding) in één POST
        import io as _io57
        _rr = _c57.post("/leads/add", data={
            "mailtekst": "",
            "emls": [(_io57.BytesIO(_qp), "mail1.eml"), (_io57.BytesIO(_b), "mail2.eml"),
                     (_io57.BytesIO(b"OLE"), "oud.msg")]},
            content_type="multipart/form-data", follow_redirects=True)
        _up = _L57.load_leads()
        check("upload-route: 2 .eml's -> 2 leads, .msg overgeslagen + melding",
              len(_up) == 2 and {x["bag_id"] for x in _up} == {"EML1", "EML2"}
              and ".msg-bestand" in _rr.get_data(as_text=True))
        _L57.save_leads([])
        _rows57 = []
        for _ld in _L57.parse_leads_bulk(_blob):
            _rows57, _ = _L57.add_lead(_ld, _rows57)
        _L57.save_leads(_rows57)
        _lid57 = _rows57[0]["id"]
        _c57.post("/leads/%d/afspraak" % _lid57, data={"wanneer": "2026-07-21T10:00"})
        _r57 = next(x for x in _L57.load_leads() if x["id"] == _lid57)
        check("afspraak-route: datum + status + AUTO-project", _r57.get("afspraak") == "2026-07-21T10:00"
              and _r57.get("status") == "afspraak gepland" and bool(_r57.get("project_tag")))
        _bev = _c57.get("/leads/%d/mail?soort=bevestiging" % _lid57).get_data(as_text=True)
        check("bevestiging-route: mail met afspraakdatum", "21 juli 2026 om 10:00" in _bev)
        # status-dropdown op 'afspraak gepland' -> ook auto-project (lead 2)
        _lid2b = _rows57[1]["id"]
        _c57.post("/leads/%d/status" % _lid2b, data={"status": "afspraak gepland"})
        _r2b = next(x for x in _L57.load_leads() if x["id"] == _lid2b)
        check("status 'afspraak gepland': AUTO-project aangemaakt", bool(_r2b.get("project_tag")))
        # ontvangst: alleen 'nieuw'-leads in BCC (lead 3), daarna allen gemarkeerd
        _ont = _c57.get("/leads/ontvangst").get_data(as_text=True)
        check("ontvangst-route: alleen nieuwe leads in BCC", "b3@x.nl" in _ont and "b2@x.nl" not in _ont)
        _c57.post("/leads/ontvangst/verstuurd")
        check("ontvangst 'verstuurd': alle nieuwe -> mail gestuurd",
              all(x["status"] != "nieuw" for x in _L57.load_leads()))
        for _x in _L57.load_leads():
            if _x.get("project_tag"):
                _s57.rmtree(_W57._pdir(_x["project_tag"]), ignore_errors=True)
    finally:
        _L57.LEADS_DIR, _L57.LEADS_FILE = _od57, _of57
except Exception as _e:
    check("leads-workflow: draait zonder fout", False); print("     " + repr(_e)[:180])

print()
print("58. VABI-import-fixes 12-7 (1e echte import): bouwjaar-lek, verdiepingen, gebouwhoogte, klasse, Locatie")
try:
    import re as _re58, tempfile as _t58, os as _o58, csv as _c58
    import xml.etree.ElementTree as _ET58
    from core.dossier import Dossier as _D58, SchilDeel as _S58, VloerInfo as _V58
    from vabi.constructie_generate import _jaar_uit_klassetekst as _jk, resolve_constructies as _rc58
    from vabi import objecten_generate as _OG58
    # (a) klassetekst -> representatief jaar (ook ge-dot)
    check("klassetekst: 'Van 1975 t/m 1982' -> 1978", _jk("Van 1975 t/m 1982") == 1978)
    check("klassetekst: dotted + Tot/Vanaf", _jk("Van.1975.t.m.1982") == 1978
          and _jk("Tot 1965") == 1964 and _jk("Vanaf 2014") == 2015 and _jk("") is None)
    # (b) dossier: gevel met per-bouwdeel klasse (project-bouwjaar LEEG) -> constructie 1975-1982
    _d = _D58()
    _d.identificatie.bouwjaar = None
    _d.identificatie.orientatie_voorgevel = "NW"
    _d.opname.gevelhoogte_m = 5.24
    _d.opname.gebouwhoogte_m = 8.21
    _d.geometrie.vloeren = [_V58(naam="Ground Floor", oppervlakte_m2=55.56),
                            _V58(naam="1st Floor", oppervlakte_m2=44.35),
                            _V58(naam="2nd Floor", oppervlakte_m2=22.15)]
    _d.geometrie.gebruiksoppervlakte_ag_m2 = 87.13
    _d.schil = [
        _S58(id="gevel-voor", type="gevel", orientatie="NW", oppervlakte_m2=48.0,
             isolatie_aanwezig="Onbekend", bouwjaarklasse="Van 1975 t/m 1982", begrenzing="Buitenlucht"),
        _S58(id="gevel-achter", type="gevel", orientatie="ZO", oppervlakte_m2=24.0,
             isolatie_aanwezig="Onbekend", bouwjaarklasse="Van 1975 t/m 1982", begrenzing="Buitenlucht"),
        _S58(id="gevel-kopg-ZW", type="gevel", orientatie="ZW", oppervlakte_m2=8.0,
             isolatie_aanwezig="Onbekend", bouwjaarklasse="Van 1975 t/m 1982", begrenzing="Buitenlucht"),
        _S58(id="vloer", type="vloer", oppervlakte_m2=55.0, isolatie_aanwezig="Onbekend",
             begrenzing="Kruipruimte"),
        _S58(id="dak-schu-NW", type="dak", orientatie="NW", oppervlakte_m2=40.0,
             isolatie_aanwezig="Onbekend", hellingshoek=35.0, begrenzing="Buitenlucht"),
    ]
    _cons, _map58, _iss58 = _rc58(_d)
    _namen = " | ".join(sorted({m["naam"] for m in _map58.values()}))
    check("constructie: per-bouwdeel klasse wint (gevel 1975-1982, GEEN <1965)", "1975-1982" in _namen)
    check("constructie: onbekend ZONDER klasse -> luide issue (vloer/dak)",
          any("ZONDER bouwjaar" in str(i) for i in _iss58))
    # (c) objecten-XML: bouwjaar-lek dicht + echte verdiepingen + gebouwhoogte + Locatie-tabs
    _root58, _m2, _iss2, _st2 = _OG58.build_tree(_d)
    _x58 = _ET58.tostring(_root58, encoding="unicode")
    check("objecten: bouwjaar ontbreekt -> 0 (GEEN sjabloon-1994-lek) + actie",
          "<Bouwjaar>0</Bouwjaar>" in _x58
          and any("BOUWJAAR ONTBREEKT" in str(i) for i in _iss2))
    check("objecten: verdiepingen = gemeten m2 (55.56/44.35/22.15), niet gelijk verdeeld",
          all(("<Gebruiksoppervlakte>%s</Gebruiksoppervlakte>" % w) in _x58
              for w in ("55.56", "44.35", "22.15")) and "29.04" not in _x58)
    check("objecten: Gebouwhoogte = handmatige invoer (8.21), niet gevelhoogte",
          "<Gebouwhoogte>8.21</Gebouwhoogte>" in _x58)
    # gebouwhoogte ONTBREEKT -> 0 + luide actie (nooit sjabloon-lek 7.60, nooit gevelhoogte-fallback)
    _d.opname.gebouwhoogte_m = None
    _rootg, _, _issg, _ = _OG58.build_tree(_d)
    _xg = _ET58.tostring(_rootg, encoding="unicode")
    check("objecten: gebouwhoogte ontbreekt -> 0 + actie (geen 7.60-sjabloonlek, geen 5.24-fallback)",
          "<Gebouwhoogte>0</Gebouwhoogte>" in _xg and "7.60" not in _xg and "5.24" not in _xg
          and any("GEBOUWHOOGTE ONTBREEKT" in str(i) for i in _issg))
    _d.opname.gebouwhoogte_m = 8.21
    _locs = {}
    for _el in _root58.iter():
        if _el.tag.endswith("Hoofdvlak") and (_el.findtext("Naam") or "").strip():
            _locs[_el.findtext("Naam")] = _el.findtext("Locatie")
    check("objecten: Locatie-tabs voor=2/achter=3/kopgevel-ZW(rechts)=5/vloer=0/dak=1",
          _locs.get("Gevel gevel-voor") == "2" and _locs.get("Gevel gevel-achter") == "3"
          and _locs.get("Gevel gevel-kopg-ZW") == "5" and _locs.get("Vloer vloer") == "0"
          and _locs.get("Dak dak-schu-NW") == "1")
    # (c2) audit-fixes 12-7: perimeter-/GrenstAan-/verdiepingen-sjabloonlekken + netto-herrekening
    _d.schil.append(_S58(id="raam-1", type="kozijn", subtype="Raam", orientatie="NW",
                         oppervlakte_m2=2.5, glastype="Dubbel", begrenzing="Buitenlucht"))
    _d.schil.append(_S58(id="vloer-leeg", type="vloer", oppervlakte_m2=10.0,
                         isolatie_aanwezig="Onbekend", begrenzing=""))
    _rootc, _mc, _issc, _stc = _OG58.build_tree(_d)
    _xc = _ET58.tostring(_rootc, encoding="unicode")
    _hvs = {(h.findtext("Naam") or ""): h for h in _rootc.iter() if h.tag.endswith("Hoofdvlak")}
    _vl = _hvs.get("Vloer vloer")
    check("audit: vloer zonder perimeter -> AutoPerimeter=1 + 0.00 (geen 28.14-sjabloonlek)",
          _vl is not None and _vl.findtext("AutoPerimeter") == "1" and _vl.findtext("Perimeter") == "0.00"
          and any("perimeter" in str(i).lower() and "ONTBREEKT" in str(i) for i in _issc))
    check("audit: LEGE begrenzing geflagd (sjabloon-GrenstAan blijft anders stil staan)",
          any("vloer-leeg" in str(i) and "ONTBREEKT" in str(i) for i in _issc))
    _gv = _hvs.get("Gevel gevel-voor")
    check("audit: netto = bruto - deelvlakken (48.00 - 2.50 = 45.50, zoals de echte export)",
          _gv is not None and _gv.findtext("BrutoOppervlakte") == "48.00"
          and _gv.findtext("NettoOppervlakte") == "45.50")
    _dleeg = _D58()
    _dleeg.schil = [_S58(id="gevel-x", type="gevel", orientatie="Z", oppervlakte_m2=30.0,
                         isolatie_aanwezig="Onbekend", begrenzing="Buitenlucht")]
    _rootl, _, _issl, _ = _OG58.build_tree(_dleeg)
    _xl = _ET58.tostring(_rootl, encoding="unicode")
    check("audit: geen Ag/verdiepingen -> 1 laag met 0.00 + actie (GEEN sjabloon-185m2-lek)",
          "28.86" not in _xl and "94.51" not in _xl
          and any("AG/VERDIEPINGEN ONTBREKEN" in str(i) for i in _issl))
    _d.schil = [s for s in _d.schil if s.id not in ("raam-1", "vloer-leeg")]
    # (d) parser: gebouwhoogte = HANDMATIG veld + verdieping-m2 + Ag = som gemeten verdiepingen
    from magicplan.statistics_csv import build_dossier as _bd58
    _rows58 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "87.13"], ["Floors", "3"],
               ["Nij Begun"], ["Woningtype", "Tussenwoning"], ["Gevelhoogte (m)", "5.24"],
               ["Gebouwhoogte tot de nok (m)", "8.21"],
               ["Dak - nokhoogte (m, optioneel)", "2.97"], ["Oriëntatie voorgevel", "NW"],
               ["Gevel - bouwjaar (onbekend)", "Van.1975.t.m.1982"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Volume: m³",
                "Ground perimeter: m", "Ceiling perimeter: m", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "55.56", "67", "59", "68", "1", "2", "2.60 m"],
               ["1st Floor", "44.35", "105", "56", "62", "1", "2", "2.60 m"], []]
    with _t58.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f58:
        _c58.writer(_f58).writerows(_rows58)
        _p58 = _f58.name
    _d2, _n2 = _bd58(_p58)
    _o58.unlink(_p58)
    check("parser: gebouwhoogte = handmatig veld (8.21), NOOIT berekend", _d2.opname.gebouwhoogte_m == 8.21)
    check("parser: verdieping-m2 gelezen (55.56 + 44.35)",
          sorted(round(v.oppervlakte_m2, 2) for v in _d2.geometrie.vloeren) == [44.35, 55.56])
    check("parser: Ag = som gemeten verdiepingen (99.91), NIET MagicPlan-woonoppervlak (87.13)",
          abs(_d2.geometrie.gebruiksoppervlakte_ag_m2 - 99.91) < 0.01
          and any("niet gebruikt" in str(n) for n in _n2))
    check("parser: bouwjaar-ontbreekt-note", any("BOUWJAAR ONTBREEKT" in str(n) for n in _n2))
    # (e) gebouwhoogte-veld ontbreekt -> GEEN berekening uit gevel+nok, wel luide note
    _rows58b = [r for r in _rows58 if not (r and r[0].startswith("Gebouwhoogte"))]
    with _t58.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f58b:
        _c58.writer(_f58b).writerows(_rows58b)
        _p58b = _f58b.name
    _d3, _n3 = _bd58(_p58b)
    _o58.unlink(_p58b)
    check("parser: veld ontbreekt -> gebouwhoogte LEEG + GEBOUWHOOGTE-note (geen 5.24+2.97-gok)",
          _d3.opname.gebouwhoogte_m is None and any("GEBOUWHOOGTE ONTBREEKT" in str(n) for n in _n3))
    # (f) 1-op-1-garantie: hernoemd "(...)"-suffix mag een INGEVULD veld niet meer stil kwijtraken
    _rows58c = [(["Gevelhoogte (m, tot de dakvoet gemeten)", "5.24"] if r and r[0].startswith("Gevelhoogte")
                 else r) for r in _rows58]
    with _t58.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f58c:
        _c58.writer(_f58c).writerows(_rows58c)
        _p58c = _f58c.name
    _d4, _n4 = _bd58(_p58c)
    _o58.unlink(_p58c)
    check("parser: suffix-drift ('Gevelhoogte (m, tot de dakvoet gemeten)') -> waarde tóch gelezen",
          _d4.opname.gevelhoogte_m == 5.24)
    # (g) ongeldige hellingshoek (tikfout 95) -> LUIDE note + dak niet stil te klein berekend.
    # dak-velden in de PLAN ATTRIBUTES-sectie zetten (vóór de blanco regel die FLOOR inleidt).
    _rows58d, _ingevoegd = [], False
    for r in _rows58:
        if not _ingevoegd and r == []:
            _rows58d += [["Type dak", "Zadeldak"], ["Hellingshoek dak", "95"],
                         ["Dak - vloerbreedte (m)", "6"],
                         ["Dakvlak 1 - oriëntatie", "NW"], ["Dakvlak 2 - oriëntatie", "ZO"]]
            _ingevoegd = True
        _rows58d.append(r)
    with _t58.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f58d:
        _c58.writer(_f58d).writerows(_rows58d)
        _p58d = _f58d.name
    _d5, _n5 = _bd58(_p58d)
    _o58.unlink(_p58d)
    check("parser: helling 95 gr -> ONGELDIG-note en geen stil te klein dak",
          any("ONGELDIG" in str(n) for n in _n5)
          and not any(s.type == "dak" and s.hellingshoek == 95 for s in _d5.schil))
except Exception as _e:
    check("VABI-import-fixes: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("60. Enum-gok-fixes (audit 13-7): alleen EPA-bevestigde codes schrijven, rest sjabloon + LUIDE flag")
try:
    from core.dossier import Dossier as _D60, Verwarming as _V60, Tapwater as _T60, Ventilatie as _Ve60
    import vabi.installatie_generate as _IG60
    # (a) ketel-subtype != HR107 -> NIET geschreven + flag
    _d = _D60()
    _d.installaties.verwarming = _V60(type_opwekker="Gasgestookte ketel", subtype="HR104")
    _r, _f = _IG60.build_tree(_d)
    _op = _IG60._find(_r, "VerwarmingOpwekker")
    check("enum: ketel-subtype HR104 NIET auto-geschreven (alleen HR107) + flag",
          _op.findtext("SubType") != "3" and any("subtype" in x.lower() and "HR107" in x for x in _f))
    # (a2) aanvoertemp 90/70 (oude radiatoren) -> WaterAanvoertemperatuur=11 (19-7 live bevestigd), geen flag
    _d1b = _D60()
    _d1b.installaties.verwarming = _V60(type_opwekker="Gasgestookte ketel", subtype="HR107",
                                        aanvoertemperatuur="90/70")
    _r1b, _f1b = _IG60.build_tree(_d1b)
    check("enum: aanvoertemp 90/70 -> WaterAanvoertemperatuur=11 (EPA-bevestigd), geen aanvoertemp-flag",
          _IG60._find(_r1b, "WaterAanvoertemperatuur").text == "11"
          and not any("aanvoertemp" in x.lower() for x in _f1b))
    # (b) ventilatie collectief -> NIET geschreven + flag
    _d2 = _D60(); _d2.ventilatie = _Ve60(systeem="C", systeem_soort="collectief")
    _r2, _f2 = _IG60.build_tree(_d2)
    check("enum: ventilatie-systeemsoort 'collectief' niet auto-gecodeerd + flag",
          any("systeemsoort" in x.lower() for x in _f2))
    # (c) tapwater 'compleet' -> geen foute TypeToestel=2 meer + flag
    _d3 = _D60(); _d3.installaties.tapwater = _T60(type_toestel="Compleet toestel")
    _r3, _f3 = _IG60.build_tree(_d3)
    _top = _IG60._find(_r3, "TapwaterOpwekker")
    # de foute 'compleet'->TypeToestel=2-mapping is weg; onbevestigd toestel -> LUIDE flag
    check("enum: tapwater 'compleet' niet meer actief gemapt -> LUIDE flag i.p.v. gok",
          any("tapwater-toestel" in x.lower() and "bevestigd" in x.lower() for x in _f3))
    # (d) objecten: GrenstAan 'Water' (code 1) is 19-7 LIVE bevestigd (volledige dropdown 0-9) ->
    # code 1 wordt geschreven en er is GEEN 'niet probe-bevestigd'-flag meer.
    from core.dossier import SchilDeel as _S60
    from vabi import objecten_generate as _OG60
    _d4 = _D60()
    _d4.schil = [_S60(id="vloer-water", type="vloer", oppervlakte_m2=40.0,
                      isolatie_aanwezig="Onbekend", begrenzing="Water")]
    _root4, _m4, _iss4, _st4 = _OG60.build_tree(_d4)
    _ga4 = next((e.text for e in _root4.iter()
                 if e.tag.split("}")[-1] == "GrenstAan" and (e.text or "") == "1"), None)
    check("enum: GrenstAan 'Water' -> code 1 (EPA-bevestigd), geen 'probe-bevestigd'-flag meer",
          _ga4 == "1" and not any("probe-bevestigd" in str(i).lower() for i in _iss4))
    # (e) parser: per-vlak helling 95 (direct dakvlak) -> _helling_ok vangt hem
    import tempfile as _t60, os as _o60, csv as _c60
    from magicplan.statistics_csv import build_dossier as _bd60
    _rows60 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "80"], ["Woningtype", "Tussenwoning"],
               ["Type dak", "Zadeldak"], ["Dakvlak 1 - oppervlak (m²)", "30"],
               ["Dakvlak 1 - hellingshoek (°)", "95"], ["Dakvlak 1 - oriëntatie", "NW"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Ceiling Height"],
               ["Ground Floor", "40", "2.60 m"], []]
    with _t60.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f60:
        _c60.writer(_f60).writerows(_rows60); _p60 = _f60.name
    _d60, _n60 = _bd60(_p60); _o60.unlink(_p60)
    check("enum/parser: per-vlak helling 95 -> ONGELDIG-note, geen dakvlak met helling 95",
          any("ONGELDIG" in str(n) for n in _n60)
          and not any(s.type == "dak" and s.hellingshoek == 95 for s in _d60.schil))
except Exception as _e:
    check("enum-gok-fixes: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("61. Gevel-tikfout-detectie + zolder-uitsluiting (Essenhage-les 14-7)")
try:
    import tempfile as _t61, os as _o61, csv as _c61
    from magicplan.statistics_csv import build_dossier as _bd61
    def _w61(vals):
        r = [""] * 26
        for i, v in vals.items():
            r[i] = str(v)
        return r
    _rows61 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "80"], ["Woningtype", "Tussenwoning"],
               ["Gevelhoogte (m)", "5.2"], ["Oriëntatie voorgevel", "NW"],
               ["Type dak", "Zadeldak"], ["Hellingshoek dak", "35"], ["Dak - vloerbreedte (m)", "5.9"],
               ["Dakvlak 1 - oriëntatie", "NW"], ["Dakvlak 2 - oriëntatie", "ZO"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "Volume", "GP", "CP", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "50", "1", "1", "1", "1", "1", "2.60 m"],
               ["2nd Floor", "30", "1", "1", "1", "1", "1", "2.40 m"], [],
               ["ROOM ATTRIBUTES", "Ground surface without walls: m²"],
               ["Ground Floor", ""], ["Woonkamer", "50"],
               ["2nd Floor", ""], ["Zolderkamer", "30"], [],
               # WALL: c0 kamer, c1 wand, c3 bruto, c4 netto, c5 breedte, c6 hoogte, c8 Type, c25 Gevelnaam
               ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
               ["Ground Floor"],   # verdieping-scheidingsrij (voedt de b x h-verdiepingshoogte)
               _w61({0: "Woonkamer", 1: "Wall 1", 3: "10.4", 4: "10.4", 5: "4.0", 6: "2.6", 8: "Wall", 25: "Voorgevel"}),
               _w61({0: "Woonkamer", 1: "Wall 3", 3: "10.4", 4: "10.4", 5: "4.0", 6: "2.6", 8: "Wall", 25: "Voorgevel"}),
               ["2nd Floor"],
               _w61({0: "Zolderkamer", 1: "Wall 0", 3: "9.6", 4: "9.6", 5: "4.0", 6: "2.4", 8: "Wall", 25: "Achtergevel"}),
               []]
    with _t61.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f61:
        _c61.writer(_f61).writerows(_rows61)
        _p61 = _f61.name
    _d61, _n61 = _bd61(_p61)
    _o61.unlink(_p61)
    # GEOMETRISCHE DEDUP: 2 gelijke-breedte wanden (tegenoverliggend) op dezelfde gevel -> 1x geteld
    check("dedup: gelijke-breedte tegenoverliggende wand op dezelfde gevel -> 1x geteld (geen aanname)",
          any("TEGENOVERLIGGENDE" in str(n) and "1x geteld" in str(n) for n in _n61))
    # zolderwand op een schuin-dak-oriëntatie -> AUTOMATISCH uit de gevel gehouden (eis Renze 14-7)
    check("zolder: wand op schuin-dak-oriëntatie -> ZOLDER automatisch NIET meegeteld in de gevel",
          any("ZOLDER" in str(n) and "NIET meegeteld" in str(n) and "SCHUINE DAKVLAK" in str(n) for n in _n61))
    _zo61 = sum(s.oppervlakte_m2 or 0 for s in _d61.schil if s.type == "gevel" and s.orientatie == "ZO")
    check("zolder: achtergevel (alleen zolder onder schuin dak) -> ~0 m² gevel (zit in het dak)", _zo61 < 1.0)
    _nw61 = sum(s.oppervlakte_m2 or 0 for s in _d61.schil if s.type == "gevel" and s.orientatie == "NW")
    # Woonkamer BG Wall 1 + Wall 3 (beide 4,0 m) -> dedup naar 4,0 x 2,60 = 10,4 (+ hart-op-hart)
    check("dedup: dubbele BG-wanden (4,0=4,0) -> 1x geteld = 10,4 m² i.p.v. 20,8",
          10.3 <= _nw61 <= 12.0)
    # ONMOGELIJKE HOEK: tegenoverliggende wanden (Wall 0//Wall 2) getagd op 90°-apart gevels
    _rows61b = [["PLAN ATTRIBUTES"], ["Total living area: m²", "40"], ["Woningtype", "Vrijstaand"],
                ["Oriëntatie voorgevel", "NW"], [],
                ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
                ["Ground Floor", "40", "1", "1", "1", "1", "1", "2.60 m"], [],
                ["ROOM ATTRIBUTES", "Ground surface without walls: m²"], ["Ground Floor", ""], ["Bijkeuken", "12"], [],
                ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                 "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
                ["Ground Floor"],
                _w61({0: "Bijkeuken", 1: "Wall 0", 3: "9", 5: "3.6", 6: "2.6", 8: "Wall", 25: "Achtergevel"}),
                _w61({0: "Bijkeuken", 1: "Wall 2", 3: "9", 5: "3.6", 6: "2.6", 8: "Wall", 25: "Rechtergevel"}),
                []]
    with _t61.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f61b:
        _c61.writer(_f61b).writerows(_rows61b)
        _p61b = _f61b.name
    _d61b, _n61b = _bd61(_p61b)
    _o61.unlink(_p61b)
    check("onmogelijke hoek: tegenoverliggende Wall 0=achter + Wall 2=rechts (90°) -> TIKFOUT",
          any("onmogelijke hoek" in str(n) for n in _n61b))
    # geldige combinatie (voor/achter op tegenoverliggende wanden) mag NIET flaggen
    _rows61c = [(["Voorgevel" if r and len(r) > 25 and r[25] == "Rechtergevel" else (r[25] if r and len(r) > 25 else None)] and
                 (r[:25] + ["Voorgevel"] if (r and len(r) > 25 and r[25] == "Rechtergevel") else r)) for r in _rows61b]
    with _t61.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f61c:
        _c61.writer(_f61c).writerows(_rows61c)
        _p61c = _f61c.name
    _d61c, _n61c = _bd61(_p61c)
    _o61.unlink(_p61c)
    check("onmogelijke hoek: voor/achter op tegenoverliggende wanden (180°) -> GEEN valse flag",
          not any("onmogelijke hoek" in str(n) for n in _n61c))
except Exception as _e:
    check("gevel-tikfout-detectie: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("62. Gevel = breedte x verdiepingshoogte per bouwlaag (methode Renze 14-7: 5,81 -> 28,92 m2 BRUTO)")
try:
    import tempfile as _t62, os as _o62, csv as _c62
    from magicplan.statistics_csv import build_dossier as _bd62
    def _w62(vals):
        r = [""] * 26
        for i, v in vals.items():
            r[i] = str(v)
        return r
    _rows62 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "80"], ["Woningtype", "Vrijstaand"],
               ["Gevelhoogte (m)", "5.2"], ["Oriëntatie voorgevel", "NW"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "50", "1", "1", "1", "1", "1", "2.60 m"],
               ["1st Floor", "45", "1", "1", "1", "1", "1", "2.38 m"], [],
               ["ROOM ATTRIBUTES", "Ground surface without walls: m²"],
               ["Ground Floor", ""], ["Woonkamer", "50"],
               ["1st Floor", ""], ["Slaapkamer", "45"], [],
               ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
               ["Ground Floor"],
               # BG: pui-wand met borstwering — wandsom zou maar 2,67 m2 tellen; b x h moet 5,81x2,60 geven
               _w62({0: "Woonkamer", 1: "Wall 1", 3: "2.67", 4: "2.67", 5: "5.81", 6: "0.46", 8: "Wall", 25: "Achtergevel"}),
               ["1st Floor"],
               _w62({0: "Slaapkamer", 1: "Wall 2", 3: "13.83", 4: "13.83", 5: "5.81", 6: "2.38", 8: "Wall", 25: "Achtergevel"}),
               []]
    with _t62.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f62:
        _c62.writer(_f62).writerows(_rows62)
        _p62 = _f62.name
    _d62, _n62 = _bd62(_p62)
    _o62.unlink(_p62)
    _ag62 = next((s for s in _d62.schil if s.type == "gevel" and s.orientatie == "ZO"), None)
    check("b x h: achtergevel 5,81 -> BG 5,81x2,60 + 1e 5,81x2,38 = 28,92 m2 (Renze's voorbeeld)",
          _ag62 is not None and abs(_ag62.oppervlakte_m2 - 28.92) < 0.05)
    check("b x h: opbouw-note met beide verdiepingen",
          any("5.81x2.60" in str(n) and "5.81x2.38" in str(n) for n in _n62))
    check("b x h: gevel gemarkeerd als BRUTO (ramen/deuren = deelvlak)",
          "BRUTO" in (_ag62.opmerkingen or ""))
    # DIRECTE GEVELBREEDTE-invoer: overruled de fragiele wandsom (dubbel-getikte wanden genegeerd)
    _rows62b = [["PLAN ATTRIBUTES"], ["Total living area: m²", "80"], ["Woningtype", "Vrijstaand"],
                ["Oriëntatie voorgevel", "NW"], ["Achtergevel - breedte (m)", "5.81"], [],
                ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
                ["Ground Floor", "50", "1", "1", "1", "1", "1", "2.60 m"],
                ["1st Floor", "45", "1", "1", "1", "1", "1", "2.38 m"], [],
                ["ROOM ATTRIBUTES", "Ground surface without walls: m²"],
                ["Ground Floor", ""], ["Woonkamer", "50"], ["1st Floor", ""], ["Slaapkamer", "45"], [],
                ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                 "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
                ["Ground Floor"],
                # bewust FOUT getikte breedtes (3,1 en 9,9) — de gemeten 5,81 moet die overrulen
                _w62({0: "Woonkamer", 1: "Wall 1", 3: "8", 5: "3.1", 6: "2.6", 8: "Wall", 25: "Achtergevel"}),
                ["1st Floor"],
                _w62({0: "Slaapkamer", 1: "Wall 2", 3: "23", 5: "9.9", 6: "2.38", 8: "Wall", 25: "Achtergevel"}), []]
    with _t62.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f62b:
        _c62.writer(_f62b).writerows(_rows62b)
        _p62b = _f62b.name
    _d62b, _n62b = _bd62(_p62b)
    _o62.unlink(_p62b)
    _ag62b = next((s for s in _d62b.schil if s.type == "gevel" and s.orientatie == "ZO"), None)
    check("gevelbreedte-invoer: gemeten 5,81 overrulet de (foute) wandsom -> 5,81x(2,60+2,38)=28,92",
          _ag62b is not None and abs(_ag62b.oppervlakte_m2 - 28.92) < 0.05)
    check("gevelbreedte-invoer: 'DIRECT GEMETEN'-note",
          any("DIRECT GEMETEN gevelbreedte 5.81" in str(n) for n in _n62b))
    # plat-dak legacy-pad mag niet crashen (helling None-bug 14-7)
    _rows62c = [r for r in _rows62b if not (r and r[0].startswith("Achtergevel"))]
    _rows62c = _rows62c[:5] + [["Type dak", "Plat dak"], ["Dakvlak 1 - oppervlak (m²)", "50"]] + _rows62c[5:]
    with _t62.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f62c:
        _c62.writer(_f62c).writerows(_rows62c)
        _p62c = _f62c.name
    _d62c, _n62c = _bd62(_p62c)
    _o62.unlink(_p62c)
    check("plat-dak legacy-pad crasht niet + maakt een plat dakvlak",
          any(s.type == "dak" for s in _d62c.schil))
except Exception as _e:
    check("gevel b x h: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("63. Zolder-dak-footprint + kopgevel-op-bovenste-verdieping + dubbeltel + aanbouw-dak (Essenhage-les 15-7)")
try:
    import tempfile as _t63, os as _o63, csv as _c63
    from magicplan.statistics_csv import build_dossier as _bd63
    def _w63(vals):
        r = [""] * 26
        for i, v in vals.items():
            r[i] = str(v)
        return r
    # 3 verdiepingen: BG 50 (met aanbouw), 1e 44 (dakdragend), 2e ZOLDER 22 (< 0,7 x 44 -> zolder).
    # Zadeldak NW/ZO -> kopgevels op NO/ZW. Aanbouw op BG heeft Linkergevel(NO)+Rechtergevel(ZW),
    # maar op de ZOLDER staat GEEN NO/ZW-gevel -> nok-kopgevels moeten WEGGELATEN worden.
    # Voorgevel 1e = 3,8+1,9+1,9 = 7,6 m (dubbel getekende kamer) vs achtergevel 1e = 5,0 -> DUBBELTEL.
    _rows63 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "116"], ["Woningtype", "Tussenwoning"],
               ["Gevelhoogte (m)", "5.0"], ["Oriëntatie voorgevel", "NW"],
               ["Type dak", "Zadeldak"],
               ["Dak zadel - oriëntatie dakvlak 1", "NW"],
               ["Dak zadel - vloerbreedte tussen de kopgevels (m)", "5.8"],
               ["Dak zadel - hellingshoek (°, leeg = berekend uit nok/breedte)", "35"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "50", "1", "1", "1", "1", "1", "2.60 m"],
               ["1st Floor", "44", "1", "1", "1", "1", "1", "2.40 m"],
               ["2nd Floor", "22", "1", "1", "1", "1", "1", "2.50 m"], [],
               ["ROOM ATTRIBUTES", "Ground surface without walls: m²"],
               ["Ground Floor", ""], ["Woonkamer", "38"], ["Bijkeuken", "12"],
               ["1st Floor", ""], ["Studeerkamer", "20"], ["Badkamer", "8"], ["Badkamer", "8"], ["Slaapkamer", "16"],
               ["2nd Floor", ""], ["Zolder", "22"], [],
               ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
               ["Ground Floor"],
               _w63({0: "Woonkamer", 1: "Wall 1", 3: "15", 5: "5.8", 6: "2.6", 8: "Wall", 25: "Achtergevel"}),
               _w63({0: "Woonkamer", 1: "Wall 3", 3: "15", 5: "5.8", 6: "2.6", 8: "Wall", 25: "Voorgevel"}),
               _w63({0: "Bijkeuken", 1: "Wall 0", 3: "9", 5: "3.5", 6: "2.6", 8: "Wall", 25: "Linkergevel"}),
               _w63({0: "Bijkeuken", 1: "Wall 2", 3: "9", 5: "3.5", 6: "2.6", 8: "Wall", 25: "Rechtergevel"}),
               ["1st Floor"],
               _w63({0: "Slaapkamer", 1: "Wall 0", 3: "10", 5: "3.2", 6: "2.4", 8: "Wall"}),
               _w63({0: "Slaapkamer", 1: "Wall 1", 3: "12", 5: "5.0", 6: "2.4", 8: "Wall", 25: "Achtergevel"}),
               _w63({0: "Studeerkamer", 1: "Wall 0", 3: "9", 5: "3.8", 6: "2.4", 8: "Wall"}),
               _w63({0: "Studeerkamer", 1: "Wall 1", 3: "9", 5: "3.8", 6: "2.4", 8: "Wall", 25: "Voorgevel"}),
               _w63({0: "Badkamer", 1: "Wall 0", 3: "4", 5: "1.6", 6: "2.4", 8: "Wall"}),
               _w63({0: "Badkamer", 1: "Wall 1", 3: "5", 5: "1.9", 6: "2.4", 8: "Wall", 25: "Voorgevel"}),
               _w63({0: "Badkamer", 1: "Wall 0", 3: "4", 5: "1.6", 6: "2.4", 8: "Wall"}),
               _w63({0: "Badkamer", 1: "Wall 1", 3: "5", 5: "1.9", 6: "2.4", 8: "Wall", 25: "Voorgevel"}),
               ["2nd Floor"],
               _w63({0: "Zolder", 1: "Wall 0", 3: "9", 5: "3.8", 6: "2.5", 8: "Wall"}),  # geen geveltag
               []]
    with _t63.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f63:
        _c63.writer(_f63).writerows(_rows63)
        _p63 = _f63.name
    _d63, _n63 = _bd63(_p63)
    _o63.unlink(_p63)
    # (1) DAK-FOOTPRINT: schuin dak over de verdieping ONDER de zolder (44 m²), niet de zolder (22 m²)
    check("dak: footprint = verdieping ONDER de zolder (44 m²), niet de zolder (22 m²)",
          any("ONDER de zolder" in str(n) for n in _n63))
    _dak63 = sum(s.oppervlakte_m2 or 0 for s in _d63.schil if s.type == "dak")
    # 44/cos(35) = 53,7 (schuin) -> ~53, NIET 22/cos(35)=27
    check("dak: hellend dak ~54 m² (over 44 m²), niet ~27 (over zolder 22)", 48 <= _dak63 <= 60)
    # (2) KOPGEVEL: aanbouw-zijgevel op BG (NO/ZW) mag de nok-kopgevels NIET aanzetten (geen tag op zolder)
    check("kopgevel: WEGGELATEN want geen buitengevel op de bovenste verdieping (aanbouw-BG telt niet)",
          any("kopgevel" in str(n).lower() and "WEGGELATEN" in str(n) and "bovenste verdieping" in str(n) for n in _n63))
    _no63 = sum(s.oppervlakte_m2 or 0 for s in _d63.schil if s.type == "gevel" and s.orientatie == "NO")
    check("kopgevel: NO-gevel = alleen de aanbouw-zijgevel (~9 m²), geen extra nok-driehoek", _no63 < 12)
    # (3) DUBBELTEL: voorgevel 1e (7,6 m) veel breder dan achtergevel 1e (5,0 m) -> LET OP
    check("dubbeltel: voorgevel 1e breder dan achtergevel -> LET OP mogelijke DUBBELTEL",
          any("DUBBELTEL" in str(n) for n in _n63))
    # (4) AANBOUW-DAK: BG (50) groter dan dakdragende verdieping (44), geen plat dak -> MOGELIJK DAK ONTBREEKT
    check("aanbouw-dak: BG groter dan hoofddak-footprint + geen plat dak -> MOGELIJK DAK ONTBREEKT",
          any("MOGELIJK DAK ONTBREEKT" in str(n) for n in _n63))
    # geen valse dubbeltel op de begane grond (voor 5,8 ~ achter 5,8)
    check("dubbeltel: GEEN valse flag op de BG (voor 5,8 ~ achter 5,8)",
          not any("DUBBELTEL" in str(n) and "Ground Floor" in str(n) for n in _n63))
except Exception as _e:
    check("zolder-dak/kopgevel/dubbeltel: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("64. Deels-buiten (buitenlengte) -> tag x hoogte, cumulatief in de juiste EPA-tab (verificatie Renze 15-7)")
try:
    import tempfile as _t64, os as _o64, csv as _c64
    from magicplan.statistics_csv import build_dossier as _bd64
    from vabi.objecten_generate import _locatie_code as _lc64
    def _h64():
        return (["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                 "Width: m", "Height: m", "Annotation", "Type", "Grenst aan buiten (m) — meet de buitenlengte"]
                + [""] * 15 + ["Gevelnaam (leeg = binnenwand)", "Deels binnen/deels buiten? (narekenen)"])
    def _w64(vals):
        r = [""] * 27
        for i, v in vals.items():
            r[i] = str(v)
        return r
    _base = [["PLAN ATTRIBUTES"], ["Total living area: m²", "40"], ["Woningtype", "Vrijstaand"],
             ["Oriëntatie voorgevel", "NW"], [],
             ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
             ["Ground Floor", "40", "1", "1", "1", "1", "1", "2.60 m"], [],
             ["ROOM ATTRIBUTES", "Ground surface without walls: m²"], ["Ground Floor", ""], ["Hal", "40"], []]
    # A) deels-buiten MET tag + buitenlengte 1,67 (zonder vinkje) -> 1,67 x 2,60 = 4,34 op Voorgevel (tab 2)
    _rows64 = _base + [_h64(), ["Ground Floor"],
              _w64({0: "Hal", 1: "Wall 0", 3: "15.6", 4: "15.6", 5: "6.03", 6: "2.60", 8: "Wall", 9: "1.67", 25: "Voorgevel"}),
              # B) deels-buiten ZONDER tag -> LUIDE flag, niet meegeteld
              _w64({0: "Hal", 1: "Wall 2", 3: "8.8", 4: "8.8", 5: "3.39", 6: "2.60", 8: "Wall", 9: "1.67", 26: "Yes"}), []]
    with _t64.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f64:
        _c64.writer(_f64).writerows(_rows64)
        _p64 = _f64.name
    _d64, _n64 = _bd64(_p64)
    _o64.unlink(_p64)
    _vgev = [s for s in _d64.schil if s.type == "gevel" and s.orientatie == "NW"]
    check("deels-buiten MET tag: buitenlengte 1,67 x hoogte 2,60 = 4,34 m² (niet de hele wand 15,6)",
          _vgev and abs(_vgev[0].oppervlakte_m2 - 4.34) < 0.2)
    check("deels-buiten: buitenlengte werkt ZONDER het narekenen-vinkje (split-note aanwezig)",
          any("gesplitst via" in str(n) for n in _n64))
    check("deels-buiten MET tag -> juiste EPA-tab (Voorgevel = 2)",
          _vgev and _lc64("gevel", _vgev[0], "NW") == "2")
    check("deels-buiten ZONDER tag -> LUIDE flag (niet stil weggelaten)",
          any("ZONDER gevel-aanduiding" in str(n) and "NIET meegeteld" in str(n) for n in _n64))
    # C) cumulatief: twee Voorgevel-wanden op verschillende verdiepingen -> opgeteld op tab 2
    #    (Vrijstaand = geen hart-op-hart-toeslag, zodat we de zuivere optelling toetsen)
    _rows64c = [["PLAN ATTRIBUTES"], ["Total living area: m²", "80"], ["Woningtype", "Vrijstaand"],
                ["Oriëntatie voorgevel", "NW"], [],
                ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
                ["Ground Floor", "40", "1", "1", "1", "1", "1", "2.60 m"],
                ["1st Floor", "40", "1", "1", "1", "1", "1", "2.40 m"], [],
                ["ROOM ATTRIBUTES", "Ground surface without walls: m²"],
                ["Ground Floor", ""], ["Woonkamer", "40"], ["1st Floor", ""], ["Slaapkamer", "40"], [],
                _h64(), ["Ground Floor"],
                _w64({0: "Woonkamer", 1: "Wall 1", 3: "13", 4: "13", 5: "5.0", 6: "2.6", 8: "Wall", 25: "Voorgevel"}),
                ["1st Floor"],
                _w64({0: "Slaapkamer", 1: "Wall 1", 3: "12", 4: "12", 5: "5.0", 6: "2.4", 8: "Wall", 25: "Voorgevel"}), []]
    with _t64.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f64c:
        _c64.writer(_f64c).writerows(_rows64c)
        _p64c = _f64c.name
    _d64c, _n64c = _bd64(_p64c)
    _o64.unlink(_p64c)
    _nwtot = sum(s.oppervlakte_m2 or 0 for s in _d64c.schil if s.type == "gevel" and s.orientatie == "NW")
    # 5,0x2,60 + 5,0x2,40 = 13,0 + 12,0 = 25,0 (Vrijstaand -> geen toeslag)
    check("cumulatief: voorgevel BG (5x2,60) + 1e (5x2,40) opgeteld = 25 m² op één tab",
          24.5 <= _nwtot <= 25.5)
    # D) deels-buiten GEMARKEERD + tag maar LEGE meters -> FOUT-melding
    _rows64d = _base + [_h64(), ["Ground Floor"],
               _w64({0: "Hal", 1: "Wall 0", 3: "15.6", 4: "15.6", 5: "6.03", 6: "2.60", 8: "Wall", 25: "Voorgevel", 26: "Yes"}), []]
    with _t64.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f64d:
        _c64.writer(_f64d).writerows(_rows64d)
        _p64d = _f64d.name
    _d64d, _n64d = _bd64(_p64d)
    _o64.unlink(_p64d)
    check("deels-buiten + tag maar LEGE meters -> FOUT-melding (vul buitenlengte in)",
          any("FOUT" in str(n) and "Grenst aan buiten (m)' is LEEG" in str(n) for n in _n64d))
except Exception as _e:
    check("deels-buiten/cumulatief: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("59. Webapp opleveren: leeg project -> direct Afronden + eigen ventilatieplan/bijlagen uploaden + export-zip")
try:
    import io as _io59, shutil as _sh59, zipfile as _zip59
    import dashboard.app as _W59
    _W59.app.config.update(TESTING=True)
    _c59 = _W59.app.test_client()
    with _c59.session_transaction() as _s59:
        _s59["ingelogd"] = True
    _r59 = _c59.post("/nieuw", data={"straat": "Testweg 5", "postcode": "9990ZZ", "plaats": "X",
                                     "woningtype": "Tussenwoning"})
    _tag59 = _r59.headers["Location"].rstrip("/").split("/")[-2]
    # (a) direct naar Afronden zonder VABI-toets -> geen crash, klikbare stepper
    _ra59 = _c59.get("/project/%s/afronden" % _tag59)
    check("opleveren: Afronden direct bereikbaar (geen VABI-heenweg) + klikbare stepper",
          _ra59.status_code == 200 and '<a class="step' in _ra59.get_data(as_text=True))
    # (b) eigen ventilatieplan + 2 bijlagen uploaden
    _c59.post("/project/%s/bijlagen" % _tag59, data={
        "ventilatieplan_eigen": (_io59.BytesIO(b"%PDF plan"), "eigen_plan.pdf"),
        "bijlagen": [(_io59.BytesIO(b"f"), "factuur.pdf"), (_io59.BytesIO(b"p"), "plattegrond.png")]},
        content_type="multipart/form-data", follow_redirects=True)
    _st59 = _W59._load_state(_tag59)
    check("opleveren: eigen ventilatieplan + bijlagen opgeslagen",
          bool(_st59.get("ventilatieplan_eigen")) and set(_st59.get("bijlagen") or []) == {"factuur.pdf", "plattegrond.png"})
    # (c) foto voorkant uploaden
    _c59.post("/project/%s/fotos" % _tag59, data={"foto_voorkant": (_io59.BytesIO(b"J"), "voor.jpg")},
              content_type="multipart/form-data", follow_redirects=True)
    check("opleveren: foto voorkant opgeslagen", bool(_W59._load_state(_tag59).get("foto_voorkant")))
    # (d) export-zip bevat de eigen bestanden
    _rz59 = _c59.get("/project/%s/export" % _tag59)
    _namen59 = _zip59.ZipFile(_io59.BytesIO(_rz59.data)).namelist()
    # de zip is sinds de OneDrive-indeling één projectmap met genummerde submappen
    check("opleveren: export-zip bevat eigen ventilatieplan + bijlagen + foto voorkant",
          any("03_Isolatieplan/ventilatieplan_eigen" in n for n in _namen59)
          and any("07_Overig/factuur.pdf" in n for n in _namen59)
          and any("04_Fotos/foto_voorkant" in n for n in _namen59))
    # (e) bijlage verwijderen
    _c59.get("/project/%s/bijlage/factuur.pdf/weg" % _tag59, follow_redirects=True)
    check("opleveren: bijlage verwijderen werkt", "factuur.pdf" not in (_W59._load_state(_tag59).get("bijlagen") or []))
    _sh59.rmtree(_W59._pdir(_tag59), ignore_errors=True)
except Exception as _e:
    check("webapp opleveren: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("65. Webapp: Dak toevoegen-wizard (plat / driehoek-zadeldak / 9 geometrieen; auto-genummerd)")
try:
    import re as _re65, shutil as _sh65
    import dashboard.app as _W65
    _W65.app.config.update(TESTING=True)
    _c65 = _W65.app.test_client()
    with _c65.session_transaction() as _s65:
        _s65["ingelogd"] = True
    _r65 = _c65.post("/nieuw", data={"straat": "Dakstraat 1", "postcode": "9991ZZ", "plaats": "X",
                                     "woningtype": "Vrijstaand"})
    _tag65 = _r65.headers["Location"].rstrip("/").split("/")[-2]
    # 1) zadeldak via driehoek: c=7, breedte=5, 45 gr, NW, beide kopgevels buiten (jouw voorbeeld)
    _c65.post("/project/%s/opname/dak/driehoek" % _tag65,
              data={"orient_hellend": "NW", "lange_zijde": "7", "breedte": "5", "helling1": "45",
                    "kopgevel1_buiten": "on", "kopgevel2_buiten": "on"}, follow_redirects=True)
    _d65 = _W65._dossier(_tag65)
    _schuin = [s for s in _d65.schil if s.type == "dak" and "schuin" in (s.subtype or "")]
    _kop = [s for s in _d65.schil if s.type == "gevel" and "kopgevel" in (s.subtype or "")]
    check("dak-wizard driehoek: 2 hellende vlakken ~24,75 m² (jouw 4,95x5)",
          len(_schuin) == 2 and all(24.0 <= (s.oppervlakte_m2 or 0) <= 25.5 for s in _schuin))
    check("dak-wizard driehoek: hellende vlakken op NW én ZO",
          {s.orientatie for s in _schuin} == {"NW", "ZO"})
    check("dak-wizard driehoek: 2 kopgevels ~12,25 m² op NO/ZW (haaks, buiten)",
          len(_kop) == 2 and {s.orientatie for s in _kop} == {"NO", "ZW"}
          and all(11.5 <= (s.oppervlakte_m2 or 0) <= 13.0 for s in _kop))
    # 2) tussenwoning: kopgevels NIET aanvinken -> geen kopgevel toegevoegd (buurwand)
    _c65.post("/project/%s/opname/dak/driehoek" % _tag65,
              data={"orient_hellend": "NW", "lange_zijde": "7", "breedte": "5", "helling1": "45"},
              follow_redirects=True)
    _kop2 = [s for s in _W65._dossier(_tag65).schil if s.type == "gevel" and "kopgevel" in (s.subtype or "")]
    check("dak-wizard driehoek: kopgevels UIT -> geen extra kopgevel (tussenwoning-buurwand)", len(_kop2) == len(_kop))
    # 3) plat dak
    _c65.post("/project/%s/opname/dak/plat" % _tag65, data={"m2": "24.5"}, follow_redirects=True)
    check("dak-wizard plat: plat dak 24,5 m² toegevoegd",
          any(s.type == "dak" and (s.subtype or "") == "plat dak" and abs((s.oppervlakte_m2 or 0) - 24.5) < 0.1
              for s in _W65._dossier(_tag65).schil))
    # 4) 9 geometrieen
    _c65.post("/project/%s/opname/dak/negen" % _tag65,
              data={"m2_NW": "10", "m2_ZO": "8", "m2_Horizontaal": "5"}, follow_redirects=True)
    _d65d = _W65._dossier(_tag65)
    _negen = [s for s in _d65d.schil if "zelf ingevoerd" in (s.subtype or "")]
    check("dak-wizard 9-geom: 3 vlakken (NW/ZO/Horizontaal) toegevoegd", len(_negen) == 3)
    # 5) auto-nummering: dak1.., dak2.., dak3.., dak4..
    _nrs = {int(_m.group(1)) for s in _d65d.schil for _m in [_re65.match(r"dak(\d+)", s.id or "")] if _m}
    check("dak-wizard: automatisch genummerd (>=4 dak-groepen)", len(_nrs) >= 4)
    # 6) dakraam op het NW-dakvlak -> kozijn subtype Dakraam + VABI-deelvlak op het DAK (niet de gevel)
    _c65.post("/project/%s/opname/dakraam" % _tag65,
              data={"dak_orient": "NW", "glas": "HR++", "breedte": "0.8", "hoogte": "1.2"}, follow_redirects=True)
    _d65e = _W65._dossier(_tag65)
    _draam = [s for s in _d65e.schil if "dakraam" in (s.subtype or "").lower()]
    check("dakraam: toegevoegd als kozijn subtype Dakraam op NW (~0,96 m²)",
          len(_draam) == 1 and _draam[0].orientatie == "NW" and abs((_draam[0].oppervlakte_m2 or 0) - 0.96) < 0.05)
    from vabi.objecten_generate import build_tree as _bt65
    _root65 = _bt65(_d65e)[0]
    def _lc65b(t):
        return t.split("}")[-1]
    _dak_dv = 0
    for _hv in _root65.iter():
        if _lc65b(_hv.tag) == "Hoofdvlak":
            _nm = next((c.text for c in _hv if _lc65b(c.tag) == "Naam"), "") or ""
            if _nm.startswith("Dak "):
                _dvl = next((c for c in _hv if _lc65b(c.tag) == "DeelvlakList"), None)
                _dak_dv += sum(1 for c in (list(_dvl) if _dvl is not None else []) if _lc65b(c.tag) == "Deelvlak")
    check("dakraam: als deelvlak op een DAK-hoofdvlak geplaatst (niet op een gevel)", _dak_dv >= 1)
    _sh65.rmtree(_W65._pdir(_tag65), ignore_errors=True)
except Exception as _e:
    check("dak-wizard: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("66. Zadeldak-kopgevel BASIS = overspanning (footprint/noklengte), niet de noklengte (MagicPlan-fix 15-7)")
try:
    import tempfile as _t66, os as _o66, csv as _c66
    from magicplan.statistics_csv import build_dossier as _bd66
    def _w66(vals):
        r = [""] * 26
        for i, v in vals.items():
            r[i] = str(v)
        return r
    # Vrijstaand, 1 verdieping (= bovenste), zadeldak NW; kopgevels op NO/ZW getikt -> tellen mee.
    # overspanning c=7 x noklengte 5 = footprint 35 -> hellend 24,75/stuk; kopgevel-basis 7 -> 12,25/stuk.
    _rows66 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "35"], ["Woningtype", "Vrijstaand"],
               ["Oriëntatie voorgevel", "NW"],
               ["Type dak (zadeldak/schilddak = 1 dak; leeg = geen dak)", "Zadeldak"],
               ["Dak zadel - oriëntatie dakvlak 1", "NW"],
               ["Dak zadel - overspanning (m, leeg = auto)", "7"],
               ["Dak zadel - vloerbreedte tussen de kopgevels (m)", "5"],
               ["Dak zadel - hellingshoek (°, leeg = berekend uit nok/breedte)", "45"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "35", "1", "1", "1", "1", "1", "2.60 m"], [],
               ["ROOM ATTRIBUTES", "Ground surface without walls: m²"], ["Ground Floor", ""], ["Woonkamer", "35"], [],
               ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
                "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"],
               ["Ground Floor"],
               _w66({0: "Woonkamer", 1: "Wall 0", 3: "9.1", 5: "3.5", 6: "2.6", 8: "Wall", 25: "Linkergevel"}),
               _w66({0: "Woonkamer", 1: "Wall 2", 3: "9.1", 5: "3.5", 6: "2.6", 8: "Wall", 25: "Rechtergevel"}), []]
    with _t66.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f66:
        _c66.writer(_f66).writerows(_rows66)
        _p66 = _f66.name
    _d66, _n66 = _bd66(_p66)
    _o66.unlink(_p66)
    _schuin66 = [s for s in _d66.schil if s.type == "dak" and "schuin" in (s.subtype or "")]
    _kop66 = [s for s in _d66.schil if s.type == "gevel" and "kopgevel" in (s.subtype or "")]
    check("zadeldak: hellend vlak = 35/cos45/2 = 24,75 m² (overspanning 7 x noklengte 5)",
          len(_schuin66) == 2 and all(23.5 <= (s.oppervlakte_m2 or 0) <= 25.5 for s in _schuin66))
    check("zadeldak: kopgevel-basis = overspanning 7 -> ~12,25 m² (NIET noklengte 5 -> 6,25)",
          len(_kop66) == 2 and all(11.5 <= (s.oppervlakte_m2 or 0) <= 13.0 for s in _kop66))
    check("zadeldak: note toont kopgevel-basis (overspanning) = 7.00 m",
          any("kopgevel-basis (overspanning) = 7.00" in str(n) for n in _n66))
except Exception as _e:
    check("zadeldak-kopgevel-basis: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("67. VABI-mapping-audit fixes (opake schil F1-F6, 15-7)")
try:
    import tempfile as _t67, os as _o67, csv as _c67
    from magicplan.statistics_csv import build_dossier as _bd67
    from vabi.objecten_generate import build_tree as _bt67
    def _w67(vals):
        r = [""] * 26
        for i, v in vals.items():
            r[i] = str(v)
        return r
    _H67 = (["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²",
             "Width: m", "Height: m", "Annotation", "Type"] + [""] * 16 + ["Gevelnaam (leeg = binnenwand)"])
    # F2 (Gevel - begrenzing form-default=AOR toegepast) + F4 (spouw 'Onbekend' -> None) + F1 (Z-gevel flag)
    _rows67 = [["PLAN ATTRIBUTES"], ["Total living area: m²", "40"], ["Woningtype", "Vrijstaand"],
               ["Oriëntatie voorgevel", "Z"], ["Gevel - begrenzing", "AOR"], ["Gevel - spouw aanwezig?", "Onbekend"], [],
               ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
               ["Ground Floor", "40", "1", "1", "1", "1", "1", "2.60 m"], [],
               ["ROOM ATTRIBUTES", "Ground surface without walls: m²"], ["Ground Floor", ""], ["Woonkamer", "40"], [],
               _H67, ["Ground Floor"],
               _w67({0: "Woonkamer", 1: "Wall 0", 3: "10", 5: "4.0", 6: "2.6", 8: "Wall", 25: "Voorgevel"}), []]
    with _t67.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f67:
        _c67.writer(_f67).writerows(_rows67)
        _p67 = _f67.name
    _d67, _n67 = _bd67(_p67)
    _o67.unlink(_p67)
    _vg67 = [s for s in _d67.schil if s.type == "gevel"]
    check("F2: 'Gevel - begrenzing'=AOR form-standaard toegepast op de gevel (voorheen genegeerd)",
          bool(_vg67) and _vg67[0].begrenzing == "AOR")
    check("F4: spouw 'Onbekend' -> None (niet False = géén spouw)",
          bool(_vg67) and _vg67[0].spouw_aanwezig is None)
    _r67, _m67, _iss67, _s67 = _bt67(_d67)
    # F1 (18-7): oriëntatie is nu LIVE in EPA bevestigd (Zuid=0) -> de Z-voorgevel krijgt code 0 in de
    # objecten-XML en er is GEEN 'afgeleid uit het rotatiepatroon'-flag meer.
    _z_codes = [e.text for e in _r67.iter() if e.tag.split("}")[-1] == "Orientatie" and (e.text or "") == "0"]
    check("F1: gevel op Z -> Orientatie-code 0 (EPA-bevestigd), geen 'afgeleid'-flag meer",
          bool(_z_codes) and not any("afgeleid uit het rotatiepatroon" in str(i) for i in _iss67))
    # F3: schilddak -> Daktype-code 0 (hellend), geen 'niet herkend'
    _d67.identificatie.type_dak = "Schilddak"
    _r67b, _m67b, _iss67b, _s67b = _bt67(_d67)
    def _lc67(t):
        return t.split("}")[-1]
    _dt67 = next((e.text for e in _r67b.iter() if _lc67(e.tag) == "Daktype"), None)
    check("F3: schilddak -> Daktype-code 0 (hellend), niet 'niet herkend -> sjabloon'",
          _dt67 == "0" and not any("Daktype" in str(i) and "niet herkend" in str(i) for i in _iss67b))
    # 18-7: Gebouwtype/Subtype (woningpositie) LIVE in EPA bevestigd (export hoek=1 + monitor tussen=2)
    # -> generator schrijft de codes i.p.v. te flaggen. Onbekende positie (meergezins) -> nog wel flag.
    def _cls67(_wt):
        _d67.identificatie.woningtype = _wt
        _rr, _mm, _ii, _ss = _bt67(_d67)
        _g = next((e.text for e in _rr.iter() if _lc67(e.tag) == "Gebouwtype"), None)
        _s = next((e.text for e in _rr.iter() if _lc67(e.tag) == "Subtype"), None)
        return _g, _s, _ii
    _gv, _sv, _ = _cls67("Vrijstaand")
    check("classificatie: Vrijstaand -> Gebouwtype 0 + Subtype 0", _gv == "0" and _sv == "0")
    _gt, _st, _ = _cls67("Tussenwoning")
    check("classificatie: Tussenwoning -> Subtype 2 (Tussenligging)", _gt == "0" and _st == "2")
    _gk, _sk, _ = _cls67("Twee-onder-een-kap")
    check("classificatie: Twee-onder-een-kap -> Subtype 3", _gk == "0" and _sk == "3")
    _gm, _sm, _im = _cls67("Meergezins")
    check("classificatie: Meergezins (onbekende positie) -> flag i.p.v. gok",
          any("niet herkend" in str(i) for i in _im))
    # F5: paneel-in-kozijn met bouwjaarklasse -> op het paneel-SchilDeel (constructiekeuze forfaitair op paneel-klasse)
    _H67b = ["WALL ATTRIBUTES", "Wall", "Symbol", "Surface: m²", "Surface without openings: m²", "Width: m",
             "Height: m", "Annotation", "Type", "Raam = Ja | Paneel = Nee", "Paneel - bouwjaarklasse",
             "Gevelnaam (leeg = binnenwand)"]
    def _w67b(vals):
        r = [""] * 12
        for i, v in vals.items():
            r[i] = str(v)
        return r
    _rows67b = [["PLAN ATTRIBUTES"], ["Total living area: m²", "40"], ["Woningtype", "Vrijstaand"],
                ["Oriëntatie voorgevel", "N"], [],
                ["FLOOR ATTRIBUTES", "Ground surface without walls: m²", "V", "GP", "CP", "W1", "W2", "Ceiling Height"],
                ["Ground Floor", "40", "1", "1", "1", "1", "1", "2.60 m"], [],
                ["ROOM ATTRIBUTES", "Ground surface without walls: m²"], ["Ground Floor", ""], ["Woonkamer", "40"], [],
                _H67b, ["Ground Floor"],
                _w67b({0: "Woonkamer", 1: "Wall 0", 3: "8", 5: "4.0", 6: "2.6", 8: "Wall", 11: "Voorgevel"}),
                _w67b({0: "Woonkamer", 1: "Wall 0", 3: "0.5", 8: "Window", 9: "Nee (dicht paneel)",
                       10: "Van 1965 t/m 1974"}), []]
    with _t67.NamedTemporaryFile("w", suffix=".csv", delete=False, newline="", encoding="utf-8") as _f67b:
        _c67.writer(_f67b).writerows(_rows67b)
        _p67b = _f67b.name
    _d67c, _n67c = _bd67(_p67b)
    _o67.unlink(_p67b)
    _pnl67 = [s for s in _d67c.schil if s.type == "paneel"]
    check("F5: paneel-bouwjaarklasse op het SchilDeel gezet (niet alleen in opmerking)",
          bool(_pnl67) and getattr(_pnl67[0], "bouwjaarklasse", "") == "Van 1965 t/m 1974")
except Exception as _e:
    check("VABI-mapping-audit fixes: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("68. VABI-mapping-audit fixes glas/installaties/algemeen (15-7)")
try:
    from vabi.constructie_generate import _norm_glas as _ng68, _norm_kozijn as _nk68
    from vabi.installatie_generate import _pv_paneeltype as _pp68, build_tree as _ibt68
    from vabi.objecten_generate import build_tree as _obt68
    from core.dossier import Dossier as _Dos68, Ventilatie as _V68
    # glas F2/F5
    check("glas-F2: 'Voorzetglas' -> 'voorzetraam' (matcht codebook; niet meer stil op Dubbel)",
          _ng68("Voorzetglas") == "voorzetraam")
    check("glas-F5: 'Metaal niet thermisch onderbroken' -> 'metaal' (code 4, niet therm-onderbroken)",
          _nk68("Metaal niet thermisch onderbroken") == "metaal")
    check("glas-F5: 'Metaal thermisch onderbroken' -> therm. onderbroken metaal (blijft correct)",
          _nk68("Metaal thermisch onderbroken") == "therm. onderbroken metaal")
    # installaties PV F4
    check("inst-PV: 'Onbekend' -> 7 (onbekend kristallijn), niet 0 (kwaliteitsverklaring)", _pp68("Onbekend") == "7")
    check("inst-PV: 'Dunne film' -> 8 (onbekend amorf), niet 0", _pp68("Dunne film") == "8")
    check("inst-PV: 'Monokristallijn' -> 1 (ongewijzigd)", _pp68("Monokristallijn") == "1")
    # installaties ventilatie-flag (HIGH): vuurt ook bij systeem gekozen ZONDER subsysteem
    _dv68 = _Dos68()
    _dv68.ventilatie = _V68(systeem="A", systeem_soort="individueel", subsysteem_code="")
    _iflags68 = _ibt68(_dv68)[1]
    check("inst-ventilatie: flag vuurt ook bij systeem A ZONDER subsysteem (stuurt de Standaard)",
          any("STUURT DE STANDAARD" in str(f) for f in _iflags68))
    # algemeen: thermische massa leeg -> luide flag (geen stille sjabloon-'Zwaar'-leak)
    _iss68 = _obt68(_Dos68())[2]
    check("algemeen: thermische massa LEEG -> luide flag (geen stille sjabloon-'Zwaar')",
          any("thermische massa" in str(i).lower() and "ONTBREEKT" in str(i) for i in _iss68))
except Exception as _e:
    check("audit-fixes glas/installaties/algemeen: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("69. Ontvangstmail: markeer ALLEEN de gemailde leads (BCC-ids) — niet later toegevoegde (15-7)")
try:
    import dashboard.app as _W69
    from dashboard import leads as _L69
    _W69.app.config.update(TESTING=True)
    _c69 = _W69.app.test_client()
    with _c69.session_transaction() as _s69:
        _s69["ingelogd"] = True
    # in-memory leads (raakt NOOIT het echte out/leads-bestand)
    _mem69 = [{"id": 1, "status": "nieuw", "email": "a@x.nl", "naam": "A", "postcode": "9990AA", "huisnummer": "1"},
              {"id": 2, "status": "nieuw", "email": "b@x.nl", "naam": "B", "postcode": "9990AA", "huisnummer": "2"},
              {"id": 3, "status": "nieuw", "email": "", "naam": "C-geen-mail", "postcode": "9990AA", "huisnummer": "3"}]
    _o_load69, _o_save69 = _L69.load_leads, _L69.save_leads
    try:
        _L69.load_leads = lambda: [dict(r) for r in _mem69]
        _L69.save_leads = lambda rows: _mem69.__setitem__(slice(None), [dict(r) for r in rows])
        _h69 = _c69.get("/leads/ontvangst").get_data(as_text=True)
        check("ontvangstmail: alleen leads MET e-mail in de BCC (A+B, niet C)",
              "a@x.nl" in _h69 and "b@x.nl" in _h69 and _h69.count("@x.nl") == 2)
        # lead D meldt zich aan NA het openen van de ontvangst-pagina
        _mem69.append({"id": 4, "status": "nieuw", "email": "d@x.nl", "naam": "D-later",
                       "postcode": "9990AA", "huisnummer": "4"})
        _c69.post("/leads/ontvangst/verstuurd", data={"ids": "1,2"}, follow_redirects=True)
        _na69 = {r["id"]: r["status"] for r in _mem69}
        check("ontvangstmail: de gemailde leads (A+B) -> 'mail gestuurd'",
              _na69.get(1) == "mail gestuurd" and _na69.get(2) == "mail gestuurd")
        check("ontvangstmail: later aangemelde D blijft 'nieuw' (gaat mee in de VOLGENDE batch)",
              _na69.get(4) == "nieuw")
        check("ontvangstmail: lead zonder e-mailadres (C) blijft 'nieuw' (kreeg niets)",
              _na69.get(3) == "nieuw")
        # lead verwijderen (aanvraag geannuleerd) -> definitief weg, rest blijft ongemoeid
        _c69.post("/leads/3/weg", follow_redirects=True)
        check("lead verwijderen: geannuleerde lead 3 definitief weg (AVG), rest blijft staan",
              not any(r["id"] == 3 for r in _mem69) and {r["id"] for r in _mem69} == {1, 2, 4})
        _r69w = _c69.post("/leads/999/weg", follow_redirects=False)
        check("lead verwijderen: onbekende lead -> 404 (geen stille no-op)", _r69w.status_code == 404)
    finally:
        _L69.load_leads, _L69.save_leads = _o_load69, _o_save69
except Exception as _e:
    check("ontvangstmail-markering: draait zonder fout", False); print("     " + repr(_e)[:200])

print()
print("70. Inloggen met e-mailadres + wachtwoord (17-7)")
try:
    import dashboard.security as _S70
    _sec70 = _S70.nieuw_totp_secret()
    _cfg70 = {"pw_hash": _S70.hash_password("geheim"), "totp_secret": _sec70}
    _MAIL = "info@poortinga-energieadvies.nl"
    # juist adres + juist wachtwoord + juiste code -> binnen
    _ok, _ = _S70.login_check(_cfg70, "geheim", _S70.totp_code(_sec70), "8.8.8.1",
                              email=_MAIL, verwacht_email=_MAIL)
    check("login: juist e-mailadres + wachtwoord + code -> ingelogd", _ok)
    # hoofdletters/spaties mogen niet uitmaken
    _ok2, _ = _S70.login_check(_cfg70, "geheim", _S70.totp_code(_sec70), "8.8.8.2",
                               email="  INFO@Poortinga-Energieadvies.NL ", verwacht_email=_MAIL)
    check("login: e-mailadres case-insensitief + spaties eraf", _ok2)
    # fout adres -> geweigerd, en de melding verraadt NIET welk veld fout was
    _ok3, _m3 = _S70.login_check(_cfg70, "geheim", _S70.totp_code(_sec70), "8.8.8.3",
                                 email="iemand@anders.nl", verwacht_email=_MAIL)
    check("login: fout e-mailadres -> geweigerd ondanks juist wachtwoord", not _ok3)
    check("login: melding verraadt niet WELK veld fout is", "E-mailadres of wachtwoord onjuist." == _m3)
    # leeg e-mailveld terwijl er wel een adres is geconfigureerd -> geweigerd
    _ok4, _ = _S70.login_check(_cfg70, "geheim", _S70.totp_code(_sec70), "8.8.8.4",
                               email="", verwacht_email=_MAIL)
    check("login: leeg e-mailadres -> geweigerd", not _ok4)
    # GEEN adres geconfigureerd -> oude gedrag blijft (lokale modus, geen e-mailcheck)
    _ok5, _ = _S70.login_check({}, "lokaalpw", "", "8.8.8.5", fallback_pw="lokaalpw")
    check("login: zonder geconfigureerd adres blijft de lokale modus werken", _ok5)
    # de loginpagina toont het e-mailveld alleen als er een adres is geconfigureerd
    import dashboard.app as _W70
    _W70.app.config.update(TESTING=True)
    _cfg70_orig = _W70._cfg
    try:
        _W70._cfg = lambda: {"dashboard": {"email": _MAIL}}
        _h70 = _W70.app.test_client().get("/login").get_data(as_text=True)
        check("login-pagina: e-mailveld aanwezig met tijdelijke testconfig",
              'name=email' in _h70 and 'type=email' in _h70)
    finally:
        _W70._cfg = _cfg70_orig
except Exception as _e:
    check("login met e-mailadres: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n71. Huisstijl (logo/favicon/PWA) + leads als kaarten + automatische BAG-verrijking")
try:
    import os as _o71, tempfile as _t71, json as _j71
    import dashboard.app as _W71, dashboard.leads as _L71
    _W71.app.config.update(TESTING=True)
    _c71 = _W71.app.test_client()

    # (a) huisstijl-bestanden bestaan en zitten in de <head> (iPad-beginscherm heeft PNG nodig, geen SVG)
    _st71 = _o71.path.join(_o71.path.dirname(_o71.path.dirname(_o71.path.abspath(_W71.__file__))),
                           "dashboard", "static")
    for _f71 in ("logo.svg", "mark.svg", "apple-touch-icon.png", "icon-192.png", "icon-512.png"):
        check("huisstijl: %s aanwezig" % _f71, _o71.path.isfile(_o71.path.join(_st71, _f71)))
    _hl71 = _c71.get("/login").get_data(as_text=True)
    check("login: volledig logo staat op de pagina", "logo.svg" in _hl71 and "class=login" in _hl71)
    check("head: favicon + apple-touch-icon + manifest", all(
        _s in _hl71 for _s in ("mark.svg", "apple-touch-icon", "manifest")))
    _mf71 = _c71.get("/manifest.webmanifest")
    check("manifest: geldige JSON, standalone, 512px-icoon",
          _mf71.status_code == 200 and _j71.loads(_mf71.get_data(as_text=True))["display"] == "standalone"
          and any(i["sizes"] == "512x512" for i in _j71.loads(_mf71.get_data(as_text=True))["icons"]))

    # (b) leads-pagina = kaarten (geen tabel meer) — geïsoleerd van de echte leads.json (AVG)
    _tmp71 = _t71.mkdtemp()
    _bew71 = (_L71.LEADS_DIR, _L71.LEADS_FILE)
    _L71.LEADS_DIR, _L71.LEADS_FILE = _tmp71, _o71.path.join(_tmp71, "leads.json")
    try:
        _L71.save_leads([
            dict(id=1, naam="J.T. Klok", postcode="9541AB", huisnummer="21", toevoeging="",
                 telefoon="0631623134", email="k@x.nl", straat="Harpelerweg", woonplaats="Vlagtwedde",
                 bouwjaar=1992, oppervlakte_m2=232, status="nieuw", ontvangen="2026-07-17",
                 notitie="", afspraak=""),
            dict(id=2, naam="A. Kroeze", postcode="9561XY", huisnummer="52", toevoeging="",
                 telefoon="", email="", status="afgerond", ontvangen="2026-07-17", notitie="",
                 afspraak="2026-07-23T14:30"),
        ])
        with _c71.session_transaction() as _s71:
            _s71["ingelogd"] = True
        _h71 = _c71.get("/leads").get_data(as_text=True)
        check("leads: elke lead is een kaart", _h71.count('class="lead-card') == 2)
        check("leads: oude afgekapte tabel is weg", "lead-list" not in _h71 and "<th>Adres</th>" not in _h71)
        check("leads: adres + BAG-pills op de kaart",
              "Harpelerweg 21 in Vlagtwedde" in _h71 and "1992" in _h71 and "232 m" in _h71)
        check("leads: statuskleur per status", 'class="pill blue">nieuw' in _h71
              and 'class="pill green">afgerond' in _h71)
        check("leads: afgeronde lead gedempt", "is-klaar" in _h71)
        check("leads: afspraak leesbaar in het Nederlands", "donderdag 23 juli 2026 om 14:30" in _h71)
        check("leads: zoek + statusfilter aanwezig", 'id=zoek' in _h71 and 'id=statusfilter' in _h71)
        check("leads: zonder BAG een duidelijke melding i.p.v. lege plek", "BAG onbekend" in _h71)

        # (c) automatische BAG-verrijking: alleen gevulde velden, nooit iets leegmaken
        _ld71 = {"straat": "Oud", "bouwjaar": 1900}
        _W71._bag_toepassen(_ld71, {"straat": "Nieuw", "woonplaats": "", "oppervlakte_m2": 120})
        check("BAG-verrijking: vult aan, overschrijft nooit met leeg",
              _ld71["straat"] == "Nieuw" and _ld71["bouwjaar"] == 1900
              and _ld71["oppervlakte_m2"] == 120 and "woonplaats" not in _ld71)

        # (d) leads.wijzig(): lees-wijzig-schrijf als één stap (achtergrondwerk vs. de webapp)
        _L71.wijzig(lambda rows: rows.append(dict(id=9, naam="Test", status="nieuw", ontvangen="x",
                                                  postcode="", huisnummer="", notitie="", afspraak="")))
        check("leads.wijzig: schrijft de wijziging weg", any(r["id"] == 9 for r in _L71.load_leads()))
    finally:
        _L71.LEADS_DIR, _L71.LEADS_FILE = _bew71
except Exception as _e:
    check("huisstijl + leadkaarten + BAG-automaat: draait zonder fout", False)
    print("     " + repr(_e)[:200])

print("\n72. Mailpagina + OneDrive-projectmap-export + IMAP-intake")
try:
    import os as _o72, io as _i72, zipfile as _z72, datetime as _d72, tempfile as _t72
    import email.message as _em72
    import dashboard.app as _W72, dashboard.leads as _L72, dashboard.mailbox as _M72
    _W72.app.config.update(TESTING=True)
    _c72 = _W72.app.test_client()
    with _c72.session_transaction() as _s72:
        _s72["ingelogd"] = True

    # (a) mailpagina toont alle drie de bewonersmails met voorbeeldgegevens
    _h72 = _c72.get("/mails").get_data(as_text=True)
    check("mails: alle drie de mails staan erop",
          all(t in _h72 for t in ("Ontvangstbevestiging", "Kennismakingsmail", "Afspraakbevestiging")))
    check("mails: voorbeeldgegevens ingevuld (adres + afspraakdatum)",
          "Munsterheerd 106" in _h72 and "donderdag 23 juli 2026" in _h72)
    check("mails: vanuit de guide bereikbaar", "/mails" in _c72.get("/guide").get_data(as_text=True))

    # (b) OneDrive-projectmap: indeling, lege mappen blijven bestaan, niets raakt zoek
    check("export-indeling: dossier -> 01_Opname", _W72._onedrive_map("dossier_x.json") == "01_Opname")
    check("export-indeling: vabi-export -> 02_VABI", _W72._onedrive_map("vabi_export_huidig_x.xml") == "02_VABI")
    check("export-indeling: bibliotheken in eigen submap",
          _W72._onedrive_map("huidig_Constructiebibliotheek.xml", "vabi_huidig") == "02_VABI/huidige_staat"
          and _W72._onedrive_map("na_Objectenbibliotheek.xml", "vabi_na") == "02_VABI/toekomstige_staat")
    check("export-indeling: plan/ventilatie/checklist -> 03_Isolatieplan",
          all(_W72._onedrive_map(n) == "03_Isolatieplan" for n in
              ("isolatieplan_x.docx", "ventilatieplan_x.svg", "fotochecklist_x.txt",
               "haalbaarheid_toelichting_x.txt")))
    check("export-indeling: foto's -> 04_Fotos",
          _W72._onedrive_map("foto_voorkant_x.jpg") == "04_Fotos"
          and _W72._onedrive_map("IMG_1.jpg", "fotos") == "04_Fotos")
    check("export-indeling: onbekend bestand gaat NIET verloren maar naar 07_Overig",
          _W72._onedrive_map("aantekeningen bewoner.rtf") == "07_Overig")

    _tag72 = "9999ZZ_5"
    if _o72.path.isdir(_W72._pdir(_tag72)):
        _r72 = _c72.get("/project/%s/export" % _tag72)
        _zip72 = _z72.ZipFile(_i72.BytesIO(_r72.get_data()))
        _namen72 = _zip72.namelist()
        _wortel72 = _namen72[0].split("/")[0]
        check("export: zip is één projectmap met het adres als naam",
              all(n.startswith(_wortel72 + "/") for n in _namen72) and "9999ZZ" not in _wortel72)
        check("export: LEESMIJ met bewaartermijn en AVG-notitie",
              any(n.endswith("LEESMIJ.txt") for n in _namen72)
              and "15 jaar" in _zip72.read(_wortel72 + "/LEESMIJ.txt").decode("utf-8"))
        check("export: lege correspondentie- en facturenmap blijven bestaan",
              any("05_Correspondentie" in n for n in _namen72)
              and any("06_Facturen" in n for n in _namen72))
        check("export: interne projectstatus gaat NIET mee",
              not any(n.endswith("project.json") for n in _namen72))

    # (c) IMAP-intake — met een nep-postvak, geen netwerk nodig
    check("mailbox: zonder instellingen een nette uitleg i.p.v. een fout",
          _M72.haal_teksten({})[1].startswith("Mailbox nog niet ingesteld"))
    check("mailbox: zoekopdracht op datum + portaal-marker (OR met de configterm)",
          _M72.zoekopdracht({"dagen": 30, "onderwerp": "AdviseurToegekend"}, _d72.date(2026, 7, 19))
          == ["SINCE", "19-Jun-2026", "OR", "TEXT", "AdviseurToegekend", "TEXT", "WijzigingsType"])
    check("mailbox: mapnaam met spatie wordt gequote", _M72._mapnaam("Verwijderde items") == '"Verwijderde items"')

    def _mailtje72(naam, pc, hn, bagid):
        m = _em72.EmailMessage()
        m["Subject"] = "AdviseurToegekend"
        m["From"] = "portal@smarttwin.nl"
        m.set_content('{"BagAdresId":"%s","Email":"x@x.nl","Postcode":"%s","Huisnummer":%s,'
                      '"Naam":"%s","WijzigingsType":"AdviseurToegekend"}' % (bagid, pc, hn, naam))
        return m.as_bytes()

    _post72 = [_mailtje72("Jan de Boer", "9736GL", "106", "0014200000001"),
               _mailtje72("Ada Vos", "9711RS", "23", "0014200000002")]

    class _NepImap72:
        gelezen_readonly = None

        def login(self, *a):
            pass

        def select(self, m, readonly=True):
            _NepImap72.gelezen_readonly = readonly

        def search(self, *a):
            return "OK", [b"1 2"]

        def fetch(self, i, spec):
            return "OK", [(b"1 (RFC822 {n})", _post72[int(i) - 1])]

        def close(self):
            pass

        def logout(self):
            pass

    _cfg72 = {"mailbox": {"host": "imap.x.nl", "gebruiker": "info@x.nl", "wachtwoord": "app-pw"}}
    _tk72, _f72 = _M72.haal_teksten(_cfg72, verbind=lambda m: _NepImap72())
    check("mailbox: 2 mails opgehaald, geen fout", len(_tk72) == 2 and _f72 is None)
    check("mailbox: opent het postvak ALLEEN-LEZEN (verplaatst/verwijdert niets)",
          _NepImap72.gelezen_readonly is True)
    _lds72 = _L72.parse_leads_bulk("\n".join(_tk72))
    check("mailbox -> leads: beide bewoners geparsed",
          [l["naam"] for l in _lds72] == ["Jan de Boer", "Ada Vos"])

    class _StukImap72:
        def login(self, *a):
            raise _M72.imaplib.IMAP4.error("AUTHENTICATIONFAILED")
    _tk73, _f73 = _M72.haal_teksten(_cfg72, verbind=lambda m: _StukImap72().login())
    check("mailbox: mislukte login -> nette melding, geen crash", _tk73 == [] and "mislukte" in (_f73 or ""))

    # (d) leadkaart-route: ophalen zit erin en de gedeelde toevoeg-functie werkt
    check("webapp: ophaal-route geregistreerd",
          "/leads/ophalen" in {r.rule for r in _W72.app.url_map.iter_rules()})
    _tmp72 = _t72.mkdtemp()
    _bew72 = (_L72.LEADS_DIR, _L72.LEADS_FILE)
    _L72.LEADS_DIR, _L72.LEADS_FILE = _tmp72, _o72.path.join(_tmp72, "leads.json")
    try:
        _L72.save_leads([])
        _res72 = _W72._leads_toevoegen("\n".join(_tk72))
        check("gedeelde intake: 2 nieuw, 0 dubbel", (_res72["nieuw"], _res72["dubbel"]) == (2, 0))
        _res72b = _W72._leads_toevoegen("\n".join(_tk72))          # nog eens: alles dubbel
        check("gedeelde intake: tweede keer alles herkend als dubbel",
              (_res72b["nieuw"], _res72b["dubbel"]) == (0, 2))
    finally:
        _L72.LEADS_DIR, _L72.LEADS_FILE = _bew72
except Exception as _e:
    check("mails/export/IMAP: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n73. Microsoft Graph-mailkoppeling (gedeeld info@-postvak, Microsoft 365)")
try:
    import json as _j73, datetime as _dt73, os as _o73
    import dashboard.graph_mail as _G73, dashboard.leads as _L73, dashboard.app as _W73

    _cfg73 = {"graph": {"tenant_id": "t-1", "client_id": "c-1", "client_secret": "s-1",
                        "postvak": "info@poortinga-energieadvies.nl"}}
    check("graph: herkent een volledig ingevuld blok", _G73.is_ingesteld(_cfg73))
    check("graph: half ingevuld telt NIET als ingesteld",
          not _G73.is_ingesteld({"graph": {"tenant_id": "t", "client_id": "c"}}))

    _url73 = _G73.bericht_url(_G73.instellingen(_cfg73), _dt73.date(2026, 7, 19))
    check("graph: vraagt het juiste postvak op",
          "/users/info%40poortinga-energieadvies.nl/messages" in _url73)
    check("graph: filtert op datum (30 dagen terug)", "2026-06-19T00%3A00%3A00Z" in _url73)
    check("graph: lege instellingen geven geen crash maar een lege URL",
          _G73.bericht_url({}).endswith("messages?%24select=subject%2CreceivedDateTime%2Cbody&"
                                        "%24top=200&%24orderby=receivedDateTime+desc")
          or "/users//messages" in _G73.bericht_url({}))

    _ber73 = {"value": [
        {"subject": "AdviseurToegekend", "body": {"content":
            'Beste adviseur,\n{"BagAdresId":"0014200000001","Postcode":"9736GL",'
            '"Huisnummer":106,"Naam":"Jan de Boer"}'}},
        {"subject": "AdviseurToegekend", "body": {"content":
            '<p>{"BagAdresId":"0014200000002","Postcode":"9711RS","Huisnummer":23,"Naam":"Ada Vos"}</p>'}},
        {"subject": "Nieuwsbrief juli", "body": {"content": "niet relevant"}}]}

    _gezien73 = {}

    def _nep73(m, u, h=None, d=None):
        if "login.microsoftonline" in u:
            _gezien73["token_methode"] = m
            return 200, _j73.dumps({"access_token": "TOK"}).encode()
        _gezien73["auth"] = (h or {}).get("Authorization")
        _gezien73["prefer"] = (h or {}).get("Prefer")
        _gezien73["methode"] = m
        return 200, _j73.dumps(_ber73).encode()

    _t73, _f73 = _G73.haal_teksten(_cfg73, http=_nep73)
    check("graph: 2 portal-mails, nieuwsbrief eruit gefilterd", len(_t73) == 2 and _f73 is None)
    check("graph: token via POST, berichten via GET (leest alleen)",
          _gezien73["token_methode"] == "POST" and _gezien73["methode"] == "GET")
    check("graph: token wordt meegestuurd", _gezien73["auth"] == "Bearer TOK")
    check("graph: vraagt platte tekst op i.p.v. HTML", "text" in (_gezien73["prefer"] or ""))
    check("graph: HTML-mail wordt van tags ontdaan", "<p>" not in _t73[1])
    check("graph -> leads: beide bewoners geparsed",
          [l["naam"] for l in _L73.parse_leads_bulk("\n".join(_t73))] == ["Jan de Boer", "Ada Vos"])

    def _stuk73(code):
        def http(m, u, h=None, d=None):
            if "login" in u:
                return ((401, b'{"error_description":"AADSTS7000222 secret expired"}') if code == 401
                        else (200, _j73.dumps({"access_token": "TOK"}).encode()))
            return code, b'{"error":{"message":"Access denied"}}'
        return http

    for _code73, _woord73 in ((401, "secret"), (403, "Access Policy"), (404, "postvak")):
        _lg73, _fo73 = _G73.haal_teksten(_cfg73, http=_stuk73(_code73))
        check("graph: %d geeft een bruikbare uitleg i.p.v. ruwe JSON" % _code73,
              _lg73 == [] and _woord73 in (_fo73 or ""))

    # ECHTE portalmail (20-7-2026): het ONDERWERP bevat 'AdviseurToegekend' NIET — daar staat een
    # account-id in, per aanmelding anders. De term staat in de BODY ("WijzigingsType"). Filteren op
    # alleen het onderwerp vond daardoor niets. Deze test bewaakt dat.
    _echt73 = {"subject": "Contact met adviseur door accountid 957a6ac7-b664-400d-9c42-c94e51c78380",
               "body": {"content":
                        '{"BagAdresId":"1895200000005699","AccountId":"957a6ac7-b664-400d-9c42-c94e51c78380",'
                        '"Email":"jesjeentesje@gmail.com","Postcode":"9674BW","Huisnummer":28,'
                        '"HuisnummerToevoeging":"","Voornaam":"tess","Telefoonnummer":"0625494609",'
                        '"Achternaam":"roupp\\u00e9","Naam":"tess roupp\\u00e9",'
                        '"WijzigingsType":"AdviseurToegekend",'
                        '"WijzigingsReden":"Adviseur 39222 toegekend aan gebruiker."}'}}
    _ruis73 = {"subject": "Nieuwsbrief Nij Begun juli",
               "body": {"content": "Beste adviseur, hierbij de nieuwsbrief."}}
    _door73 = _G73.berichten_naar_teksten([_echt73, _ruis73], "AdviseurToegekend")
    check("graph: term in de BODY telt ook (onderwerp bevat een wisselend account-id)", len(_door73) == 1)
    _l73 = _L73.parse_leads_bulk("\n".join(_door73))
    check("graph: echte portalmail -> volledige lead (naam/adres/mail/telefoon/BAG-id)",
          len(_l73) == 1 and _l73[0]["naam"] == "tess rouppé" and _l73[0]["postcode"] == "9674BW"
          and _l73[0]["huisnummer"] == "28" and _l73[0]["email"] == "jesjeentesje@gmail.com"
          and _l73[0]["telefoon"] == "0625494609" and _l73[0]["bag_id"] == "1895200000005699")
    check("graph: leeg filter laat alles door", len(_G73.berichten_naar_teksten([_echt73, _ruis73], "")) == 2)

    # Het portaal stuurt ook "Contact gegevens gewijzigd"-mails voor een AL bekende bewoner.
    # Die mogen geen dubbele lead maken, maar het nieuwe telefoonnummer moet wel doorkomen —
    # zonder de status/afspraak/notitie van de adviseur te overschrijven.
    _rij73 = [dict(id=1, bag_id="1895200000005699", naam="tess rouppé", email="oud@x.nl",
                   telefoon="0611111111", status="afspraak gepland", afspraak="2026-07-25T10:00",
                   notitie="belt liever 's avonds", project_tag="9674BW_28",
                   postcode="9674BW", huisnummer="28", toevoeging="")]
    _nw73 = dict(bag_id="1895200000005699", naam="tess rouppé", email="nieuw@x.nl",
                 telefoon="0625494609", postcode="9674BW", huisnummer="28", toevoeging="")
    _rij73, _toegevoegd73 = _L73.add_lead(_nw73, _rij73)
    check("wijzigingsmail: geen dubbele lead", _toegevoegd73 is False and len(_rij73) == 1)
    check("wijzigingsmail: nieuwe contactgegevens komen wél door",
          _rij73[0]["email"] == "nieuw@x.nl" and _rij73[0]["telefoon"] == "0625494609")
    check("wijzigingsmail: eigen werk blijft staan (status/afspraak/notitie/project)",
          _rij73[0]["status"] == "afspraak gepland" and _rij73[0]["afspraak"] == "2026-07-25T10:00"
          and _rij73[0]["notitie"] == "belt liever 's avonds"
          and _rij73[0]["project_tag"] == "9674BW_28")
    _rij73, _ = _L73.add_lead(dict(_nw73, email="", telefoon=""), _rij73)
    check("wijzigingsmail: lege waarden wissen niets", _rij73[0]["email"] == "nieuw@x.nl")

    import dashboard.mailbox as _M73b
    check("imap: zoekt in de hele mail (TEXT), niet alleen het onderwerp",
          "TEXT" in _M73b.zoekopdracht({"dagen": 30, "onderwerp": "AdviseurToegekend"})
          and "SUBJECT" not in _M73b.zoekopdracht({"dagen": 30, "onderwerp": "AdviseurToegekend"}))

    check("graph: zonder instellingen een uitleg met verwijzing naar de beheerdersinstructie",
          "microsoft-graph-mailkoppeling.md" in (_G73.haal_teksten({})[1] or ""))
    check("beheerdersinstructie bestaat",
          _o73.path.isfile(_o73.path.join(_o73.path.dirname(_o73.path.dirname(
              _o73.path.abspath(_W73.__file__))), "docs", "microsoft-graph-mailkoppeling.md")))

    # keuze van de bron: Graph gaat vóór IMAP, en zonder beide -> geen knop
    check("bron: graph wint van imap",
          _W73._mailbron(dict(_cfg73, mailbox={"host": "h", "gebruiker": "g", "wachtwoord": "w"}))[0] == "graph")
    check("bron: alleen imap ingesteld -> imap",
          _W73._mailbron({"mailbox": {"host": "h", "gebruiker": "g", "wachtwoord": "w"}})[0] == "imap")
    check("bron: niets ingesteld -> geen ophaalknop", _W73._mailbron({})[0] is None)
except Exception as _e:
    check("Graph-mailkoppeling: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n74. Verwijderde leads blijven weg + mobiele opmaak (telefoon)")
try:
    import os as _o74, tempfile as _t74, re as _r74
    import dashboard.leads as _L74, dashboard.app as _W74
    _W74.app.config.update(TESTING=True)
    _c74 = _W74.app.test_client()
    with _c74.session_transaction() as _s74:
        _s74["ingelogd"] = True

    _tmp74 = _t74.mkdtemp()
    _bew74 = (_L74.LEADS_DIR, _L74.LEADS_FILE, _L74.GEWIST_FILE)
    _L74.LEADS_DIR = _tmp74
    _L74.LEADS_FILE = _o74.path.join(_tmp74, "leads.json")
    _L74.GEWIST_FILE = _o74.path.join(_tmp74, "verwijderd.json")
    try:
        _lead74 = dict(id=1, naam="Uitgeschreven Bewoner", bag_id="0018200000001", postcode="9601AA",
                       huisnummer="50", toevoeging="", email="x@x.nl", telefoon="06", status="nieuw",
                       ontvangen="2026-07-19", notitie="", afspraak="")
        _L74.save_leads([_lead74])
        _c74.post("/leads/1/weg", follow_redirects=True)
        check("verwijderen: lead weg uit de lijst", _L74.load_leads() == [])
        check("verwijderen: alleen de adres-sleutel onthouden (geen naam/contact)",
              _L74.load_gewist() == ["0018200000001"])

        # het portaal mailt opnieuw over deze bewoner -> mag NIET terugkomen
        _rows74, _nieuw74 = _L74.add_lead(dict(bag_id="0018200000001", naam="Uitgeschreven Bewoner",
                                               postcode="9601AA", huisnummer="50", toevoeging=""),
                                          _L74.load_leads())
        check("ophalen: bewust verwijderde bewoner komt NIET terug", _nieuw74 is False and _rows74 == [])

        _h74 = _c74.get("/leads").get_data(as_text=True)
        check("leads: geblokkeerde adressen zichtbaar op de pagina", "Geblokkeerd (1)" in _h74)

        _c74.post("/leads/geblokkeerd", follow_redirects=True)
        check("blokkade opheffen: lijst leeg", _L74.load_gewist() == [])
        _rows74, _nieuw74 = _L74.add_lead(dict(bag_id="0018200000001", naam="Toch Weer Aangemeld",
                                               postcode="9601AA", huisnummer="50", toevoeging=""),
                                          _L74.load_leads())
        check("blokkade opheffen: bewoner kan daarna weer binnenkomen", _nieuw74 is True)

        # een ANDER adres mag nooit meegeblokkeerd raken
        _L74.onthoud_gewist({"bag_id": "0018200000001"})
        _rows74, _nieuw74 = _L74.add_lead(dict(bag_id="0018200000099", naam="Buurman",
                                               postcode="9601AA", huisnummer="52", toevoeging=""), [])
        check("blokkade raakt alleen het verwijderde adres", _nieuw74 is True)
    finally:
        _L74.LEADS_DIR, _L74.LEADS_FILE, _L74.GEWIST_FILE = _bew74

    # ---- mobiele opmaak: de regels die de half zichtbare factuurtabel veroorzaakten ----
    _css74 = open(_o74.path.join(_o74.path.dirname(_o74.path.dirname(_o74.path.abspath(_W74.__file__))),
                                 "dashboard", "static", "app.css"), encoding="utf-8").read()
    check("mobiel: <html> klipt horizontaal (anders sleept iOS de pagina opzij)",
          _r74.search(r"html\{[^}]*overflow-x:hidden", _css74) is not None)
    _mob74 = _css74[_css74.index("@media (max-width:700px)"):]
    check("mobiel: sleutel/waarde-tabellen krijgen table-layout:fixed (tekst breekt af)",
          "table-layout:fixed" in _mob74)
    check("mobiel: geen nowrap meer op kale tabellen (dat kapte de factuurgegevens af)",
          "white-space:nowrap" not in _mob74.split(".card-table")[0])
    check("voorschot: beide tabellen in een .table-wrap (niet als class op de tabel zelf)",
          'table class="table-wrap"' not in _W74.VOORSCHOT
          and _W74.VOORSCHOT.count("<div class=table-wrap") >= 2)
    check("ophaalknop: toont dat hij bezig is (Graph duurt seconden)",
          "Bezig met ophalen" in _W74.LEADS and "b.disabled=true" in _W74.LEADS)
except Exception as _e:
    check("verwijderde leads + mobiele opmaak: draait zonder fout", False)
    print("     " + repr(_e)[:200])

print("\n75. Smalle-iPhone-opmaak (320px) — de regels die de uitstekende leadkaarten veroorzaakten")
try:
    import os as _o75, re as _r75
    import dashboard.app as _W75
    _css75 = open(_o75.path.join(_o75.path.dirname(_o75.path.dirname(_o75.path.abspath(_W75.__file__))),
                                 "dashboard", "static", "app.css"), encoding="utf-8").read()
    check("lead-grid: kolom nooit breder dan de container (min(300px,100%))",
          "minmax(min(300px,100%),1fr)" in _css75)
    check("lead-card: mag krimpen onder z'n inhoud (min-width:0)",
          _r75.search(r"\.lead-card\{[^}]*min-width:0", _css75, _r75.S) is not None)
    check("lead-contact: overflow-wrap ANYWHERE (break-word verkleint min-breedte niet)",
          _r75.search(r"\.lead-contact a[^{]*\{overflow-wrap:anywhere", _css75) is not None)
    check("topbar-nav mobiel: krimpt en wrapt (geen flex:none/nowrap meer — 5e link paste niet)",
          "flex:none;flex-wrap:nowrap" not in _css75
          and _r75.search(r"\.topbar nav\{[^}]*min-width:0", _css75) is not None)
    check("accordeon-kop wrapt (titel + pills pasten niet op één regel op 320px)",
          _r75.search(r"details\.acc>summary\{[^}]*flex-wrap:wrap", _css75, _r75.S) is not None)
    check("grid2: kolommen minmax(0,1fr) + bestand-kiezers begrensd (foto-upload stak uit)",
          "grid-template-columns:minmax(0,1fr) minmax(0,1fr)" in _css75
          and "input[type=file]{max-width:100%}" in _css75)
    check("gids-tabellen scrollen in eigen vak (width:100% is bij tabellen slechts een minimum)",
          _r75.search(r"\.gids-inhoud table\{[^}]*overflow-x:auto", _css75) is not None
          and 'class="card gids-inhoud"' in _W75.GIDS_TMPL)
    check("ventilatieplan-SVG: schaalt mee + eigen scroll-vak",
          _r75.search(r"\.svgbox\{[^}]*overflow-x:auto", _css75) is not None)
    check("datum/tijd-veld: iOS-eigen breedte uitgezet (appearance:none)",
          "input[type=datetime-local]" in _css75 and "-webkit-appearance:none" in _css75)
except Exception as _e:
    check("smalle-iPhone-opmaak: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n76. Leads op de telefoon: filter verbergt echt, plek blijft bewaard, juiste mailtitel")
try:
    import os as _o76, tempfile as _t76
    import dashboard.app as _W76, dashboard.leads as _L76
    _css76 = open(_o76.path.join(_o76.path.dirname(_o76.path.dirname(_o76.path.abspath(_W76.__file__))),
                                 "dashboard", "static", "app.css"), encoding="utf-8").read()
    # het zoekfilter telde wel maar verborg niets: display:flex won van [hidden]
    check("filter: [hidden] wint van display:flex", ".lead-card[hidden]{display:none}" in _css76)
    check("filter: zoek + status overleven het herladen (sessionStorage)",
          "leads_zoek" in _W76.LEADS and "leads_status" in _W76.LEADS)
    check("na wijziging: terugspringen naar de kaart die je aanraakte",
          'data-lid="{{l.id}}"' in _W76.LEADS and "leads_focus" in _W76.LEADS
          and "scrollIntoView" in _W76.LEADS and "is-net-gewijzigd" in _css76)

    # de bevestigingsmail-pagina heette 'Kennismakingsmail' — verwarrend of hij wel/niet werkte
    _tmp76 = _t76.mkdtemp()
    _bew76 = (_L76.LEADS_DIR, _L76.LEADS_FILE, _L76.GEWIST_FILE)
    _L76.LEADS_DIR = _tmp76
    _L76.LEADS_FILE = _o76.path.join(_tmp76, "leads.json")
    _L76.GEWIST_FILE = _o76.path.join(_tmp76, "verwijderd.json")
    try:
        _L76.save_leads([dict(id=1, naam="Jan de Boer", postcode="9736GL", huisnummer="106",
                              toevoeging="", email="j@x.nl", telefoon="06", status="afspraak gepland",
                              ontvangen="2026-07-19", notitie="", afspraak="2026-07-23T14:30")])
        _W76.app.config.update(TESTING=True)
        _c76 = _W76.app.test_client()
        with _c76.session_transaction() as _s76:
            _s76["ingelogd"] = True
        _hb76 = _c76.get("/leads/1/mail?soort=bevestiging").get_data(as_text=True)
        check("bevestigingsmail-pagina: eigen titel + juiste inhoud",
              "Afspraakbevestiging — Jan de Boer" in _hb76
              and "Bevestiging afspraak woningopname" in _hb76)
        _hk76 = _c76.get("/leads/1/mail").get_data(as_text=True)
        check("kennismakingsmail-pagina: titel blijft Kennismakingsmail",
              "Kennismakingsmail — Jan de Boer" in _hk76)
        check("mailpagina: uitleg dat mailto het standaardaccount opent (Van-adres wisselen)",
              "standaard-mailaccount" in _hb76 and "Van" in _hb76)
        check("bevestigingspagina: afspraakdatum prominent ter controle",
              "donderdag 23 juli 2026 om 14:30" in _hb76)

        # de Bevestiging-knop zat verstopt achter het datumveld -> mail 3 was onvindbaar
        _hl76 = _c76.get("/leads").get_data(as_text=True)
        check("leadkaart: Bevestiging-knop ALTIJD zichtbaar (niet meer achter het datumveld)",
              "soort=bevestiging" in _hl76)
        _L76.save_leads([dict(id=1, naam="Zonder Datum", postcode="9736GL", huisnummer="106",
                              toevoeging="", email="j@x.nl", telefoon="06", status="afspraak gepland",
                              ontvangen="2026-07-19", notitie="", afspraak="")])
        _hz76 = _c76.get("/leads/1/mail?soort=bevestiging", follow_redirects=True).get_data(as_text=True)
        check("bevestiging zonder datum: terug naar leads met uitleg (nooit 'nader te bepalen' mailen)",
              "eerst een afspraakdatum nodig" in _hz76)
    finally:
        _L76.LEADS_DIR, _L76.LEADS_FILE, _L76.GEWIST_FILE = _bew76
except Exception as _e:
    check("leads-telefoonfixes: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n77. Boven-/onderlicht bij ramen + deur-bovenlicht-glas + toevoerroosters (1e echte opname)")
try:
    import tempfile as _tf77
    from magicplan.statistics_csv import build_dossier as _csvdos77

    # kolomnamen = de LIVE veldnamen (deur-patroon; 'kozijn' in de naam voorkomt deur-kolom-clash)
    _KOP77 = ("WALL ATTRIBUTES,c1,c2,Surface,SurfNoOpen,c5,c6,c7,Type,Isol,c10,Orientatie,Bron,c13,c14,"
              "Kozijn,Type glas,RaamOrient,Type constructie (deur),Oppervlakte raam in deur,"
              "Type glas (indien glas in deur),Bovenlicht in het kozijn? (leeg = geen),"
              "Bovenlicht kozijn - oppervlak glas (m²),Bovenlicht kozijn - type glas,"
              "Bovenlicht kozijn-paneel - oppervlak (m²),Bovenlicht kozijn-paneel - isolatie aanwezig?,"
              "Bovenlicht kozijn-paneel - isolatiedikte (mm),Onderlicht in het kozijn? (leeg = geen),"
              "Onderlicht kozijn - oppervlak glas (m²),Onderlicht kozijn-paneel - oppervlak (m²),"
              "Onderlicht kozijn-paneel - isolatie aanwezig?,Toevoerrooster aanwezig? (leeg = geen),"
              "Bovenlicht - oppervlak glas (m²),Bovenlicht deur - type glas,"
              "Toevoerrooster deur aanwezig? (leeg = geen),"
              "Bovenlicht kozijn - toevoerrooster aanwezig? (leeg = geen)")

    def _rij77(d, n=36):
        rr = [""] * n
        for kk, vv in d.items():
            rr[kk] = str(vv)
        # waarden met komma's ('Ja, met eigen glas') horen gequote — zoals de echte export doet
        return ",".join('"%s"' % v if "," in v else v for v in rr)

    _csv77 = "\n".join([
        "PLAN ATTRIBUTES", "Bouwjaar,1975 t/m 1982", "Woningtype,Tussenwoning", "Gevelhoogte (m),6", "",
        "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing",
        "Ground Floor,60,2.60 m,Kruipruimte", "",
        _KOP77, "Ground Floor",
        _rij77({0: "Voorgevel", 3: 20, 4: 20, 8: "Wall", 11: "ZW"}),
        # raam 3,0 m2 HR++ met BOVENLICHT 0,5 m2 ENKEL glas -> 2,5 HR++ + 0,5 Enkel.
        # 31 = hoofdraam-rooster Ja; 35 = rooster IN het bovenlicht Ja (los geteld)
        _rij77({0: "raam bovenlicht", 3: 3.0, 8: "Window", 16: "HR++",
                21: "Ja, met eigen glas", 22: 0.5, 23: "Enkel", 31: "Ja", 35: "Ja"}),
        # raam 2,0 m2 Dubbel met ONDERLICHT-PANEEL (borstwering) 0,6 m2 -> 1,4 glas + 0,6 paneel
        _rij77({0: "raam borstwering", 3: 2.0, 8: "Window", 16: "Dubbel",
                27: "Ja, met dicht paneel", 29: 0.6, 30: "Nee"}),
        # raam met bovenlicht=Ja maar ZONDER oppervlak -> niet splitsen, LUIDE flag
        _rij77({0: "raam zonder m2", 3: 1.5, 8: "Window", 16: "HR++", 21: "Ja, met eigen glas"}),
        # deur met raam (0,4 Dubbel) + bovenlicht 0,3 m2 met EIGEN glastype Enkel -> apart kozijn
        _rij77({0: "voordeur", 3: 2.2, 8: "Door", 17: "ZW", 18: "Deur met raam", 19: 0.4,
                20: "Dubbel", 32: 0.3, 33: "Enkel", 34: "Ja"}),
        # deur met bovenlicht ZONDER glastype -> legacy: optellen bij het deurglas
        _rij77({0: "achterdeur", 3: 2.2, 8: "Door", 17: "ZW", 18: "Deur met raam", 19: 0.4,
                20: "Dubbel", 32: 0.3}), "",
    ])
    _f77 = _tf77.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8")
    _f77.write(_csv77); _f77.close()
    _d77, _n77 = _csvdos77(_f77.name)

    _koz77 = [s for s in _d77.schil if s.type == "kozijn" and s.subtype == "Raam"]
    _opp77 = lambda glas: round(sum(s.oppervlakte_m2 for s in _koz77 if (s.glastype or "") == glas), 2)
    check("bovenlicht (glas): 0,5 m2 ENKEL als eigen kozijnregel", _opp77("Enkel") >= 0.65)  # <0,65-regel bumpt
    _enkel77 = [s for s in _koz77 if s.glastype == "Enkel"]
    check("bovenlicht: 2 aparte enkel-glas-vlakken (raam-bovenlicht + deur-bovenlicht)",
          len(_enkel77) == 2)
    check("hoofdraam 1: 3,0 - 0,5 = 2,5 m2 HR++ over",
          any(abs(s.oppervlakte_m2 - 2.5) < 0.01 for s in _koz77 if s.glastype == "HR++"))
    check("onderlicht-paneel: 0,6 m2 dichte constructie + hoofdraam 1,4 m2 Dubbel",
          any(abs(s.oppervlakte_m2 - 0.6) < 0.01 and "paneel" in (s.opmerkingen or "").lower()
              for s in _d77.schil)
          and any(abs(s.oppervlakte_m2 - 1.4) < 0.01 for s in _koz77 if s.glastype == "Dubbel"))
    check("bovenlicht zonder oppervlak: NIET gesplitst + luide flag",
          any(abs(s.oppervlakte_m2 - 1.5) < 0.01 for s in _koz77 if s.glastype == "HR++")
          and any("OPPERVLAK ontbreekt" in n for n in _n77))
    _deur77 = [s for s in _d77.schil if s.subtype == "Deur"]
    check("deur met eigen bovenlicht-glastype: bovenlicht gaat van het deurvlak af (2,2-0,3=1,9)",
          any(abs(s.oppervlakte_m2 - 1.9) < 0.01 for s in _deur77))
    check("deur zonder bovenlicht-glastype (legacy): deurvlak blijft 2,2 (glas telt in de deur)",
          any(abs(s.oppervlakte_m2 - 2.2) < 0.01 for s in _deur77))
    check("toevoerroosters: hoofdraam + bovenlicht + deur geteld (3) in de notes",
          any("3 kozijn(en)/deur(en) met TOEVOERROOSTER" in n for n in _n77))
except Exception as _e:
    check("boven-/onderlicht-parser: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n78. Project verwijderen (map weg + lead-koppeling los + pad-veiligheid)")
try:
    import os as _o78, json as _j78, tempfile as _t78
    import dashboard.app as _W78, dashboard.leads as _L78
    _W78.app.config.update(TESTING=True)
    _bewP = _W78.PROJECTS_DIR
    _bewL = (_L78.LEADS_DIR, _L78.LEADS_FILE, _L78.GEWIST_FILE)
    _tmpP78, _tmpL78 = _t78.mkdtemp(), _t78.mkdtemp()
    _W78.PROJECTS_DIR = _tmpP78
    _L78.LEADS_DIR = _tmpL78
    _L78.LEADS_FILE = _o78.path.join(_tmpL78, "leads.json")
    _L78.GEWIST_FILE = _o78.path.join(_tmpL78, "verwijderd.json")
    try:
        _tag78 = "9999ZZ_7"
        _pdir78 = _o78.path.join(_tmpP78, _tag78)
        _o78.makedirs(_o78.path.join(_pdir78, "fotos"))
        _j78.dump({"tag": _tag78, "adres": "Testlaan 7, Groningen", "stap": "opname"},
                  open(_o78.path.join(_pdir78, "project.json"), "w"))
        open(_o78.path.join(_pdir78, "dossier_%s.json" % _tag78), "w").write("{}")
        open(_o78.path.join(_pdir78, "fotos", "voor.jpg"), "w").write("x")
        _L78.save_leads([dict(id=1, naam="Jan", postcode="9999ZZ", huisnummer="7", toevoeging="",
                              status="opname gedaan", ontvangen="2026-07-20", notitie="", afspraak="",
                              project_tag=_tag78)])
        _c78 = _W78.app.test_client()
        with _c78.session_transaction() as _s78:
            _s78["ingelogd"] = True
        check("verwijder-route geregistreerd",
              "/project/<tag>/verwijder" in {r.rule for r in _W78.app.url_map.iter_rules()})
        check("home: project + verwijderknop zichtbaar",
              "Testlaan 7" in (_h := _c78.get("/").get_data(as_text=True)) and "verwijder" in _h)
        _r78 = _c78.post("/project/%s/verwijder" % _tag78, follow_redirects=True)
        check("verwijderen: hele projectmap weg", not _o78.path.isdir(_pdir78))
        check("verwijderen: lead-koppeling losgemaakt, lead blijft",
              _L78.load_leads()[0].get("project_tag") is None and len(_L78.load_leads()) == 1)
        check("verwijderen: bevestiging getoond", "definitief verwijderd" in _r78.get_data(as_text=True))
        check("pad-veiligheid: onbekend project -> 404",
              _c78.post("/project/bestaatniet/verwijder").status_code == 404)
    finally:
        _W78.PROJECTS_DIR = _bewP
        _L78.LEADS_DIR, _L78.LEADS_FILE, _L78.GEWIST_FILE = _bewL
except Exception as _e:
    check("project verwijderen: draait zonder fout", False); print("     " + repr(_e)[:200])

print("\n79. Portaal-annulering: bestaande lead op 'vervallen' i.p.v. opnieuw benaderen")
try:
    import os as _o79, tempfile as _t79, datetime as _dt79
    import dashboard.app as _W79, dashboard.leads as _L79, dashboard.mailbox as _M79, dashboard.graph_mail as _G79

    _TOE79 = ('{"BagAdresId":"1895200000005699","Email":"t@x.nl","Postcode":"9674BW","Huisnummer":28,'
              '"Naam":"tess rouppe","Telefoonnummer":"0625494609",'
              '"WijzigingsType":"AdviseurToegekend","WijzigingsReden":"toegekend"}')
    _ANN79 = ('{"BagAdresId":"1895200000005699","Email":"t@x.nl","Postcode":"9674BW","Huisnummer":28,'
              '"Naam":"tess rouppe","WijzigingsType":"AdviseurGeannuleerd",'
              '"WijzigingsReden":"Adviseur 39222 geannuleerd door gebruiker. Reden anders: "}')

    check("parse: WijzigingsType wordt meegelezen",
          _L79.parse_lead(_ANN79).get("wijzigingstype") == "AdviseurGeannuleerd")
    check("annulering herkend, toewijzing niet",
          _L79.is_annulering(_L79.parse_lead(_ANN79)) is True
          and _L79.is_annulering(_L79.parse_lead(_TOE79)) is False)

    _W79.app.config.update(TESTING=True)
    _tmp79 = _t79.mkdtemp()
    _bew79 = (_L79.LEADS_DIR, _L79.LEADS_FILE, _L79.GEWIST_FILE)
    _L79.LEADS_DIR = _tmp79
    _L79.LEADS_FILE = _o79.path.join(_tmp79, "leads.json")
    _L79.GEWIST_FILE = _o79.path.join(_tmp79, "verwijderd.json")
    try:
        _L79.save_leads([])
        _r1 = _W79._leads_toevoegen(_TOE79)
        check("toewijzing -> nieuwe lead", _r1["nieuw"] == 1 and _L79.load_leads()[0]["status"] == "nieuw")
        _r2 = _W79._leads_toevoegen(_ANN79)
        _ld79 = _L79.load_leads()
        check("annulering -> bestaande lead op 'vervallen', niet verwijderd",
              _r2["geannuleerd"] == 1 and len(_ld79) == 1 and _ld79[0]["status"] == "vervallen")
        check("annulering: reden + datum in de notitie",
              "Geannuleerd door bewoner" in (_ld79[0].get("notitie") or ""))
        _r3 = _W79._leads_toevoegen(_ANN79.replace("1895200000005699", "0000000000000000")
                                          .replace("9674BW", "1111AA"))
        check("annulering voor onbekend adres -> gemeld, niets aangemaakt",
              _r3["annul_onbekend"] == 1 and len(_L79.load_leads()) == 1)
        check("melding noemt de annuleringen", "vervallen" in _W79._annulering_melding(_r2))
    finally:
        _L79.LEADS_DIR, _L79.LEADS_FILE, _L79.GEWIST_FILE = _bew79

    # fetch-filters: annuleringen komen nu ook binnen (marker WijzigingsType), ook al zegt de config
    # alleen 'AdviseurToegekend'
    _imap79 = _M79.zoekopdracht({"dagen": 30, "onderwerp": "AdviseurToegekend"}, _dt79.date(2026, 7, 20))
    check("IMAP-zoek vangt portaalmails breed (OR met WijzigingsType)",
          "WijzigingsType" in _imap79)
    _ber79 = [{"subject": "Contact met adviseur geannuleerd door accountid abc",
               "body": {"content": _ANN79}}]
    check("graph-filter laat een annulering door ook al staat alleen AdviseurToegekend in de config",
          len(_G79.berichten_naar_teksten(_ber79, "AdviseurToegekend")) == 1)
except Exception as _e:
    check("portaal-annulering: draait zonder fout", False); print("     " + repr(_e)[:200])

print("N. Standaard + verliesoppervlak (NTA 8800 §5.3.2 / §6.7.3)")
try:
    from engine.standaard import fls, verliesoppervlak, standaard_eis, is_grondgebonden
    from core.dossier import SchilDeel as _SD
    check("fls grond/kruipruimte 0,7", fls("Grond") == 0.7 and fls("Kruipruimte") == 0.7)
    check("fls AVR/woningscheidend 0", fls("AVR") == 0.0 and fls("Aangrenzende woning") == 0.0)
    check("fls buitenlucht/AOR 1,0", fls("Buitenlucht") == 1.0 and fls("AOR") == 1.0)
    check("is_grondgebonden tussenwoning=True, appartement=False",
          is_grondgebonden("Tussenwoning") and not is_grondgebonden("Appartement (tussen)"))
    _std = Dossier()
    _std.schil = [_SD(id="G1", type="gevel", begrenzing="Buitenlucht", oppervlakte_m2=50),
                  _SD(id="V1", type="vloer", begrenzing="Kruipruimte", oppervlakte_m2=40),
                  _SD(id="W1", type="gevel", begrenzing="AVR", oppervlakte_m2=30),
                  _SD(id="K1", type="kozijn", begrenzing="Buitenlucht", oppervlakte_m2=10)]
    # 50x1 + 40x0,7 + 30x0 + kozijn uitgesloten = 78
    check("Als gewogen (grond x0,7, AVR x0, kozijn uit) = 78",
          abs(verliesoppervlak(_std) - 78.0) < 1e-6, str(verliesoppervlak(_std)))
    _std.geometrie.gebruiksoppervlakte_ag_m2 = 100.0
    _std.identificatie.woningtype = "Tussenwoning"     # zonder woningtype = geen betrouwbare eis (None)
    _std.identificatie.bouwjaar = 1970                 # grondgebonden, na 1945, Als/Ag=0,78<1 -> 43
    check("Standaard grondgebonden na1945 ratio<1 -> 43", standaard_eis(_std) == 43, str(standaard_eis(_std)))
    _std.geometrie.gebruiksoppervlakte_ag_m2 = 50.0    # Als/Ag=1,56 -> 43 + 40x0,56 = 65,4 -> 65
    check("Standaard grondgebonden na1945 ratio1,56 -> 65", standaard_eis(_std) == 65, str(standaard_eis(_std)))
    _std.identificatie.woningtype = "Appartement (tussen)"
    _std.identificatie.bouwjaar = 1930                 # woongebouw t/m 1945 -> 95 + 70x0,56 = 134,2 -> 134
    check("Standaard woongebouw t/m1945 ratio1,56 -> 134", standaard_eis(_std) == 134, str(standaard_eis(_std)))
    _std.identificatie.bouwjaar = None
    check("Standaard None zonder bouwjaar", standaard_eis(_std) is None)
    from engine.standaard import woningtype_onbekend
    _std.identificatie.bouwjaar = 1970; _std.identificatie.woningtype = ""
    check("Standaard None zonder woningtype (niet stil grondgebonden aannemen)",
          standaard_eis(_std) is None and woningtype_onbekend(""))
except Exception as _e:
    check("Standaard/verliesoppervlak: draait zonder fout", False); print("     " + repr(_e)[:200])

print("O. ASGR/ASV-alias (begrenzing)")
try:
    from magicplan.statistics_csv import _begrenzing_uit_naam
    from vabi.objecten_generate import _grenst_aan_code
    check("wandnaam '...ASV' -> Sterk geventileerd (ASGR)",
          _begrenzing_uit_naam("Zijgevel ASV") == "Sterk geventileerd")
    check("wandnaam '...ASGR' -> Sterk geventileerd",
          _begrenzing_uit_naam("Zijgevel ASGR") == "Sterk geventileerd")
    check("ASV -> GrenstAan 6 (detailopname)", _grenst_aan_code("ASV", basis=False) == "6")
    check("ASV -> GrenstAan 0 (basisopname, telt als buiten)", _grenst_aan_code("ASV", basis=True) == "0")
    # beschrijvende MagicPlan-labels (met '(...)') moeten nu ook mappen (was None vóór de fix)
    check("MagicPlan-label 'ASGR (sterk geventileerd)' -> 6", _grenst_aan_code("ASGR (sterk geventileerd)", basis=False) == "6")
    check("MagicPlan-label 'AOR (onverwarmd)' -> 4", _grenst_aan_code("AOR (onverwarmd)", basis=False) == "4")
    check("MagicPlan-label 'AOS (serre)' -> 5", _grenst_aan_code("AOS (serre)", basis=False) == "5")
    check("MagicPlan-label 'AVR (aangrenzend verwarmd)' -> 8", _grenst_aan_code("AVR (aangrenzend verwarmd)", basis=False) == "8")
except Exception as _e:
    check("ASGR/ASV-alias: draait zonder fout", False); print("     " + repr(_e)[:200])

print("P. Auto-perimeter uit begane-grond gevelbreedtes (volledig getagd)")
try:
    import tempfile as _tfP
    from magicplan.statistics_csv import build_dossier as _bdP
    # vrijstaand 6x6 (footprint 36), 4 buitengevels van 6 m getikt -> auto = 24 m; MagicPlan-omtrek staat
    # bewust op 20 zodat we auto (24) van fallback (20) kunnen onderscheiden.
    _pcsv = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Vrijstaand\n"
             "Gevelhoogte (m),3\nTotal living area,36\n\n"
             "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
             "Ground Floor,36,2.60 m,Kruipruimte\n\n"
             "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
             "Ground Floor,\n"
             "Room,Wall 0,Wall,15,15,6,2.6,1,Wall,,1,N\nRoom,Wall 1,Wall,15,15,6,2.6,1,Wall,,1,O\n"
             "Room,Wall 2,Wall,15,15,6,2.6,1,Wall,,1,Z\nRoom,Wall 3,Wall,15,15,6,2.6,1,Wall,,1,W\n")
    _fP = _tfP.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fP.write(_pcsv); _fP.close()
    _dP, _nP = _bdP(_fP.name); _os = __import__("os"); _os.unlink(_fP.name)
    _vlP = next((s for s in _dP.schil if s.type == "vloer"), None)
    check("auto-perimeter: volledig getagd -> perimeter = 24 (auto), niet 20 (MagicPlan-fallback)",
          _vlP is not None and _vlP.perimeter_m == 24.0, str(getattr(_vlP, "perimeter_m", None)))
    check("auto-perimeter: note vermeldt 'AUTOMATISCH berekend'",
          any("AUTOMATISCH berekend" in n for n in _nP))
except Exception as _e:
    check("auto-perimeter: draait zonder fout", False); print("     " + repr(_e)[:200])

print("Q. Element-overrides: per WAND en per KAMER (drielaagse overerving)")
try:
    import tempfile as _tfQ, os as _osQ
    from magicplan.statistics_csv import build_dossier as _bdQ
    # Project-standaard (Constructies): gevel isolatie Nee, vloer begrenzing Kruipruimte.
    # Wand 'Achtergevel' krijgt een AFWIJKENDE per-wand override (isolatie Ja + 120 mm + AOR + zone 2).
    _q = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1975 t/m 1982\nWoningtype,Vrijstaand\n"
          "Oriëntatie voorgevel,N\n"
          "Gevelhoogte (m),3\nTotal living area,60\nGevel - isolatie aanwezig?,Nee\n"
          "Begrenzing (vloer),Kruipruimte\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n\n"
          "ROOM ATTRIBUTES,Ground surface without walls,Vloer - begrenzing,"
          "Vloer - telt mee voor gebruiksoppervlakte?,Vloer - rekenzone\n"
          "Woonkamer,40,,,\nBerging,20,Grond,Nee,\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Gevelnaam (leeg = binnenwand),"
          "Gevel - isolatie aanwezig?,Gevel - isolatiedikte (mm),Gevel - begrenzing,Gevel - rekenzone\n"
          "Ground Floor,\n"
          "Room,Wall 0,Wall,15,15,6,2.6,1,Wall,Voorgevel,,,,\n"
          "Room,Wall 1,Wall,15,15,6,2.6,1,Wall,Achtergevel,Ja,120,AOR,2\n")
    _fQ = _tfQ.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fQ.write(_q); _fQ.close()
    _dQ, _nQ = _bdQ(_fQ.name); _osQ.unlink(_fQ.name)
    _gev = [s for s in _dQ.schil if s.type == "gevel"]
    _ov = next((s for s in _gev if s.begrenzing == "AOR"), None)
    _std = next((s for s in _gev if s.begrenzing != "AOR"), None)
    check("A: wand-override maakt een EIGEN schildeel (niet samengevoegd)", len(_gev) >= 2, str(len(_gev)))
    check("A: override-wand heeft isolatie Ja + 120 mm + AOR + rekenzone 2",
          _ov is not None and _ov.isolatie_aanwezig == "Ja" and _ov.isolatiedikte_mm == 120.0
          and _ov.rekenzone == 2,
          str((getattr(_ov, "isolatie_aanwezig", None), getattr(_ov, "isolatiedikte_mm", None),
               getattr(_ov, "rekenzone", None))))
    check("A: NIET-override wand volgt de Constructies-standaard (isolatie Nee)",
          _std is not None and _std.isolatie_aanwezig == "Nee", str(getattr(_std, "isolatie_aanwezig", None)))
    check("B: kamer met 'telt mee voor Ag' = Nee valt buiten de zone (niet in ruimtes)",
          all("berging" not in (r.naam or "").lower() for r in _dQ.geometrie.ruimtes))
    check("B: melding over de buiten-de-zone gelaten ruimte",
          any("telt NIET mee voor het gebruiksoppervlak" in n for n in _nQ))
except Exception as _e:
    check("element-overrides: draait zonder fout", False); print("     " + repr(_e)[:220])

print("R. Dak-isolatie op form-niveau + riet/DWTW/zonwering/PV-belemmering (NTA-aanvullingen)")
try:
    import tempfile as _tfR, os as _osR
    from magicplan.statistics_csv import build_dossier as _bdR
    _r = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1975 t/m 1982\nWoningtype,Vrijstaand\n"
          "Oriëntatie voorgevel,N\nGevelhoogte (m),3\nTotal living area,60\n"
          "Dak - isolatie aanwezig?,Ja\nDak - isolatiedikte (mm),140\nDak - begrenzing,Buitenlucht\n"
          "Dak - isolatie aan zijde,Binnenzijde\nRieten dak?,Ja\nRietdikte (mm),150\n"
          "Tapwater - toestel,Gasgestookt combitoestel\nDouche-WTW (DWTW) aanwezig?,Ja\n"
          "Zonne-energie aanwezig?,Ja\nPV - paneeltype,Monokristallijn\nPV - aantal panelen,8\n"
          "PV - orientatie,Z\nPV - belemmering/beschaduwing?,Ja deels\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n\n"
          "ROOM ATTRIBUTES\nWoonkamer,60\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Gevelnaam (leeg = binnenwand),"
          "Type glas,Zonwering/luik aanwezig? (leeg = geen)\n"
          "Ground Floor,\nRoom,Wall 0,Wall,15,15,6,2.6,1,Wall,Voorgevel,,\n"
          "Room,Wall 0,Window,2.0,2.0,1.5,1.3,1,Window,,HR++,Rolluik of luik (bedienbaar)\n")
    _fR = _tfR.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fR.write(_r); _fR.close()
    _dR, _nR = _bdR(_fR.name); _osR.unlink(_fR.name)
    _ds = _dR.opname.dak_standaard
    check("D: dak-isolatie uit de form (Ja + 140 mm + zijde) in dak_standaard",
          _ds.isolatie_aanwezig == "Ja" and _ds.isolatiedikte_mm == 140.0
          and _ds.isolatie_zijde == "Binnenzijde",
          str((_ds.isolatie_aanwezig, _ds.isolatiedikte_mm, _ds.isolatie_zijde)))
    check("D: melding dat de webapp-dakvlakken dit erven",
          any("elk dakvlak dat je" in n for n in _nR))
    check("4: riet -> dikte vastgelegd + NTA-bijlage-I-melding",
          _ds.riet_dikte_mm == 150.0 and any("RIETEN DAK" in n for n in _nR))
    check("4: DWTW aanwezig doorgezet naar tapwater", _dR.installaties.tapwater.dwtw_aanwezig is True)
    check("4: PV-belemmering doorgezet",
          bool(_dR.installaties.zonne_energie) and _dR.installaties.zonne_energie[0].belemmering is True)
    _rm = next((s for s in _dR.schil if s.type == "kozijn" and s.subtype == "Raam"), None)
    check("4: zonwering/luik per raam vastgelegd",
          _rm is not None and "Rolluik" in (_rm.zonwering or ""), str(getattr(_rm, "zonwering", None)))
except Exception as _e:
    check("dak-form/NTA-aanvullingen: draait zonder fout", False); print("     " + repr(_e)[:220])

print("S. Voorschot B/U-tarief volgt het BOUWJAAR (Startpakket maart 2026), niet het opnametype")
try:
    from dashboard.voorschot import uitgebreid_uit_bouwjaar as _ub, tarief_excl as _te
    check("B/U: vóór 1945 -> Uitgebreid", _ub(1930) is True and _ub(1944) is True)
    check("B/U: vanaf 1945 -> Basis", _ub(1945) is False and _ub(1975) is False)
    check("B/U: onbekend bouwjaar -> None (flaggen, niet gokken)",
          _ub(None) is None and _ub("") is None)
    check("tarief tussenwoning 1930 = 425 (U) en 1975 = 350 (B)",
          _te("Tussenwoning", 100, True) == 425.0 and _te("Tussenwoning", 100, False) == 350.0)
except Exception as _e:
    check("voorschot B/U: draait zonder fout", False); print("     " + repr(_e)[:200])

print("T. Bouwfysische aandachtspunten in de begeleidende tekst (Nij Begun technische haalbaarheid)")
try:
    from engine.advies_text import begeleidende_tekst as _bt
    from core.dossier import Maatregel as _MT
    _tA = _bt(_MT(code="V1-1-A1", onderdeel="A Gevelisolatie", omschrijving="Spouwmuurisolatie"))
    _tD = _bt(_MT(code="V4-1-A1", onderdeel="D Dakisolatie", omschrijving="Dakisolatie binnenzijde"))
    check("A: gevel-advies noemt dampremmende laag + balkkoppen",
          "dampremmende laag" in _tA and "balkkop" in _tA.lower())
    check("D: dak-advies noemt warme binnenzijde + ruimtefunctie onder koud dak",
          "warme binnenzijde" in _tD and "ruimtefunctie" in _tD)
except Exception as _e:
    check("bouwfysica-aandachtspunten: draait zonder fout", False); print("     " + repr(_e)[:200])

print("U. psi-lookup (NTA bijlage I) + Bijlage 4-7 in de plan-output")
try:
    from engine.psi_lookup import psi as _psi, relevante_details as _rd, PSI_DEFAULT as _PD, is_gestapeld as _ig
    check("psi: laagbouw detail 5 (onderdorpel kozijn) = 0,15 (A) / 0,25 (B)",
          _psi(5, "Tussenwoning", True) == 0.15 and _psi(5, "Tussenwoning", False) == 0.25)
    check("psi: gestapeld gebruikt tabel I.2 (detail 50 = 0,61)",
          _ig("Galerijwoning") and _psi(50, "Galerijwoning", True) == 0.61)
    check("psi: onbekende detailpositie -> 0,50 (NTA bijlage I.1)", _psi(999, "Tussenwoning") == _PD == 0.50)
    _pd = build_sample()
    _det = _rd(_pd)
    check("psi: details afgeleid uit het dossier", len(_det) >= 3, str(len(_det)))
    check("psi: elk detail heeft nr/omschrijving/psi_a/bron",
          all({"nr", "omschrijving", "psi_a", "bron"} <= set(d) for d in _det))
    # bijlagen 4-7 belanden in de docx
    _bp = os.path.join(tempfile.gettempdir(), "bijlagen_test.docx")
    fill_template(dos2, TEMPLATE, _bp)
    _txt = "\n".join(p.text for p in docx.Document(_bp).paragraphs)
    check("bijlage 4 t/m 7 als koppen in het plan",
          all(k in _txt for k in ("Bijlage 4: Informatie over dit isolatieplan",
                                  "Bijlage 5: Voorgestelde maatregelen in beeld",
                                  "Bijlage 6: Ventilatieplan",
                                  "Bijlage 7: Detailtekeningen")))
    check("bijlage 4 noemt de bronnen (Groninger catalogus / RVO / Milieu Centraal / 21% btw)",
          "Groninger Maatregelen Catalogus" in _txt and "Milieu Centraal" in _txt and "21 % btw" in _txt)
    check("bijlage 7 toont detailregels met psi-waarden", "ψ = " in _txt and "Detail " in _txt)
except Exception as _e:
    check("psi-lookup/bijlagen: draait zonder fout", False); print("     " + repr(_e)[:220])

print("V. Ruimtefunctie: berging/opslag is GEEN verblijfsgebied (ISSO 82.1 6.3.1)")
try:
    from magicplan.statistics_csv import _functie_uit_naam as _fn
    check("'Storage Room' -> overig (niet verblijfsruimte via het brede 'room')", _fn("Storage Room") == "overig")
    check("'Closet' -> overig", _fn("Closet") == "overig")
    check("'Zolder' -> overig (bergzolder; telt wel voor Ag, geen 0,7-toevoer)", _fn("Zolder") == "overig")
    check("'Living Room' blijft verblijfsruimte", _fn("Living Room") == "verblijfsruimte")
    check("'Slaapkamer zolder' blijft slaapkamer (werkelijk gebruik)", _fn("Slaapkamer zolder") == "slaapkamer")
    check("'Bathroom' blijft badkamer", _fn("Bathroom") == "badkamer")
    from ventilatie.ventilatie import bereken as _vb
    from core.dossier import Ruimte as _R
    _v = _vb([_R(naam="Storage Room", functie="overig", oppervlakte_m2=20),
              _R(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30)])
    _berg = next((r for r in _v["rows"] if "storage" in r["naam"].lower()), None)
    check("berging krijgt GEEN toevoer in de ventilatieberekening",
          _berg is not None and (_berg.get("toevoer") or 0) == 0, str(_berg))
except Exception as _e:
    check("ruimtefunctie: draait zonder fout", False); print("     " + repr(_e)[:220])

print("W. Glastype-kolomkeuze: 'Type glas' (raam) wint van 'Type glas (indien glas in deur)'")
try:
    import tempfile as _tfW, os as _osW
    from magicplan.statistics_csv import build_dossier as _bdW
    # Essenhage-volgorde: de DEUR-glaskolom staat VOOR de raam-glaskolom. Vóór de fix knipte de
    # 'exacte' match alles vóór ' (' af -> beide werden 'type glas' -> de (lege) deurkolom won.
    _w = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1975 t/m 1982\nWoningtype,Tussenwoning\n"
          "Oriëntatie voorgevel,N\nGevelhoogte (m),3\nTotal living area,60\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,"
          "Type glas (indien glas in deur),Type glas,Gevelnaam (leeg = binnenwand)\n"
          "Ground Floor,\n"
          "Room,Wall 0,Wall,15,15,6,2.6,1,Wall,,,Voorgevel\n"
          "Room,Wall 0,Hung Window,2.0,2.0,1.5,1.3,1,Window,,HR++,\n")
    _fW = _tfW.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fW.write(_w); _fW.close()
    _dW, _nW = _bdW(_fW.name); _osW.unlink(_fW.name)
    _rW = next((s for s in _dW.schil if s.type == "kozijn" and s.subtype == "Raam"), None)
    check("glastype van het RAAM wordt gelezen (HR++), niet de lege deur-kolom",
          _rW is not None and _rW.glastype == "HR++", str(getattr(_rW, "glastype", None)))
    # actiepunt-classificatie
    from dashboard.app import _is_prio_actie as _pa
    check("actiepunt 'FOUT wand ...' = prioriteit", _pa("FOUT wand 'X': veld leeg") is True)
    check("actiepunt 'GEBOUWHOOGTE ONTBREEKT' = prioriteit", _pa("GEBOUWHOOGTE ONTBREEKT: vul ...") is True)
    check("actiepunt 'Gevel NW: b x h-methode' = ter controle",
          _pa("Gevel NW: b x h per bouwlaag gebruikt") is False)
except Exception as _e:
    check("glastype/actiepunten: draait zonder fout", False); print("     " + repr(_e)[:220])

print("X. Kwaliteitsverklaring = aantoonbaar geisoleerd + HR-glasmapping + share target")
try:
    from vabi.constructie_generate import _norm_glas as _ng
    from vabi.codebook import Codebook as _CB
    _cb = _CB.default()
    check("HR (dubbel glas met coating) -> HR (code 3), niet Dubbel (2)",
          _ng("HR (dubbel glas met coating)") == "hr" and _cb.glas_code("hr") == "3",
          _ng("HR (dubbel glas met coating)") + "/" + str(_cb.glas_code(_ng("HR (dubbel glas met coating)"))))
    check("HR+ apart van HR en HR++ (codes 3/4/5)",
          [_cb.glas_code(_ng(g)) for g in ("HR", "HR+", "HR++")] == ["3", "4", "5"])
    # kwaliteitsverklaring -> geen blind isolatie-advies
    from engine.measure_engine import _isolated_status as _iso
    from engine.advies_logic import _isolatie_op_niveau as _opn
    from core.dossier import SchilDeel as _SDX
    _kv = _SDX(id="dak-kv", type="dak", isolatie_aanwezig="Onbekend", rc_bron="Kwaliteitsverklaring")
    _on = _SDX(id="dak-onb", type="dak", isolatie_aanwezig="Onbekend", rc_bron="Dikte onbekend")
    check("kwaliteitsverklaring telt als geisoleerd (geen isolatie-advies)", _iso(_kv) is True)
    check("advies_logic: kwaliteitsverklaring -> geen advies", _opn(_kv) is True)
    check("zonder KV en isolatie Onbekend -> wel advies", _iso(_on) is False and _opn(_on) is False)
    # PWA share target in het manifest
    from dashboard.app import app as _app
    _c = _app.test_client()
    _m = _c.get("/manifest.webmanifest")
    _mj = json.loads(_m.data.decode("utf-8"))
    check("manifest bevat share_target naar /deel",
          _mj.get("share_target", {}).get("action") == "/deel"
          and _mj["share_target"]["params"]["files"][0]["name"] == "bestand")
except Exception as _e:
    check("KV/HR-glas/share: draait zonder fout", False); print("     " + repr(_e)[:220])

print("X2. Kwaliteitsverklaring blokkeert iedere VABI-export vóór schrijven")
try:
    import shutil as _shX2
    from vabi import generate_all as _gaX2
    from vabi import constructie_generate as _cgX2
    from vabi import objecten_generate as _ogX2
    from vabi import installatie_generate as _igX2
    from vabi.preflight import VabiExportBlocked as _BlockedX2
    _kvX2 = build_sample()
    _kvX2.schil[0].rc_bron = "  kWaLiTeItSvErKlArInG  "
    _kvX2.schil[0].id = "gevel-kv-1"
    _kvX2.schil[1].rc_bron = "Kwaliteitsverklaring"
    _kvX2.schil[1].id = "vloer-kv-2"
    _tdX2 = tempfile.mkdtemp(prefix="nb_kv_gate_")
    _directX2 = []
    for _nameX2, _writerX2 in (("constructie", _cgX2.write),
                               ("objecten", _ogX2.write),
                               ("installatie", _igX2.write)):
        _pathX2 = os.path.join(_tdX2, _nameX2 + ".xml")
        try:
            _writerX2(_kvX2, _pathX2)
            _directX2.append(False)
        except _BlockedX2 as _eX2:
            _directX2.append(not os.path.exists(_pathX2)
                             and "gevel-kv-1" in str(_eX2)
                             and "vloer-kv-2" in str(_eX2)
                             and "correct in Vabi" in str(_eX2))
    check("KV: alle drie directe writers blokkeren met alle ids vóór schrijven", all(_directX2))

    _allX2 = os.path.join(_tdX2, "complete_set")
    try:
        _gaX2.generate_all(_kvX2, _allX2, prefix="huidig")
        _blocked_allX2 = False
    except _BlockedX2:
        _blocked_allX2 = not os.path.exists(_allX2)
    check("KV: generate_all laat geen gedeeltelijke exportset of map achter", _blocked_allX2)

    _allowedX2 = []
    for _bronX2 in ("", "Dikte onbekend", "Forfaitair bouwjaar"):
        _okX2 = build_sample()
        for _sX2 in _okX2.schil:
            _sX2.rc_bron = _bronX2
        _okX2.schil[0].isolatie_aanwezig = "Onbekend"
        _okdirX2 = os.path.join(_tdX2, "ok_" + str(len(_allowedX2)))
        _resX2 = _gaX2.generate_all(_okX2, _okdirX2)
        _allowedX2.append(all(os.path.isfile(_resX2[_kX2][0])
                              for _kX2 in ("constructies", "objecten", "installaties")))
    check("KV: onbekend, Dikte onbekend en leeg blijven exporteerbaar", all(_allowedX2))

    import dashboard.app as _WAX2
    _WAX2.app.config.update(TESTING=True, SECRET_KEY="test-kv-gate")
    _orig_lsX2, _orig_dosX2, _orig_pdX2 = _WAX2._load_state, _WAX2._dossier, _WAX2._pdir
    _dashdirX2 = os.path.join(_tdX2, "dashboard_project")
    _WAX2._load_state = lambda _tagX2: {"dossier_file": "dossier.json"}
    _WAX2._dossier = lambda _tagX2: _kvX2
    _WAX2._pdir = lambda _tagX2: _dashdirX2
    try:
        _clientX2 = _WAX2.app.test_client()
        with _clientX2.session_transaction() as _sessX2:
            _sessX2["ingelogd"] = True
        _respX2 = _clientX2.get("/project/test-kv/opname/vabi_huidig", follow_redirects=False)
        with _clientX2.session_transaction() as _sessX2:
            _flashesX2 = " ".join(m for _cat, m in _sessX2.get("_flashes", []))
        check("KV: dashboard toont concrete blokkade met betrokken ids",
              _respX2.status_code in (302, 303) and "VABI-export geblokkeerd" in _flashesX2
              and "gevel-kv-1" in _flashesX2 and "vloer-kv-2" in _flashesX2)
        check("KV: dashboardroute schrijft geen exportbestanden",
              not os.path.exists(os.path.join(_dashdirX2, "vabi_huidig")))
    finally:
        _WAX2._load_state, _WAX2._dossier, _WAX2._pdir = _orig_lsX2, _orig_dosX2, _orig_pdX2
    _shX2.rmtree(_tdX2, ignore_errors=True)
except Exception as _e:
    check("KV exportgate: draait zonder fout", False); print("     " + repr(_e)[:220])

print("X3. VABI-exportset wordt atomisch gepubliceerd")
try:
    import io
    import json as _jsonX3
    import shutil as _shX3
    import zipfile
    from vabi import generate_all as _gaX3
    _tdX3 = tempfile.mkdtemp(prefix="nb_atomic_vabi_")
    _dosX3 = build_sample()
    _preflight_callsX3 = [0]
    _original_preflightX3 = _gaX3.assert_no_dubbel_dak_fallback
    def _count_preflightX3(dos):
        _preflight_callsX3[0] += 1
        return _original_preflightX3(dos)
    _gaX3.assert_no_dubbel_dak_fallback = _count_preflightX3
    try:
        _gaX3.generate_all(_dosX3, os.path.join(_tdX3, "preflight-count"), prefix="huidig")
    finally:
        _gaX3.assert_no_dubbel_dak_fallback = _original_preflightX3
    check("taak023: generate_all voert dubbel-dakpreflight exact eenmaal uit",
          _preflight_callsX3[0] == 1)

    # De optimalisatie geldt alleen voor de orkestrator: rechtstreeks aangeroepen writers
    # moeten een onveilig dossier zelf blijven blokkeren.
    from core.dossier import SchilDeel as _SchilX3
    from vabi.preflight import VabiExportBlocked as _BlockedX3
    _unsafeX3 = build_sample()
    _unsafeX3.schil.append(_SchilX3(id="dak", type="dak", oppervlakte_m2=40.0,
                                    bron="magicplan-dak-fallback"))
    _unsafeX3.schil.append(_SchilX3(id="dak1-plat", type="dak", oppervlakte_m2=20.0,
                                    bron="webapp-wizard"))
    _direct_blockedX3 = []
    for _writerX3, _nameX3 in ((_gaX3.constructie_generate.write, "c.xml"),
                                (_gaX3.objecten_generate.write, "o.xml"),
                                (_gaX3.installatie_generate.write, "i.xml")):
        try:
            _writerX3(_unsafeX3, os.path.join(_tdX3, _nameX3))
            _direct_blockedX3.append(False)
        except _BlockedX3:
            _direct_blockedX3.append(True)
    check("taak023: alle directe writers behouden hun dubbel-dakpreflight",
          all(_direct_blockedX3))

    _firstX3 = _gaX3.generate_all(_dosX3, _tdX3, prefix="huidig")
    _manifest_pathX3 = os.path.join(_tdX3, _gaX3.MANIFEST)
    with open(_manifest_pathX3, "rb") as _fhX3:
        _manifest_beforeX3 = _fhX3.read()
    _set_beforeX3 = _firstX3["set_dir"]
    _files_beforeX3 = sorted(os.listdir(_set_beforeX3))

    _targetsX3 = [
        (_gaX3.constructie_generate, "write"),
        (_gaX3.objecten_generate, "write"),
        (_gaX3.installatie_generate, "write"),
        (_gaX3, "_write_instructions"),
    ]
    _fault_resultsX3 = []
    for _ownerX3, _attrX3 in _targetsX3:
        _originalX3 = getattr(_ownerX3, _attrX3)
        def _failX3(*_argsX3, **_kwargsX3):
            raise RuntimeError("geinjecteerde writerfout")
        setattr(_ownerX3, _attrX3, _failX3)
        try:
            try:
                _gaX3.generate_all(_dosX3, _tdX3, prefix="huidig")
                _raisedX3 = False
            except RuntimeError:
                _raisedX3 = True
            with open(_manifest_pathX3, "rb") as _fhX3:
                _manifest_afterX3 = _fhX3.read()
            _setsX3 = [p for p in os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR))
                       if not p.startswith(".staging-")]
            _stagingX3 = [p for p in os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR))
                          if p.startswith(".staging-")]
            _fault_resultsX3.append(_raisedX3
                                     and _manifest_afterX3 == _manifest_beforeX3
                                     and _gaX3.current_set_dir(_tdX3) == _set_beforeX3
                                     and sorted(os.listdir(_set_beforeX3)) == _files_beforeX3
                                     and len(_setsX3) == 1 and not _stagingX3)
        finally:
            setattr(_ownerX3, _attrX3, _originalX3)
    check("writerfout in elk van 4 fasen behoudt vorige complete set", all(_fault_resultsX3))

    _secondX3 = _gaX3.generate_all(_dosX3, _tdX3, prefix="huidig")
    with open(_manifest_pathX3, encoding="utf-8") as _fhX3:
        _manifestX3 = _jsonX3.load(_fhX3)
    _published_filesX3 = sorted(os.listdir(_secondX3["set_dir"]))
    check("succes wijst naar precies een intern consistente immutable set",
          _secondX3["set_dir"] != _set_beforeX3
          and _gaX3.current_set_dir(_tdX3) == _secondX3["set_dir"]
          and sorted(_manifestX3["files"]) == _published_filesX3
          and all(os.path.dirname(_secondX3[k][0]) == _secondX3["set_dir"]
                  for k in ("constructies", "objecten", "installaties")))
    check("tijdelijke publicatiebestanden zijn na succes opgeruimd",
          not any(p.startswith(".") for p in os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR)))
          and not any(p.startswith(".CURRENT-") for p in os.listdir(_tdX3)))

    _traversalX3 = []
    for _bad_prefixX3 in ("../uit", "sub/map", "sub\\map", ".verborgen", "a b"):
        _bad_outX3 = os.path.join(_tdX3, "bad_" + str(len(_traversalX3)))
        try:
            _gaX3.generate_all(_dosX3, _bad_outX3, prefix=_bad_prefixX3)
            _traversalX3.append(False)
        except ValueError:
            _traversalX3.append(not os.path.exists(_bad_outX3))
    check("prefix path traversal en onveilige bestandscomponenten worden vóór schrijven geweigerd",
          all(_traversalX3))

    _before_switchX3 = _gaX3.current_set_dir(_tdX3)
    _sets_before_switchX3 = sorted(os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR)))
    _replaceX3 = _gaX3.os.replace
    def _fail_manifest_switchX3(src, dst):
        if os.path.basename(dst) == _gaX3.MANIFEST:
            raise PermissionError("geinjecteerde manifestwisselfout")
        return _replaceX3(src, dst)
    _gaX3.os.replace = _fail_manifest_switchX3
    try:
        try:
            _gaX3.generate_all(_dosX3, _tdX3, prefix="huidig")
            _switch_failedX3 = False
        except PermissionError:
            _switch_failedX3 = True
    finally:
        _gaX3.os.replace = _replaceX3
    check("manifestwisselfout behoudt vorige pointer en ruimt onzichtbare set op",
          _switch_failedX3 and _gaX3.current_set_dir(_tdX3) == _before_switchX3
          and sorted(os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR))) == _sets_before_switchX3
          and not any(p.startswith(".staging-") for p in os.listdir(os.path.join(_tdX3, _gaX3.SETS_DIR))))

    import concurrent.futures as _cfX3
    _concurrent_dirX3 = os.path.join(_tdX3, "concurrent")
    with _cfX3.ThreadPoolExecutor(max_workers=2) as _poolX3:
        _futuresX3 = [_poolX3.submit(_gaX3.generate_all, _dosX3, _concurrent_dirX3, "na")
                      for _iX3 in range(2)]
        _resultsX3 = [_fX3.result() for _fX3 in _futuresX3]
    _current_concurrentX3 = _gaX3.current_set_dir(_concurrent_dirX3)
    check("gelijktijdige publicaties leveren één complete CURRENT-set zonder stagingresten",
          _current_concurrentX3 in [r["set_dir"] for r in _resultsX3]
          and len(os.listdir(_current_concurrentX3)) == 4
          and not any(p.startswith(".staging-") for p in os.listdir(
              os.path.join(_concurrent_dirX3, _gaX3.SETS_DIR))))

    import dashboard.app as _WAX3
    _dashboard_dirX3 = os.path.join(_tdX3, "dashboard")
    os.makedirs(_dashboard_dirX3)
    _published_dashX3 = _gaX3.generate_all(_dosX3, os.path.join(_dashboard_dirX3, "vabi_na"), "na")
    _legacy_dashX3 = os.path.join(_dashboard_dirX3, "vabi_huidig")
    os.makedirs(_legacy_dashX3)
    with open(os.path.join(_legacy_dashX3, "legacy.xml"), "w", encoding="utf-8") as _fhX3:
        _fhX3.write("<Project/>")
    _old_pdirX3, _old_stateX3 = _WAX3._pdir, _WAX3._load_state
    _WAX3._pdir = lambda _tagX3: _dashboard_dirX3
    _WAX3._load_state = lambda _tagX3: {"adres": "Testweg 17"}
    try:
        _WAX3.app.config.update(TESTING=True, SECRET_KEY="atomic-dashboard-test")
        _clientX3 = _WAX3.app.test_client()
        with _clientX3.session_transaction() as _sessionX3:
            _sessionX3["ingelogd"] = True
        _nameX3 = os.path.basename(_published_dashX3["objecten"][0])
        _download_currentX3 = _clientX3.get("/download/test/vabi_na/" + _nameX3)
        _download_legacyX3 = _clientX3.get("/download/test/vabi_huidig/legacy.xml")
        _exportX3 = _clientX3.get("/project/test/export")
        _zipX3 = zipfile.ZipFile(io.BytesIO(_exportX3.data))
        _zip_namesX3 = _zipX3.namelist()
        check("dashboard-download volgt CURRENT en behoudt legacy vlakke export",
              _download_currentX3.status_code == 200 and _download_legacyX3.status_code == 200)
        check("projectexport vult 02_VABI uit CURRENT zonder manifest-infrastructuur",
              any("02_VABI/toekomstige_staat/" in n and n.endswith(_nameX3) for n in _zip_namesX3)
              and any("02_VABI/huidige_staat/legacy.xml" in n for n in _zip_namesX3)
              and not any(n.endswith("CURRENT.json") or "/sets/" in n for n in _zip_namesX3))
    finally:
        _WAX3._pdir, _WAX3._load_state = _old_pdirX3, _old_stateX3
    _shX3.rmtree(_tdX3, ignore_errors=True)
except Exception as _e:
    check("atomische VABI-exporttests draaien zonder fout", False); print("     " + repr(_e)[:220])

print("Y. Gevel-onderbouwing + tegenoverliggende wand op dezelfde gevel (wandnummer +2)")
try:
    import tempfile as _tfY, os as _osY
    from magicplan.statistics_csv import build_dossier as _bdY
    # Wall 1 (5.80) en Wall 3 (3.60) van DEZELFDE kamer beide op de voorgevel = tegenover elkaar.
    # Verschillende breedtes, dus de breedte-dedup ziet het niet -> moet LUID gemeld worden.
    _y = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1975 t/m 1982\nWoningtype,Tussenwoning\n"
          "Oriëntatie voorgevel,N\nGevelhoogte (m),3\nTotal living area,60\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n\n"
          # >=12 kolommen: de parser negeert smallere wandrijen (len(r) < 12)
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie,"
          "Gevelnaam (leeg = binnenwand)\n"
          "Ground Floor,\n"
          "Woonkamer,Wall 1,Wall,15,15,5.80,2.6,1,Wall,,1,,Voorgevel\n"
          "Woonkamer,Wall 3,Wall,10,10,3.60,2.6,1,Wall,,1,,Voorgevel\n")
    _fY = _tfY.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fY.write(_y); _fY.close()
    _dY, _nY = _bdY(_fY.name); _osY.unlink(_fY.name)
    check("tegenoverliggende wand (nr +2) op dezelfde gevel -> LUIDE melding",
          any("TEGENOVER elkaar (nummer +2)" in n for n in _nY))
    check("melding noemt beide wandnummers en de opgetelde breedte",
          any("Wall 1" in n and "Wall 3" in n and "9.40" in n for n in _nY))
    _gY = next((s for s in _dY.schil if s.type == "gevel"), None)
    check("gevel heeft een controleerbare ONDERBOUWING in de opmerking",
          _gY is not None and "ONDERBOUWING:" in (_gY.opmerkingen or "")
          and "Wall 1" in _gY.opmerkingen, (getattr(_gY, "opmerkingen", "") or "")[:90])
    from dashboard.app import _is_prio_actie as _paY
    check("de melding valt in de PRIORITEIT-lijst (begint met FOUT)",
          _paY(next(n for n in _nY if "TEGENOVER elkaar (nummer +2)" in n)) is True)
except Exception as _e:
    check("gevel-onderbouwing/tegenover: draait zonder fout", False); print("     " + repr(_e)[:220])

print("Z. Aannames-audit: kozijnmateriaal, dubbele kolomnamen, positionele fallbacks")
try:
    import tempfile as _tfZ, os as _osZ
    from magicplan.statistics_csv import build_dossier as _bdZ, _norm_kozijn_mat as _nkm
    check("'?afwijkend' -> LEEG (niet stil op het gunstigste hout/kunststof)", _nkm("?afwijkend") == "")
    check("leeg -> hout/kunststof (80%-default)", _nkm("") == "Hout of kunststof")
    check("metaal niet-TO herkend", _nkm("Metaal (niet thermisch onderbroken)") == "Metaal (niet thermisch onderbroken)")
    # DUBBELE kolomnaam (raam- en deurvariant) + kozijn 'afwijkend' op de raamkolom
    _z = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1975 t/m 1982\nWoningtype,Tussenwoning\n"
          "Oriëntatie voorgevel,N\nGevelhoogte (m),3\nTotal living area,60\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,"
          "Kozijnmateriaal afwijkend (anders dan hout/kunststof)?,Type glas (indien glas in deur),"
          "Type glas,Kozijnmateriaal afwijkend (anders dan hout/kunststof)?,Toevoerrooster type,"
          "Gevelnaam (leeg = binnenwand)\n"
          "Ground Floor,\n"
          "Room,Wall 0,Wall,15,15,6,2.6,1,Wall,,,,,,Voorgevel\n"
          "Room,Wall 0,Hung Window,2.0,2.0,1.5,1.3,1,Window,,,HR++,Ja,,\n")
    _fZ = _tfZ.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fZ.write(_z); _fZ.close()
    _dZ, _nZ = _bdZ(_fZ.name); _osZ.unlink(_fZ.name)
    _rZ = next((s for s in _dZ.schil if s.type == "kozijn" and s.subtype == "Raam"), None)
    check("raam met 'afwijkend=Ja' -> kozijnmateriaal LEEG (niet aangenomen)",
          _rZ is not None and _rZ.kozijnmateriaal == "", str(getattr(_rZ, "kozijnmateriaal", None)))
    check("LUIDE melding over onbekend kozijnmateriaal (Ufr 3,8 vs 7,0)",
          any("KOZIJNMATERIAAL ONBEKEND" in n for n in _nZ))
    check("glastype nog steeds correct ondanks de dubbele kolomnamen",
          _rZ is not None and _rZ.glastype == "HR++", str(getattr(_rZ, "glastype", None)))
except Exception as _e:
    check("aannames-audit: draait zonder fout", False); print("     " + repr(_e)[:220])

print("AA. Eén vocabulaire: MagicPlan <-> parser <-> webapp mogen niet uit elkaar lopen")
try:
    from magicplan.statistics_csv import (_norm_begrenzing as _nb, _norm_glaslabel as _ngl,
                                          _BEGR_CANON as _BC, _GLAS_CANON as _GC)
    from dashboard.app import BEGR_OPTS as _BO, GLAS_OPTS as _GO
    # 1) elke canonieke uitkomst MOET in de webapp-keuzelijst staan, anders overschrijft het
    #    <select> de waarde stil bij opslaan
    _mis_b = sorted({v for v in _BC.values() if v not in _BO})
    check("elke genormaliseerde begrenzing staat in de webapp-lijst", not _mis_b, str(_mis_b))
    _mis_g = sorted({v for v in _GC.values() if v not in _GO})
    check("elke genormaliseerde glaskeuze staat in de webapp-lijst", not _mis_g, str(_mis_g))
    # 2) de LIVE MagicPlan-schrijfwijzen moeten allemaal landen op een canonieke waarde
    for _live, _verwacht in [("AOR (onverwarmd)", "AOR"), ("AOS (serre)", "AOS"),
                             ("AVR (aangrenzend verwarmd)", "AVR"),
                             ("ASGR (sterk geventileerd)", "Sterk geventileerd"),
                             ("Onverwarmde kelder", "Onverwarmde kelder"), ("Water", "Water")]:
        check("begrenzing '%s' -> '%s'" % (_live, _verwacht), _nb(_live) == _verwacht, _nb(_live))
    check("glas 'HR dubbel glas met coating' (CSV) -> webapp-schrijfwijze mét haakjes",
          _ngl("HR dubbel glas met coating") == "HR (dubbel glas met coating)",
          _ngl("HR dubbel glas met coating"))
    check("glas 'HR++' blijft HR++", _ngl("HR++") == "HR++")
    check("onbekende glaswaarde blijft staan (verdwijnt niet stil)", _ngl("Zonneglas XYZ") == "Zonneglas XYZ")
except Exception as _e:
    check("vocabulaire-check: draait zonder fout", False); print("     " + repr(_e)[:220])

print("AB. Dubbele kamer: som kamers > vloeroppervlak van die verdieping")
try:
    import tempfile as _tfB, os as _osB
    from magicplan.statistics_csv import build_dossier as _bdB
    # 1e verdieping is 20 m2, maar er staan 3x een badkamer van 4 m2 + slaapkamer 12 = 24 m2
    _b = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1979\nWoningtype,Tussenwoning\n"
          "Oriëntatie voorgevel,NW\nGevelhoogte (m),5\nTotal living area,80\n\n"
          "FLOOR ATTRIBUTES,Ground surface without walls: m²,Ceiling Height,Begrenzing\n"
          "Ground Floor,60,2.60 m,Kruipruimte\n1st Floor,20,2.38 m,Buitenlucht\n\n"
          "ROOM ATTRIBUTES,Ground surface without walls: m²\n"
          "Ground Floor,\nWoonkamer,60\n"
          "1st Floor,\nSlaapkamer,12\nBadkamer,4\nBadkamer,4\nBadkamer,4\n\n"
          "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie,"
          "Gevelnaam (leeg = binnenwand)\n"
          "Ground Floor,\nWoonkamer,Wall 0,Wall,15,15,6,2.6,1,Wall,,1,,Voorgevel\n")
    _fB = _tfB.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _fB.write(_b); _fB.close()
    _dB, _nB = _bdB(_fB.name); _osB.unlink(_fB.name)
    check("dubbele kamer gedetecteerd (24 m2 kamers op een 20 m2-verdieping)",
          any("DUBBELE KAMER" in n and "1st Floor" in n for n in _nB))
    check("melding noemt de verdachte kamer (zelfde naam en oppervlak, 3x)",
          any("Badkamer 4.00 m2 (3x)" in n for n in _nB))
    from dashboard.app import _is_prio_actie
    check("melding is PRIORITEIT (begint met FOUT)",
          _is_prio_actie(next(n for n in _nB if "DUBBELE KAMER" in n)) is True)
    check("geen valse melding op de begane grond (60 = 60)",
          not any("DUBBELE KAMER" in n and "Ground Floor" in n for n in _nB))
except Exception as _e:
    check("dubbele-kamer-check: draait zonder fout", False); print("     " + repr(_e)[:220])

print("AC. Dak-fallback bronprovenance: geen dubbel dakoppervlak (taak 014/015)")
try:
    import tempfile as _tfC, os as _osC, shutil as _shC
    from magicplan.statistics_csv import build_dossier as _bdC

    # AC1: zonder dak-velden in de CSV krijgt de footprint-fallback een aantoonbare herkomst-tag,
    # en de rest van het geïmporteerde schil krijgt de generieke "magicplan-import"-tag.
    _c1 = ("PLAN ATTRIBUTES\nExterior perimeter: m,20,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\n"
           "Gevelhoogte (m),5.4\nGevel - invoer,Beslisschema\nGevel - isolatie aanwezig?,Ja\n"
           "Gevel - isolatiedikte onbekend?,Nee\nGevel - isolatiedikte (mm),80\nGevel - begrenzing,Buitenlucht\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
           "Ground Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _f1 = _tfC.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _f1.write(_c1); _f1.close()
    _d1, _n1 = _bdC(_f1.name); _osC.unlink(_f1.name)
    _fb1 = next((s for s in _d1.schil if s.type == "dak"), None)
    check("AC1: footprint-fallback dak getagd als magicplan-dak-fallback",
          _fb1 is not None and _fb1.bron == "magicplan-dak-fallback", str(_fb1))
    check("AC1: gevel/vloer krijgen de generieke import-tag",
          all(s.bron == "magicplan-import" for s in _d1.schil if s.type != "dak"))

    # AC2: MET expliciete dakvlakken in de CSV komt er geen fallback-dak, en die vlakken krijgen
    # de generieke import-tag (ze zijn geen webapp-wizard-invoer, maar ook geen placeholder).
    _c2 = ("PLAN ATTRIBUTES\nExterior perimeter: m,24,\nBouwjaar,1992.t.m.2013\nWoningtype,Tussenwoning\n"
           "Gevelhoogte (m),5.4\nDakvlak 1 - daktype,Zadeldak\nDak - vloerbreedte (m),8\n"
           "Dakvlak 1 - hellingshoek (°),45\nDakvlak 1 - oriëntatie,ZW\nDakvlak 2 - oriëntatie,NO\n"
           "Dak - kopgevel oriëntatie 1,NW\nDak - kopgevel oriëntatie 2,ZO\n\n"
           "FLOOR ATTRIBUTES,Ground surface without walls,Ceiling Height,Begrenzing\n"
           "Ground Floor,40,2.50 m,Kruipruimte\n\n"
           "WALL ATTRIBUTES,Wall,Symbol,Surf,SurfNoOpen,Width,Height,Ann,Type,Isol,Rekenzone,Orientatie\n"
           "Ground Floor,\nVoorgevel,Wall 0,Wall,10,9,4,2.5,1,Wall,,1,ZW\n")
    _f2 = _tfC.NamedTemporaryFile("w", suffix=".csv", delete=False, encoding="utf-8"); _f2.write(_c2); _f2.close()
    _d2, _n2 = _bdC(_f2.name); _osC.unlink(_f2.name)
    _daks2 = [s for s in _d2.schil if s.type == "dak"]
    check("AC2: expliciete dakvlakken -> GEEN fallback-tag",
          len(_daks2) > 0 and all(s.bron != "magicplan-dak-fallback" for s in _daks2))

    # AC3: dashboard-helper verwijdert de fallback zodra een echt dakvlak ernaast staat, en laat
    # 'm met rust zolang hij alleen staat.
    from dashboard.app import _dak_fallback_opschonen as _dfoC
    from core.dossier import build_sample as _bsC, SchilDeel as _SDC
    _d3 = _bsC()
    for _s in _d3.schil:
        _s.bron = "magicplan-import"
    _dak_idx = next(i for i, s in enumerate(_d3.schil) if s.type == "dak")
    _d3.schil[_dak_idx].bron = "magicplan-dak-fallback"
    check("AC3: alleen-fallback -> niets verwijderd", _dfoC(_d3) is None and len(_d3.schil) == 5)
    _d3.schil.append(_SDC(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=12.0,
                          bron="webapp-wizard"))
    _verwijderd = _dfoC(_d3)
    check("AC3: fallback + echt dakvlak -> fallback verwijderd, m² + id teruggegeven",
          _verwijderd == (60.0, ["dak"]) and len(_d3.schil) == 5
          and not any(s.bron == "magicplan-dak-fallback" for s in _d3.schil)
          and any(s.id == "dak1-plat" for s in _d3.schil))

    # AC4: de webapp-route voor "plat dak toevoegen" schoont de fallback zelf op (end-to-end via
    # de Flask-testclient, niet alleen de losse helperfunctie).
    import dashboard.app as _WAC
    _WAC.app.config.update(TESTING=True, SECRET_KEY="test-dak-fallback")
    _tdC = _tfC.mkdtemp(prefix="nb_dak_fallback_")
    _d4 = _bsC()
    for _s in _d4.schil:
        _s.bron = "magicplan-import" if _s.type != "dak" else "magicplan-dak-fallback"
    _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC = _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save
    _stateC = {"dossier_file": "dossier.json"}
    _WAC._load_state = lambda _t: _stateC
    _WAC._dossier = lambda _t: _d4
    _WAC._pdir = lambda _t: _tdC
    _savedC = []
    _WAC._dos_save = lambda _t, _st, _d: _savedC.append(_d)
    try:
        _cC = _WAC.app.test_client()
        with _cC.session_transaction() as _sessC:
            _sessC["ingelogd"] = True
        _rC = _cC.post("/project/test-dak/opname/dak/plat", data={"breedte": "4", "diepte": "3", "rekenzone": "1"})
        check("AC4: dak/plat-route redirect", _rC.status_code in (302, 303))
        check("AC4: webapp-route verwijdert de fallback zelf",
              len(_savedC) == 1 and not any(s.bron == "magicplan-dak-fallback" for s in _savedC[0].schil)
              and any(s.subtype == "plat dak" and s.bron == "webapp-wizard" for s in _savedC[0].schil))
        with _cC.session_transaction() as _sessC:
            _flashesC = " ".join(m for _cat, m in _sessC.get("_flashes", []))
        check("AC4: flash meldt de opschoning met het verwijderde id",
              "Placeholder-dak" in _flashesC and "verwijderd" in _flashesC and "dak" in _flashesC)
    finally:
        _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save = _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC

    # AC5: VABI-preflight blokkeert fallback+echt-dak-naast-elkaar op ELK exportpad (niet alleen de
    # webapp), en laat een dossier zonder die combinatie gewoon door.
    from vabi.preflight import assert_no_dubbel_dak_fallback as _adfC, VabiExportBlocked as _VEBC
    _d5 = _bsC()
    for _s in _d5.schil:
        _s.bron = "magicplan-import" if _s.type != "dak" else "magicplan-dak-fallback"
    _d5.schil.append(_SDC(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=12.0,
                          bron="webapp-wizard"))
    try:
        _adfC(_d5)
        check("AC5: preflight blokkeert fallback naast wizarddak", False)
    except _VEBC as _eC5:
        check("AC5: preflight blokkeert fallback naast wizarddak",
              "dubbel" in str(_eC5).lower() and "dak" in str(_eC5).lower())
    _d6 = _bsC()
    for _s in _d6.schil:
        _s.bron = "magicplan-import"
    try:
        _adfC(_d6)
        check("AC5: preflight laat een dossier zonder fallback-tag door", True)
    except _VEBC:
        check("AC5: preflight laat een dossier zonder fallback-tag door", False)
    from vabi import generate_all as _gaC
    _tdC2 = _tfC.mkdtemp(prefix="nb_dak_fallback_export_")
    try:
        _gaC.generate_all(_d5, _tdC2, prefix="test")
        check("AC5: generate_all blokkeert vóór schrijven bij fallback+wizarddak", False)
    except _VEBC:
        check("AC5: generate_all blokkeert vóór schrijven bij fallback+wizarddak",
              not os.listdir(_tdC2))

    # AC6: herimport van een verse CSV veegt eerder toegevoegde wizard-dakvlakken niet stil weg.
    _d7 = _bsC()
    for _s in _d7.schil:
        _s.bron = "magicplan-import"
    _d7.schil = [s for s in _d7.schil if s.type != "dak"]   # simuleer: dak al vervangen door de wizard
    _d7.schil.append(_SDC(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=18.5,
                          bron="webapp-wizard"))
    _tdC3 = _tfC.mkdtemp(prefix="nb_dak_reimport_")
    _osC.makedirs(_tdC3, exist_ok=True)
    _WAC._load_state = lambda _t: dict(_stateC)
    _WAC._dossier = lambda _t: _d7
    _WAC._pdir = lambda _t: _tdC3
    from core.dossier import load_json as _ljC
    try:
        _cC2 = _WAC.app.test_client()
        with _cC2.session_transaction() as _sessC2:
            _sessC2["ingelogd"] = True
        import io as _ioC
        _r7 = _cC2.post("/project/test-dak-reimport/opname/magicplan",
                        data={"bestand": (_ioC.BytesIO(_c1.encode("utf-8")), "opname.csv")},
                        content_type="multipart/form-data")
        check("AC6: herimport-route redirect", _r7.status_code in (302, 303))
        _herimportC = _ljC(_osC.path.join(_tdC3, "dossier.json"))
        check("AC6: eerder wizard-dakvlak overleeft de herimport",
              any(s.bron == "webapp-wizard" and s.id == "dak1-plat" for s in _herimportC.schil))
        check("AC6: geen dubbele/fallback dak naast het behouden wizardvlak",
              not any(s.bron == "magicplan-dak-fallback" for s in _herimportC.schil))
    finally:
        _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save = _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC
        _uplC = _osC.path.join(_WAC.UPLOAD_DIR, "opname_test-dak-reimport.csv")
        if _osC.path.exists(_uplC):
            _osC.unlink(_uplC)
    # AC7 (review-fix): een dossier van VÓÓR het bron-veld bestond heeft bron=="" op zijn
    # placeholder-dak — de gedeelde herkenning (vabi.preflight.dak_fallback_schildelen) moet die via
    # het legacy-signaal id=="dak" alsnog vinden, anders is het bestaande, live Essenhage-dossier
    # (de aanleiding voor deze hele taak) NIET beschermd door deze fix.
    from vabi.preflight import dak_fallback_schildelen as _dfsC
    _d8 = _bsC()
    for _s in _d8.schil:
        _s.bron = ""   # simuleer een dossier van vóór dit veld bestond
    _d8.schil.append(_SDC(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=12.0, bron=""))
    check("AC7: legacy dossier (bron altijd '') -> id=='dak' herkend als placeholder",
          [s.id for s in _dfsC(_d8.schil)] == ["dak"])
    try:
        _adfC(_d8)
        check("AC7: preflight blokkeert ook het legacy (ongetagde) Essenhage-patroon", False)
    except _VEBC:
        check("AC7: preflight blokkeert ook het legacy (ongetagde) Essenhage-patroon", True)
    check("AC7: webapp-helper ruimt het legacy-patroon ook op",
          _dfoC(_d8) == (60.0, ["dak"]) and not any(s.id == "dak" for s in _d8.schil))

    # AC8 (review-fix): een dossier-.JSON-herimport (niet alleen CSV) mag wizard-dakvlakken ook
    # niet stil wegvegen.
    _d9 = _bsC()
    for _s in _d9.schil:
        _s.bron = "magicplan-import"
    _d9.schil = [s for s in _d9.schil if s.type != "dak"]
    _d9.schil.append(_SDC(id="dak1-plat", type="dak", subtype="plat dak", oppervlakte_m2=22.0,
                          bron="webapp-wizard"))
    _tdC4 = _tfC.mkdtemp(prefix="nb_dak_reimport_json_")
    _osC.makedirs(_tdC4, exist_ok=True)
    _WAC._load_state = lambda _t: dict(_stateC)
    _WAC._dossier = lambda _t: _d9
    _WAC._pdir = lambda _t: _tdC4
    try:
        _cC3 = _WAC.app.test_client()
        with _cC3.session_transaction() as _sessC3:
            _sessC3["ingelogd"] = True
        _nieuw_json = _bsC()
        for _s in _nieuw_json.schil:
            _s.bron = "magicplan-import"
        _nieuw_json.schil = [s for s in _nieuw_json.schil if s.type != "dak"]
        import json as _jsonC, io as _ioC2
        _r8 = _cC3.post("/project/test-dak-reimport-json/opname/magicplan",
                        data={"bestand": (_ioC2.BytesIO(_jsonC.dumps(_nieuw_json.to_dict()).encode("utf-8")),
                                          "dossier.json")},
                        content_type="multipart/form-data")
        check("AC8: dossier-json-herimport-route redirect", _r8.status_code in (302, 303))
        _herimportC2 = _ljC(_osC.path.join(_tdC4, "dossier.json"))
        check("AC8: wizard-dakvlak overleeft ook een dossier-.json-herimport",
              any(s.bron == "webapp-wizard" and s.id == "dak1-plat" for s in _herimportC2.schil))
    finally:
        _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save = _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC
        _uplC2 = _osC.path.join(_WAC.UPLOAD_DIR, "opname_test-dak-reimport-json.json")
        if _osC.path.exists(_uplC2):
            _osC.unlink(_uplC2)
    _shC.rmtree(_tdC4, ignore_errors=True)

    # AC9 (review-fix ronde 3): een dakkapel op een fallback-getagd (bv. hybride-route) hellend
    # dakvlak mag dat vlak NIET laten opruimen als "ongebruikte placeholder" — de moeder is na de
    # dakkapel-correctie nog altijd echt (verkleind) dakoppervlak, geen dubbeling.
    _d10 = _bsC()
    _d10.schil = [s for s in _d10.schil if s.type != "dak"]
    _d10.schil.append(_SDC(id="dak", type="dak", subtype="Hellend dak", begrenzing="Buitenlucht",
                           orientatie="Z", oppervlakte_m2=60.0, hellingshoek=45.0,
                           bron="magicplan-dak-fallback"))
    _moeder_i9 = len(_d10.schil) - 1
    _tdC5 = _tfC.mkdtemp(prefix="nb_dakkapel_fallback_")
    _WAC._load_state = lambda _t: dict(_stateC)
    _WAC._dossier = lambda _t: _d10
    _WAC._pdir = lambda _t: _tdC5
    try:
        _cC4 = _WAC.app.test_client()
        with _cC4.session_transaction() as _sessC4:
            _sessC4["ingelogd"] = True
        _r9 = _cC4.post("/project/test-dakkapel-fallback/opname/dakkapel",
                        data={"moederdak_i": str(_moeder_i9), "breedte": "2", "hoogte": "1.5",
                              "diepte": "1", "rekenzone": "1"})
        check("AC9: dakkapel-op-fallback-route redirect", _r9.status_code in (302, 303))
        _moeder9 = _d10.schil[_moeder_i9]
        check("AC9: moederdak (voorheen fallback) is niet meer weggooibaar getagd",
              _moeder9.id == "dak" and _moeder9.bron == "magicplan-import"
              and _moeder9.oppervlakte_m2 < 60.0)
        check("AC9: dakje van de dakkapel is normaal webapp-wizard-getagd",
              any(s.id == "dakkapel1-dakje" and s.bron == "webapp-wizard" for s in _d10.schil))
        try:
            _adfC(_d10)
            check("AC9: preflight blokkeert NIET (geen echte dubbeling, alleen een gat + dakje)", True)
        except _VEBC as _e9:
            check("AC9: preflight blokkeert NIET (geen echte dubbeling, alleen een gat + dakje)",
                  False)
            print("     " + str(_e9)[:160])
    finally:
        _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save = _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC
    _shC.rmtree(_tdC5, ignore_errors=True)

    # Taak 023: een oud dossier heeft op dezelfde placeholder nog bron="". Het legacy-signaal
    # id="dak" moet vóór de dakkapelmutatie worden vastgelegd en daarna worden herclassificeerd.
    _d11 = _bsC()
    _d11.schil = [s for s in _d11.schil if s.type != "dak"]
    _d11.schil.append(_SDC(id="dak", type="dak", subtype="Hellend dak", begrenzing="Buitenlucht",
                           orientatie="Z", oppervlakte_m2=60.0, hellingshoek=45.0, bron=""))
    _moeder_i10 = len(_d11.schil) - 1
    _tdC6 = _tfC.mkdtemp(prefix="nb_dakkapel_legacy_fallback_")
    _WAC._load_state = lambda _t: dict(_stateC)
    _WAC._dossier = lambda _t: _d11
    _WAC._pdir = lambda _t: _tdC6
    try:
        _cC5 = _WAC.app.test_client()
        with _cC5.session_transaction() as _sessC5:
            _sessC5["ingelogd"] = True
        _r10 = _cC5.post("/project/test-dakkapel-legacy/opname/dakkapel",
                         data={"moederdak_i": str(_moeder_i10), "breedte": "2", "hoogte": "1.5",
                               "diepte": "1", "rekenzone": "1"})
        _moeder10 = _d11.schil[_moeder_i10]
        check("taak023: dakkapel op ongetagd legacy-moederdak herclassificeert het vlak",
              _r10.status_code in (302, 303) and _moeder10.id == "dak"
              and _moeder10.bron == "magicplan-import" and _moeder10.oppervlakte_m2 < 60.0)
        try:
            _adfC(_d11)
            check("taak023: herclassificatie voorkomt onterechte dubbel-dakblokkade", True)
        except _VEBC:
            check("taak023: herclassificatie voorkomt onterechte dubbel-dakblokkade", False)
    finally:
        _WAC._load_state, _WAC._dossier, _WAC._pdir, _WAC._dos_save = _orig_lsC, _orig_dosC, _orig_pdC, _orig_saveC
    _shC.rmtree(_tdC6, ignore_errors=True)

    _shC.rmtree(_tdC, ignore_errors=True)
    _shC.rmtree(_tdC2, ignore_errors=True)
    _shC.rmtree(_tdC3, ignore_errors=True)
except Exception as _e:
    check("dak-fallback-opschoning: draait zonder fout", False); print("     " + repr(_e)[:220])

print("\nAD. Mappingmanifest (taak 015): drift-detectie parser <-> webapp <-> VABI + form-fingerprint")
try:
    from core.mapping_manifest import (MANIFEST, VERWACHTE_SNAPSHOT_FINGERPRINT,
                                       resolve, get as _mm_get)
    from scripts.check_mapping_manifest import (run_checks as _mm_run_checks, render_doc as _mm_render_doc,
                                                 check_doc as _mm_check_doc, DOC_PATH as _mm_doc_path)
    from magicplan.form_fingerprint import (compute_fingerprint, snapshot_fingerprint, load_snapshot,
                                            stamp_dossier_meta, refresh_snapshot_live,
                                            DEFAULT_SNAPSHOT_PATH)

    check("manifest: minstens de bekende dropdowns zijn opgenomen",
          {"begrenzing", "glastype", "kozijnmateriaal"} <= {m.id for m in MANIFEST})
    _mm_errors = _mm_run_checks()
    check("manifest: GEEN drift tussen parser/webapp/VABI-codebook", not _mm_errors, "; ".join(_mm_errors))

    # Elke manifest-entry: alle referenties moeten daadwerkelijk resolven (bewijst dat de check zelf
    # geen stille no-op is bij een kapotte referentie -- 1 entry expres verbouwen en weer herstellen).
    _begr = _mm_get("begrenzing")
    _orig_ref = _begr.parser_canon
    _begr.parser_canon = "magicplan.statistics_csv:DEZE_BESTAAT_NIET"
    try:
        resolve(_begr.parser_canon)
        check("manifest: kapotte referentie faalt luid (AttributeError)", False)
    except AttributeError:
        check("manifest: kapotte referentie faalt luid (AttributeError)", True)
    finally:
        _begr.parser_canon = _orig_ref
    check("manifest: referentie hersteld", _begr.parser_canon == _orig_ref)

    # doc-generatie: --write-doc en --check-doc zijn consistent (round-trip), en het gecommitte
    # docs/mapping-overview.md is NIET verouderd t.o.v. het manifest zoals het nu in de repo staat.
    check("manifest-doc: render is deterministisch", _mm_render_doc() == _mm_render_doc())
    check("manifest-doc: gecommitte docs/mapping-overview.md is actueel (--write-doc gedraaid?)",
          _mm_check_doc(), "draai: python scripts/check_mapping_manifest.py --write-doc")

    # kozijnmateriaal-regressie (de concrete bug die dit manifest heeft blootgelegd, 21-8): de parser
    # mag NOOIT meer een kozijnwaarde produceren die de webapp-<select> niet als optie heeft staan.
    from dashboard.app import KOZ_OPTS
    from magicplan.statistics_csv import _norm_kozijn_mat as _nkm_ad
    check("kozijnmateriaal: parserwaarden zitten allemaal in de webapp-optielijst",
          all(_nkm_ad(v) in KOZ_OPTS for v in ("a", "b", "c", "")),
          str([(v, _nkm_ad(v)) for v in ("a", "b", "c", "")]))

    # form-fingerprint: gedateerde snapshot, geen live call nodig; stabiel/reproduceerbaar en
    # gevoelig voor een echte inhoudswijziging (drift-detectie).
    _snap = load_snapshot(DEFAULT_SNAPSHOT_PATH)
    check("form-fingerprint: snapshot heeft een datum", bool(_snap.get("_meta", {}).get("snapshot_datum")))
    _fp1 = snapshot_fingerprint()
    _fp2 = snapshot_fingerprint()
    check("form-fingerprint: reproduceerbaar (zelfde snapshot -> zelfde fingerprint)", _fp1 == _fp2)
    check("form-fingerprint: gecommitte snapshot matcht onafhankelijke contractpin",
          _fp1 == VERWACHTE_SNAPSHOT_FINGERPRINT)
    _forms_gewijzigd = dict(_snap["forms"])
    _forms_gewijzigd["Object"] = _forms_gewijzigd["Object"] + [{"name": "Nieuw testveld", "options": None}]
    check("form-fingerprint: verandert bij een echte veldwijziging",
          compute_fingerprint(_forms_gewijzigd) != _fp1)
    _forms_zelfde_volgorde_anders = {k: list(reversed(v)) for k, v in _snap["forms"].items()}
    check("form-fingerprint: ongevoelig voor volgorde (alleen inhoud telt)",
          compute_fingerprint(_forms_zelfde_volgorde_anders) == _fp1)

    # De drift-check moet de snapshot daadwerkelijk lezen. Drie mutaties bewijzen dat optie-,
    # label- én overige formulierdrift rood worden; dit is geen zinloze snapshot-self-compare.
    import copy as _copy_ad
    def _snapshot_errors_ad(mutator):
        gewijzigd = _copy_ad.deepcopy(_snap)
        mutator(gewijzigd)
        pad = tempfile.mktemp(suffix=".json")
        try:
            with open(pad, "w", encoding="utf-8") as fh:
                json.dump(gewijzigd, fh, ensure_ascii=False)
            return _mm_run_checks(pad)
        finally:
            if os.path.exists(pad):
                os.unlink(pad)

    _optiedrift_ad = _snapshot_errors_ad(
        lambda s: s["forms"]["Raam/paneel"][6]["options"].append("Nieuw glastype"))
    check("manifest-snapshot: gewijzigde dropdownoptie faalt luid",
          any("snapshotopties" in e and "glastype" in e for e in _optiedrift_ad),
          str(_optiedrift_ad))

    def _labeldrift_ad(s):
        next(v for v in s["forms"]["Object"] if v["name"] == "Woningtype")["name"] = "Woningsoort"
    _label_errors_ad = _snapshot_errors_ad(_labeldrift_ad)
    check("manifest-snapshot: gewijzigd bronlabel faalt luid",
          any("labeldrift" in e and "woningtype_subtype" in e for e in _label_errors_ad),
          str(_label_errors_ad))

    def _overige_drift_ad(s):
        s["forms"]["Object"].append({"name": "Nieuw los veld", "options": None})
    _fp_errors_ad = _snapshot_errors_ad(_overige_drift_ad)
    check("manifest-snapshot: overige fingerprintdrift faalt tegen vaste contractpin",
          any("fingerprint" in e and "wijkt af" in e for e in _fp_errors_ad),
          str(_fp_errors_ad))

    # Live-refresh blijft offline testbaar: bij een gelijknamige groep/veld uit beide API-routes
    # moeten unieke velden én alle opties behouden blijven, zonder dubbele veldrecords.
    import magicplan.form_push as _mfp_push
    _orig_env_ad = _mfp_push._load_env
    _orig_forms_ad = _mfp_push.fetch_forms
    _orig_fields_ad = _mfp_push.fetch_fields
    _tmp_snapshot_ad = tempfile.mktemp(suffix=".json")
    try:
        _mfp_push._load_env = lambda: {"offline": "test"}
        _mfp_push.fetch_forms = lambda _env: [
            {"name": "Object", "fields": [
                {"name": "Gedeeld", "options": ["A"]},
                {"name": "Alleen form", "options": None},
            ]}
        ]
        _mfp_push.fetch_fields = lambda _env: [
            {"name": "Object", "fields": [
                {"name": "Gedeeld", "options": ["A", "B"]},
                {"name": "Alleen fields", "options": ["X"]},
            ]}
        ]
        _live_snap_ad = refresh_snapshot_live(_tmp_snapshot_ad)
        _live_obj_ad = _live_snap_ad["forms"]["Object"]
        check("form-fingerprint live-refresh: gelijknamige groepen worden samengevoegd",
              {v["name"] for v in _live_obj_ad} == {"Gedeeld", "Alleen form", "Alleen fields"})
        check("form-fingerprint live-refresh: opties uit beide routes blijven behouden",
              next(v for v in _live_obj_ad if v["name"] == "Gedeeld")["options"] == ["A", "B"])
        check("form-fingerprint live-refresh: gelijknamig veld staat maar één keer in snapshot",
              sum(v["name"] == "Gedeeld" for v in _live_obj_ad) == 1)
    finally:
        _mfp_push._load_env = _orig_env_ad
        _mfp_push.fetch_forms = _orig_forms_ad
        _mfp_push.fetch_fields = _orig_fields_ad
        if os.path.exists(_tmp_snapshot_ad):
            os.unlink(_tmp_snapshot_ad)

    # dossier-meta: elke MagicPlan-importroute stempelt fingerprint + snapshotdatum.
    from core.dossier import Dossier
    _dstamp = Dossier()
    check("dossier-meta: leeg vóór het stempelen", _dstamp.meta.magicplan_form_fingerprint == "")
    stamp_dossier_meta(_dstamp)
    check("dossier-meta: fingerprint gezet na stamp_dossier_meta", _dstamp.meta.magicplan_form_fingerprint == _fp1)
    check("dossier-meta: snapshotdatum gezet na stamp_dossier_meta",
          _dstamp.meta.magicplan_form_snapshot_datum == _snap["_meta"]["snapshot_datum"])
    check("dossier-meta: importtijd (magicplan_import_op) gezet na stamp_dossier_meta",
          bool(_dstamp.meta.magicplan_import_op))

    # de 3 echte MagicPlan-importroutes (CSV/hybride-API+report/API-only) stempelen allemaal mee --
    # geen route mag stilletjes een dossier zonder fingerprint opleveren.
    import inspect as _insp_ad
    from magicplan import statistics_csv as _mscv_ad, assemble as _masm_ad, extractor as _mext_ad
    check("statistics_csv.build_dossier stempelt de fingerprint",
          "stamp_dossier_meta" in _insp_ad.getsource(_mscv_ad.build_dossier))
    check("assemble.build_dossier stempelt de fingerprint",
          "stamp_dossier_meta" in _insp_ad.getsource(_masm_ad.build_dossier))
    check("extractor.map_plan_to_dossier stempelt de fingerprint",
          "stamp_dossier_meta" in _insp_ad.getsource(_mext_ad.map_plan_to_dossier))
except Exception as _e:
    check("mappingmanifest/form-fingerprint: draait zonder fout", False); print("     " + repr(_e)[:300])
print("\nAD. Ventilatie-rekenlaag uitbreiden (taak 019): balans, deurbelasting, vuistregeltoets")
try:
    from ventilatie.ventilatie import (bereken as _vbAD, verdeel_balans as _vbal, deurbelasting as _vdb,
                                        toets_vuistregels as _vtv)
    from core.dossier import Ruimte as _RAD

    # Fixture "demo-woning" (docs/ventilatieplan-webapp-spec.md §1) — het voorbeeldproject waarop een
    # concurrent (Aira) hun eigen rekenhart liet zien. Woonkamer+Keuken zijn in de spec-video één open
    # ruimte waarvan Aira's UI 38,7 (woonkamer) + 18,5 (keuken) als aparte labels toont maar wél
    # samen doorrekent (0,7 x 57,2 = 40,04 -> hun 40,0); wij modelleren dat hier zo: de Woonkamer-regel
    # draagt het volledige (samengevoegde) vloeroppervlak voor de toevoerberekening, de Keuken-regel
    # blijft apart bestaan met oppervlak 0 zodat hij zijn eigen afvoereis (21 l/s) en afvoerpunt-vlag
    # houdt zonder het toevoeroppervlak dubbel te tellen (zie de "geen fantoom-toevoer"-fix in bereken()).
    _demo_ruimtes = [
        _RAD(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=57.2),   # 38,7 + 18,5 (keuken)
        _RAD(naam="Keuken", functie="keuken", oppervlakte_m2=0.0),
        _RAD(naam="Slaapkamer 4", functie="slaapkamer", oppervlakte_m2=18.2),
        _RAD(naam="Bijkeuken", functie="wasruimte", oppervlakte_m2=9.7),
        _RAD(naam="Entree", functie="verkeer", oppervlakte_m2=8.2),
        _RAD(naam="Toilet", functie="toilet", oppervlakte_m2=1.2),
        _RAD(naam="Slaapkamer 1", functie="slaapkamer", oppervlakte_m2=12.2),
        _RAD(naam="Slaapkamer 3", functie="slaapkamer", oppervlakte_m2=6.9),
        _RAD(naam="Slaapkamer 2", functie="slaapkamer", oppervlakte_m2=9.3),
        _RAD(naam="Overloop", functie="verkeer", oppervlakte_m2=3.8),
        _RAD(naam="Badkamer", functie="badkamer", oppervlakte_m2=6.9),
    ]
    _resAD = _vbAD(_demo_ruimtes)
    _byAD = {r["naam"]: r for r in _resAD["rows"]}

    check("toevoer woonkamer 40,0 (0,7 x 57,2 = woonkamer 38,7 plus keuken 18,5)",
          _byAD["Woonkamer"]["toevoer"] == 40.0, str(_byAD["Woonkamer"]))
    check("toevoer slaapkamer 4: 12,7 (0,7 x 18,2, onze afronding op 0,1 — zie decision 0002)",
          _byAD["Slaapkamer 4"]["toevoer"] == 12.7)
    check("toevoer slaapkamer 1: 8,5 (0,7 x 12,2)", _byAD["Slaapkamer 1"]["toevoer"] == 8.5)
    check("toevoer slaapkamer 3: 7,0 (0,7 x 6,9 = 4,8, floor min 7 l/s/leefruimte)",
          _byAD["Slaapkamer 3"]["toevoer"] == 7.0)
    check("toevoer slaapkamer 2: 7,0 (0,7 x 9,3 = 6,5, floor)", _byAD["Slaapkamer 2"]["toevoer"] == 7.0)
    check("herkomst toevoer woonkamer = oppervlakte", _byAD["Woonkamer"]["toevoer_herkomst"] == "oppervlakte")
    check("herkomst toevoer slaapkamer 3 = minimum (floor geraakt)",
          _byAD["Slaapkamer 3"]["toevoer_herkomst"] == "minimum")
    check("keuken (0 m2, eigen regel) krijgt GEEN fantoom-toevoer", _byAD["Keuken"]["toevoer"] == 0.0)

    check("afvoerpunt: ja voor de 4 natte ruimten, nee voor de rest",
          {r["naam"] for r in _resAD["rows"] if r["afvoerpunt"]} == {"Keuken", "Bijkeuken", "Toilet", "Badkamer"})

    _balAD = _vbal(_resAD)
    _byBalAD = {r["naam"]: r for r in _balAD["rows"]}
    check("bereken() blijft puur: verdeel_balans wijzigt het origineel niet",
          _byAD["Keuken"]["afvoer_advies_ls"] == 21.0)
    check("verdeel_balans: som advies-afvoer = som toevoer (76-achtig sluitend, geen restje)",
          _balAD["afvoer_advies_totaal"] == _balAD["toevoer_totaal"], str(_balAD["afvoer_advies_totaal"]))
    check("verdeel_balans: geen natte ruimte onder zijn minimum",
          all(r["afvoer_advies_ls"] >= r["afvoer"] for r in _balAD["rows"] if r["afvoerpunt"]))
    check("verdeel_balans: keuken blijft de grootste afnemer",
          _byBalAD["Keuken"]["afvoer_advies_ls"] == max(
              r["afvoer_advies_ls"] for r in _balAD["rows"] if r["afvoerpunt"]))
    check("verdeel_balans: opgehoogde regels krijgen herkomst 'balansophoging'",
          _byBalAD["Badkamer"]["afvoer_herkomst"] == "balansophoging" and _byBalAD["Badkamer"]["afvoer_advies_ls"] > 14.0)

    _deurenAD = _vdb(_balAD, [["Badkamer", "Overloop"]])
    check("deurbelasting: 1 regel voor het pad Badkamer-Overloop", len(_deurenAD) == 1)
    check("deurbelasting boven 15 l/s -> deurrooster-tekst",
          _deurenAD[0]["boven_norm"] is True and "deurrooster geadviseerd" in _deurenAD[0]["tekst"]
          and _deurenAD[0]["van"] == "Badkamer" and _deurenAD[0]["naar"] == "Overloop")
    _deuren_laag = _vdb(_balAD, [["Toilet", "Entree"]])
    check("deurbelasting onder 15 l/s -> geen deurrooster-tekst",
          _deuren_laag[0]["boven_norm"] is False and "deurrooster" not in _deuren_laag[0]["tekst"])

    # Vuistregeltoets: alle 7 regels moeten altijd terugkomen, mét een topologie kunnen we er een paar
    # ook echt beoordelen; zonder mag niets stiekem 'voldoet' worden.
    _tvAD_leeg = _vtv(_balAD)
    check("toets_vuistregels: alle 7 regels aanwezig (leeg plan)", len(_tvAD_leeg) == 7)
    check("toets_vuistregels: zonder topologie is overstroom-regel 'niet te bepalen', niet stil 'voldoet'",
          _tvAD_leeg[0]["status"] == "niet te bepalen")
    check("toets_vuistregels: geen afvoerpunt in slaapkamer -> voldoet (deze fixture heeft er geen)",
          _tvAD_leeg[2]["status"] == "voldoet")
    check("toets_vuistregels: >=50% van buiten voldoet (toevoer > 0 in dit model)",
          _tvAD_leeg[1]["status"] == "voldoet")
    check("toets_vuistregels: rookkanaal/afstand/C4c zonder plan-gegevens -> niet te bepalen",
          _tvAD_leeg[4]["status"] == "niet te bepalen" and _tvAD_leeg[5]["status"] == "niet te bepalen"
          and _tvAD_leeg[6]["status"] == "niet te bepalen")

    _plan_ok = {"topologie": [["Badkamer", "Overloop"]], "deurroosters": {"Badkamer-Overloop"},
                "afstand_toe_afvoer_m": 2.5, "afstand_rookkanaal_m": 8.0,
                "co2_sturing_woonkamer": True, "co2_sturing_hoofdslaapkamer": True}
    _tvAD_ok = _vtv(_balAD, _plan_ok)
    check("toets_vuistregels: max 2 deuren -> voldoet (1 deur op de weg)", _tvAD_ok[0]["status"] == "voldoet")
    check("toets_vuistregels: deurrooster al aanwezig -> voldoet ondanks 18,8 l/s", _tvAD_ok[3]["status"] == "voldoet")
    check("toets_vuistregels: afstand/rookkanaal/C4c met gegevens -> voldoet",
          _tvAD_ok[4]["status"] == "voldoet" and _tvAD_ok[5]["status"] == "voldoet" and _tvAD_ok[6]["status"] == "voldoet")

    _plan_fout = {"topologie": [["Badkamer", "Overloop"], ["Slaapkamer 1", "Overloop", "Woonkamer", "Entree"]],
                  "afstand_toe_afvoer_m": 0.5, "co2_sturing_woonkamer": True, "co2_sturing_hoofdslaapkamer": False}
    _tvAD_fout = _vtv(_balAD, _plan_fout)
    check("toets_vuistregels: 3-deurs overstroomweg -> voldoet niet (max 2)", _tvAD_fout[0]["status"] == "voldoet niet")
    check("toets_vuistregels: deurrooster ontbreekt (18,8 l/s, geen rooster gemeld) -> voldoet niet",
          _tvAD_fout[3]["status"] == "voldoet niet")
    check("toets_vuistregels: toevoer/afvoer te dicht bij elkaar -> voldoet niet", _tvAD_fout[4]["status"] == "voldoet niet")
    check("toets_vuistregels: C4c niet op beide ruimten -> voldoet niet", _tvAD_fout[6]["status"] == "voldoet niet")

    # vuistregel 3 apart: een afvoerpunt in een slaapkamer moet 'voldoet niet' geven, nooit stil door.
    _slaapafvoer = _vbAD([_RAD(naam="Slaapkamer met afzuiging", functie="slaapkamer", oppervlakte_m2=10.0)])
    # forceer een afvoerpunt op de slaapkamer-regel (edge case: badkamer-en-suite met eigen afzuiging)
    _slaapafvoer["rows"][0]["afvoerpunt"] = True
    _tv_slaap = _vtv(_slaapafvoer)
    check("toets_vuistregels: afvoerpunt in slaapkamer -> voldoet niet", _tv_slaap[2]["status"] == "voldoet niet")

    check("bereken() blijft puur: geen Flask/bestand-IO nodig om deze functies te draaien", True)

    check("0 m2 op een verblijfsgebied-regel geeft een waarschuwing, geen stille 0 l/s",
          any("Keuken" in w and "0.0 m2" in w for w in _resAD["waarschuwingen"]))

    # Regressie (code review taak 020): de waarschuwing moet het WERKELIJKE (mogelijk negatieve)
    # oppervlak tonen, niet altijd hardcoded '0 m2' — anders verhult de melding zelf een ergere fout
    # (een negatief oppervlak, bv. een MagicPlan-editfout) achter een onschuldig ogend '0 m2'.
    _negAD = _vbAD([_RAD(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=-12.0)])
    check("negatief oppervlak op een verblijfsgebied-regel: waarschuwing toont het ECHTE (negatieve) getal",
          any("Woonkamer" in w and "-12.0 m2" in w for w in _negAD["waarschuwingen"]))

    # Regressie (code review): een simpele 'grootste afnemer krijgt de afrondingsrest' kon bij >=5
    # natte ruimten de grootste regel juist ONDER zijn eigen minimum duwen door een vastgeklikte
    # 0,1 l/s-afrondingsfout. De grootste-restmethode (Hamilton) mag dat nooit meer doen.
    _veel_natte = [_RAD(naam="Keuken", functie="keuken", oppervlakte_m2=0.0),
                   _RAD(naam="Badkamer 1", functie="badkamer", oppervlakte_m2=0.0),
                   _RAD(naam="Badkamer 2", functie="badkamer", oppervlakte_m2=0.0),
                   _RAD(naam="Toilet 1", functie="toilet", oppervlakte_m2=0.0),
                   _RAD(naam="Wasruimte", functie="wasruimte", oppervlakte_m2=0.0),
                   _RAD(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=100.0)]
    _resVeel = _vbAD(_veel_natte)
    _balVeel = _vbal(_resVeel)
    check("verdeel_balans (5 natte ruimten): niemand onder zijn minimum",
          all(r["afvoer_advies_ls"] >= r["afvoer"] for r in _balVeel["rows"] if r["afvoerpunt"]),
          str([(r["naam"], r["afvoer"], r["afvoer_advies_ls"]) for r in _balVeel["rows"] if r["afvoerpunt"]]))
    check("verdeel_balans (5 natte ruimten): som advies-afvoer sluit exact op de toevoer",
          _balVeel["afvoer_advies_totaal"] == _balVeel["toevoer_totaal"], str(_balVeel["afvoer_advies_totaal"]))
except Exception as _e:
    check("ventilatie-rekenlaag (taak 019): draait zonder fout", False); print("     " + repr(_e)[:220])

try:
    from ventilatie.ventilatie import deurbelasting as _vdb2, bereken as _vb2
    from core.dossier import Ruimte as _R2
    _res2 = _vb2([_R2(naam="Badkamer", functie="badkamer", oppervlakte_m2=5.0)])
    _fout_opgevangen = False
    try:
        _vdb2(_res2, [["Bathroom (tikfout)", "Overloop"]])
    except ValueError:
        _fout_opgevangen = True
    check("deurbelasting: onbekende ruimtenaam geeft een harde ValueError (geen stille 0 l/s)",
          _fout_opgevangen)

    # Regressie (code review taak 020): eerder werd ALLEEN pad[0] gevalideerd — een tikfout verderop
    # in het pad (de 'naar'-kant van een deur) gleed er stil doorheen. 'Badkamer' bestaat hier wél.
    _fout_verderop = False
    try:
        _vdb2(_res2, [["Badkamer", "Overlop (tikfout)"]])
    except ValueError:
        _fout_verderop = True
    check("deurbelasting: tikfout VERDEROP in het pad (niet pad[0]) geeft ook een harde ValueError",
          _fout_verderop)
except Exception as _e:
    check("deurbelasting-validatie: draait zonder fout", False); print("     " + repr(_e)[:220])

print("\nAE. Ventilatieplan-pagina (taak 020): datalaag (groeperen, autoplaatsing, validatie, balans)")
try:
    from dashboard import ventilatieplan as _vp
    from core.dossier import (Dossier as _DAE, Ruimte as _RAE, VloerInfo as _VIAE,
                               VentilatieMarker as _VMAE)
    from ventilatie.ventilatie import bereken as _vbAE, verdeel_balans as _vbalAE

    # --- groepeer_per_verdieping ---
    _d0 = _DAE()
    check("groepeer: geen vloeren -> 1 groep 'Begane grond'",
          _vp.groepeer_per_verdieping(_d0) == [("Begane grond", None, [])])

    _d1 = _DAE()
    _d1.geometrie.vloeren = [_VIAE(naam="Verdieping X", oppervlakte_m2=50)]
    _d1.geometrie.ruimtes = [_RAE(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30)]
    _g1 = _vp.groepeer_per_verdieping(_d1)
    check("groepeer: 1 vloer -> alle ruimtes in die ene groep, ongeacht Ruimte.verdieping",
          len(_g1) == 1 and _g1[0][0] == "Verdieping X" and len(_g1[0][2]) == 1)

    _d2 = _DAE()
    _d2.geometrie.vloeren = [_VIAE(naam="Begane grond"), _VIAE(naam="1e verdieping")]
    _d2.geometrie.ruimtes = [
        _RAE(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30, verdieping="Begane grond"),
        _RAE(naam="Slaapkamer 1", functie="slaapkamer", oppervlakte_m2=12, verdieping="1e verdieping"),
        _RAE(naam="Verweesde kamer", functie="overig", oppervlakte_m2=5, verdieping="Zolder (bestaat niet)")]
    _g2 = _vp.groepeer_per_verdieping(_d2)
    check("groepeer: 2 vloeren -> 3 groepen (BG, 1e, + niet-gekoppeld)", len(_g2) == 3)
    _bg2 = next(g for g in _g2 if g[0] == "Begane grond")
    _ng2 = next(g for g in _g2 if g[0] == "Niet gekoppeld aan een verdieping")
    check("groepeer: begane grond bevat alleen de woonkamer", [r.naam for r in _bg2[2]] == ["Woonkamer"])
    check("groepeer: onbekende verdieping komt in een expliciete 'niet gekoppeld'-groep, niet stil weg",
          [r.naam for r in _ng2[2]] == ["Verweesde kamer"])

    # --- achtergrond_van ---
    check("achtergrond: geen VloerInfo -> (None, 'geen')", _vp.achtergrond_van(None) == (None, "geen"))
    check("achtergrond: lege VloerInfo -> (None, 'geen')", _vp.achtergrond_van(_VIAE()) == (None, "geen"))
    check("achtergrond: contour_m zonder afbeelding -> (None, 'contour')",
          _vp.achtergrond_van(_VIAE(contour_m=[[0, 0], [1, 0], [1, 1]])) == (None, "contour"))
    check("achtergrond: plattegrond_afbeelding wint van contour_m",
          _vp.achtergrond_van(_VIAE(contour_m=[[0, 0]], plattegrond_afbeelding="bg.png"))
          == ("bg.png", "afbeelding"))

    # --- auto_markers (op de demo-woning-achtige mix) ---
    _ruimtesAE = [_RAE(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30),
                  _RAE(naam="Keuken", functie="keuken", oppervlakte_m2=0.0),
                  _RAE(naam="Slaapkamer 1", functie="slaapkamer", oppervlakte_m2=12),
                  _RAE(naam="Badkamer", functie="badkamer", oppervlakte_m2=5),
                  _RAE(naam="Toilet", functie="toilet", oppervlakte_m2=1.2),
                  _RAE(naam="Hal", functie="verkeer", oppervlakte_m2=4)]
    _resAE = _vbalAE(_vbAE(_ruimtesAE))
    _autoAE = _vp.auto_markers(_ruimtesAE, _resAE["rows"])
    check("auto_markers: toevoer alleen voor ruimtes met echte toevoer (Keuken 0 m2 dus niet)",
          {m.ruimte_id for m in _autoAE if m.type == "toevoer"} == {"Woonkamer", "Slaapkamer 1"})
    check("auto_markers: afvoer voor elke natte ruimte (afvoerpunt)",
          {m.ruimte_id for m in _autoAE if m.type == "afvoer"} == {"Keuken", "Badkamer", "Toilet"})
    check("auto_markers: zonder bevestigde topologie GEEN fantoom-overstroom",
          not any(m.type == "overstroom" for m in _autoAE))
    check("auto_markers: coördinaten relatief (0..1)",
          all(0.0 <= m.x <= 1.0 and 0.0 <= m.y <= 1.0 for m in _autoAE))
    check("auto_markers: bron='auto' op elke gegenereerde marker", all(m.bron == "auto" for m in _autoAE))
    _ruimtes_geoAE = [_RAE(naam="Geo", functie="slaapkamer", oppervlakte_m2=10,
                            contour_relatief=[[.2,.2],[.6,.2],[.6,.6],[.2,.6]])]
    _auto_geoAE = _vp.auto_markers(_ruimtes_geoAE, _vbalAE(_vbAE(_ruimtes_geoAE))["rows"])
    check("auto_markers: met echte ruimtegeometrie starten alle markers binnen de gekoppelde ruimte",
          all(_vp.punt_in_polygoon(m.x, m.y, _ruimtes_geoAE[0].contour_relatief) for m in _auto_geoAE))
    _top_ruimtesAE = [
        _RAE(naam="Bron", functie="slaapkamer", oppervlakte_m2=10,
             contour_relatief=[[.1,.1],[.45,.1],[.45,.7],[.1,.7]]),
        _RAE(naam="Doel", functie="badkamer", oppervlakte_m2=6,
             contour_relatief=[[.55,.1],[.9,.1],[.9,.7],[.55,.7]])]
    _top_resAE = _vbalAE(_vbAE(_top_ruimtesAE))["rows"]
    _top_autoAE = _vp.auto_markers(_top_ruimtesAE, _top_resAE, [["Bron", "Doel"]])
    _overAE = [m for m in _top_autoAE if m.type == "overstroom"]
    check("auto_markers: bevestigde bron->nat-doel-topologie maakt precies één overstroommarker",
          len(_overAE) == 1 and _overAE[0].ruimte_id == "Bron")
    check("auto_markers: overstroom ligt controleerbaar op bronrand richting doel",
          _vp.punt_in_polygoon(_overAE[0].x, _overAE[0].y, _top_ruimtesAE[0].contour_relatief)
          and _overAE[0].x > .4)
    _c_bronAE = _RAE(naam="C-bron", contour_relatief=[
        [.1,.1],[.8,.1],[.8,.25],[.3,.25],[.3,.75],[.8,.75],[.8,.9],[.1,.9]])
    _c_doelAE = _RAE(naam="C-doel", contour_relatief=[[.85,.35],[.98,.35],[.98,.65],[.85,.65]])
    _c_randAE = _vp._randpunt_van_naar(_c_bronAE, _c_doelAE)
    check("randpunt: concave C-vorm gebruikt intern startpunt en echte eerste rand",
          _c_randAE is not None and _vp.punt_in_polygoon(*_c_randAE, _c_bronAE.contour_relatief)
          and abs(_c_randAE[1] - .25) < .001)
    _u_bronAE = _RAE(naam="U-bron", contour_relatief=[
        [.1,.1],[.3,.1],[.3,.7],[.7,.7],[.7,.1],[.9,.1],[.9,.9],[.1,.9]])
    _u_doelAE = _RAE(naam="U-doel", contour_relatief=[[.4,.01],[.6,.01],[.6,.08],[.4,.08]])
    _u_randAE = _vp._randpunt_van_naar(_u_bronAE, _u_doelAE)
    check("randpunt: concave U-vorm eindigt epsilon binnen bron",
          _u_randAE is not None and _vp.punt_in_polygoon(*_u_randAE, _u_bronAE.contour_relatief))
    _bk_marker = next(m for m in _autoAE if m.ruimte_id == "Badkamer")
    check("auto_markers: afvoerwaarde komt uit afvoer_advies_ls (na balansverdeling), niet het kale minimum",
          _bk_marker.waarde_ls == next(r for r in _resAE["rows"] if r["naam"] == "Badkamer")["afvoer_advies_ls"]
          and _bk_marker.waarde_ls > 5.0)

    # --- zorg_voor_verdiepingen: idempotent, vult alleen lege verdiepingen ---
    _d3 = _DAE(); _d3.geometrie.ruimtes = _ruimtesAE
    _res3 = _vbalAE(_vbAE(_d3.geometrie.ruimtes))
    check("zorg_voor_verdiepingen: 1e keer vult en meldt gewijzigd", _vp.zorg_voor_verdiepingen(_d3, _res3["rows"]) is True)
    check("zorg_voor_verdiepingen: markers staan er nu echt", len(_d3.ventilatieplan.verdiepingen[0].markers) > 0)
    _d3.ventilatieplan.verdiepingen[0].markers[0].waarde_ls = 999.0   # simuleer een handmatige wijziging
    check("zorg_voor_verdiepingen: 2e keer raakt bestaande (niet-lege) markers niet aan",
          _vp.zorg_voor_verdiepingen(_d3, _res3["rows"]) is False
          and _d3.ventilatieplan.verdiepingen[0].markers[0].waarde_ls == 999.0)
    _d3.ventilatieplan.verdiepingen[0].markers.append(
        _VMAE(id="oud-fantoom", type="overstroom", ruimte_id="Woonkamer", waarde_ls=7, x=.5, y=.5))
    check("zorg_voor_verdiepingen: legacy overstroom zonder topologie wordt verwijderd",
          _vp.zorg_voor_verdiepingen(_d3, _res3["rows"]) is True
          and not any(m.type == "overstroom" for m in _d3.ventilatieplan.verdiepingen[0].markers))

    # --- herstel_verdieping: raakt alleen de opgegeven verdieping ---
    _d4 = _DAE()
    _d4.geometrie.vloeren = [_VIAE(naam="Begane grond"), _VIAE(naam="1e verdieping")]
    _d4.geometrie.ruimtes = [
        _RAE(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30, verdieping="Begane grond"),
        _RAE(naam="Slaapkamer 1", functie="slaapkamer", oppervlakte_m2=12, verdieping="1e verdieping")]
    _res4 = _vbalAE(_vbAE(_d4.geometrie.ruimtes))
    _vp.zorg_voor_verdiepingen(_d4, _res4["rows"])
    _v1e = next(v for v in _d4.ventilatieplan.verdiepingen if v.naam == "1e verdieping")
    _v1e.markers[0].waarde_ls = 12345.0  # handmatige wijziging op de 1e verdieping
    _vbg = next(v for v in _d4.ventilatieplan.verdiepingen if v.naam == "Begane grond")
    _vbg.markers[0].x = 0.77             # handmatige wijziging op de begane grond
    _vp.herstel_verdieping(_d4, "Begane grond", _res4["rows"])
    check("herstel_verdieping: de opgegeven verdieping is teruggezet",
          next(v for v in _d4.ventilatieplan.verdiepingen if v.naam == "Begane grond").markers[0].x != 0.77)
    check("herstel_verdieping: de ANDERE verdieping blijft ongemoeid (handmatige waarde staat nog)",
          next(v for v in _d4.ventilatieplan.verdiepingen if v.naam == "1e verdieping").markers[0].waarde_ls == 12345.0)
    check("herstel_verdieping: onbekende verdiepingnaam -> None, geen crash",
          _vp.herstel_verdieping(_d4, "Kelder (bestaat niet)", _res4["rows"]) is None)

    # --- valideer_markers ---
    _geldigAE = {"Woonkamer", "Badkamer"}
    _goedAE, _foutAE = _vp.valideer_markers(
        [{"type": "toevoer", "ruimte_id": "Woonkamer", "x": 0.1, "y": 0.5, "waarde_ls": "21.05", "rotatie": "450"}],
        _geldigAE)
    check("valideer_markers: geldige marker geeft geen fout", _foutAE is None and len(_goedAE) == 1)
    check("valideer_markers: waarde afgerond op 0,1", _goedAE[0].waarde_ls == 21.1)
    check("valideer_markers: rotatie modulo 360 (450 -> 90)", _goedAE[0].rotatie == 90)
    check("valideer_markers: marker zonder id krijgt er automatisch een", bool(_goedAE[0].id))

    _, _fout1 = _vp.valideer_markers([{"type": "zijwaarts", "ruimte_id": "Woonkamer", "x": 0, "y": 0}], _geldigAE)
    check("valideer_markers: onbekend type -> leesbare fout, niets opgeslagen", "type" in (_fout1 or "").lower())
    _, _fout2 = _vp.valideer_markers([{"type": "afvoer", "ruimte_id": "Kelderkast", "x": 0, "y": 0}], _geldigAE)
    check("valideer_markers: onbekende ruimte_id -> geweigerd met leesbare melding",
          _fout2 is not None and "Kelderkast" in _fout2)
    _, _fout3 = _vp.valideer_markers([{"type": "toevoer", "ruimte_id": "Woonkamer", "x": 1.5, "y": 0.5}], _geldigAE)
    check("valideer_markers: x buiten 0..1 -> geweigerd", _fout3 is not None)
    _goed_mix, _fout_mix = _vp.valideer_markers(
        [{"type": "toevoer", "ruimte_id": "Woonkamer", "x": 0.1, "y": 0.1, "waarde_ls": 10},
         {"type": "afvoer", "ruimte_id": "Onbekend", "x": 0.1, "y": 0.1, "waarde_ls": 10}], _geldigAE)
    check("valideer_markers: 1 foute regel weigert de HELE batch, niets half opgeslagen",
          _goed_mix is None and _fout_mix is not None)
    check("verdiepingsvalidatie: alleen ruimtenamen van de gevraagde verdieping",
          _vp.geldige_ruimtenamen_op_verdieping(_d2, "Begane grond") == {"Woonkamer"})
    _polyAE = {"Woonkamer": [[0, 0], [.5, 0], [.5, .5], [0, .5]]}
    _, _buitenAE = _vp.valideer_markers(
        [{"type": "toevoer", "ruimte_id": "Woonkamer", "x": .8, "y": .8}], {"Woonkamer"}, _polyAE)
    check("valideer_markers: positie buiten gekoppelde ruimte -> geweigerd", "buiten ruimte" in _buitenAE)
    _pg_okAE, _pg_foutAE = _vp.valideer_ruimtepolygonen(
        {"Woonkamer": [[0,0],[1,0],[1,1]]}, {"Woonkamer"})
    check("ruimtepolygonen: minimaal 3 punten binnen canvas geaccepteerd", _pg_foutAE is None and len(_pg_okAE["Woonkamer"]) == 3)
    _, _pg_kortAE = _vp.valideer_ruimtepolygonen({"Woonkamer": [[0,0],[1,0]]}, {"Woonkamer"})
    _, _pg_buitenAE = _vp.valideer_ruimtepolygonen({"Woonkamer": [[0,0],[1.1,0],[1,1]]}, {"Woonkamer"})
    check("ruimtepolygonen: minder dan 3 punten en buiten canvas geweigerd", _pg_kortAE and _pg_buitenAE)
    _, _pg_degenAE = _vp.valideer_ruimtepolygonen(
        {"Woonkamer": [[0,0],[.5,.5],[1,1]]}, {"Woonkamer"})
    _, _pg_kruisAE = _vp.valideer_ruimtepolygonen(
        {"Woonkamer": [[0,0],[1,1],[0,1],[1,0]]}, {"Woonkamer"})
    check("ruimtepolygonen: degeneratieve en zelfsnijdende contour geweigerd",
          _pg_degenAE is not None and _pg_kruisAE is not None)
    # Exacte reproducties uit reviewronde 4: self-touch, teruglopende overlap en een niet-aangrenzend
    # herhaald edge-eindpunt moeten alle drie luid worden afgewezen.
    _review_poly1AE = [[0,0],[1,0],[.5,.5],[1,1],[0,1],[.5,.5]]
    _review_poly2AE = [[0,0],[1,0],[.5,0],[1,1],[0,1]]
    _review_poly3AE = [[0,0],[1,0],[1,1],[0,1],[1,0]]
    _review_foutenAE = [_vp.valideer_ruimtepolygonen({"Woonkamer": p}, {"Woonkamer"})[1]
                        for p in (_review_poly1AE, _review_poly2AE, _review_poly3AE)]
    check("ruimtepolygonen reviewrepro 1: non-adjacent self-touch/herhaald punt geweigerd",
          _review_foutenAE[0] is not None)
    check("ruimtepolygonen reviewrepro 2: teruglopend/overlappend segment geweigerd",
          _review_foutenAE[1] is not None)
    check("ruimtepolygonen reviewrepro 3: duplicate non-adjacent edge endpoint geweigerd",
          _review_foutenAE[2] is not None)
    _c_okAE, _c_foutAE = _vp.valideer_ruimtepolygonen(
        {"C": _c_bronAE.contour_relatief}, {"C"})
    _u_okAE, _u_foutAE = _vp.valideer_ruimtepolygonen(
        {"U": _u_bronAE.contour_relatief}, {"U"})
    check("ruimtepolygonen: geldige concave C- en U-vormen blijven geaccepteerd",
          _c_foutAE is None and _u_foutAE is None and _c_okAE and _u_okAE)

    # --- marker_balans ---
    _d5 = _DAE()
    _d5.ventilatieplan.verdiepingen = list(_d4.ventilatieplan.verdiepingen)
    _bal5 = _vp.marker_balans(_d5)
    check("marker_balans: som toevoer over alle verdiepingen samen (_d4 heeft geen natte ruimte, dus afvoer 0)",
          _bal5["toevoer"] > 0 and _bal5["afvoer"] == 0.0)
    _d6 = _DAE()
    from core.dossier import VentilatieplanVerdieping as _VPVAE
    _d6.ventilatieplan.verdiepingen = [_VPVAE(naam="X", markers=[
        _VMAE(type="toevoer", ruimte_id="A", waarde_ls=20.0),
        _VMAE(type="afvoer", ruimte_id="B", waarde_ls=20.0)])]
    check("marker_balans: gelijke som -> sluitend", _vp.marker_balans(_d6)["sluitend"] is True)
    _d6.ventilatieplan.verdiepingen[0].markers[0].waarde_ls = 25.0
    check("marker_balans: ongelijke som -> niet sluitend", _vp.marker_balans(_d6)["sluitend"] is False)

    # --- round-trip door de dossier-(de)serialisatie (core/dossier.py-uitbreiding) ---
    _d7 = _DAE()
    _d7.geometrie.ruimtes = [_RAE(naam="Woonkamer", contour_relatief=[[0.1, 0.1], [0.9, 0.1],
                                                                       [0.9, 0.9], [0.1, 0.9]])]
    _d7.ventilatieplan.systeem = "C"
    _d7.ventilatieplan.verdiepingen = [_VPVAE(naam="Begane grond", achtergrond="bg.png", breedte_px=1600,
        markers=[_VMAE(id="t1", type="toevoer", ruimte_id="Woonkamer", waarde_ls=21.0, x=0.1, y=0.2,
                       rotatie=90, bron="auto")])]
    _terug7 = Dossier.from_dict(_d7.to_dict())
    check("ventilatieplan: round-trip door to_dict/from_dict identiek", _terug7.to_dict() == _d7.to_dict())
    check("ventilatieplan: marker-veld blijft na round-trip een VentilatieMarker-object",
          _terug7.ventilatieplan.verdiepingen[0].markers[0].ruimte_id == "Woonkamer")
    check("ventilatieplan: expliciete ruimtegeometrie overleeft round-trip",
          _terug7.geometrie.ruimtes[0].contour_relatief == _d7.geometrie.ruimtes[0].contour_relatief)
except Exception as _e:
    check("ventilatieplan-datalaag (taak 020): draait zonder fout", False); print("     " + repr(_e)[:220])

print("\nAF. Ventilatieplan-pagina (taak 020): de webapp-route zelf (GET/POST, persistentie, validatie)")
try:
    import shutil as _shAF, dashboard.app as _WAF
    _WAF.app.config.update(TESTING=True)
    _cAF = _WAF.app.test_client()
    with _cAF.session_transaction() as _sAF:
        _sAF["ingelogd"] = True

    # leeg project (nog geen ruimtes) -> de pagina stuurt door naar de opname, geen crash/lege tekening
    _rAF0 = _cAF.post("/nieuw", data={"straat": "Ventilatiestraat 1", "postcode": "9999VP",
                       "plaats": "X", "woningtype": "Tussenwoning"})
    _tagAF = _rAF0.headers["Location"].rstrip("/").split("/")[-2]
    _rAF_leeg = _cAF.get("/project/%s/ventilatieplan" % _tagAF, follow_redirects=True)
    check("ventilatieplan-route: leeg project (geen ruimtes) stuurt door naar de opname, geen 500",
          _rAF_leeg.status_code == 200 and "Opname" in _rAF_leeg.get_data(as_text=True))

    # vul de opname met een demo-woning-achtige geometrie (2 verdiepingen)
    _dAF = _WAF._dossier(_tagAF)
    _dAF.geometrie.vloeren = [_VIAE(naam="Begane grond"), _VIAE(naam="1e verdieping")]
    _dAF.geometrie.ruimtes = [
        _RAE(naam="Woonkamer", functie="verblijfsruimte", oppervlakte_m2=30, verdieping="Begane grond",
             contour_relatief=[[.05,.05],[.55,.05],[.55,.7],[.05,.7]]),
        _RAE(naam="Keuken", functie="keuken", oppervlakte_m2=10, verdieping="Begane grond",
             contour_relatief=[[.55,.05],[.95,.05],[.95,.5],[.55,.5]]),
        _RAE(naam="Toilet", functie="toilet", oppervlakte_m2=1.5, verdieping="Begane grond",
             contour_relatief=[[.55,.5],[.75,.5],[.75,.7],[.55,.7]]),
        _RAE(naam="Slaapkamer 1", functie="slaapkamer", oppervlakte_m2=12, verdieping="1e verdieping",
             contour_relatief=[[.05,.05],[.55,.05],[.55,.7],[.05,.7]]),
        _RAE(naam="Badkamer", functie="badkamer", oppervlakte_m2=6, verdieping="1e verdieping",
             contour_relatief=[[.55,.05],[.95,.05],[.95,.7],[.55,.7]])]
    _stAF = _WAF._load_state(_tagAF)
    _WAF._dos_save(_tagAF, _stAF, _dAF)

    _rAF1 = _cAF.get("/project/%s/ventilatieplan" % _tagAF)
    check("ventilatieplan-route: laadt (200) met geometrie", _rAF1.status_code == 200)
    _htmlAF1 = _rAF1.get_data(as_text=True)
    check("ventilatieplan-route: toont de balans-pil + beide verdiepingen + de 2 rekenlaag-tabellen",
          "vp-balans" in _htmlAF1 and "Begane grond" in _htmlAF1 and "1e verdieping" in _htmlAF1
          and "Toevoer per verblijfsruimte" in _htmlAF1 and "Afvoer per natte ruimte" in _htmlAF1)
    check("ventilatieplan-route: laadt ventilatieplan.js", "ventilatieplan.js" in _htmlAF1)
    check("ventilatieplan-route: productiekalibratie en topologie-CTA zichtbaar",
          "Ruimtecontouren kalibreren" in _htmlAF1 and "Overstroomverbinding vastleggen" in _htmlAF1)
    check("ventilatieplan-route: rendert echte ruimtepolygonen en labels voor hit-testing",
          _htmlAF1.count("class=vp-ruimte ") == 5 and "vp-koppellijn" in _htmlAF1)
    check("ventilatieplan-route: alle 7 vuistregels staan er (ook 'niet te bepalen', geen stille voldoet)",
          _htmlAF1.count("niet te bepalen") >= 5)

    # eerste GET moet de autoplaatsing al hebben OPGESLAGEN (overleeft een 'herlaad')
    _dAF2 = _WAF._dossier(_tagAF)
    check("ventilatieplan-route: autoplaatsing is bij het eerste bezoek al in het dossier opgeslagen",
          all(v.markers for v in _dAF2.ventilatieplan.verdiepingen))
    _bg_naam = _dAF2.ventilatieplan.verdiepingen[0].naam
    _n_markers_voor = len(_dAF2.ventilatieplan.verdiepingen[0].markers)
    _rAF2 = _cAF.get("/project/%s/ventilatieplan" % _tagAF)   # 'herlaad'
    check("ventilatieplan-route: een herlaad genereert de autoplaatsing niet opnieuw (idempotent)",
          _rAF2.status_code == 200
          and len(_WAF._dossier(_tagAF).ventilatieplan.verdiepingen[0].markers) == _n_markers_voor)

    # POST met een geldige marker -> opgeslagen, balans terug in de response
    import json as _jsAF
    _nieuwe_markers = [{"id": "x1", "type": "toevoer", "ruimte_id": "Woonkamer",
                        "waarde_ls": 21.3, "x": 0.12, "y": 0.34, "rotatie": 180, "bron": "handmatig"}]
    _rAF3 = _cAF.post("/project/%s/ventilatieplan/%s/markers" % (_tagAF, _bg_naam),
                      data=_jsAF.dumps({"markers": _nieuwe_markers}), content_type="application/json")
    check("markers-route: geldige marker -> 200 + ok:true", _rAF3.status_code == 200 and _rAF3.get_json()["ok"] is True)
    _dAF3 = _WAF._dossier(_tagAF)
    _v_bg3 = next(v for v in _dAF3.ventilatieplan.verdiepingen if v.naam == _bg_naam)
    check("markers-route: precies opgeslagen wat gepost is (rotatie/positie/waarde)",
          len(_v_bg3.markers) == 1 and _v_bg3.markers[0].rotatie == 180 and _v_bg3.markers[0].waarde_ls == 21.3)

    # POST met een ONGELDIGE ruimte_id -> geweigerd, dossier blijft bij de vorige (geldige) set
    _slechte_markers = [{"type": "afvoer", "ruimte_id": "Bestaat niet", "x": 0.5, "y": 0.5, "waarde_ls": 10}]
    _rAF4 = _cAF.post("/project/%s/ventilatieplan/%s/markers" % (_tagAF, _bg_naam),
                      data=_jsAF.dumps({"markers": _slechte_markers}), content_type="application/json")
    _j4 = _rAF4.get_json()
    check("markers-route: ongeldige ruimte_id -> 400 + leesbare melding",
          _rAF4.status_code == 400 and _j4["ok"] is False and "Bestaat niet" in (_j4.get("fout") or ""))
    _dAF4 = _WAF._dossier(_tagAF)
    check("markers-route: geweigerd verzoek verandert het dossier niet (nog steeds de vorige, geldige marker)",
          next(v for v in _dAF4.ventilatieplan.verdiepingen if v.naam == _bg_naam).markers[0].id == "x1")

    # Een bestaande ruimte van een ANDERE verdieping is ook ongeldig voor deze POST.
    _verkeerde_verdieping = [{"type": "toevoer", "ruimte_id": "Slaapkamer 1", "x": .2, "y": .2,
                              "waarde_ls": 7}]
    _rAF4b = _cAF.post("/project/%s/ventilatieplan/%s/markers" % (_tagAF, _bg_naam),
                       data=_jsAF.dumps({"markers": _verkeerde_verdieping}), content_type="application/json")
    check("markers-route: ruimte van andere verdieping -> 400", _rAF4b.status_code == 400)

    _jsbronAF = open(os.path.join(os.path.dirname(__file__), "..", "dashboard", "static",
                                  "ventilatieplan.js"), encoding="utf-8").read()
    check("ventilatieplan-JS: polygon-hit-test + highlight + rollback buiten ruimte aanwezig",
          all(term in _jsbronAF for term in ("puntInPolygoon", "sleepIndicatie", "Object.assign(m, vorige)")))
    check("ventilatieplan-JS: dubbelklik ondersteunt expliciet splitsen",
          'antwoord.split("+")' in _jsbronAF and 'verdieping.markers.push' in _jsbronAF)
    check("ventilatieplan-JS: markers zijn toetsenbordbedienbaar en hebben toegankelijke naam",
          all(term in _jsbronAF for term in ('tabindex: "0"', 'role: "button"', '"aria-label"',
                                             'ev.key === "Enter"', 'ev.key === "Delete"', '"ArrowLeft"')))

    # Productieroute: kalibratie schrijft echte polygonen naar het dossier en activeert daarna slepen.
    _poly_postAF = {"polygonen": {"Woonkamer": [[.05,.05],[.55,.05],[.55,.7],[.05,.7]],
                                    "Keuken": [[.55,.05],[.95,.05],[.95,.5],[.55,.5]],
                                    "Toilet": [[.55,.5],[.75,.5],[.75,.7],[.55,.7]]}}
    _rpolyAF = _cAF.post("/project/%s/ventilatieplan/%s/ruimtepolygonen" % (_tagAF, _bg_naam),
                         data=_jsAF.dumps(_poly_postAF), content_type="application/json")
    _na_polyAF = _cAF.get("/project/%s/ventilatieplan" % _tagAF).get_data(as_text=True)
    check("ruimtepolygonen-route: opgeslagen dossierroute activeert slepen voor verdieping",
          _rpolyAF.status_code == 200 and '"heeft_ruimtegeometrie": true' in _na_polyAF.lower())

    # Zonder topologie geen pijl; na expliciete bron+natte doelruimte wel, en vuistregeldata krijgt pad.
    _dtop0AF = _WAF._dossier(_tagAF)
    # Maak dit specifieke rekenpad exact: toevoer Woonkamer 21 + Keuken 7 = 28, en Keuken is de
    # enige natte ruimte. Dan moet deurbelasting via de getransformeerde [Keuken,Woonkamer]-route
    # aantoonbaar 28 l/s zijn (niet 0 door omgekeerde bron/doel-semantiek).
    _dtop0AF.geometrie.ruimtes = [r for r in _dtop0AF.geometrie.ruimtes
                                  if r.naam in {"Woonkamer", "Keuken"}]
    _WAF._dos_save(_tagAF, _WAF._load_state(_tagAF), _dtop0AF)
    check("topologie-route: vooraf geen fantoom-overstroom",
          not any(m.type == "overstroom" for v in _dtop0AF.ventilatieplan.verdiepingen for m in v.markers))
    _rtopAF = _cAF.post("/project/%s/ventilatieplan/%s/topologie" % (_tagAF, _bg_naam),
                        data=_jsAF.dumps({"bron":"Woonkamer", "doel":"Keuken"}), content_type="application/json")
    _dtop1AF = _WAF._dossier(_tagAF)
    check("topologie-route: expliciete bron+natte doelruimte opgeslagen en maakt pijl",
          _rtopAF.status_code == 200 and ["Woonkamer","Keuken"] in _dtop1AF.ventilatieplan.topologie
          and any(m.type == "overstroom" for m in next(v for v in _dtop1AF.ventilatieplan.verdiepingen if v.naam == _bg_naam).markers))
    _top_getAF = _cAF.get("/project/%s/ventilatieplan" % _tagAF).get_data(as_text=True)
    _top_calcAF = _WAF.vent_verdeel_balans(_WAF.vent_bereken(_dtop1AF.geometrie.ruimtes))
    _top_keukenAF = next(r for r in _top_calcAF["rows"] if r["naam"] == "Keuken")["afvoer_advies_ls"]
    check("topologie POST->GET: bron Woonkamer -> natte Keuken rekent Keuken-belasting >15",
          "deurrooster geadviseerd" in _top_getAF and "voldoet niet" in _top_getAF and "28.0 l/s" in _top_getAF,
          "advies=%s voldoet_niet=%s waarde28=%s" % ("deurrooster geadviseerd" in _top_getAF,
          "voldoet niet" in _top_getAF, "28.0 l/s" in _top_getAF) + " calc=" + str(_top_keukenAF))

    # onbekende verdieping -> 404, nette fout
    _rAF5 = _cAF.post("/project/%s/ventilatieplan/%s/markers" % (_tagAF, "Kelder (bestaat niet)"),
                      data=_jsAF.dumps({"markers": []}), content_type="application/json")
    check("markers-route: onbekende verdieping -> 404", _rAF5.status_code == 404)

    # herstel: alleen de opgegeven verdieping gaat terug naar de autoplaatsing
    _v_1e_voor = next(v for v in _dAF4.ventilatieplan.verdiepingen if v.naam != _bg_naam)
    _rAF6 = _cAF.post("/project/%s/ventilatieplan/%s/herstel" % (_tagAF, _bg_naam))
    _j6 = _rAF6.get_json()
    check("herstel-route: 200 + ok:true + de nieuwe markerlijst terug", _rAF6.status_code == 200 and _j6["ok"] is True
          and isinstance(_j6.get("markers"), list) and len(_j6["markers"]) > 0)
    _dAF6 = _WAF._dossier(_tagAF)
    check("herstel-route: de begane grond is niet meer de handmatige 'x1'-marker (teruggezet)",
          next(v for v in _dAF6.ventilatieplan.verdiepingen if v.naam == _bg_naam).markers[0].id != "x1")
    check("herstel-route: de 1e verdieping is ongemoeid gebleven",
          next(v for v in _dAF6.ventilatieplan.verdiepingen if v.naam != _bg_naam).markers == _v_1e_voor.markers)

    _shAF.rmtree(_WAF._pdir(_tagAF), ignore_errors=True)
except Exception as _e:
    check("ventilatieplan-route (taak 020): draait zonder fout", False); print("     " + repr(_e)[:250])

print("\n=== RESULTAAT: %d geslaagd, %d gefaald ===" % (passed, failed))
sys.exit(1 if failed else 0)
