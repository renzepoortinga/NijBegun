"""
Vult het officiele Nij Begun isolatieplan-Word-template vanuit een canoniek dossier.
Exporteert fill(dossier, template_path, out_path) zodat de orchestrator (run.py) het kan aanroepen.
"""
import os, sys, argparse, copy, re
import docx
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.dossier import load_json  # noqa: E402

def eur(x):
    return "" if x is None else "EUR " + ("%.2f" % x).replace(".", ",")
def num_nl(x):
    return "" if x is None else str(x).replace(".", ",")

def replace_in_paragraph(p, mapping):
    full = "".join(run.text for run in p.runs); new = full
    for k, v in mapping.items():
        if k in new: new = new.replace(k, v)
    if new != full:
        for run in p.runs: run.text = ""
        if p.runs: p.runs[0].text = new
        else: p.add_run(new)

def replace_everywhere(doc, mapping):
    for p in doc.paragraphs: replace_in_paragraph(p, mapping)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: replace_in_paragraph(p, mapping)

def set_cell(cell, text): cell.text = "" if text is None else str(text)
def find_table(doc, pred):
    for t in doc.tables:
        if pred(t): return t
    return None
def first_cell(row): return row.cells[0].text.strip()

def insert_row_after(row):
    new_tr = copy.deepcopy(row._tr); row._tr.addnext(new_tr)
    from docx.table import _Row
    return _Row(new_tr, row._parent)

def fill_kv_table(table, kv):
    for row in table.rows:
        cells = row.cells
        for li, vi in ((0, 1), (2, 3)):
            if len(cells) > vi:
                label = cells[li].text.strip().lower(); cur = cells[vi].text.strip()
                fillable = (cur == "" or "keuze maken" in cur.lower() or cur == "U/B/L")
                for key, val in kv.items():
                    if key in label and fillable:
                        set_cell(cells[vi], val); break

def onderdeel_letter(text):
    m = re.search(r'Onderdeel\s+([A-E])', text); return m.group(1) if m else None

def fill_maatregelen(table, per_letter):
    rows = list(table.rows); i = 0
    while i < len(rows):
        letter = onderdeel_letter(first_cell(rows[i]))
        if letter and letter in per_letter:
            ms = per_letter[letter]; target_idx = i + 1
            if target_idx < len(rows):
                anchor = rows[target_idx]
                for j, m in enumerate(ms):
                    r = anchor if j == 0 else insert_row_after(anchor)
                    if j > 0: anchor = r
                    cells = r.cells
                    vals = [m["omschrijving"], m["code"], m["rc_u_doel"], m["materiaal"],
                            num_nl(m["oppervlakte_m2"]) + " m2", eur(m["kosten"])]
                    for ci in range(min(len(cells), len(vals))): set_cell(cells[ci], vals[ci])
                rows = list(table.rows)
        i += 1

def fill_warmteverlies(table, ber, renovatiejaar):
    for row in table.rows:
        lbl = first_cell(row).lower()
        if len(row.cells) < 2: continue
        if "standaard voor woning" in lbl: set_cell(row.cells[1], num_nl(ber.standaard_eis_kwh_m2))
        elif "maatregelen zijn toegepast" in lbl: set_cell(row.cells[1], num_nl(ber.kwh_m2_na_maatregelen))
        elif "renovatiejaar" in lbl: set_cell(row.cells[1], str(renovatiejaar) if renovatiejaar else "Nee")


def _summary(schils, is_glas, rc_threshold):
    if not schils:
        return ("", "")
    rcs = [s.rc_huidig for s in schils if s.rc_huidig is not None]
    us = [s.u_huidig for s in schils if s.u_huidig is not None]
    statuses = set(s.isolatie_aanwezig for s in schils)
    if (statuses & {"Wel", "Ja"}) and not (statuses & {"Niet", "Nee", "Gedeeltelijk", "Onbekend", ""}):
        niveau = "Wel geisoleerd"
    elif rcs:
        if min(rcs) >= rc_threshold: niveau = "Wel geisoleerd"
        elif max(rcs) >= rc_threshold: niveau = "Gedeeltelijk geisoleerd"
        else: niveau = "Niet/onbekend geisoleerd"
    else:
        niveau = "Niet/onbekend"
    if is_glas and us:
        waarde = ("U %.1f W/m2K" % us[0]) if len(set(us)) == 1 else ("U %.1f-%.1f W/m2K" % (min(us), max(us)))
        glas = sorted({s.glastype for s in schils if s.glastype})
        if glas: waarde += " (" + ", ".join(glas) + ")"
    elif rcs:
        waarde = ("Rc %.2f m2.K/W" % rcs[0]) if len(set(rcs)) == 1 else ("Rc %.2f-%.2f m2.K/W" % (min(rcs), max(rcs)))
    else:
        waarde = ""
    return (niveau, waarde)

def _pick(schils, sub_any=None, sub_none=None, glas_any=None, glas_none=None, hellend=None):
    """Filter schildelen op subtype-/glastype-trefwoorden en (voor daken) hellend vs plat."""
    out = []
    for s in schils:
        sub = (s.subtype or "").lower()
        gl = (s.glastype or "").lower()
        if sub_any and not any(k in sub for k in sub_any):
            continue
        if sub_none and any(k in sub for k in sub_none):
            continue
        if glas_any and not any(k in gl for k in glas_any):
            continue
        if glas_none and any(k in gl for k in glas_none):
            continue
        if hellend is True and not ((s.hellingshoek or 0) > 0):
            continue
        if hellend is False and ((s.hellingshoek or 0) > 0):
            continue
        out.append(s)
    return out


def _staat_subsets(dos):
    """Schildeel-subsets per template-regel van sectie 3 (Huidige woningstaat V1-V6)."""
    S = dos.schil
    gevels = [s for s in S if s.type == "gevel"]
    vloeren = [s for s in S if s.type == "vloer"]
    daken = [s for s in S if s.type == "dak"]
    kozijnen = [s for s in S if s.type == "kozijn"]
    panelen = [s for s in S if s.type == "paneel"] + _pick(kozijnen, sub_any=("paneel",))
    _niet_raam = ("deur", "paneel", "dakraam", "tuimel", "vierpans")
    return {
        "gevel": gevels,
        "raam": _pick(kozijnen, sub_none=_niet_raam, glas_none=("voorzet", "achterzet", "vacu")),
        "voorzet": _pick(kozijnen, glas_any=("voorzet", "achterzet")),
        "vacuum": _pick(kozijnen, glas_any=("vacu",)),
        "deur": _pick(kozijnen, sub_any=("deur",)),
        "paneel": panelen,
        "kozijn_alle": kozijnen,
        "vloer_bg": _pick(vloeren, sub_none=("zolder", "vliering")),
        "vloer_zolder": _pick(vloeren, sub_any=("zolder", "vliering")),
        "dak_hellend": _pick(daken, hellend=True),
        "dak_plat": _pick(daken, hellend=False),
        "dakraam": _pick(kozijnen, sub_any=("dakraam", "tuimel")),
    }


def huidige_staat_gaten(dos):
    """Template-regels uit sectie 3 die de tool NIET kan vullen omdat de OPNAME de gegevens mist ->
    de adviseur vult ze bewust zelf in. Vul je de opnamevelden (gevel-/dak-isolatiezijde,
    bodemisolatie, kierdichting) wél in, dan verdwijnen deze meldingen vanzelf."""
    sub = _staat_subsets(dos)
    o = dos.opname
    gaten = []
    if not (o.gevel_isolatie_zijde or "").strip():
        gaten.append("V1 Gevelisolatie: aan welke zijde (spouw / binnenzijde / buitenzijde)? "
                     "-> MagicPlan-veld 'Gevel - isolatie aan zijde'")
    if (sub["dak_hellend"] or sub["dak_plat"]) and not (o.dak_isolatie_zijde or "").strip():
        gaten.append("V4 Dakisolatie: binnen- of buitenzijde? -> MagicPlan-veld 'Dak - isolatie aan zijde'")
    if not (o.bodemisolatie or "").strip():
        gaten.append("V3 Bodemisolatie (kruipruimte) -> MagicPlan-veld 'Bodemisolatie kruipruimte'")
    if not (o.kierdichting or "").strip() and not (o.qv10_gemeten or o.qv10_waarde):
        gaten.append("V6 Kierdichting: geen qv10 en geen omschrijving -> MagicPlan-veld 'Kierdichting'")
    gaten.append("V4 Vierpansraam in dakvlak — niet in de opname; vul zelf in indien aanwezig")
    return gaten


def fill_huidige_staat(doc, dos):
    """Vult sectie 3 'Huidige woningstaat' (V1-V6). Elke regel krijgt (isolatieniveau, isolatiewaarde)
    uit de bijbehorende schildelen. Regels waarvoor de opname geen gegevens heeft blijven leeg — het
    template staat dat expliciet toe — en komen terug via huidige_staat_gaten()."""
    sub = _staat_subsets(dos)
    vent_letter = ""
    syss = (dos.ventilatie.systeem or "")
    for L in ("A", "B", "C", "D"):
        if syss.strip().upper().startswith(L) or (L + ":") in syss:
            vent_letter = L; break

    def _fillable(txt):
        t = txt.strip().lower()
        return t == "" or t.startswith("wel/niet") or "in het geval van" in t or "gedeeltelijk ge" in t

    def put(row, niveau, waarde):
        c = row.cells
        if len(c) > 1 and _fillable(c[1].text): set_cell(c[1], niveau)
        if len(c) > 2 and _fillable(c[2].text): set_cell(c[2], waarde)

    def put_set(row, key, is_glas, drempel):
        s = sub.get(key) or []
        if s:
            n, w = _summary(s, is_glas, drempel)
            put(row, n, w)

    o = dos.opname
    zijde_g = (o.gevel_isolatie_zijde or "").strip().lower()      # spouw | binnenzijde | buitenzijde
    zijde_d = (o.dak_isolatie_zijde or "").strip().lower()        # binnenzijde | buitenzijde
    for t in doc.tables:
        for row in t.rows:
            lbl = first_cell(row).lower()
            # --- V1 Gevel: de gevelstaat gaat naar de regel die bij de isolatiezijde hoort ---
            if "spouwmuur isolatie" in lbl:
                if "binnenzijde" not in zijde_g and "buitenzijde" not in zijde_g:
                    put_set(row, "gevel", False, 2.5)     # leeg/spouw -> spouwmuur-regel
            elif "gevel isolatie" in lbl and "binnenzijde" in lbl:
                if "binnenzijde" in zijde_g:
                    put_set(row, "gevel", False, 2.5)
            elif "gevel isolatie" in lbl and "buitenzijde" in lbl:
                if "buitenzijde" in zijde_g:
                    put_set(row, "gevel", False, 2.5)
            # --- V2 Glas en kozijnen ---
            elif "beglazing (enkel" in lbl or (lbl.startswith("beglazing") and "zet" not in lbl):
                put_set(row, "raam", True, 1.6)
            elif "zet beglazing" in lbl or "voor- of achterzet" in lbl:
                put_set(row, "voorzet", True, 1.6)
            elif lbl.startswith("deuren"):
                put_set(row, "deur", True, 1.6)
            elif "panelen in kozijn" in lbl:
                put_set(row, "paneel", False, 2.5)
            elif lbl.startswith("kozijnen"):
                mats = sorted({s.kozijnmateriaal for s in sub["kozijn_alle"] if s.kozijnmateriaal})
                if mats:
                    put(row, ", ".join(mats), "")
            elif "vacu" in lbl:
                put_set(row, "vacuum", True, 1.6)
            # --- V3 Vloeren ---
            elif "onderzijde begane grond" in lbl:
                put_set(row, "vloer_bg", False, 2.5)
            elif "zolder" in lbl or "vliering" in lbl:
                put_set(row, "vloer_zolder", False, 2.5)
            elif "bodemisolatie" in lbl:
                bi = (o.bodemisolatie or "").strip()
                if bi and len(row.cells) > 1 and _fillable(row.cells[1].text):
                    set_cell(row.cells[1], bi)
            # --- V4 Daken: hellend/plat naar de regel die bij de isolatiezijde hoort ---
            elif "hellend dak" in lbl:
                rij_buiten, zij_buiten = ("buitenzijde" in lbl), ("buitenzijde" in zijde_d)
                if (zijde_d and zij_buiten == rij_buiten) or (not zijde_d and not rij_buiten):
                    put_set(row, "dak_hellend", False, 3.5)   # leeg -> binnenzijde-regel
            elif "plat dak" in lbl:
                if zijde_d and ("buitenzijde" in zijde_d) == ("buitenzijde" in lbl):
                    put_set(row, "dak_plat", False, 3.5)
            elif "tuimelvenster" in lbl or "dakraam" in lbl:
                put_set(row, "dakraam", True, 1.6)
            # --- V5 Ventilatie / V6 Kierdichting / warmteverlies ---
            elif "huidig warmteverlies in de" in lbl:
                if len(row.cells) > 1 and row.cells[1].text.strip() == "":
                    set_cell(row.cells[1], num_nl(dos.berekening.kwh_m2_huidig))
            elif "kierdichting" in lbl:
                qv = o.qv10_waarde
                txt = (o.kierdichting or "").strip() or (
                    ("qv10 = %s dm3/s.m2 (%s)" % (num_nl(qv), "gemeten" if o.qv10_gemeten else "forfaitair"))
                    if qv else "")
                if txt and len(row.cells) > 1 and _fillable(row.cells[1].text):
                    set_cell(row.cells[1], txt)
            elif vent_letter and lbl.startswith("type " + vent_letter.lower()):
                if len(row.cells) > 1 and row.cells[1].text.strip() == "":
                    set_cell(row.cells[1], "Aanwezig (huidige situatie): " + syss)


def _clone_table_after(table):
    """Kloon een prijsopbouw-blok en plaats het (met lege tussenalinea) erna."""
    from docx.table import Table
    from docx.oxml.ns import qn
    new_tbl = copy.deepcopy(table._tbl)
    table._tbl.addnext(new_tbl)
    new_tbl.addprevious(new_tbl.makeelement(qn("w:p"), {}))  # anders voegt Word tabellen samen
    return Table(new_tbl, table._parent)

def _cat_data_rows(table):
    """{1: rij, 2: rij, 3: rij, 'subtotaal': rij} voor een prijsopbouw-blok."""
    rows = list(table.rows); out = {}
    for ri, row in enumerate(rows):
        c1 = row.cells[1].text.lower() if len(row.cells) > 1 else ""
        for cat in (1, 2, 3):
            if ("categorie %d" % cat) in c1 and ri + 1 < len(rows):
                out[cat] = rows[ri + 1]
        if "subtotaal" in first_cell(row).lower():
            out["subtotaal"] = row
    return out

def _fill_prijsregel(row, code, oms, prijs, eenheid, hoeveelheid, kosten):
    hoev = (num_nl(round(hoeveelheid, 2)) + " " + (eenheid or "m2")) if hoeveelheid else (eenheid or "")
    vals = [code, (oms or "")[:40], eur(prijs), hoev, eur(kosten)]
    for ci in range(min(len(row.cells), len(vals))):
        set_cell(row.cells[ci], vals[ci])

def fill_prijsopbouw(doc, dos):
    blocks = [t for t in doc.tables if t.rows and first_cell(t.rows[0]).lower().startswith("maatregel:")]
    ms = dos.maatregelen
    while blocks and len(blocks) < len(ms):     # blok per maatregel (kloon de rest)
        blocks.append(_clone_table_after(blocks[-1]))
    for bi, t in enumerate(blocks):
        if bi >= len(ms):                        # ongebruikt blok leegmaken
            set_cell(t.rows[0].cells[0], "Maatregel: -"); continue
        m = ms[bi]
        set_cell(t.rows[0].cells[0], "Maatregel: " + (m.omschrijving or m.code)[:60])
        cr = _cat_data_rows(t)
        if 1 in cr:
            _fill_prijsregel(cr[1], m.code, m.omschrijving, m.prijs_per_eenheid,
                             m.eenheid, m.oppervlakte_m2, m.kosten)
        subtot = m.kosten or 0
        for cat in (2, 3):
            sp = next((s for s in m.subposten if s.categorie == cat), None)
            if sp and cat in cr:
                _fill_prijsregel(cr[cat], sp.code, sp.omschrijving, sp.prijs_per_eenheid,
                                 sp.eenheid, sp.hoeveelheid, sp.kosten)
                subtot += sp.kosten or 0
        if "subtotaal" in cr and len(cr["subtotaal"].cells) > 1:
            set_cell(cr["subtotaal"].cells[-1], eur(round(subtot, 2)))

def fill(dos, template_path, out_path):
    """Vul het template vanuit dossier-object. Retourneer (aantal_maatregelen, totaal)."""
    doc = docx.Document(template_path)
    idf, adv, opn, ber = dos.identificatie, dos.adviseur, dos.opname, dos.berekening
    woning = idf.woningtype or ""
    replace_everywhere(doc, {
        "< Straatnaam >": (idf.straat + " " + idf.huisnummer).strip(),
        "< Postcode >": idf.postcode, "<Plaatsnaam>": idf.plaats,
        "< voor en achternaam isolatieadviseur >": adv.naam,
        "< bedrijfsnaam invullen >": adv.bedrijf,
        "<telefoonnummer waarop isolatieadviseur bereikbaar is>": adv.telefoon,
        "< Keuze maken >": woning})
    t_geg = find_table(doc, lambda t: t.rows and "1. Gegevens" in t.rows[0].cells[0].text)
    if t_geg:
        fill_kv_table(t_geg, {
            "adres": (idf.straat + " " + idf.huisnummer).strip(),
            "postcode": (idf.postcode + " " + idf.plaats).strip(),
            "bouwjaar": str(idf.bouwjaar or ""), "type woning": woning,
            "type advies": opn.type_advies, "opnamedatum": opn.opnamedatum,
            "rapportagedatum": opn.rapportagedatum, "organisatie": adv.bedrijf,
            "gebruikte isolat": "Vabi EPA-W (NTA 8800)", "naam adviseur": adv.naam})
    per_letter = {}
    for m in dos.maatregelen:
        per_letter.setdefault((m.onderdeel or " ")[0], []).append({
            "omschrijving": m.omschrijving, "code": m.code, "rc_u_doel": m.rc_u_doel,
            "materiaal": m.materiaal, "oppervlakte_m2": m.oppervlakte_m2, "kosten": m.kosten})
    for t in doc.tables:
        if any(onderdeel_letter(first_cell(r)) for r in t.rows): fill_maatregelen(t, per_letter)
    totaal = sum((m.kosten or 0) for m in dos.maatregelen)
    for t in doc.tables:
        for row in t.rows:
            if "totale kosten" in first_cell(row).lower() and len(row.cells) > 1:
                set_cell(row.cells[-1], eur(totaal))
    t_wv = find_table(doc, lambda t: t.rows and "Warmteverlies" in t.rows[0].cells[0].text)
    if t_wv: fill_warmteverlies(t_wv, ber, idf.renovatiejaar)
    fill_huidige_staat(doc, dos)
    fill_prijsopbouw(doc, dos)
    voeg_bijlagen_4_7_toe(doc, dos)
    os.makedirs(os.path.dirname(out_path), exist_ok=True); doc.save(out_path)
    return sum(len(v) for v in per_letter.values()), totaal


# Standaardtekst van bijlage 4, letterlijk naar de strekking van de goedgekeurde voorbeeldplannen
# (Voorbeeldplannen/Voorbeeld isolatieplan bouwjaar 1970.pdf).
_BIJLAGE4 = [
    ("Over het isolatieplan",
     "Voor de berekeningen in dit plan gaan we uit van een standaard energieverbruik en standaard "
     "bewonersgedrag. Het werkelijke energiegebruik kan daarvan afwijken. Voor de investeringskosten "
     "gebruiken we de Groninger Maatregelen Catalogus; waar bedragen ontbreken hanteren we de "
     "kostenkengetallen van kostenkengetallen.rvo.nl. Voor de energietarieven volgen we Milieu "
     "Centraal (per kwartaal bijgewerkt). Kosten en besparingen kunnen dus afwijken van de "
     "werkelijkheid. Alle bedragen zijn inclusief 21 % btw."),
    ("Isoleren en koudebruggen",
     "Door te isoleren kunnen (grotere) koudebruggen ontstaan. Een koudebrug is een plek in de "
     "buitenschil waar warmte makkelijker ontsnapt dan via de omliggende constructie: de isolatie is "
     "daar onderbroken, dunner of afwezig, of materialen met een hoge warmtegeleiding (beton, staal) "
     "voeren de warmte af. De oppervlaktetemperatuur aan de binnenzijde daalt, waardoor condensatie "
     "kan optreden met schimmel en vochtplekken tot gevolg. Dat tast het binnenklimaat aan en kan op "
     "termijn de constructie beschadigen. In bijlage 7 staan de aansluitdetails die bij de "
     "voorgestelde maatregelen aandacht vragen."),
]


def voeg_bijlagen_4_7_toe(doc, dos):
    """Bijlage 4 t/m 7 ACHTER de bestaande inhoud zetten (de goedgekeurde voorbeeldplannen hebben
    bijlage 1 t/m 7; het officiële template bevat 1 t/m 3). Puur additief — de M29-lay-out van het
    template zelf blijft ongemoeid (Beoordelingsformulier-eis)."""
    from engine.psi_lookup import relevante_details

    def kop(tekst):
        p = doc.add_paragraph()
        r = p.add_run(tekst)
        r.bold = True
        return p

    doc.add_page_break()
    kop("Bijlage 4: Informatie over dit isolatieplan")
    for titel, tekst in _BIJLAGE4:
        kop(titel)
        doc.add_paragraph(tekst)

    doc.add_page_break()
    kop("Bijlage 5: Voorgestelde maatregelen in beeld")
    doc.add_paragraph(
        "Bij het verduurzamen van een bestaande woning ontstaan risico's op koudebruggen. Hieronder "
        "staat per bouwdeel welke maatregel is voorgesteld en waar volgens de adviseur aandacht nodig "
        "is. Voeg hier de schets van de woning toe waarop de maatregelen en de potentiële "
        "koudebruggen zijn aangegeven (te vervangen ramen blauw gearceerd).")
    for m in (getattr(dos, "maatregelen", None) or []):
        doc.add_paragraph("%s — %s%s" % (m.onderdeel or "", m.omschrijving or "",
                                         (" (%s)" % m.code) if m.code else ""), style=None)
    if not (getattr(dos, "maatregelen", None) or []):
        doc.add_paragraph("(nog geen maatregelen gekozen)")

    doc.add_page_break()
    kop("Bijlage 6: Ventilatieplan")
    doc.add_paragraph(
        "Het ventilatieplan (berekening en plattegrond) is als aparte bijlage bijgevoegd. Toevoer "
        "0,7 dm³/s per m² verblijfsgebied, afvoer keuken 21 / bad 14 / toilet 7 dm³/s, toe- en afvoer "
        "in balans met overstroom tussen de ruimten.")

    doc.add_page_break()
    kop("Bijlage 7: Detailtekeningen")
    details = relevante_details(dos)
    doc.add_paragraph(
        "De onderstaande aansluitdetails komen in deze woning voor en vragen bij uitvoering aandacht. "
        "De nummering volgt de ISSO-Referentiedetails; de ψ-waarden zijn de forfaitaire waarden uit "
        "NTA 8800 bijlage I (kolom A = aansluiting volgens de aanvullende voorwaarden uitgevoerd, "
        "kolom B = niet). Voeg per detail de bijbehorende detailtekening toe.")
    for d in details:
        doc.add_paragraph(
            "Detail %d — %s | ψ = %.2f W/(m·K) bij goede uitvoering, %.2f W/(m·K) zonder | "
            "voorwaarde: %s" % (d["nr"], d["omschrijving"], d["psi_a"], d["psi_b"], d["voorwaarde"]))
    if not details:
        doc.add_paragraph("(geen details afgeleid uit de opname — vul ze handmatig aan)")

def main():
    here = os.path.dirname(os.path.abspath(__file__)); root = os.path.dirname(here)
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--dossier", default=os.path.join(root, "sample_dossier_priced.json"))
    ap.add_argument("--out", default=os.path.join(root, "out", "isolatieplan.docx"))
    a = ap.parse_args()
    dos = load_json(a.dossier)
    n, totaal = fill(dos, a.template, a.out)
    print("OK: " + a.out)
    print("  maatregelen ingevuld: %d | totale kosten: %s" % (n, eur(totaal)))

if __name__ == "__main__":
    main()
